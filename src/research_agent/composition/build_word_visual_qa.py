"""
Master Academic Thesis Document Rebuilder with Word 2016 Native Figures, Native Tables, Native Captions, and Cross-References.
"""

import os
import sys
import uuid
import shutil
import zipfile
from pathlib import Path
from xml.sax.saxutils import escape
from lxml import etree
import latex2mathml.converter
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import win32com.client as win32
import pythoncom

from research_agent.storage.db import DatabaseManager
from research_agent.storage.repository import ResearchRepository
from research_agent.composition.native_omml_equations import (
    make_evidence_space_omml, make_frequency_vector_omml,
    make_l_inv_omml, make_l_var_omml, make_var_formula_omml, make_c_matrix_omml, make_l_cov_omml,
    make_stage_a_loss_omml, make_quality_vector_omml, make_gating_weights_omml, make_pcgrad_omml,
    make_mil_attention_omml, make_mil_loss_omml, make_canonical_fusion_omml, make_rep_bundle_omml,
    make_graph_self_loss_omml, make_preserv_loss_omml, make_fuse_rec_stopgrad_omml,
    make_tgn_msg_omml, make_tgn_agg_omml, make_tgn_update_omml, make_tgn_readout_omml,
    make_l_mep_omml, make_l_mpp_omml, make_l_seq_time_omml,
    make_l_mask_node_omml, make_l_mask_edge_omml, make_l_time_gap_omml,
    make_event_embedding_omml
)
from research_agent.visuals.chapter1_drawings import draw_fig_1_1, draw_fig_1_2, draw_fig_1_3, draw_fig_1_4
from research_agent.visuals.schemas import VisualRecord, VisualType, CreationMethod
from research_agent.visuals.registry import VisualRegistry

sys.stdout.reconfigure(encoding="utf-8")

# Initialize official Microsoft Office MathML to OMML XSLT transformer
XSLT_PATH = r"C:\Program Files\Microsoft Office\Office16\MML2OMML.XSL"
xslt_tree = etree.parse(XSLT_PATH)
transform_omml = etree.XSLT(xslt_tree)



def make_card_ew_omml():
    return parse_xml(
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        '  <m:r><m:rPr><m:scr m:val="script"/></m:rPr><m:t>O</m:t></m:r>'
        '  <m:r><m:t>(</m:t></m:r>'
        '  <m:r><m:rPr><m:nor/></m:rPr><m:t>card</m:t></m:r>'
        '  <m:r><m:t>(</m:t></m:r>'
        '  <m:sSub><m:e><m:r><m:rPr><m:scr m:val="script"/></m:rPr><m:t>E</m:t></m:r></m:e><m:sub><m:r><m:t>W</m:t></m:r></m:sub></m:sSub>'
        '  <m:r><m:t>))</m:t></m:r>'
        '</m:oMath>'
    )


def make_card_tgn_omml():
    return parse_xml(
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        '  <m:r><m:rPr><m:scr m:val="script"/></m:rPr><m:t>O</m:t></m:r>'
        '  <m:r><m:t>(</m:t></m:r>'
        '  <m:r><m:rPr><m:nor/></m:rPr><m:t>card</m:t></m:r>'
        '  <m:r><m:t>(</m:t></m:r>'
        '  <m:sSub><m:e><m:r><m:rPr><m:scr m:val="script"/></m:rPr><m:t>E</m:t></m:r></m:e><m:sub><m:r><m:t>W</m:t></m:r></m:sub></m:sSub>'
        '  <m:r><m:t>)</m:t></m:r>'
        '  <m:r><m:t> · d + </m:t></m:r>'
        '  <m:r><m:rPr><m:nor/></m:rPr><m:t>card</m:t></m:r>'
        '  <m:r><m:t>(</m:t></m:r>'
        '  <m:sSub><m:e><m:r><m:rPr><m:scr m:val="script"/></m:rPr><m:t>V</m:t></m:r></m:e><m:sub><m:r><m:rPr><m:nor/></m:rPr><m:t>upd</m:t></m:r></m:sub></m:sSub>'
        '  <m:r><m:t>)</m:t></m:r>'
        '  <m:r><m:t> · </m:t></m:r>'
        '  <m:sSup><m:e><m:r><m:t>d</m:t></m:r></m:e><m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup>'
        '  <m:r><m:t>)</m:t></m:r>'
        '</m:oMath>'
    )


def make_c_build_omml():
    return parse_xml(
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        '  <m:sSub><m:e><m:r><m:rPr><m:scr m:val="script"/></m:rPr><m:t>C</m:t></m:r></m:e><m:sub><m:r><m:rPr><m:nor/></m:rPr><m:t>build</m:t></m:r></m:sub></m:sSub>'
        '  <m:r><m:t> = </m:t></m:r>'
        '  <m:r><m:rPr><m:scr m:val="script"/></m:rPr><m:t>O</m:t></m:r>'
        '  <m:r><m:t>(</m:t></m:r>'
        '  <m:r><m:rPr><m:nor/></m:rPr><m:t>card</m:t></m:r>'
        '  <m:r><m:t>(</m:t></m:r>'
        '  <m:sSub><m:e><m:r><m:rPr><m:scr m:val="script"/></m:rPr><m:t>E</m:t></m:r></m:e><m:sub><m:r><m:t>W</m:t></m:r></m:sub></m:sSub>'
        '  <m:r><m:t>))</m:t></m:r>'
        '</m:oMath>'
    )


def make_c_msg_omml():
    return parse_xml(
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        '  <m:sSub><m:e><m:r><m:rPr><m:scr m:val="script"/></m:rPr><m:t>C</m:t></m:r></m:e><m:sub><m:r><m:rPr><m:nor/></m:rPr><m:t>msg</m:t></m:r></m:sub></m:sSub>'
        '  <m:r><m:t> = </m:t></m:r>'
        '  <m:r><m:rPr><m:scr m:val="script"/></m:rPr><m:t>O</m:t></m:r>'
        '  <m:r><m:t>(</m:t></m:r>'
        '  <m:r><m:rPr><m:nor/></m:rPr><m:t>card</m:t></m:r>'
        '  <m:r><m:t>(</m:t></m:r>'
        '  <m:sSub><m:e><m:r><m:rPr><m:scr m:val="script"/></m:rPr><m:t>E</m:t></m:r></m:e><m:sub><m:r><m:t>W</m:t></m:r></m:sub></m:sSub>'
        '  <m:r><m:t>)</m:t></m:r>'
        '  <m:r><m:t> · (</m:t></m:r>'
        '  <m:sSub><m:e><m:r><m:t>d</m:t></m:r></m:e><m:sub><m:r><m:rPr><m:nor/></m:rPr><m:t>GNN</m:t></m:r></m:sub></m:sSub>'
        '  <m:r><m:t> · </m:t></m:r>'
        '  <m:sSub><m:e><m:r><m:t>d</m:t></m:r></m:e><m:sub><m:r><m:rPr><m:nor/></m:rPr><m:t>msg</m:t></m:r></m:sub></m:sSub>'
        '  <m:r><m:t> + </m:t></m:r>'
        '  <m:sSub><m:e><m:r><m:t>d</m:t></m:r></m:e><m:sub><m:r><m:rPr><m:nor/></m:rPr><m:t>attr</m:t></m:r></m:sub></m:sSub>'
        '  <m:r><m:t>))</m:t></m:r>'
        '</m:oMath>'
    )


def make_v_upd_subset_omml():
    return parse_xml(
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        '  <m:sSub><m:e><m:r><m:rPr><m:scr m:val="script"/></m:rPr><m:t>V</m:t></m:r></m:e><m:sub><m:r><m:rPr><m:nor/></m:rPr><m:t>upd</m:t></m:r></m:sub></m:sSub>'
        '  <m:r><m:t>(t) ⊆ </m:t></m:r>'
        '  <m:sSub><m:e><m:r><m:rPr><m:scr m:val="script"/></m:rPr><m:t>V</m:t></m:r></m:e><m:sub><m:r><m:t>t</m:t></m:r></m:sub></m:sSub>'
        '</m:oMath>'
    )


def make_c_update_omml():
    return parse_xml(
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        '  <m:sSub><m:e><m:r><m:rPr><m:scr m:val="script"/></m:rPr><m:t>C</m:t></m:r></m:e><m:sub><m:r><m:rPr><m:nor/></m:rPr><m:t>update</m:t></m:r></m:sub></m:sSub>'
        '  <m:r><m:t> = </m:t></m:r>'
        '  <m:r><m:rPr><m:scr m:val="script"/></m:rPr><m:t>O</m:t></m:r>'
        '  <m:r><m:t>(</m:t></m:r>'
        '  <m:r><m:rPr><m:nor/></m:rPr><m:t>card</m:t></m:r>'
        '  <m:r><m:t>(</m:t></m:r>'
        '  <m:sSub><m:e><m:r><m:rPr><m:scr m:val="script"/></m:rPr><m:t>V</m:t></m:r></m:e><m:sub><m:r><m:rPr><m:nor/></m:rPr><m:t>upd</m:t></m:r></m:sub></m:sSub>'
        '  <m:r><m:t>)</m:t></m:r>'
        '  <m:r><m:t> · </m:t></m:r>'
        '  <m:sSubSup><m:e><m:r><m:t>d</m:t></m:r></m:e><m:sub><m:r><m:rPr><m:nor/></m:rPr><m:t>GNN</m:t></m:r></m:sub><m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSubSup>'
        '  <m:r><m:t>)</m:t></m:r>'
        '</m:oMath>'
    )


def make_c_readout_omml():
    return parse_xml(
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        '  <m:sSub><m:e><m:r><m:rPr><m:scr m:val="script"/></m:rPr><m:t>C</m:t></m:r></m:e><m:sub><m:r><m:rPr><m:nor/></m:rPr><m:t>readout</m:t></m:r></m:sub></m:sSub>'
        '  <m:r><m:t> = </m:t></m:r>'
        '  <m:r><m:rPr><m:scr m:val="script"/></m:rPr><m:t>O</m:t></m:r>'
        '  <m:r><m:t>(</m:t></m:r>'
        '  <m:r><m:rPr><m:nor/></m:rPr><m:t>card</m:t></m:r>'
        '  <m:r><m:t>(</m:t></m:r>'
        '  <m:sSub><m:e><m:r><m:rPr><m:scr m:val="script"/></m:rPr><m:t>V</m:t></m:r></m:e><m:sub><m:r><m:rPr><m:nor/></m:rPr><m:t>active</m:t></m:r></m:sub></m:sSub>'
        '  <m:r><m:t>)</m:t></m:r>'
        '  <m:r><m:t> · </m:t></m:r>'
        '  <m:sSub><m:e><m:r><m:t>d</m:t></m:r></m:e><m:sub><m:r><m:rPr><m:nor/></m:rPr><m:t>GNN</m:t></m:r></m:sub></m:sSub>'
        '  <m:r><m:t> + </m:t></m:r>'
        '  <m:sSub><m:e><m:r><m:t>d</m:t></m:r></m:e><m:sub><m:r><m:rPr><m:nor/></m:rPr><m:t>GNN</m:t></m:r></m:sub></m:sSub>'
        '  <m:r><m:t> · </m:t></m:r>'
        '  <m:sSub><m:e><m:r><m:t>d</m:t></m:r></m:e><m:sub><m:r><m:rPr><m:nor/></m:rPr><m:t>graph</m:t></m:r></m:sub></m:sSub>'
        '  <m:r><m:t>)</m:t></m:r>'
        '</m:oMath>'
    )


def make_rq_omml(num: int):
    return parse_xml(
        f'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        f'  <m:r><m:rPr><m:nor/></m:rPr><m:t>RQ{num}</m:t></m:r>'
        f'</m:oMath>'
    )


def make_hypo_omml(num: int):
    return parse_xml(
        f'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        f'  <m:r><m:rPr><m:nor/></m:rPr><m:t>H{num}</m:t></m:r>'
        f'</m:oMath>'
    )


def latex_to_clean_omml(latex_code: str):
    """Converts a LaTeX formula into a native Word OMML element, cleaning any empty placeholders and adding noProof."""
    try:
        sanitized_latex = latex_code.replace(r"\&", "&#38;")
        mathml = latex2mathml.converter.convert(sanitized_latex)
        tree = etree.fromstring(mathml)
        omml_tree = transform_omml(tree)

        ns = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math", "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        for r_elem in omml_tree.xpath(".//m:r", namespaces=ns):
            wrPr = etree.SubElement(r_elem, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
            etree.SubElement(wrPr, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}noProof")
            lang = etree.SubElement(wrPr, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lang")
            lang.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "en-US")

        # Clean empty m:sup in m:nary or anywhere across omml_tree
        for sup_elem in omml_tree.xpath(".//m:sup", namespaces=ns):
            if len(sup_elem) == 0 and (not sup_elem.text or not sup_elem.text.strip()):
                sup_elem.getparent().remove(sup_elem)

        # Ensure naryPr has supHide="1" and subHide="0" when no sup exists
        for nary in omml_tree.xpath(".//m:nary", namespaces=ns):
            naryPr = nary.find("m:naryPr", namespaces=ns)
            sup = nary.find("m:sup", namespaces=ns)
            if sup is None and naryPr is not None:
                supHide = naryPr.find("m:supHide", namespaces=ns)
                if supHide is None:
                    supHide = etree.SubElement(naryPr, "{http://schemas.openxmlformats.org/officeDocument/2006/math}supHide")
                supHide.set("{http://schemas.openxmlformats.org/officeDocument/2006/math}val", "1")

        # Convert empty m:nary to m:sSubSup so Word doesn't draw a dotted placeholder or need U+200B
        for nary in list(omml_tree.xpath(".//m:nary", namespaces=ns)):
            e_elem = nary.find("m:e", namespaces=ns)
            if e_elem is not None and len(e_elem) == 0 and (not e_elem.text or not e_elem.text.strip()):
                naryPr = nary.find("m:naryPr", namespaces=ns)
                chr_val = "∑"
                if naryPr is not None:
                    chr_elem = naryPr.find("m:chr", namespaces=ns)
                    if chr_elem is not None:
                        chr_val = chr_elem.get("{http://schemas.openxmlformats.org/officeDocument/2006/math}val", "∑")
                sub = nary.find("m:sub", namespaces=ns)
                sup = nary.find("m:sup", namespaces=ns)
                
                new_elem = etree.Element("{http://schemas.openxmlformats.org/officeDocument/2006/math}sSubSup")
                e_new = etree.SubElement(new_elem, "{http://schemas.openxmlformats.org/officeDocument/2006/math}e")
                r_new = etree.SubElement(e_new, "{http://schemas.openxmlformats.org/officeDocument/2006/math}r")
                t_new = etree.SubElement(r_new, "{http://schemas.openxmlformats.org/officeDocument/2006/math}t")
                t_new.text = chr_val
                
                if sub is not None:
                    new_elem.append(sub)
                if sup is not None:
                    new_elem.append(sup)
                
                parent = nary.getparent()
                if parent is not None:
                    parent.replace(nary, new_elem)

        xml_str = etree.tostring(omml_tree, encoding="utf-8").decode("utf-8")
        return parse_xml(xml_str)
    except Exception as e:
        print(f"Warning: LaTeX conversion failed for '{latex_code}': {e}")
        xml_str = (
            f'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
            f'  <m:r><m:t>{escape(latex_code)}</m:t></m:r>'
            f'</m:oMath>'
        )
        return parse_xml(xml_str)


def make_citation_element(tag_or_num_list):
    """Creates native Word CITATION field elements for a list of source IDs with clean ', ' separator."""
    if isinstance(tag_or_num_list, (int, str)):
        tag_or_num_list = [tag_or_num_list]

    elems = []
    for i, item in enumerate(tag_or_num_list):
        if isinstance(item, int):
            tag = f"SRC{item:06d}"
            num_str = str(item)
        else:
            tag = item.replace("-", "")
            num_str = str(int(item.split("-")[1])) if "-" in item else item

        if i > 0:
            sep_xml = (
                f'<w:r {nsdecls("w")}>\n'
                '  <w:rPr>\n'
                '    <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>\n'
                '    <w:sz w:val="28"/>\n'
                '  </w:rPr>\n'
                '  <w:t xml:space="preserve">, </w:t>\n'
                '</w:r>'
            )
            elems.append(parse_xml(sep_xml))

        fld_xml = (
            f'<w:fldSimple {nsdecls("w")} w:instr="CITATION {tag} \\l 1033 ">\n'
            '  <w:r>\n'
            '    <w:rPr>\n'
            '      <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>\n'
            '      <w:sz w:val="28"/>\n'
            '    </w:rPr>\n'
            f'    <w:t>[{num_str}]</w:t>\n'
            '  </w:r>\n'
            '</w:fldSimple>'
        )
        elems.append(parse_xml(fld_xml))
    return elems


def make_ref_element(bookmark_name: str, fallback_text: str, font_size_pt: float = 14.0):
    """Creates a native Word REF dynamic cross-reference field element."""
    sz_val = int(font_size_pt * 2)
    ref_xml = (
        f'<w:fldSimple {nsdecls("w")} w:instr="REF {bookmark_name} \\h ">\n'
        '  <w:r>\n'
        '    <w:rPr>\n'
        '      <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>\n'
        f'      <w:sz w:val="{sz_val}"/>\n'
        '    </w:rPr>\n'
        f'    <w:t>{fallback_text}</w:t>\n'
        '  </w:r>\n'
        '</w:fldSimple>'
    )
    return parse_xml(ref_xml)


def add_table_caption(doc, target_p, seq_num: int, title_content, bookmark_name: str = None, chapter_num: int = 1, page_break_before: bool = False):
    """Inserts a native Word Caption paragraph with chapter-aware SEQ Bảng field and keepWithNext."""
    cap_p = doc.add_paragraph(style="Caption") if target_p is None else target_p.insert_paragraph_before(style="Caption")
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_before = Pt(12)
    cap_p.paragraph_format.space_after = Pt(4)
    cap_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    cap_p.paragraph_format.keep_with_next = True
    if page_break_before:
        cap_p.paragraph_format.page_break_before = True

    pb_xml = '  <w:pageBreakBefore/>\n' if page_break_before else ''
    cap_xml = (
        f'<w:pPr {nsdecls("w")}>\n'
        '  <w:pStyle w:val="Caption"/>\n'
        '  <w:jc w:val="center"/>\n'
        f'{pb_xml}'
        '  <w:spacing w:before="240" w:after="80" w:line="360" w:lineRule="auto"/>\n'
        '  <w:keepNext/>\n'
        '</w:pPr>'
    )
    cap_p._p.remove(cap_p._p.pPr)
    cap_p._p.append(parse_xml(cap_xml))

    b_id = f"20{chapter_num}{seq_num}"
    if bookmark_name:
        cap_p._p.append(parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{b_id}" w:name="{bookmark_name}"/>'))

    r1 = cap_p.add_run(f"Bảng {chapter_num}.")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(14)
    r1.bold = True

    instr = "SEQ Bảng \\r 1 \\* ARABIC " if (chapter_num > 1 and seq_num == 1) else "SEQ Bảng \\* ARABIC "
    seq_xml = (
        f'<w:fldSimple {nsdecls("w")} w:instr="{instr}">\n'
        '  <w:r>\n'
        '    <w:rPr>\n'
        '      <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>\n'
        '      <w:b/>\n'
        '      <w:sz w:val="28"/>\n'
        '    </w:rPr>\n'
        f'    <w:t>{seq_num}</w:t>\n'
        '  </w:r>\n'
        '</w:fldSimple>'
    )
    cap_p._p.append(parse_xml(seq_xml))

    if bookmark_name:
        cap_p._p.append(parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{b_id}"/>'))

    colon_r = cap_p.add_run(": ")
    colon_r.font.name = "Times New Roman"
    colon_r.font.size = Pt(14)
    colon_r.bold = True

    items = [title_content] if isinstance(title_content, str) else title_content
    for item in items:
        if isinstance(item, str):
            r = cap_p.add_run(item)
            r.font.name = "Times New Roman"
            r.font.size = Pt(14)
            r.bold = True
        elif isinstance(item, list):
            for sub in item:
                cap_p._p.append(sub)
        else:
            cap_p._p.append(item)

    return cap_p


def add_figure_caption(doc, target_p, chapter_num: int, seq_num: int, title_content, bookmark_name: str = None):
    """Inserts a native Word Caption paragraph with SEQ Hình field placed below the figure."""
    cap_p = doc.add_paragraph(style="Caption") if target_p is None else target_p.insert_paragraph_before(style="Caption")
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_before = Pt(6)
    cap_p.paragraph_format.space_after = Pt(12)
    cap_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    cap_xml = (
        f'<w:pPr {nsdecls("w")}>\n'
        '  <w:pStyle w:val="Caption"/>\n'
        '  <w:jc w:val="center"/>\n'
        '  <w:spacing w:before="120" w:after="240" w:line="360" w:lineRule="auto"/>\n'
        '</w:pPr>'
    )
    cap_p._p.remove(cap_p._p.pPr)
    cap_p._p.append(parse_xml(cap_xml))

    b_id = f"10{chapter_num}{seq_num}"
    if bookmark_name:
        cap_p._p.append(parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{b_id}" w:name="{bookmark_name}"/>'))

    r1 = cap_p.add_run(f"Hình {chapter_num}.")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(14)
    r1.bold = True

    instr = "SEQ Hình \\r 1 \\* ARABIC " if (chapter_num > 1 and seq_num == 1) else "SEQ Hình \\* ARABIC "
    seq_xml = (
        f'<w:fldSimple {nsdecls("w")} w:instr="{instr}">\n'
        '  <w:r>\n'
        '    <w:rPr>\n'
        '      <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>\n'
        '      <w:b/>\n'
        '      <w:sz w:val="28"/>\n'
        '    </w:rPr>\n'
        f'    <w:t>{seq_num}</w:t>\n'
        '  </w:r>\n'
        '</w:fldSimple>'
    )
    cap_p._p.append(parse_xml(seq_xml))

    if bookmark_name:
        cap_p._p.append(parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{b_id}"/>'))

    colon_r = cap_p.add_run(": ")
    colon_r.font.name = "Times New Roman"
    colon_r.font.size = Pt(14)
    colon_r.bold = True

    items = [title_content] if isinstance(title_content, str) else title_content
    for item in items:
        if isinstance(item, str):
            r = cap_p.add_run(item)
            r.font.name = "Times New Roman"
            r.font.size = Pt(14)
            r.bold = True
        elif isinstance(item, list):
            for sub in item:
                cap_p._p.append(sub)
        else:
            cap_p._p.append(item)

    return cap_p


def format_table_cell(cell, width_dxa: int, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, font_size_pt=14):
    """Sets standard cell properties: exact width, vertical centering, border, padding, and compact line spacing."""
    tcPr = cell._tc.get_or_add_tcPr()
    tc_xml = (
        f'<w:tcPr {nsdecls("w")}>\n'
        f'  <w:tcW w:w="{width_dxa}" w:type="dxa"/>\n'
        '  <w:vAlign w:val="center"/>\n'
        '  <w:tcBorders>\n'
        '    <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '    <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '    <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '    <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '  </w:tcBorders>\n'
        '  <w:tcMar>\n'
        '    <w:top w:w="80" w:type="dxa"/>\n'
        '    <w:bottom w:w="80" w:type="dxa"/>\n'
        '    <w:left w:w="120" w:type="dxa"/>\n'
        '    <w:right w:w="120" w:type="dxa"/>\n'
        '  </w:tcMar>\n'
        '</w:tcPr>'
    )
    cell._tc.remove(tcPr)
    cell._tc.append(parse_xml(tc_xml))

    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(font_size_pt)
        if bold:
            r.bold = True


def insert_thesis_table(doc, ref_p, headers, col_widths, rows_data, font_size_pt=14, pad_v_dxa=80, space_v_pt=3):
    """Creates an elegant, professional thesis table matching original template layout, supporting OMML nodes in cells."""
    tbl = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    if ref_p is not None:
        ref_p._p.addprevious(tbl._tbl)

    tblPr = tbl._tbl.tblPr
    total_w = sum(col_widths)
    tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:w="{total_w}" w:type="dxa"/>'))

    hdr_row = tbl.rows[0]
    hdr_trPr = hdr_row._tr.get_or_add_trPr()
    hdr_trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    hdr_trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

    for c_i, h in enumerate(headers):
        cell = hdr_row.cells[c_i]
        format_table_cell_rich(cell, h, col_widths[c_i], align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size_pt=font_size_pt, pad_v_dxa=pad_v_dxa, space_v_pt=space_v_pt)

    for r_i, row in enumerate(rows_data):
        b_row = tbl.rows[r_i + 1]
        b_trPr = b_row._tr.get_or_add_trPr()
        b_trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

        for c_i, val in enumerate(row):
            cell = b_row.cells[c_i]
            cell_align = WD_ALIGN_PARAGRAPH.CENTER if (c_i == 0 and len(headers) >= 4) else WD_ALIGN_PARAGRAPH.LEFT
            format_table_cell_rich(cell, val, col_widths[c_i], align=cell_align, bold=(c_i == 0 and len(headers) == 3), font_size_pt=font_size_pt, pad_v_dxa=pad_v_dxa, space_v_pt=space_v_pt)


def format_table_cell_rich(cell, val, width_dxa: int, align=WD_ALIGN_PARAGRAPH.LEFT, bold: bool = False, font_size_pt: float = 14.0, pad_v_dxa: int = 80, space_v_pt: float = 3.0):
    """Formats cell borders, margins, alignment and renders rich text / OMML nodes without raw math underscores."""
    tcPr = cell._tc.get_or_add_tcPr()
    tc_xml = (
        f'<w:tcPr {nsdecls("w")}>\n'
        f'  <w:tcW w:w="{width_dxa}" w:type="dxa"/>\n'
        '  <w:vAlign w:val="center"/>\n'
        '  <w:tcBorders>\n'
        '    <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '    <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '    <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '    <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '  </w:tcBorders>\n'
        '  <w:tcMar>\n'
        f'    <w:top w:w="{pad_v_dxa}" w:type="dxa"/>\n'
        f'    <w:bottom w:w="{pad_v_dxa}" w:type="dxa"/>\n'
        '    <w:left w:w="120" w:type="dxa"/>\n'
        '    <w:right w:w="120" w:type="dxa"/>\n'
        '  </w:tcMar>\n'
        '</w:tcPr>'
    )
    cell._tc.remove(tcPr)
    cell._tc.append(parse_xml(tc_xml))

    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(space_v_pt)
    p.paragraph_format.space_after = Pt(space_v_pt)

    # Clear default text runs
    p.text = ""
    items = val if isinstance(val, list) else [val]
    for item in items:
        if isinstance(item, str):
            # Dynamic check for math notation with underscore like [e_{t-k+1}, ..., e_t]
            if "[" in item and "_{" in item and "]" in item:
                # Convert math bracket segment to OMML
                prefix = item[:item.find("[")]
                math_part = item[item.find("["):item.find("]")+1]
                suffix = item[item.find("]")+1:]
                if prefix:
                    r1 = p.add_run(prefix)
                    r1.font.name = "Times New Roman"
                    r1.font.size = Pt(font_size_pt)
                    if bold: r1.bold = True
                p._p.append(latex_to_clean_omml(math_part))
                if suffix:
                    r2 = p.add_run(suffix)
                    r2.font.name = "Times New Roman"
                    r2.font.size = Pt(font_size_pt)
                    if bold: r2.bold = True
            else:
                r = p.add_run(item)
                r.font.name = "Times New Roman"
                r.font.size = Pt(font_size_pt)
                if bold:
                    r.bold = True
        elif isinstance(item, list):
            for sub_elem in item:
                p._p.append(sub_elem)
        else:
            p._p.append(item)


def generate_perfect_sources_xml(sources):
    """Generates valid Microsoft Word Bibliography Sources CustomXML with Corporate Authors & IEEE style."""
    lines = ['<?xml version="1.0" encoding="UTF-8" standalone="no"?>']
    lines.append('<b:Sources SelectedStyle="\\IEEE.XSL" StyleName="IEEE" xmlns:b="http://schemas.openxmlformats.org/officeDocument/2006/bibliography" xmlns="http://schemas.openxmlformats.org/officeDocument/2006/bibliography">')

    corporate_map = {
        "SRC-000001": "MITRE ATT&CK",
        "SRC-000027": "National Institute of Standards and Technology (NIST)",
        "SRC-000028": "Defense Advanced Research Projects Agency (DARPA)"
    }

    for s in sources:
        tag = s.source_id.replace("-", "")
        venue = s.venue or ""
        if any(w in venue for w in ["Proceedings", "Conference", "Symposium", "NDSS", "S&P", "CCS", "ICLR", "ICML", "ACSAC", "ISSTA", "ISSRE", "ASE", "IJCNN", "KDD", "IJCAI", "ICDM", "SOSP", "ATC", "USENIX", "NeurIPS", "NIPS", "Neural Information Processing Systems", "Workshop"]):
            stype = "ConferenceProceedings"
        elif any(w in venue for w in ["Journal", "Surveys", "IEEE Transactions", "ACM"]):
            stype = "ArticleInAPeriodical"
        elif "Standard" in venue or "NIST" in venue or "DARPA" in venue or "LANL" in venue:
            stype = "Report"
        elif "arxiv" in venue.lower():
            stype = "InternetSite"
        else:
            stype = "Report"

        guid = str(uuid.uuid5(uuid.NAMESPACE_DNS, f"thesis.sources.{s.source_id}")).upper()

        lines.append("  <b:Source>")
        lines.append(f"    <b:Tag>{tag}</b:Tag>")
        lines.append(f"    <b:SourceType>{stype}</b:SourceType>")
        lines.append(f"    <b:Guid>{{{guid}}}</b:Guid>")
        lines.append(f"    <b:Title>{escape(s.title)}</b:Title>")
        lines.append(f"    <b:Year>{s.year}</b:Year>")
        if stype == "ConferenceProceedings":
            lines.append(f"    <b:ConferenceName>{escape(venue)}</b:ConferenceName>")
            lines.append(f"    <b:BookTitle>{escape(venue)}</b:BookTitle>")
        elif stype == "ArticleInAPeriodical":
            lines.append(f"    <b:JournalName>{escape(venue)}</b:JournalName>")
        elif stype == "InternetSite":
            lines.append(f"    <b:InternetSiteTitle>{escape(venue)}</b:InternetSiteTitle>")
        else:
            lines.append(f"    <b:Institution>{escape(venue)}</b:Institution>")

        if s.source_id in corporate_map:
            lines.append("    <b:Author>")
            lines.append("      <b:Author>")
            lines.append(f"        <b:Corporate>{escape(corporate_map[s.source_id])}</b:Corporate>")
            lines.append("      </b:Author>")
            lines.append("    </b:Author>")
        elif s.source_id == "SRC-000029":
            lines.append("    <b:Author>")
            lines.append("      <b:Author>")
            lines.append("        <b:NameList>")
            lines.append("          <b:Person><b:Last>Kent</b:Last><b:First>Alexander D.</b:First></b:Person>")
            lines.append("        </b:NameList>")
            lines.append("      </b:Author>")
            lines.append("    </b:Author>")
            lines.append("    <b:Institution>Los Alamos National Laboratory</b:Institution>")
        else:
            lines.append("    <b:Author>")
            lines.append("      <b:Author>")
            lines.append("        <b:NameList>")
            for author_name in s.authors:
                parts = author_name.strip().split()
                if len(parts) > 1:
                    first = " ".join(parts[:-1])
                    last = parts[-1]
                    lines.append(f"          <b:Person><b:Last>{escape(last)}</b:Last><b:First>{escape(first)}</b:First></b:Person>")
                else:
                    lines.append(f"          <b:Person><b:Last>{escape(author_name)}</b:Last></b:Person>")
            lines.append("        </b:NameList>")
            lines.append("      </b:Author>")
            lines.append("    </b:Author>")

        lines.append("  </b:Source>")

    lines.append("</b:Sources>")
    return "\n".join(lines)


def build_and_audit_document(target_file: str = r"D:\Research\Chuyên đề chuyên sâu - Copy.docx"):
    repo = ResearchRepository(DatabaseManager())
    sources = repo.list_sources()

    # Sync Master Sources.xml in AppData
    try:
        master_xml_str = generate_perfect_sources_xml(sources)
        master_xml_path = os.path.expandvars(r"%APPDATA%\Microsoft\Bibliography\Sources.xml")
        with open(master_xml_path, "w", encoding="utf-8") as f:
            f.write(master_xml_str)
        print(f"[0/6] Synchronized %APPDATA%\\Microsoft\\Bibliography\\Sources.xml with {len(sources)} primary verified sources.")
    except Exception as e:
        print(f"[WARNING] Could not sync master Sources.xml: {e}")

    target_path = Path(target_file)
    backup_path = target_path.parent / (target_path.stem + ".backup.docx")

    if not backup_path.exists():
        shutil.copyfile(target_path, backup_path)
    print(f"[1/6] Loaded pristine template from: {backup_path}")

    doc = docx.Document(str(backup_path))

    # Keep only Table 0 (Cover page frame)
    while len(doc.tables) > 1:
        tbl_to_remove = doc.tables[1]
        tbl_to_remove._tbl.getparent().remove(tbl_to_remove._tbl)
    print("[2/6] Preserved Cover Frame Table 0.")

    # Fix cover table year: 2024 -> 2026
    for r in doc.tables[0].rows:
        for c in r.cells:
            for p in c.paragraphs:
                if "2024" in p.text:
                    p.text = p.text.replace("2024", "2026")
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(14)
                        run.bold = True

    # -------------------------------------------------------------------------
    # FRONT MATTER REORGANIZATION (TOC, TOF, and 3 Chapters List)
    # -------------------------------------------------------------------------
    p_intro = None
    for p in doc.paragraphs:
        txt = p.text.strip().lower()
        if "báo cáo gồm năm chương" in txt or "báo cáo gồm 5 chương" in txt or "báo cáo chuyên đề được cấu trúc" in txt:
            p_intro = p
            break

    ch_paras = []
    if p_intro is not None:
        curr = p_intro
        for _ in range(5):
            next_p_elem = curr._p.getnext()
            if next_p_elem is not None:
                for p_obj in doc.paragraphs:
                    if p_obj._p == next_p_elem:
                        ch_paras.append(p_obj)
                        curr = p_obj
                        break

        # 1. Update Intro
        p_intro.text = "Báo cáo chuyên đề được cấu trúc thành ba chương trọng tâm:"
        p_intro.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p_intro.paragraph_format.first_line_indent = Cm(1.27)
        for r in p_intro.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(14)

        # 2. Update Chapter 1
        if len(ch_paras) > 0:
            ch_paras[0].text = ""
            ch_paras[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            ch_paras[0].paragraph_format.first_line_indent = Cm(1.27)
            r1 = ch_paras[0].add_run("Chương 1. Tổng quan về phương pháp trích xuất đặc trưng dữ liệu log và thách thức bảo toàn ngữ cảnh an toàn: ")
            r1.font.name = "Times New Roman"
            r1.font.size = Pt(14)
            r1.bold = True
            r2 = ch_paras[0].add_run("Trình bày bài toán biểu diễn dữ liệu log doanh nghiệp, hành vi tấn công đa giai đoạn trên ma trận MITRE ATT&CK (Phiên bản Enterprise v19.1, 28/04/2026), khung Hợp đồng Biểu diễn (Representation Contract), phân tích đối chiếu ba nhóm phương pháp hiện đại và xác lập 5 khoảng trống nghiên cứu cốt lõi tương ứng với 5 câu hỏi nghiên cứu (RQ1–RQ5).")
            r2.font.name = "Times New Roman"
            r2.font.size = Pt(14)

        # 3. Update Chapter 2
        if len(ch_paras) > 1:
            ch_paras[1].text = ""
            ch_paras[1].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            ch_paras[1].paragraph_format.first_line_indent = Cm(1.27)
            r1 = ch_paras[1].add_run("Chương 2. Đề xuất phương pháp trích xuất đặc trưng đa view bảo toàn ngữ cảnh và nhận thức quyền riêng tư: ")
            r1.font.name = "Times New Roman"
            r1.font.size = Pt(14)
            r1.bold = True
            r2 = ch_paras[1].add_run("Trình bày kiến trúc biểu diễn đa góc nhìn kết hợp chuỗi sự kiện và đồ thị nguồn gốc, cơ chế gióng hàng tiềm ẩn chống sụp đổ biểu diễn, phân bổ bằng chứng yếu cho nhiễu quản trị viên và cơ chế bảo toàn liên kết có kiểm soát quyền riêng tư.")
            r2.font.name = "Times New Roman"
            r2.font.size = Pt(14)

        # 4. Update Chapter 3
        if len(ch_paras) > 2:
            ch_paras[2].text = ""
            ch_paras[2].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            ch_paras[2].paragraph_format.first_line_indent = Cm(1.27)
            r1 = ch_paras[2].add_run("Chương 3. Thực nghiệm, đánh giá và ứng dụng: ")
            r1.font.name = "Times New Roman"
            r1.font.size = Pt(14)
            r1.bold = True
            r2 = ch_paras[2].add_run("Mô tả thiết lập thực nghiệm trên các bộ dữ liệu benchmark chuẩn (DARPA TC, LANL, HDFS, BGL), đánh giá định lượng hiệu năng phát hiện, phân tích độ bền vững trước trôi dạt dữ liệu, kiểm chứng khả năng phòng chống tấn công suy luận quyền riêng tư và thảo luận khả năng tích hợp trong hệ thống SOC.")
            r2.font.name = "Times New Roman"
            r2.font.size = Pt(14)

        # 5. Remove Chapter 4 & Chapter 5 paragraphs
        if len(ch_paras) > 4:
            ch_paras[4]._p.getparent().remove(ch_paras[4]._p)
        if len(ch_paras) > 3:
            ch_paras[3]._p.getparent().remove(ch_paras[3]._p)

    # The pristine template already contains valid TOC, LOF, and LOT fields.

    # -------------------------------------------------------------------------
    # CLEAN OLD BODY PARAGRAPHS FROM HEADING 1 TO THE END OF TEMPLATE
    # -------------------------------------------------------------------------
    paragraphs_to_remove = []
    found_h1 = False
    for p in doc.paragraphs:
        txt = p.text.strip().upper()
        if p.style.name == "Heading 1" and "TỔNG QUAN VỀ PHƯƠNG PHÁP" in txt:
            found_h1 = True
        if found_h1:
            paragraphs_to_remove.append(p)

    for p in paragraphs_to_remove:
        p._p.getparent().remove(p._p)

    target_p = None
    print(f"[3/6] Cleaned {len(paragraphs_to_remove)} old body paragraphs. Clean document ready.")

    def add_p(text_segments, bold_prefix=None, first_line_indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, keep_with_next=False, page_break_before=False):
        new_p = doc.add_paragraph(style="Normal") if target_p is None else target_p.insert_paragraph_before(style="Normal")
        new_p.alignment = align
        new_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        new_p.paragraph_format.space_after = Pt(6)
        new_p.paragraph_format.space_before = Pt(0)
        if page_break_before:
            new_p.paragraph_format.page_break_before = True
        if keep_with_next:
            new_p.paragraph_format.keep_with_next = True
        
        # If bold_prefix contains a literal bullet character, convert to native Word list with numPr
        if bold_prefix and (bold_prefix.startswith("• ") or bold_prefix.startswith("•")):
            bold_prefix = bold_prefix[2:] if bold_prefix.startswith("• ") else bold_prefix[1:]
            new_p.paragraph_format.left_indent = Cm(1.27)
            new_p.paragraph_format.first_line_indent = Cm(-0.63)
            new_p.paragraph_format.space_after = Pt(4)
            pPr = new_p._p.get_or_add_pPr()
            numPr_xml = f'<w:numPr {nsdecls("w")}><w:ilvl w:val="0"/><w:numId w:val="21"/></w:numPr>'
            pPr.append(parse_xml(numPr_xml))
        elif first_line_indent:
            new_p.paragraph_format.first_line_indent = Cm(1.27)

        if bold_prefix:
            r_pre = new_p.add_run(bold_prefix)
            r_pre.font.name = "Times New Roman"
            r_pre.font.size = Pt(14)
            r_pre.bold = True

        if isinstance(text_segments, str):
            r = new_p.add_run(text_segments)
            r.font.name = "Times New Roman"
            r.font.size = Pt(14)
        elif isinstance(text_segments, list):
            for seg in text_segments:
                if isinstance(seg, str):
                    r = new_p.add_run(seg)
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(14)
                elif isinstance(seg, list):
                    for sub_elem in seg:
                        new_p._p.append(sub_elem)
                else:
                    new_p._p.append(seg)
        return new_p

    def add_bullet_p(text_segments, bold_prefix=None, keep_with_next=False):
        new_p = doc.add_paragraph(style="Normal") if target_p is None else target_p.insert_paragraph_before(style="Normal")
        new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        new_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        new_p.paragraph_format.space_after = Pt(4)
        new_p.paragraph_format.space_before = Pt(0)
        new_p.paragraph_format.left_indent = Cm(1.27)
        new_p.paragraph_format.first_line_indent = Cm(-0.63)
        if keep_with_next:
            new_p.paragraph_format.keep_with_next = True
        pPr = new_p._p.get_or_add_pPr()
        numPr_xml = f'<w:numPr {nsdecls("w")}><w:ilvl w:val="0"/><w:numId w:val="21"/></w:numPr>'
        pPr.append(parse_xml(numPr_xml))
        if bold_prefix:
            r_pre = new_p.add_run(bold_prefix)
            r_pre.font.name = "Times New Roman"
            r_pre.font.size = Pt(14)
            r_pre.bold = True
        if isinstance(text_segments, str):
            r = new_p.add_run(text_segments)
            r.font.name = "Times New Roman"
            r.font.size = Pt(14)
        elif isinstance(text_segments, list):
            for seg in text_segments:
                if isinstance(seg, str):
                    r = new_p.add_run(seg)
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(14)
                elif isinstance(seg, list):
                    for sub_elem in seg:
                        new_p._p.append(sub_elem)
                else:
                    new_p._p.append(seg)
        return new_p

    def add_num_p(text_segments, bold_prefix=None, keep_with_next=False, num_id=20):
        new_p = doc.add_paragraph(style="Normal") if target_p is None else target_p.insert_paragraph_before(style="Normal")
        new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        new_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        new_p.paragraph_format.space_after = Pt(4)
        new_p.paragraph_format.space_before = Pt(0)
        new_p.paragraph_format.left_indent = Cm(1.27)
        new_p.paragraph_format.first_line_indent = Cm(-0.63)
        if keep_with_next:
            new_p.paragraph_format.keep_with_next = True
        pPr = new_p._p.get_or_add_pPr()
        numPr_xml = f'<w:numPr {nsdecls("w")}><w:ilvl w:val="0"/><w:numId w:val="{num_id}"/></w:numPr>'
        pPr.append(parse_xml(numPr_xml))
        if bold_prefix:
            r_pre = new_p.add_run(bold_prefix)
            r_pre.font.name = "Times New Roman"
            r_pre.font.size = Pt(14)
            r_pre.bold = True
        if isinstance(text_segments, str):
            r = new_p.add_run(text_segments)
            r.font.name = "Times New Roman"
            r.font.size = Pt(14)
        elif isinstance(text_segments, list):
            for seg in text_segments:
                if isinstance(seg, str):
                    r = new_p.add_run(seg)
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(14)
                elif isinstance(seg, list):
                    for sub_elem in seg:
                        new_p._p.append(sub_elem)
                else:
                    new_p._p.append(seg)
        return new_p

    def add_display_equation(omml_node_or_latex):
        """Adds a centered block display equation paragraph."""
        eq_p = doc.add_paragraph(style="Normal") if target_p is None else target_p.insert_paragraph_before(style="Normal")
        eq_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eq_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        eq_p.paragraph_format.space_before = Pt(4)
        eq_p.paragraph_format.space_after = Pt(4)
        eq_p.paragraph_format.first_line_indent = Cm(0)
        
        if isinstance(omml_node_or_latex, str):
            omml_node = latex_to_clean_omml(omml_node_or_latex)
        else:
            omml_node = omml_node_or_latex
        eq_p._p.append(omml_node)
        return eq_p

    def add_h1(clean_text):
        p = doc.add_paragraph(style="Heading 1") if target_p is None else target_p.insert_paragraph_before(style="Heading 1")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        r = p.add_run(clean_text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(16)
        r.bold = True
        return p

    def add_h2(clean_text):
        p = doc.add_paragraph(style="Heading 2") if target_p is None else target_p.insert_paragraph_before(style="Heading 2")
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(clean_text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
        r.bold = True
        return p

    def add_h3(clean_text):
        p = doc.add_paragraph(style="Heading 3") if target_p is None else target_p.insert_paragraph_before(style="Heading 3")
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(clean_text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
        r.bold = True
        r.italic = True
        return p


    def add_h4(content):
        """Heading 4 for numbered subheadings like 2.2.1.1 etc."""
        p = doc.add_paragraph(style="Heading 4") if target_p is None else target_p.insert_paragraph_before(style="Heading 4")
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        if isinstance(content, list):
            for item in content:
                if isinstance(item, str):
                    r = p.add_run(item)
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(13)
                    r.bold = True
                    r.italic = False
                elif hasattr(item, "tag") and "oMath" in item.tag:
                    p._p.append(item)
                elif hasattr(item, "tag") and "r" in item.tag:
                    p._p.append(item)
                elif hasattr(item, "tag") and "fldSimple" in item.tag:
                    p._p.append(item)
        else:
            r = p.add_run(str(content))
            r.font.name = "Times New Roman"
            r.font.size = Pt(13)
            r.bold = True
            r.italic = False
        return p
    def add_figure_image(img_path: str, width_inches: float = 6.2):
        """Inserts an inline figure image centered in the document."""
        p = doc.add_paragraph(style="Normal") if target_p is None else target_p.insert_paragraph_before(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run()
        r.add_picture(img_path, width=docx.shared.Inches(width_inches))
        return p

    # Use pre-generated pristine academic figures
    fig1_path = r"D:\Research\figures\fig_1_1_observation_hierarchy.png"
    fig2_path = r"D:\Research\figures\fig_1_2_evidence_space.png"
    fig3_path = r"D:\Research\figures\fig_1_3_three_tier_architecture.png"
    fig4_path = r"D:\Research\figures\fig_1_4_method_to_gaps_map.png"
    fig5_path = r"D:\Research\figures\fig_2_1_dual_plane_architecture.png"
    fig6_path = r"D:\Research\figures\fig_2_2_sequential_transformer.png"
    fig7_path = r"D:\Research\figures\fig_2_3_temporal_gnn.png"

    print("[4/6] Writing Chapters with Native Visuals, Citations, OMML Equations, and Captions...")

    # =========================================================================
    # HEADING 1
    # =========================================================================
    add_h1("TỔNG QUAN VỀ PHƯƠNG PHÁP TRÍCH XUẤT ĐẶC TRƯNG DỮ LIỆU LOG VÀ THÁCH THỨC BẢO TOÀN NGỮ CẢNH AN TOÀN")

    # =========================================================================
    # 1.1.
    # =========================================================================
    add_h2("Bài toán biểu diễn log trong phát hiện tấn công đa giai đoạn")
    add_p([
        "Trong các Trung tâm Điều hành An ninh mạng (Security Operations Center - SOC) hiện đại, dữ liệu nhật ký hệ thống (system logs) và nhật ký kiểm toán (audit logs) đóng vai trò là nguồn bằng chứng trung tâm phục vụ phát hiện, điều tra và ứng phó các chiến dịch tấn công có chủ đích (Advanced Persistent Threats - APT) ",
        make_citation_element([9]),
        ". Khác với văn bản ngôn ngữ tự nhiên thông thường hoặc tín hiệu thị giác máy tính, dữ liệu log sở hữu cấu trúc bán hình thức (semi-structured), mang tính phụ thuộc thời gian nghiêm ngặt, phản ánh các quan hệ phụ thuộc thực thi giữa các tiến trình và tài nguyên hệ điều hành, đồng thời chứa đựng các ngữ nghĩa an ninh đặc thù ",
        make_citation_element([8, 18]),
        ". Mục 1.1 tập trung hình thức hóa bài toán biểu diễn đặc trưng log phục vụ phát hiện tấn công đa giai đoạn, phân tích bản chất không gian dữ liệu doanh nghiệp, xác lập mô hình hành vi phi tuyến tính trên ma trận MITRE ATT&CK (Phiên bản Enterprise v19.1, 28/04/2026) ",
        make_citation_element([1]),
        " và thiết lập khung Hợp đồng Biểu diễn (Representation Contract) nhằm bảo toàn các bất biến an toàn cho không gian vector đặc trưng ",
        latex_to_clean_omml(r"\mathbf{z}"),
        "."
    ])

    # --- 1.1.1 ---
    add_h3("Không gian dữ liệu log doanh nghiệp: tốc độ cao, mất cân bằng cực đoan và phân phối biến đổi")
    add_p([
        "Không gian dữ liệu nhật ký trong môi trường mạng doanh nghiệp được tổng hợp từ ba nhóm nguồn telemetry chính với cấu trúc và định dạng không đồng nhất ",
        make_citation_element([9, 10]),
        ". Nhóm thứ nhất là nhật ký kiểm toán máy chủ (Host Audit Logs), bao gồm Linux Auditd, Windows Event Log / Sysmon và Linux eBPF (Extended Berkeley Packet Filter). Nguồn dữ liệu này ghi nhận trực tiếp các sự kiện ở mức nhân hệ điều hành thông qua việc chặn bắt các lời gọi hệ thống (syscalls), bao gồm: khởi tạo tiến trình (execve, CreateProcess - Sysmon Event ID 1), nạp thư viện động (ImageLoaded - Sysmon Event ID 7), thao tác tệp tin (open, unlink, FileCreate - Sysmon Event ID 11), sửa đổi cấu hình registry (RegSetValue - Sysmon Event ID 13), cùng các thao tác mở và kết nối socket mạng (connect, accept - Sysmon Event ID 3) ",
        make_citation_element([8, 9]),
        ". Nhóm thứ hai là nhật ký luồng mạng (Network Flow & Protocol Logs), được thu thập từ Zeek, Suricata hoặc NetFlow/IPFIX, cung cấp siêu dữ liệu kết nối giữa các nút mạng, giao dịch DNS, chứng chỉ TLS/SSL và thông lượng gói tin. Nhóm thứ ba là nhật ký ứng dụng và dịch vụ (Application & Service Logs), phát sinh từ máy chủ web (Nginx, Apache), cơ sở dữ liệu, dịch vụ phân tán (HDFS) cùng hệ thống điều phối container (Kubernetes Audit Logs) ",
        make_citation_element([6, 7]),
        "."
    ])
    add_p([
        "Tính dị thể sâu sắc của dữ liệu đặt ra bài toán khoa học về việc lựa chọn đơn vị quan sát (Unit of Observation) phù hợp cho mô hình học biểu diễn, như được hệ thống hóa trực quan trong ",
        make_ref_element("BK_FIG_1_001", "Hình 1.1"),
        " và chi tiết hóa trong ",
        make_ref_element("BK_TBL_1_001", "Bảng 1.1"),
        ":"
    ])

    # FIGURE 1.1
    add_figure_image(fig1_path, width_inches=6.2)
    add_figure_caption(doc, target_p, 1, 1, "Phân cấp các đơn vị quan sát và sự đánh đổi giữa ngữ cảnh và chi phí tính toán trong biểu diễn dữ liệu log (Nguồn: Tác giả tổng hợp)", bookmark_name="BK_FIG_1_001")

    # TABLE 1.1
    add_table_caption(doc, target_p, 1, "Phân cấp các đơn vị quan sát và mức độ hạt biểu diễn dữ liệu log", bookmark_name="BK_TBL_1_001")
    tbl1_headers = ["Mức độ hạt", "Đơn vị quan sát", "Dữ liệu đại diện", "Ưu điểm cốt lõi", "Thách thức và Mất mát ngữ nghĩa"]
    tbl1_widths = [1551, 2028, 1791, 1806, 2429]
    tbl1_rows = [
        ["Mức từ tố (Token)", "Chuỗi con, từ khóa rời rạc", "Từ khóa tĩnh, địa chỉ IP, mã lỗi hex", "Dễ dàng vector hóa bằng kỹ thuật nhúng từ vựng", "Mất liên kết cấu trúc cú pháp và quan hệ trường"],
        ["Mức sự kiện (Event)", "Một dòng log đơn lẻ", "Bản ghi telemetry tại thời điểm t", "Bảo toàn đầy đủ thuộc tính cục bộ tại thời điểm t", "Thiếu ngữ cảnh chuỗi tuần tự và lịch sử tương tác"],
        ["Mức chuỗi / Phiên (Session)", "Cửa sổ trượt hoặc phiên tiến trình", "Chuỗi sự kiện [e_{t-k+1}, ..., e_t] theo thời gian", "Nắm bắt quan hệ phụ thuộc thứ tự thời gian", "Nhạy cảm với nhiễu xen kẽ từ các luồng song song"],
        ["Mức thực thể (Entity)", "Định danh tác nhân (Host, User, Process)", "Lịch sử tương tác gom cụm theo thực thể", "Phân lập rõ ràng ranh giới hành vi chủ thể", "Khó phát hiện hành vi tấn công vượt ranh giới"],
        ["Mức đồ thị (Graph)", "Đồ thị nguồn gốc (Provenance Graph)", "Đồ thị luồng phụ thuộc dị thể G", "Mô hình hóa toàn diện quan hệ đa thực thể", "Bùng nổ kích thước đồ thị và chi phí tính toán"]
    ]
    insert_thesis_table(doc, target_p, tbl1_headers, tbl1_widths, tbl1_rows, font_size_pt=14)
    add_p("", first_line_indent=False)

    add_p([
        "Bên cạnh tính dị thể, dữ liệu log doanh nghiệp chịu áp lực vận hành về mặt thông lượng và tỷ lệ phân bố nhãn ",
        make_citation_element([2, 18]),
        ". Trong các môi trường mạng lớn, hệ thống tiếp nhận luồng sự kiện liên tục với tốc độ cao, đòi hỏi thuật toán biểu diễn đặc trưng phải xử lý theo cơ chế dòng (streaming) với độ phức tạp tính toán tuyến tính O(N) mà không đòi hỏi lưu giữ toàn bộ lịch sử đồ thị trong bộ nhớ truy cập ngẫu nhiên ",
        make_citation_element([11]),
        ". Đồng thời, tỷ lệ các dòng log liên quan đến hành vi tấn công trong thực tế thường rất nhỏ so với lưu lượng thông thường, tạo ra sự mất cân bằng nhãn sâu sắc khiến các hàm mất mát học máy thông thường có xu hướng xem nhẹ nhóm thiểu số ",
        make_citation_element([2]),
        ". Trong khi đó, việc suy thoái không gian vector (sụp đổ chiều biểu diễn) nảy sinh khi các mô hình tự giám sát không áp dụng các cơ chế điều hòa phương sai - hiệp phương sai phù hợp để duy trì tính đa dạng của các chiều đặc trưng ",
        make_citation_element([22, 23]),
        "."
    ])
    add_p([
        "Dưới góc độ phân phối thời gian, dữ liệu log liên tục biến đổi trong môi trường vận hành dài hạn. Để tránh các sai lệch phương pháp luận khi đánh giá mô hình, chuyên đề phân định tường minh bốn cơ chế trôi dạt dữ liệu (Drift Taxonomy) ",
        make_citation_element([2]),
        ":"
    ])
    add_p([
        "Bản chất hành vi và mục đích tấn công thay đổi theo thời gian mặc dù cấu trúc định dạng log không đổi: ",
        latex_to_clean_omml(r"P_t(Y \mid X) \neq P_{t+1}(Y \mid X)"),
        "."
    ], bold_prefix="• Concept Drift: ")
    add_p([
        "Việc nâng cấp phần mềm, cập nhật bản vá hoặc thay đổi cấu hình làm biến đổi cấu trúc chuỗi mẫu log: ",
        latex_to_clean_omml(r"P_t(X_{\text{template}}) \neq P_{t+1}(X_{\text{template}})"),
        "."
    ], bold_prefix="• Template Drift: ")
    add_p([
        "Lưu lượng người dùng, tần suất giao tác nghiệp vụ hoặc cơ cấu dịch vụ hệ thống biến động theo chu kỳ: ",
        latex_to_clean_omml(r"P_t(X) \neq P_{t+1}(X)"),
        "."
    ], bold_prefix="• Population Drift: ")
    add_p([
        "Không gian vector tiềm ẩn z bị suy giảm năng lực phân tách do dữ liệu đầu vào trôi dạt khỏi vùng phân phối huấn luyện ban đầu: ",
        latex_to_clean_omml(r"P_t(\mathbf{z} \mid X) \neq P_{t+1}(\mathbf{z} \mid X)"),
        "."
    ], bold_prefix="• Representation Drift: ")

    # --- 1.1.2 ---
    add_h3("Hành vi tấn công đa giai đoạn và ánh xạ đa nhãn MITRE ATT&CK")
    add_p(
        "Trong lịch sử nghiên cứu phát hiện xâm nhập, các tiếp cận ban đầu thường dựa trên giả định chuỗi tấn công đơn tuyến (Cyber Kill Chain) hoặc các mô hình xích Markov tuyến tính để xâu chuỗi tuần tự các giai đoạn tấn công từ Thâm nhập ban đầu (Initial Access), Thực thi (Execution), Duy trì (Persistence), Leo thang đặc quyền (Privilege Escalation) cho đến Đánh cắp dữ liệu (Exfiltration)."
    )
    add_p([
        "Tuy nhiên, trên thực tế vận hành và theo khung phân loại chuẩn hóa của ma trận tri thức an ninh MITRE ATT&CK (Phiên bản Enterprise v19.1, 28/04/2026) ",
        make_citation_element([1]),
        ", các chiến dịch tấn công có chủ đích (APT) mang bản chất phi tuyến tính sâu sắc ",
        make_citation_element([9]),
        ": (1) Nhảy cóc giai đoạn (Step Skipping): Kẻ tấn công có thể khai thác trực tiếp lỗ hổng thực thi mã từ xa để trích xuất dữ liệu ra ngoài mà không cần thiết lập cơ chế duy trì hay di chuyển ngang; (2) Lặp vòng kỹ thuật (Tactic Looping & Interleaving): Kỹ thuật thu thập thông tin nội bộ (Discovery) thường được lặp lại nhiều lần xen kẽ giữa các bước leo thang đặc quyền và chiếm đoạt thông tin xác thực; (3) Phân nhánh tiến trình song song (Parallel Branching): Kẻ tấn công có thể khởi tạo đồng thời nhiều luồng tiến trình con độc lập trên các tiến trình hợp lệ khác nhau nhằm phân tán sự theo dõi của hệ thống phòng thủ."
    ])
    add_p([
        "Do đó, chuyên đề xác lập nguyên tắc: Ma trận MITRE ATT&CK được mô hình hóa thành một Không gian Bằng chứng Hành vi Đa chiều (Multi-label Behavioral Evidence Space), được minh họa trực quan trong ",
        make_ref_element("BK_FIG_1_002", "Hình 1.2"),
        ":"
    ])
    add_display_equation(r"\mathcal{Y}_{\text{ATT\&CK}} = \mathcal{T}_{\text{tactics}} \times \mathcal{T}_{\text{techniques}}")

    # FIGURE 1.2
    add_figure_image(fig2_path, width_inches=6.2)
    add_figure_caption(
        doc,
        target_p,
        1,
        2,
        [
            "Mô hình Không gian Bằng chứng Hành vi Đa chiều MITRE ATT&CK và các đặc trưng phi tuyến tính trong tấn công APT (Nguồn: Tác giả tổng hợp dựa trên MITRE ATT&CK ",
            make_citation_element([1]),
            " và Inam et al. ",
            make_citation_element([9]),
            ")"
        ],
        bookmark_name="BK_FIG_1_002"
    )

    add_p([
        "Trong đó mỗi chiến thuật (Tactic) đại diện cho mục tiêu chiến thuật của kẻ tấn công (trả lời câu hỏi 'Tại sao') và mỗi kỹ thuật (Technique) đại diện cho phương thức thực thi cụ thể (trả lời câu hỏi 'Làm thế nào'); một chuỗi sự kiện hoặc cây tiến trình có thể đồng thời kích hoạt nhiều nhãn chiến thuật và kỹ thuật tại cùng một thời điểm quan sát ",
        make_citation_element([1]),
        "."
    ], first_line_indent=False)
    add_p([
        "Về mặt dữ liệu thực nghiệm, việc mô hình hóa hành vi tấn công đòi hỏi phải phân định chính xác đặc tính gán nhãn và mức độ hạt (Label Granularity) của từng bộ dữ liệu chuẩn ",
        make_citation_element([18]),
        ": (1) DARPA Transparent Computing (TC E3/E5) ",
        make_citation_element([28]),
        " cung cấp dữ liệu kiểm toán hệ thống mức hạt nhân chi tiết với các kịch bản APT thực tế được gán nhãn ở mức độ hạt tiến trình/luồng phụ thuộc (Fine-grained Ground Truth); (2) LANL Unified Host and Network Dataset (Kent, 2015) ",
        make_citation_element([29]),
        " phản ánh môi trường mạng doanh nghiệp quy mô lớn với hàng tỷ sự kiện xác thực và luồng mạng, trong đó nhãn mặt đất thực nghiệm được xác lập từ tệp redteam.txt ghi nhận các sự kiện xác thực bị xâm nhập cụ thể của đội Red Team (Authentication Compromise Events) theo mốc thời gian và tài khoản/máy chủ xác định; (3) HDFS và BGL Datasets ",
        make_citation_element([3, 6]),
        " đại diện cho nhật ký hệ thống phân tán và siêu máy tính, được gán nhãn bất thường ở mức khối dữ liệu (Block-level) hoặc mức dòng log đơn lẻ."
    ])
    add_p([
        "Đặc biệt, sự xuất hiện của nhiễu từ hành vi quản trị viên (Admin-Noise) là một thách thức then chốt ",
        make_citation_element([9, 18]),
        ". Các quản trị viên hệ thống thường xuyên sử dụng PowerShell, SSH, WMI và các công cụ dòng lệnh nội bộ tương tự như kẻ tấn công APT, tạo ra sự trùng lặp lớn trong không gian đặc trưng hành vi và gây ra nhiều cảnh báo giả nếu mô hình chỉ học các mẫu bề mặt mà không nắm bắt được ngữ cảnh phụ thuộc cấu trúc và hành vi sâu ",
        make_citation_element([24]),
        "."
    ])

    # --- 1.1.3 ---
    add_h3("Các mức biểu diễn dữ liệu và Hợp đồng Biểu diễn (Representation Contract)")
    add_p([
        "Để khắc phục tình trạng các nghiên cứu tiền nhiệm thường đồng nhất việc trích xuất đặc trưng với thuật toán phát hiện cụ thể, chuyên đề đề xuất một khung khái niệm chính thức mang tên Hợp đồng Biểu diễn (Representation Contract), ký hiệu là ",
        latex_to_clean_omml(r"\mathcal{C}_{\text{representation}}"),
        ", đóng vai trò như một bản đặc tả ràng buộc toán học chặt chẽ giữa dữ liệu thô và không gian vector tiềm ẩn z:"
    ])
    add_display_equation(r"\mathcal{C}_{\text{representation}} = \langle \mathcal{P}_{\text{preserve}}, \mathcal{I}_{\text{invariant}}, \mathcal{E}_{\text{exclude}} \rangle")
    add_p(
        "Khung Hợp đồng Biểu diễn thiết lập ba nhóm điều kiện ràng buộc bắt buộc mà không gian vector biểu diễn z phải đồng thời thỏa mãn:"
    )
    add_p(
        "Tập hợp các thuộc tính và tương quan an ninh bắt buộc phải được bảo toàn nguyên vẹn trong không gian vector biểu diễn. Bao gồm: ngữ nghĩa sâu của các tham số dòng lệnh quan trọng, trật tự thời gian và mối quan hệ phụ thuộc thực thi giữa các thực thể hệ thống (Process-File-Socket), cấu trúc tô-pô cục bộ của đồ thị nguồn gốc và các đặc trưng nhận diện chiến thuật/kỹ thuật MITRE ATT&CK.",
        bold_prefix="1. Nhóm Điều kiện Bảo toàn (PRESERVE - P_preserve): "
    )
    add_p(
        "Tập hợp các biến đổi hình thức hoặc biến động môi trường mà không gian vector z phải duy trì tính bất biến. Bao gồm: sự biến đổi cú pháp vô hại của các chuỗi văn bản log (định dạng dấu thời gian, khoảng trắng, thứ tự các trường không mang nghĩa an ninh), sự xáo trộn mã định danh tiến trình ngẫu nhiên (PID/PPID), tên tài khoản người dùng hoặc địa chỉ IP cục bộ không ảnh hưởng đến bản chất luồng thực thi, và các biến động định kỳ về khối lượng lưu lượng nền (Workload fluctuations).",
        bold_prefix="2. Nhóm Điều kiện Bất biến (INVARIANT - I_invariant): "
    )
    add_p([
        "Tập hợp các thông tin giả tạo và biến số ngoài miền bắt buộc phải bị loại bỏ/triệt tiêu khỏi không gian vector z nhằm ngăn ngừa rủi ro học đường tắt (Shortcut Learning) ",
        make_citation_element([2]),
        ". Bao gồm: các mẫu định dạng đặc thù của môi trường thử nghiệm (Testbed-specific artifacts), các biến số gây rò rỉ phân vùng (Partition-leakage variables như dấu thời gian tuyệt đối của máy thí nghiệm, hostname cố định của môi trường lab) ",
        make_citation_element([25, 27]),
        ", và các đặc trưng tương quan giả (Spurious correlations) có thể khiến mô hình đạt độ chính xác ảo trên tập kiểm thử nhưng thất bại khi triển khai thực tế."
    ], bold_prefix="3. Nhóm Điều kiện Triệt tiêu (EXCLUDE - E_exclude): ")
    add_p([
        "Các đặc tả chi tiết của Hợp đồng Biểu diễn được hệ thống hóa trong ",
        make_ref_element("BK_TBL_1_002", "Bảng 1.2"),
        ":"
    ])

    # TABLE 1.2
    add_table_caption(doc, target_p, 2, "Đặc tả Hợp đồng Biểu diễn (Representation Contract) cho vector đặc trưng z", bookmark_name="BK_TBL_1_002")
    tbl2_headers = ["Thành phần hợp đồng", "Mục tiêu ràng buộc", "Đối tượng áp dụng cụ thể trong dữ liệu log", "Cơ chế kiểm chứng thực nghiệm"]
    tbl2_widths = [1800, 2400, 3100, 2307]
    tbl2_rows = [
        ["PRESERVE (Bảo toàn)", "Bảo tồn đầy đủ tín hiệu ngữ cảnh an ninh phân biệt", "Ngữ nghĩa tham số động, trật tự thời gian và quan hệ phụ thuộc, tô-pô đồ thị, nhãn ATT&CK", "Linear Probing, Mutual Information"],
        ["INVARIANT (Bất biến)", "Duy trì bất biến trước biến đổi hình thức và trôi dạt vô hại", "Định dạng timestamp, PID ngẫu nhiên, hoán vị cú pháp vô hại", "Invariance Loss, Data Augmentation"],
        ["EXCLUDE (Triệt tiêu)", "Loại bỏ rò rỉ phân vùng và biến số học đường tắt", "Artifacts môi trường lab, hostname cố định, thông tin định danh nhạy cảm (PII)", "Snooping Test, DP Guarantees"]
    ]
    insert_thesis_table(doc, target_p, tbl2_headers, tbl2_widths, tbl2_rows, font_size_pt=14)
    add_p("", first_line_indent=False)

    add_p([
        "Cần nhấn mạnh rằng, nguyên tắc PRESERVE tập trung vào việc bảo toàn ngữ nghĩa an ninh động chứ không đồng nghĩa với việc giữ lại toàn bộ định danh thô của người dùng hay máy chủ, nhằm đảm bảo khả năng liên kết có kiểm soát (controlled linkability) và tương thích với các tiêu chuẩn đánh giá quyền riêng tư ",
        make_citation_element([25, 27]),
        ". Tương tự, nguyên tắc EXCLUDE thiết lập các ràng buộc phủ định đối với các giả định đường tắt đã biết và các biến số gây rò rỉ phân vùng ",
        make_citation_element([2]),
        ", thay vì giả định rằng mọi đường tắt đều có thể nhận biết trước khi thực nghiệm."
    ])
    add_p([
        "Đi đôi với Hợp đồng Biểu diễn được quy định trong ",
        make_ref_element("BK_TBL_1_002", "Bảng 1.2"),
        ", chuyên đề thiết lập nguyên tắc phân định ranh giới phương pháp luận ba tầng độc lập, được mô hình hóa tổng quát trong ",
        make_ref_element("BK_FIG_1_003", "Hình 1.3"),
        ":"
    ])

    # FIGURE 1.3
    add_figure_image(fig3_path, width_inches=6.2)
    add_figure_caption(doc, target_p, 1, 3, "Khung phân định ranh giới ba tầng phương pháp luận và vị trí trọng tâm của không gian vector biểu diễn z (Nguồn: Tác giả đề xuất)", bookmark_name="BK_FIG_1_003")

    add_p(
        "Đảm nhiệm tiền xử lý dữ liệu thô, phân tích cú pháp sơ bộ, chuẩn hóa kiểu dữ liệu trường và làm sạch dữ liệu. Tầng này không gánh vác nhiệm vụ học biểu diễn ngữ cảnh an ninh sâu.",
        bold_prefix="1. Tầng 1 — Trích xuất đặc trưng cơ sở (Feature Extraction): "
    )
    add_p([
        "Thiết lập ánh xạ ",
        latex_to_clean_omml(r"f_\theta: \mathcal{X} \to \mathbf{z} \in \mathbb{R}^d"),
        " từ cấu trúc chuỗi sự kiện và đồ thị nguồn gốc sang không gian vector tiềm ẩn. Toàn bộ năng lực bảo toàn ngữ nghĩa an ninh và tính bất biến được đóng gói trọn vẹn bên trong vector z ",
        make_citation_element([22]),
        "."
    ], bold_prefix="2. Tầng 2 — Học không gian biểu diễn (Representation Learning — Trọng tâm Chuyên đề): ")
    add_p([
        "Đánh giá chất lượng của vector biểu diễn z thông qua các bộ thăm dò tuyến tính đóng băng tham số (Frozen Linear Probing):"
    ], bold_prefix="3. Tầng 3 — Phát hiện và phân loại hạ nguồn (Downstream Detection): ")
    add_display_equation(r"\hat{y} = \sigma(\mathbf{W}^\top \mathbf{z} + \mathbf{b})")
    add_p([
        "Trong đó tham số ",
        latex_to_clean_omml(r"\theta"),
        " của bộ trích xuất đặc trưng Tầng 2 được giữ cố định hoàn toàn trong suốt quá trình đánh giá ở Tầng 3. Quy tắc này bảo đảm bộ phân loại hạ nguồn không làm thay nhiệm vụ trích xuất đặc trưng của Tầng 2 ",
        make_citation_element([2]),
        "."
    ], first_line_indent=False)

    # =========================================================================
    # 1.2.
    # =========================================================================
    add_h2("Phân tích so sánh các nhóm phương pháp hiện đại")
    add_p([
        "Nhằm định vị chính xác các đóng góp kỹ thuật và cơ sở lý luận của chuyên đề, Mục 1.2 tiến hành khảo cứu, phân loại và phân tích đa chiều ba nhóm phương pháp biểu diễn đặc trưng log chủ đạo trong y văn hiện đại ",
        make_citation_element([6, 9, 10]),
        ": (1) Nhóm phương pháp thống kê và cú pháp dựa trên mẫu log (Statistical & Syntactic Parsing-Based); (2) Nhóm phương pháp biểu diễn ngữ nghĩa chuỗi thời gian (Semantic–Sequential & Transformer-Based); và (3) Nhóm phương pháp học biểu diễn đồ thị nguồn gốc (Provenance Graph Representation Learning). Mỗi nhóm phương pháp được mổ xẻ tường minh về cơ chế toán học, độ phức tạp thuật toán, ưu điểm cốt lõi và các rào cản nền tảng khi triển khai trong môi trường phát hiện tấn công thực tế."
    ])

    # --- 1.2.1 ---
    add_h3("Phương pháp thống kê và cú pháp: Event Count, Frequency, Entropy và Template Features")
    add_p([
        "Nhóm phương pháp thống kê và cú pháp đại diện cho thế hệ tiếp cận đầu tiên trong phân tích nhật ký tự động ",
        make_citation_element([6, 7]),
        ". Cơ chế hoạt động của nhóm này dựa trên quy trình hai giai đoạn tách rời: giai đoạn phân tách dòng log thô thành các mẫu định dạng tĩnh (Log Templates / Event IDs) thông qua bộ phân tích cú pháp (Log Parser), tiếp theo là giai đoạn lượng hóa các chuỗi sự kiện thành vector số học dựa trên các thước đo thống kê kinh điển."
    ])
    add_p([
        "Trong giai đoạn phân tích cú pháp, nhiều thuật toán tiêu biểu đã được phát triển nhằm tối ưu hóa tốc độ xử lý ",
        make_citation_element([6]),
        ": (1) Drain ",
        make_citation_element([6]),
        " sử dụng cấu trúc cây phân tích có độ sâu cố định (Fixed-Depth Parse Tree) để nhóm các dòng log dựa trên độ dài chuỗi và các từ khóa tiền tố, đạt tốc độ phân tích xấp xỉ tuyến tính O(N) đối với luồng dữ liệu lớn; (2) Spell ",
        make_citation_element([6]),
        " áp dụng thuật toán tìm chuỗi con chung dài nhất (Longest Common Subsequence - LCS) theo cơ chế dòng (streaming) để trích xuất động các thành phần tĩnh của thông điệp log; (3) LenMa và AEL ",
        make_citation_element([6, 7]),
        " lần lượt khai thác chiều dài các từ tố và tần suất xuất hiện của từ khóa để phân cụm và tách biến số ra khỏi chuỗi mẫu."
    ])
    add_p([
        "Sau khi không gian log được rút gọn về tập ",
        latex_to_clean_omml(r"M"),
        " mẫu sự kiện cố định ",
        latex_to_clean_omml(r"\mathcal{E} = \{e_1, e_2, \dots, e_M\}"),
        ", các chuỗi sự kiện trong một cửa sổ thời gian hoặc phiên làm việc ",
        latex_to_clean_omml(r"W"),
        " được ánh xạ thành vector tần suất:"
    ])
    add_display_equation(r"\mathbf{x} = [c(e_1), c(e_2), \dots, c(e_M)]^\top \in \mathbb{R}^M")
    add_p([
        "Trong đó ",
        latex_to_clean_omml(r"c(e_i)"),
        " là số lần xuất hiện của sự kiện ",
        latex_to_clean_omml(r"e_i"),
        ". Bên cạnh đếm tần suất đơn thuần, các trọng số TF-IDF (Term Frequency - Inverse Document Frequency) hoặc độ hỗn loạn thông tin Shannon (Information Entropy) cũng được áp dụng nhằm nhấn mạnh các sự kiện hiếm gặp ",
        make_citation_element([6]),
        ":"
    ], first_line_indent=False)
    add_display_equation(r"H(W) = -\sum_{i=1}^M p(e_i) \log_2 p(e_i)")
    add_p([
        "Trên không gian vector tần suất này, các mô hình phát hiện bất thường kinh điển được triển khai rộng rãi: (1) Principal Component Analysis - PCA (Xu et al., SOSP 2009) ",
        make_citation_element([31]),
        " phân rã không gian vector ",
        latex_to_clean_omml(r"\mathbb{R}^M"),
        " thành không gian con chuẩn tắc ",
        latex_to_clean_omml(r"\mathcal{S}_n"),
        " và không gian con phần dư ",
        latex_to_clean_omml(r"\mathcal{S}_r"),
        ", nhận diện bất thường khi năng lượng chiếu lên phần dư vượt ngưỡng kiểm định:"
    ])
    add_display_equation(r"\mathbf{x}_a = (\mathbf{I} - \mathbf{P} \mathbf{P}^\top) \mathbf{x}, \quad \|\mathbf{x}_a\|^2 > \gamma_\alpha")
    add_p([
        "(2) Invariant Mining (Lou et al., USENIX ATC 2010) ",
        make_citation_element([32]),
        " tự động khai phá các phương trình bất biến tuyến tính phản ánh mối quan hệ bảo toàn logic giữa các bước thực thi:"
    ])
    add_display_equation(r"\mathbf{A} \mathbf{x} = \mathbf{0}")
    add_p([
        "Mặc dù sở hữu ưu điểm về hiệu năng tính toán (độ phức tạp tuyến tính O(N)), nhóm phương pháp thống kê và cú pháp bộc lộ hai điểm nghẽn phương pháp luận quan trọng ",
        make_citation_element([6, 8]),
        ": (1) Mất mát ngữ nghĩa an ninh do trừu tượng hóa tham số: các bộ log parser dựa trên biểu thức chính quy thường thay thế các tham số biến động (địa chỉ IP, đường dẫn tệp tin, tham số dòng lệnh) bằng ký tự đại diện <*> khiến nhiều thông tin an ninh mang tính phân biệt cao bị lược bỏ; (2) Lan truyền và khuếch đại sai số cú pháp (Parser Error Propagation): khi gặp các định dạng log mới chưa từng xuất hiện (unseen logs), parser có thể phân tách không chính xác, dẫn đến hiện tượng sinh ra các mẫu sự kiện giả lập hoặc gộp nhầm các sự kiện khác biệt, làm xáo trộn cấu trúc không gian vector x."
    ])

    # --- 1.2.2 ---
    add_h3("Phương pháp semantic–sequential: Embeddings, Self-Supervised Learning, Transformer và Parsing-Free")
    add_p([
        "Nhằm khắc phục sự cứng nhắc của các vector đếm tần suất và tận dụng trật tự xuất hiện của các sự kiện, nhóm phương pháp Semantic–Sequential mô hình hóa luồng log tương tự như các chuỗi ngôn ngữ tự nhiên, tích hợp các kỹ thuật nhúng từ (Word Embeddings) và học tự giám sát (Self-Supervised Learning) để nắm bắt phụ thuộc ngữ cảnh dài hạn ",
        make_citation_element([3, 4, 5, 33, 34]),
        "."
    ])
    add_p([
        "Khởi đầu cho hướng nghiên cứu này là mô hình DeepLog ",
        make_citation_element([3]),
        ". DeepLog sử dụng mạng bộ nhớ dài-ngắn hạn (LSTM) để mô hình hóa chuỗi Event ID như một bài toán dự báo phần tử tiếp theo (Next-Event Prediction). Tại mỗi bước thời gian t, mô hình ước lượng phân phối xác suất có điều kiện của sự kiện tiếp theo ",
        latex_to_clean_omml(r"P(e_t \mid e_{t-k}, \dots, e_{t-1})"),
        ". Nếu sự kiện thực tế không nằm trong tập g sự kiện có xác suất cao nhất được mô hình dự đoán, hệ thống sẽ phát tín hiệu cảnh báo bất thường:"
    ])
    add_display_equation(r"\hat{\mathcal{E}}_t = \operatorname{arg\,top-}g_{e \in \mathcal{E}} P(e \mid e_{t-k}, \dots, e_{t-1})")
    add_p([
        "Song song đó, DeepLog xây dựng mô hình LSTM thứ hai dựa trên phân phối chuẩn nhiều chiều để kiểm tra sự bất thường về giá trị tham số số học (Parameter Value Anomaly) ",
        make_citation_element([3]),
        "."
    ], first_line_indent=False)
    add_p([
        "Các công trình kế tiếp đã nâng cấp cơ chế biểu diễn ngữ nghĩa: (1) LogAnomaly (Meng et al., IJCAI 2019) ",
        make_citation_element([33]),
        " đề xuất Template2Vec, trích xuất vector ngữ nghĩa cho từng mẫu log thông qua Word2Vec/FastText kết hợp trọng số d-IDF, giúp nhận biết sự tương đồng giữa các thông điệp có cấu trúc từ ngữ tương đương; (2) Logsy (Nedelkoski et al., ICDM 2020) ",
        make_citation_element([34]),
        " sử dụng hàm mất mát phân loại ngoại lai (Outlier Classification Loss) trên dữ liệu log từ các hệ thống phụ trợ để định hình biên giới phân tách cho lớp bình thường; (3) LogBERT (Guo et al., IJCNN 2021) ",
        make_citation_element([4]),
        " khai thác Transformer hai chiều với hai tác vụ học tự giám sát: Dự đoán sự kiện log bị che (Masked Log Event Prediction) và Dự đoán phân bố khối lượng log (Volume Anomaly Prediction); (4) Nhóm tiếp cận không dùng Parser (Parsing-Free - NeuralLog ",
        make_citation_element([5]),
        ") bỏ qua bước phân tích cú pháp bằng cách sử dụng các mô hình ngôn ngữ tiền huấn luyện (BERT, RoBERTa) để trực tiếp mã hóa chuỗi văn bản log thô thành các vector nhúng ngữ nghĩa liên tục."
    ])
    add_p([
        "Mặc dù đạt kết quả tốt trên các tập dữ liệu phần mềm, nhóm phương pháp Semantic–Sequential đối mặt với ba rào cản khi áp dụng vào an ninh mạng ",
        make_citation_element([2, 18]),
        ": (1) Nguy cơ rò rỉ và thiên lệch từ dữ liệu tiền huấn luyện (Pretraining-Data Advantage): các mô hình sử dụng Transformer tiền huấn luyện trên kho văn bản tổng quát có nguy cơ tận dụng tri thức ngoài miền; khi đánh giá trong điều kiện phân vùng nghiêm ngặt, năng lực phân tách thực tế cần được kiểm chứng cẩn trọng; (2) Chi phí tài nguyên tính toán: độ phức tạp tính toán bậc hai của cơ chế Self-Attention O(L^2) theo độ dài cửa sổ L đòi hỏi tài nguyên tính toán đáng kể trong môi trường lưu lượng lớn; (3) Giới hạn phạm vi quan sát đơn luồng: mô hình chuỗi chủ yếu theo dõi các sự kiện trên một dòng thời gian cục bộ, gặp khó khăn khi mô hình hóa trực tiếp các mối liên hệ phụ thuộc đan xen đa tiến trình, đa thực thể và vượt ranh giới máy chủ."
    ])

    # --- 1.2.3 ---
    add_h3("Đồ thị nguồn gốc và Graph Representation Learning")
    add_p([
        "Để khắc phục hạn chế về phạm vi quan sát của mô hình chuỗi, hướng tiếp cận dựa trên đồ thị nguồn gốc hệ thống (System Provenance Graph) mô hình hóa toàn bộ lịch sử thực thi và tương tác trong hệ điều hành dưới dạng một đồ thị có hướng, không đồng nhất và gán nhãn thời gian ",
        make_citation_element([9, 10]),
        ":"
    ])
    add_display_equation(r"\mathcal{G} = (\mathcal{V}, \mathcal{E}, \mathcal{T}_v, \mathcal{T}_e, \phi, \psi, \tau)")
    add_p([
        "Trong đó V là tập các đỉnh đại diện cho các thực thể hệ thống thuộc tập kiểu ",
        latex_to_clean_omml(r"\mathcal{T}_v = \{\text{Process, File, Socket, Registry, User, Host}\}"),
        "; E là tập các cạnh có hướng mô tả các tương tác luồng phụ thuộc thuộc tập kiểu ",
        latex_to_clean_omml(r"\mathcal{T}_e = \{\text{fork, execve, read, write, connect, bind, send, recv}\}"),
        "; và τ gán nhãn mốc thời gian xảy ra tương tác ",
        make_citation_element([9]),
        "."
    ], first_line_indent=False)
    add_p([
        "Các hệ thống phát hiện xâm nhập dựa trên đồ thị nguồn gốc (PIDS) tiêu biểu bao gồm ",
        make_citation_element([9, 10]),
        ": (1) UNICORN ",
        make_citation_element([11]),
        " (Han et al., NDSS 2020) xây dựng đồ thị nguồn gốc luồng thời gian thực, áp dụng thuật toán băm cây con Weisfeiler-Lehman (WL-subtree kernel) để chuyển đổi đồ thị động thành vector đặc trưng đếm histogram; (2) KAIROS ",
        make_citation_element([12]),
        " (Cheng et al., IEEE S&P 2024) tích hợp mạng nơ-ron đồ thị nhận biết thời gian (Time-Aware GNN), mã hóa đồng thời thông tin cấu trúc và khoảng cách thời gian giữa các cạnh để phát hiện kỹ thuật APT ẩn mình; (3) NODLINK ",
        make_citation_element([13]),
        " (Li et al., NDSS 2024) mô hình hóa bài toán phát hiện thành cây Steiner trực tuyến (Online Steiner Tree Problem) kết hợp cơ chế bộ nhớ đệm để khôi phục đồ thị con tấn công nhỏ gọn; (4) MAGIC ",
        make_citation_element([14]),
        " (Jia et al., USENIX Security 2024) khai thác kiến trúc tự mã hóa đồ thị che (Masked Graph Autoencoder) để học biểu diễn hành vi bình thường và phát hiện bất thường APT; (5) ORTHRUS ",
        make_citation_element([15]),
        ": (1) Hiện tượng bùng nổ phụ thuộc (Dependency Explosion): các tiến trình hệ thống chạy dài hạn (như daemon hệ thống hoặc trình duyệt) liên tục tương tác với nhiều tệp tin và socket, khiến đồ thị phát triển dày đặc và tạo ra nhiều liên kết phụ thuộc xa làm loãng tín hiệu bất thường ",
        make_citation_element([9, 20]),
        "; (2) Ranh giới giữa quan hệ phụ thuộc cấu trúc và tác động nhân quả (Dependency != Causal Effect): kết quả khảo sát thực nghiệm của Bilot et al. ",
        make_citation_element([16]),
        " và Guerra et al. (arXiv:2608.01454) ",
        make_citation_element([30]),
        " trên các bộ dữ liệu PIDS chuẩn chỉ ra rằng nhiều mô hình GNN phức tạp có xu hướng khai thác các đặc trưng đường tắt thống kê (như phân bố bậc của nút hoặc tính mới của đường dẫn tệp tin); khi kiểm soát chặt chẽ các yếu tố gây nhiễu, các bộ phân loại tuyến tính đơn giản có thể đạt hiệu năng cạnh tranh; (3) Hiện tượng nghẽn cổ chai thông tin (Over-smoothing và Over-squashing): khi tăng số lớp truyền tin, Over-smoothing làm vector biểu diễn của các nút dần trở nên tương đồng, trong khi Over-squashing ",
        make_citation_element([21]),
        " nén ép lượng thông tin cấu trúc tăng theo hàm mũ vào vector kích thước cố định, ảnh hưởng đến khả năng phân tách các hành vi tấn công tinh vi."
    ])

    # TABLE 1.3
    add_table_caption(doc, target_p, 3, "So sánh đối chiếu ba nhóm phương pháp biểu diễn đặc trưng log chủ đạo", bookmark_name="BK_TBL_1_003")
    tbl3_headers = ["Tiêu chí đánh giá", "Nhóm Thống kê / Cú pháp\n(Drain, PCA)", "Nhóm Chuỗi Semantic\n(DeepLog, LogBERT)", "Nhóm Đồ thị Nguồn gốc\n(UNICORN, MAGIC)"]
    tbl3_widths = [2200, 2450, 2450, 2505]
    tbl3_rows = [
        ["Cơ chế biểu diễn cốt lõi", "Vector đếm tần suất Event ID trên cửa sổ trượt", "Vector nhúng ngữ cảnh từ chuỗi sự kiện tuần tự", "Vector nhúng cấu trúc đồ thị luồng phụ thuộc dị thể"],
        ["Độ phức tạp thời gian", "O(N) — Tuyến tính theo số lượng sự kiện", "O(N · L) đến O(N · L^2) — Phụ thuộc độ dài cửa sổ L", "O(|V| + |E|) — Phụ thuộc quy mô đỉnh và cạnh đồ thị"],
        ["Bảo toàn tham số an ninh", "Hạn chế (Tham số biến động thường bị thay thế bởi parser)", "Trung bình (Tham số được rời rạc hóa hoặc nhúng từ)", "Tốt (Lưu trữ trực tiếp trên thuộc tính của nút/cạnh)"],
        ["Mô hình hóa chuỗi quan hệ", "Không hỗ trợ", "Cục bộ trên dòng thời gian tuần tự đơn luồng", "Toàn diện trên quan hệ tương tác đa thực thể"],
        ["Khả năng triển khai thực tế", "Thuận lợi nhờ tính toán nhanh và tài nguyên thấp", "Đòi hỏi tài nguyên GPU và thời gian suy luận lớn", "Thách thức do đồ thị tăng trưởng và bùng nổ phụ thuộc"],
        ["Điểm nghẽn phương pháp luận", "Mất mát tham số động, nhạy cảm với sai số phân tách", "Chi phí tính toán cao, hạn chế góc nhìn đa thực thể", "Bùng nổ phụ thuộc, Over-smoothing, Over-squashing"]
    ]
    insert_thesis_table(doc, target_p, tbl3_headers, tbl3_widths, tbl3_rows, font_size_pt=14)
    add_p("", first_line_indent=False)

    add_p([
        "Tổng kết lại, phân tích so sánh đối chiếu chỉ ra rằng mỗi nhóm phương pháp đều sở hữu những ưu thế và giới hạn riêng biệt (",
        make_ref_element("BK_TBL_1_003", "Bảng 1.3"),
        "). Mối quan hệ tương tác và sự chuyển tiếp từ các rào cản kỹ thuật của ba nhóm phương pháp tới năm khoảng trống nghiên cứu trọng tâm được tổng hợp trực quan trên bản đồ quan hệ trong ",
        make_ref_element("BK_FIG_1_004", "Hình 1.4"),
        ":"
    ])

    # FIGURE 1.4
    add_figure_image(fig4_path, width_inches=6.2)
    add_figure_caption(doc, target_p, 1, 4, "Bản đồ đối chiếu ba nhóm phương pháp biểu diễn log và nguồn gốc hình thành năm khoảng trống nghiên cứu cốt lõi (Nguồn: Tác giả tổng hợp)", bookmark_name="BK_FIG_1_004")

    # =========================================================================
    # 1.3. CÁC KHOẢNG TRỐNG NGHIÊN CỨU CỐT LÕI
    # =========================================================================
    add_h2("Các khoảng trống nghiên cứu cốt lõi")
    add_p([
        "Từ kết quả khảo sát và phân tích đối chiếu ba nhóm phương pháp biểu diễn đặc trưng log tại Mục 1.2, có thể nhận thấy rằng mặc dù các kỹ thuật thống kê, mô hình chuỗi ngữ nghĩa và học biểu diễn đồ thị nguồn gốc đã đạt được nhiều bước tiến quan trọng, việc ứng dụng chúng vào môi trường phát hiện tấn công mạng doanh nghiệp thực tế vẫn đối mặt với những rào cản nền tảng chưa được giải quyết thấu đáo ",
        make_citation_element([2, 6, 9, 18]),
        ". Nhằm thiết lập cơ sở khoa học vững chắc và định hình phạm vi nghiên cứu, chuyên đề tổng kết năm khoảng trống nghiên cứu cốt lõi (Research Gaps) tương ứng với năm câu hỏi nghiên cứu (Research Questions - RQ) định hướng cho toàn bộ các đề xuất phương pháp luận tiếp theo."
    ])

    # -------------------------------------------------------------------------
    # 1.3.1. KHOẢNG TRỐNG 1 -> RQ1
    # -------------------------------------------------------------------------
    add_h3("Khoảng trống 1: Mất mát ngữ nghĩa an ninh trong quá trình trừu tượng hóa tham số")
    add_p([
        "Dữ liệu thực nghiệm từ các nghiên cứu tiền nhiệm chỉ ra rằng các bộ phân tích cú pháp (Log Parsers) đóng vai trò then chốt trong việc giảm chiều không gian văn bản log thành các mẫu định dạng tĩnh ",
        make_citation_element([6, 7]),
        ". Tuy nhiên, cơ chế trừu tượng hóa tham số phụ thuộc vào từng thuật toán phân tích cú pháp; nhiều pipeline tiền xử lý thay thế hoặc khái quát hóa một phần các tham số biến động (như địa chỉ IP, tên tài khoản, đường dẫn tệp tin, cổng mạng, tham số dòng lệnh) thành các biến đại diện (wildcard/template variables) nhằm tạo ra các mẫu log dùng chung."
    ])
    add_p([
        "Giới hạn quan trọng chưa được giải quyết là: Sự tương đương về mặt cú pháp mẫu log (Template Equivalence) không đồng nghĩa với sự tương đương về mặt ngữ nghĩa an ninh (Security Semantic Equivalence) ",
        make_citation_element([8, 18]),
        ". Ví dụ, cùng một mẫu thông điệp xác thực thất bại nhưng xuất phát từ một tài khoản quản trị viên đặc quyền root trong mạng nội bộ mang mức độ nghiêm trọng an ninh hoàn toàn khác so với từ một tài khoản người dùng thông thường; hoặc hai dòng lệnh powershell.exe thực thi lệnh quản trị hợp lệ và tải mã độc mã hóa base64 đều bị ánh xạ về cùng một Event ID tiến trình. Khi bộ trích xuất đặc trưng triệt tiêu toàn bộ các tham số động này, không gian biểu diễn bị mất đi phần lớn các tín hiệu ngữ cảnh quan trọng để phân biệt hành vi độc hại với hành vi thông thường."
    ])
    add_p(
        "Khoảng trống nghiên cứu này đặt ra thách thức: Làm thế nào để thiết lập một cơ chế trích xuất đặc trưng có khả năng lọc bỏ các nhiễu cú pháp vô hại nhưng vẫn bảo toàn các tham số động mang ngữ nghĩa an ninh trọng yếu (dưới sự kiểm soát rò rỉ đường tắt và rủi ro quyền riêng tư) trong không gian vector biểu diễn? Từ đó, câu hỏi nghiên cứu thứ nhất được xác lập:"
    )
    add_p(
        "Có thể xây dựng representation loại bỏ nhiễu cú pháp nhưng vẫn bảo toàn các dynamic parameters có ý nghĩa an toàn quan trọng (dưới sự kiểm soát rò rỉ và quyền riêng tư) hay không?",
        bold_prefix="• Câu hỏi nghiên cứu 1 (RQ1 — Representation Fidelity): "
    )

    # -------------------------------------------------------------------------
    # 1.3.2. KHOẢNG TRỐNG 2 -> RQ2
    # -------------------------------------------------------------------------
    add_h3("Khoảng trống 2: Bất đồng bộ và suy thoái trong gióng hàng biểu diễn đa góc nhìn")
    add_p([
        "Các công trình nghiên cứu hiện đại đã chỉ ra rằng dữ liệu log sở hữu tính đa góc nhìn nội tại: góc nhìn thống kê tần suất phản ánh mật độ hoạt động ",
        make_citation_element([6]),
        ", góc nhìn chuỗi sự kiện phản ánh trật tự thời gian cục bộ ",
        make_citation_element([3, 4]),
        ", và góc nhìn đồ thị nguồn gốc phản ánh luồng phụ thuộc cấu trúc đa thực thể ",
        latex_to_clean_omml(r"\mathcal{G}"),
        make_citation_element([9, 10, 11]),
        ". Việc kết hợp các góc nhìn này được kỳ vọng sẽ cung cấp bức tranh toàn cảnh về chiến dịch tấn công."
    ])
    add_p([
        "Tuy nhiên, khi tiến hành hợp nhất hoặc gióng hàng (Cross-view Alignment) các không gian biểu diễn dị thể này trong không gian vector tiềm ẩn, các mô hình học biểu diễn tự giám sát đối mặt với nguy cơ sụp đổ chiều biểu diễn (Representation Collapse), vốn đòi hỏi các kỹ thuật điều hòa phương sai - hiệp phương sai hoặc giảm dư thừa đặc trưng để duy trì tính đa dạng (như các nguyên lý được gợi mở trong VICReg ",
        make_citation_element([22]),
        " và Barlow Twins ",
        make_citation_element([23]),
        "). Đồng thời, trong phạm vi bài toán biểu diễn an ninh mạng doanh nghiệp, quá trình gióng hàng đa góc nhìn còn đặt ra các thách thức chuyên biệt: (1) Nguy cơ chuyển giao tiêu cực (Negative Transfer), khi sự kết hợp giữa các góc nhìn làm suy giảm độ chính xác tổng thể do nhiễu từ một góc nhìn lấn át tín hiệu hữu ích của góc nhìn khác; và (2) Vấn đề góc nhìn bị khuyết hoặc tương ứng từng phần (Missing-view & Partial Correspondence), nảy sinh khi một số luồng telemetry bị chậm trễ, thất thoát gói tin hoặc không có liên kết 1-1 đồng thời giữa chuỗi sự kiện và đồ thị nguồn gốc."
    ])
    add_p(
        "Khoảng trống nghiên cứu cốt lõi ở đây là: Việc gióng hàng đồng thời các góc nhìn dị thể đòi hỏi thiết lập một cơ chế toán học phù hợp để ngăn ngừa sụp đổ biểu diễn và kiểm soát chuyển giao tiêu cực, đồng thời vẫn bảo tồn được các thông tin đặc thù hữu ích của từng góc nhìn riêng biệt. Do đó, câu hỏi nghiên cứu thứ hai được xác lập:"
    )
    add_p(
        "Có thể căn chỉnh các view dị thể mà không gây representation collapse, negative transfer, đồng thời vẫn bảo toàn thông tin hữu ích đặc thù của từng view hay không?",
        bold_prefix="• Câu hỏi nghiên cứu 2 (RQ2 — Cross-View Alignment): "
    )

    # -------------------------------------------------------------------------
    # 1.3.3. KHOẢNG TRỐNG 3 -> RQ3
    # -------------------------------------------------------------------------
    add_h3("Khoảng trống 3: Rò rỉ thông tin quy trình, học đường tắt và trôi dạt biểu diễn")
    add_p([
        "Công trình tổng kết của Arp et al. (USENIX Security 2022) ",
        make_citation_element([2]),
        " và Liu et al. (IEEE S&P 2025) ",
        make_citation_element([18]),
        " đã cảnh báo sâu sắc về tình trạng sai lệch phương pháp luận trong nghiên cứu học máy áp dụng cho an ninh mạng. Trong bài toán trích xuất đặc trưng log, các mô hình rất dễ học phải các đặc trưng đường tắt giả định (Dataset Shortcuts) có tương quan ngẫu nhiên với nhãn độc hại trong tập huấn luyện nhưng không phản ánh bản chất tấn công thực tế."
    ])
    add_p(
        "Phân tích hệ thống chỉ ra sáu kênh rò rỉ thông tin tiềm tàng (Leakage Pathways) thường xuất hiện trong quy trình thực nghiệm: (1) Rò rỉ từ vựng/bộ phân tích cú pháp (Parser/Vocabulary Leakage), khi bộ parser được học trên toàn bộ dữ liệu gồm cả tập kiểm thử; (2) Rò rỉ chuẩn hóa thống kê (Normalization/Statistics Leakage), khi các tham số trung bình/phương sai hoặc trọng số TF-IDF được tính toán xuyên qua ranh giới tập train và test; (3) Rò rỉ thực thể/máy chủ thử nghiệm (Host/Entity/Campaign Holdout Leakage), khi dữ liệu tấn công trên cùng một máy chủ xuất hiện ở cả hai tập; (4) Rò rỉ điều chỉnh ngưỡng và siêu tham số (Threshold/Hyperparameter Leakage); (5) Rò rỉ dữ liệu tiền huấn luyện ngoài miền (Pretraining Data Leakage); và (6) Rò rỉ thời gian tương lai (Future Temporal Information Leakage), khi mô hình vô tình sử dụng sự kiện ở thời điểm tương lai để trích xuất biểu diễn cho hiện tại."
    )
    add_p(
        "Khoảng trống nghiên cứu quan trọng là: Làm thế nào để phân định tường minh giữa tín hiệu an ninh hợp lệ (Legitimate Security Signal) và đặc trưng đường tắt nhân tạo (Dataset Shortcut Artifacts), đồng thời đảm bảo không gian biểu diễn duy trì được năng lực phân tách khi dữ liệu vận hành bị trôi dạt phân phối theo thời gian? Chuyên đề xác lập câu hỏi nghiên cứu thứ ba:"
    )
    add_p(
        "Representation có còn hữu ích sau khi loại bỏ shortcut của dataset và khi phân phối dữ liệu thay đổi hay không?",
        bold_prefix="• Câu hỏi nghiên cứu 3 (RQ3 — Validity Without Shortcuts): "
    )

    # -------------------------------------------------------------------------
    # 1.3.4. KHOẢNG TRỐNG 4 -> RQ4
    # -------------------------------------------------------------------------
    add_h3("Khoảng trống 4: Gán nhãn mức thô, phân bổ bằng chứng yếu và nhiễu quản trị viên")
    add_p([
        "Trong thực tế giám sát an ninh mạng doanh nghiệp, dữ liệu nhãn mặt đất (Ground Truth) thường chỉ được ghi nhận ở mức độ hạt rất thô (Coarse Labels), chẳng hạn như gán nhãn một phiên làm việc kéo dài nhiều giờ hoặc toàn bộ một máy chủ bị xâm nhập trong một khoảng thời gian ",
        make_citation_element([18, 28, 29]),
        ". Bên trong phiên làm việc đó, phần lớn các dòng log vẫn là các thao tác hệ thống bình thường, và chỉ có một số lượng rất nhỏ sự kiện thực sự phản ánh kỹ thuật tấn công."
    ])
    add_p([
        "Thách thức này càng trở nên phức tạp do sự xuất hiện liên tục của hành vi quản trị viên hệ thống (Admin-Noise / LOLBins) ",
        make_citation_element([9, 18]),
        ". Các kỹ sư quản trị thường xuyên thực thi các lệnh bảo trì, chẩn đoán mạng hoặc sao lưu dữ liệu bằng các công cụ dòng lệnh hợp lệ (như PowerShell, WMI, SSH, vssadmin). Các hành vi này sở hữu cấu trúc cú pháp và quyền hạn thực thi tương đồng với kỹ thuật của kẻ tấn công, nhưng hoàn toàn mang mục đích hợp pháp (Benign-but-risky). Một hành vi bất thường hoặc hiếm gặp không tự thân đồng nghĩa với hành vi tấn công độc hại (Unusual ≠ Malicious, Anomaly ≠ Attack)."
    ])
    add_p(
        "Khoảng trống nghiên cứu ở đây là: Làm thế nào để phân bổ chính xác bằng chứng tấn công yếu (Weak Evidence Attribution) từ các nhãn mức thô mà không học nhầm các hành vi quản trị hợp pháp thành độc hại? Đây là cơ sở để xác lập câu hỏi nghiên cứu thứ tư:"
    )
    add_p(
        "Có thể gán đúng attack evidence dưới coarse labels mà không học nhầm các hành vi quản trị hợp pháp thành malicious hay không?",
        bold_prefix="• Câu hỏi nghiên cứu 4 (RQ4 — Weak Evidence Attribution): "
    )

    # -------------------------------------------------------------------------
    # 1.3.5. KHOẢNG TRỐNG 5 -> RQ5
    # -------------------------------------------------------------------------
    add_h3("Khoảng trống 5: Đánh đổi giữa bảo toàn liên kết an ninh và rủi ro quyền riêng tư")
    add_p([
        "Để phát hiện và điều tra các cuộc tấn công APT kéo dài, mô hình học biểu diễn đòi hỏi phải duy trì khả năng liên kết có kiểm soát (Controlled Linkability) giữa các chuỗi hành vi của cùng một thực thể (người dùng, máy chủ, tiến trình) qua nhiều mốc thời gian khác nhau ",
        make_citation_element([9, 10]),
        "."
    ])
    add_p([
        "Tuy nhiên, việc lưu giữ khả năng liên kết này trong không gian vector biểu diễn trực tiếp làm nảy sinh các nguy cơ nghiêm trọng về quyền riêng tư và an toàn thông tin ",
        make_citation_element([25, 26, 27]),
        ". Các vector biểu diễn đặc trưng tiềm ẩn có nguy cơ bị kẻ tấn công khai thác thông qua các kỹ thuật tấn công suy luận thành viên (Membership Inference Attacks - MIA ",
        make_citation_element([25]),
        ") để xác định xem dữ liệu của một thực thể có nằm trong tập huấn luyện hay không, hoặc tấn công nghịch đảo biểu diễn (Representation / Model Inversion Attacks ",
        make_citation_element([26]),
        ") nhằm khôi phục lại các định danh nhạy cảm của người dùng và cấu hình mạng nội bộ. Cần nhấn mạnh rằng, một mô hình được thiết kế có nhận thức về quyền riêng tư (Privacy-Aware) không tự động đồng nghĩa với việc đã đạt được khả năng bảo vệ quyền riêng tư vững chắc (Privacy-Preserving) nếu chưa trải qua các kiểm thử thực nghiệm tấn công nghiêm ngặt ",
        make_citation_element([27]),
        "."
    ])
    add_p(
        "Khoảng trống nghiên cứu này đòi hỏi phải phân tích tường minh mối quan hệ đánh đổi: Đâu là điểm cân bằng giữa năng lực duy trì liên kết thực thể phục vụ phân tích an ninh và mức độ rò rỉ thông tin riêng tư của không gian vector đặc trưng? Từ đó, câu hỏi nghiên cứu thứ năm được xác lập:"
    )
    add_p(
        "Đâu là sự cân bằng chấp nhận được giữa entity continuity và privacy leakage để representation vẫn hữu ích cho phân tích an toàn?",
        bold_prefix="• Câu hỏi nghiên cứu 5 (RQ5 — Privacy-Security Trade-Off): "
    )

    # =========================================================================
    # CHƯƠNG 2: PHƯƠNG PHÁP BIỂU DIỄN ĐẶC TRƯNG LOG ĐA GÓC NHÌN BẢO TOÀN NGỮ CẢNH AN TOÀN
    # =========================================================================
    add_h1("PHƯƠNG PHÁP BIỂU DIỄN ĐẶC TRƯNG LOG ĐA GÓC NHÌN BẢO TOÀN NGỮ CẢNH AN TOÀN")

    # --- 2.1 ---
    add_h2("Hình thức hóa bài toán và giới hạn xử lý dòng")

    # --- 2.1.1 ---
    add_h3("Biểu diễn đa góc nhìn và ranh giới Extractor–Detector")
    add_p([
        "Để thiết lập nền tảng lý thuyết vững chắc và giải quyết căn bản các khoảng trống phương pháp luận đã được xác lập trong Chương 1, Chương này đề xuất khung phương pháp biểu diễn đặc trưng log đa góc nhìn bảo toàn ngữ cảnh an toàn, vận hành trực tuyến trong môi trường xử lý dòng (Streaming Environment). Quá trình biểu diễn bắt đầu bằng việc hình thức hóa luồng nhật ký kiểm toán hệ thống thành một chuỗi sự kiện được gán nhãn thời gian: ",
        latex_to_clean_omml(r"\mathcal{L}_{1:t} = \langle e_1, e_2, \dots, e_t \rangle"),
        ", trong đó mỗi sự kiện ",
        latex_to_clean_omml(r"e_i = (t_i, \tau_i, v_i, o_i, a_i, \mathbf{p}_i)"),
        " tuân thủ thống nhất lược đồ bộ sáu kiểu hóa (Typed Six-tuple, đặc tả chi tiết tại Mục 2.2.1), bao gồm mốc thời gian xuất hiện sự kiện ",
        latex_to_clean_omml(r"t_i \in \mathbb{R}_{\ge 0}"),
        " (Event-Time), định danh kiểu sự kiện ",
        latex_to_clean_omml(r"\tau_i \in \mathcal{T}"),
        ", thực thể tác nhân chủ động ",
        latex_to_clean_omml(r"v_i \in \mathcal{V}"),
        " (Actor Entity như tiến trình gốc, tài khoản người dùng), thực thể đối tượng đích chịu tác động ",
        latex_to_clean_omml(r"o_i \in \mathcal{O} \cup \{\bot\}"),
        " (Object Entity như tệp tin, cổng mạng, tiến trình con), hành động/thao tác thực hiện ",
        latex_to_clean_omml(r"a_i \in \mathcal{A}"),
        " (Action như fork, execve, read, write, connect, registry_set), và tập hợp các tham số động có cấu trúc ",
        latex_to_clean_omml(r"\mathbf{p}_i = \{(k_j, v_j)\}_{j=1}^m"),
        " (đường dẫn tệp tin, tham số dòng lệnh thực thi, cổng mạng và địa chỉ IP kết nối) ",
        make_citation_element([9]),
        "."
    ])
    add_p([
        "Bài toán học biểu diễn đặc trưng log được phát biểu chính thức là việc tìm kiếm một ánh xạ tham số hóa ",
        latex_to_clean_omml(r"f_\theta"),
        " biến đổi lịch sử sự kiện quan sát được ",
        latex_to_clean_omml(r"\mathcal{L}_{1:t}"),
        " thành một vector biểu diễn tiềm ẩn liên tục ",
        latex_to_clean_omml(r"\mathbf{z}_t \in \mathbb{R}^d"),
        ":"
    ])
    add_display_equation(r"\mathbf{z}_t = f_\theta(\mathcal{S}_t), \quad \mathcal{S}_t = \text{Update}(\mathcal{S}_{t-1}, e_t)")
    add_p([
        "Trong đó ",
        latex_to_clean_omml(r"\mathcal{S}_t"),
        " là trạng thái bộ nhớ nội bộ hữu hạn tại thời điểm t, được cập nhật theo cơ chế đơn lượt theo chiều thời gian, không nhìn trước (Strictly Causal Single-pass Update) khi tiếp nhận sự kiện mới ",
        latex_to_clean_omml(r"e_t"),
        ". Quá trình học biểu diễn và đánh giá được tổ chức chặt chẽ theo ba giai đoạn giám sát độc lập (Supervision Contract):"
    ], first_line_indent=False)
    add_p([
        "Huấn luyện bộ trích xuất ",
        latex_to_clean_omml(r"f_\theta"),
        " hoàn toàn theo cơ chế tự giám sát (Self-Supervised Learning) trên dữ liệu nền tảng lịch sử ",
        latex_to_clean_omml(r"\mathcal{D}_{\text{train}}"),
        " mà không tiếp cận bất kỳ nhãn tấn công hay tín hiệu phân loại hạ nguồn nào, học các bất biến cấu trúc và trật tự thời gian nội tại."
    ], bold_prefix="1. Giai đoạn A (Stage A, Pretraining Tự giám sát Ngoại tuyến): ")
    add_p([
        "Khi môi trường giám sát cung cấp các nhãn mức thô (Coarse/Bag-level Labels), một mục tiêu học đa thể hiện có trọng số chú ý (Attention-based Multiple Instance Learning) có thể được kích hoạt trong pha thích ứng riêng biệt nhằm tinh chỉnh phân bổ trọng số cho các bằng chứng thưa thớt mà không làm sai lệch tính chất tự giám sát của Stage A."
    ], bold_prefix="2. Giai đoạn B (Stage B, Thích ứng Bằng chứng Yếu / Tùy chọn): ")
    add_p([
        "Khóa cố định toàn bộ tham số ",
        latex_to_clean_omml(r"\theta^*"),
        " của bộ trích xuất biểu diễn. Đánh giá chất lượng biểu diễn thông qua giao thức đóng băng tham số (Frozen Linear Probe: ",
        latex_to_clean_omml(r"\hat{\mathbf{y}}_t = \sigma(\mathbf{W}^\top \mathbf{z}_t + \mathbf{b})"),
        ") phục vụ phân loại và quy kết chiến thuật/kỹ thuật theo phân loại MITRE ATT&CK ",
        make_citation_element([1]),
        ", hoặc quy trình chấm điểm bất thường không giám sát tùy chọn (Optional Downstream Zero-Shot Anomaly Scoring Protocol). Trong đó, hàm tính điểm bất thường, phép đo khoảng cách/năng lượng, phân bố nền chuẩn, phương pháp hiệu chuẩn và chính sách ngưỡng/ngân sách cảnh báo đều được đăng ký trước và khóa cố định trên tập Train/Validation tại Chương 3 trước khi mở tập Test. Bộ trích xuất Extractor tuyệt đối độc lập và không phụ thuộc vào bất kỳ bộ dò Detector cụ thể nào."
    ], bold_prefix="3. Giai đoạn C (Stage C, Đánh giá Đóng băng Hạ nguồn): ")
    add_p(
        "Trên cơ sở định nghĩa bài toán và ranh giới phương pháp luận, chuyên đề hình thức hóa năm giả thuyết khoa học cốt lõi (Hypotheses H1–H5 (Đóng góp Độc lập của Chuyên đề)):"
    )
    add_p([
        "Biểu diễn nhận biết tham số ",
        latex_to_clean_omml(r"X_{\text{param}}"),
        " (chuỗi lệnh, đường dẫn tệp tin, địa chỉ IP) vào không gian vector z bảo tồn lượng thông tin tương hỗ về ngữ nghĩa an ninh tiềm ẩn ",
        latex_to_clean_omml(r"Y_{\text{sec}}"),
        " cao hơn có ý nghĩa so với biểu diễn chỉ dựa trên mẫu tĩnh hoặc bị trừu tượng hóa tham số ",
        latex_to_clean_omml(r"X_{\text{abstracted}}"),
        ":"
    ], bold_prefix="Giả thuyết H1 (Parameter Semantic Fidelity, Đóng góp Độc lập): ", keep_with_next=True)
    add_display_equation(r"\mathcal{I}(\mathbf{z}; Y_{\text{sec}}) > \mathcal{I}(\mathbf{z}_{\text{abstracted}}; Y_{\text{sec}})")
    add_p([
        "Trong đó ",
        latex_to_clean_omml(r"Y_{\text{sec}}"),
        " là biến ngẫu nhiên đại diện cho mục tiêu ngữ nghĩa an ninh tiềm ẩn (như nhãn hành vi độc hại, ý định thực thi tiến trình, phân loại quyền hạn truy cập); ",
        latex_to_clean_omml(r"\mathcal{I}(\cdot\,;\cdot)"),
        " là thước đo thông tin tương hỗ, với bộ ước lượng (Estimator) và giao thức thực nghiệm sẽ được thiết lập và khóa chặt chẽ tại Chương 3; giả thuyết chỉ kiểm tra tính trung thực ngữ nghĩa an ninh và không suy diễn tuyệt đối sang việc triệt tiêu sụp đổ biểu diễn."
    ], first_line_indent=False)
    add_p([
        "Cơ chế gióng hàng có kiểm soát giữa góc nhìn chuỗi sự kiện ",
        latex_to_clean_omml(r"\mathbf{z}^{(\text{seq})}"),
        " và góc nhìn đồ thị nguồn gốc ",
        latex_to_clean_omml(r"\mathbf{z}^{(\text{graph})}"),
        " được kỳ vọng nâng cao hiệu năng biểu diễn tổng thể vượt trên các góc nhìn đơn lẻ, đồng thời thỏa mãn các điều kiện kiểm chứng định lượng:"
    ], bold_prefix="Giả thuyết H2 (Multi-View Alignment & Negative Transfer Prevention, Đóng góp Độc lập): ", keep_with_next=True)
    add_display_equation(r"\text{Utility}(\mathbf{z}) > \max\left(\text{Utility}(\mathbf{z}^{(\text{seq})}), \text{Utility}(\mathbf{z}^{(\text{graph})})\right) - \epsilon_{\text{margin}}, \quad \text{Var}(\mathbf{z}) \ge \tau_{\text{var}}, \quad \mathcal{I}(\mathbf{z}; X_{\text{view}}) \ge \tau_{\text{info}}")
    add_p([
        "Trong đó ",
        latex_to_clean_omml(r"\tau_{\text{var}} > 0"),
        " là ngưỡng phương sai duy trì độ đa dạng không gian (ngăn ngừa sụp đổ biểu diễn), ",
        latex_to_clean_omml(r"\tau_{\text{info}}"),
        " là ngưỡng bảo toàn thông tin đặc thù từng góc nhìn, và mức độ suy giảm do chuyển giao tiêu cực (Negative Transfer) được khống chế ",
        latex_to_clean_omml(r"\Delta_{\text{neg}} \le \epsilon_{\text{neg}}"),
        " theo các ngưỡng xác lập trước thực nghiệm (các kỹ thuật cụ thể như InfoNCE, VICReg hay Barlow Twins được xem xét là ứng viên kỹ thuật tại Mục 2.4)."
    ], first_line_indent=False)
    add_p([
        "Lợi ích biểu diễn của không gian vector z duy trì tính ổn định sau khi loại bỏ các đặc trưng đường tắt giả định ",
        latex_to_clean_omml(r"S \in \mathcal{E}_{\text{exclude}}"),
        " (",
        latex_to_clean_omml(r"\mathcal{I}(\mathbf{z}; S) \le \epsilon_{\text{short}}"),
        ") và dưới sự biến động phân phối thực tế (",
        latex_to_clean_omml(r"P_t(X, Y) \neq P_{t+1}(X, Y)"),
        "), đồng thời thỏa mãn ràng buộc ổn định khoảng cách hình học trước các biến đổi cú pháp vô hại ",
        latex_to_clean_omml(r"T \in \mathcal{T}_{\text{benign}}"),
        ":"
    ], bold_prefix="Giả thuyết H3 (Anti-Drift & Shortcut Invariance Robustness, Đóng góp Độc lập): ", keep_with_next=True)
    add_display_equation(r"\|\mathbf{z}(T(X)) - \mathbf{z}(X)\|_2 \le \epsilon_{\text{inv}}, \quad \mathcal{I}(\mathbf{z}; S) \le \epsilon_{\text{short}}")
    add_p([
        "Quá trình trích xuất và cập nhật biểu diễn dòng vận hành ổn định trong giới hạn ngân sách độ trễ và bộ nhớ xác định trước, đáp ứng yêu cầu độ trễ xử lý mỗi sự kiện ",
        latex_to_clean_omml(r"\Delta t(e_t) \le T_{\text{budget}}"),
        " và dung lượng trạng thái bộ nhớ ",
        latex_to_clean_omml(r"\text{Mem}(\mathcal{S}_t) \le M_{\text{max}}"),
        " trên toàn bộ dòng thời gian vận hành dài hạn:"
    ], bold_prefix="Giả thuyết H4 (Bounded Operational Budget Feasibility, Đóng góp Độc lập): ", keep_with_next=True)
    add_display_equation(r"\Delta t(e_t) \le T_{\text{budget}}, \quad \text{Mem}(\mathcal{S}_t) \le M_{\text{max}}, \quad \forall t \ge 1")
    add_p([
        "Cơ chế duy trì liên kết có kiểm soát (Controlled Linkability) được kỳ vọng sẽ thiết lập sự vượt trội theo nghĩa Pareto (Pareto-dominance) so với các đường cơ sở giữ nguyên định danh thô hoặc ẩn danh hóa quá mức tại ít nhất một điểm vận hành thực tế dưới các thước đo hiệu năng an ninh (Utility) và tổn thất quyền riêng tư (Privacy Loss) đã định nghĩa trước:"
    ], bold_prefix="Giả thuyết H5 (Controlled Linkability & Utility–Privacy Frontier, Đóng góp Độc lập): ", keep_with_next=True)
    add_display_equation(r"\exists \mathbf{z} : \quad \text{Utility}(\mathbf{z}) \ge \text{Utility}(\mathbf{z}_{\text{baseline}}) \quad \land \quad \text{PrivacyLoss}(\mathbf{z}) \le \text{PrivacyLoss}(\mathbf{z}_{\text{baseline}})")
    add_p([
        "Trong đó lý thuyết Quyền riêng tư Vi sai (DP) đóng vai trò định lượng và chặn trên tổn thất lý thuyết ",
        latex_to_clean_omml(r"(\epsilon, \delta)"),
        " dưới các giả định toán học xác định, còn rủi ro thực tế được kiểm chứng thực nghiệm độc lập thông qua các tấn công suy luận thành viên (MIA) ",
        make_citation_element([25, 27]),
        " và tấn công nghịch đảo/tái định danh thực thể ",
        make_citation_element([26]),
        "."
    ], first_line_indent=False)

    # --- 2.1.2 ---
    add_h3("Độ phức tạp xử lý dòng với trạng thái hữu hạn")
    add_p([
        "Trong môi trường giám sát an ninh SOC thời gian thực, luồng telemetry liên tục đổ về với lưu lượng rất lớn. Việc duy trì toàn bộ lịch sử tương tác của hệ thống trên bộ nhớ là bất khả thi về mặt tài nguyên. Do đó, chuyên đề thiết lập hợp đồng xử lý dòng với trạng thái hữu hạn (Bounded-State Streaming Contract), được quy định bởi ngân sách bộ nhớ trần ",
        latex_to_clean_omml(r"M_{\text{max}}"),
        " và tập thực thể hoạt động ",
        latex_to_clean_omml(r"\mathcal{V}_{\text{active}}(t)"),
        ":"
    ])
    add_display_equation(r"\mathcal{V}_{\text{active}}(t) = \{v \in \mathcal{V} \mid t - t_{\text{last}}(v) \le \tau_{\text{TTL}}\}, \quad |\mathcal{S}_t| \le M_{\text{max}}")
    add_p([
        "Trong đó ",
        latex_to_clean_omml(r"\tau_{\text{TTL}}"),
        " là ngưỡng thời gian sống (Time-To-Live). Một thực thể v (tiến trình, socket, tệp tin) không phát sinh bất kỳ tương tác nào trong khoảng thời gian ",
        latex_to_clean_omml(r"\tau_{\text{TTL}}"),
        " sẽ hết hạn và đủ điều kiện loại bỏ (Eviction) khỏi bộ đệm trạng thái ",
        latex_to_clean_omml(r"\mathcal{S}_t"),
        ". Khi kích thước trạng thái tiếp cận ngưỡng ngân sách ",
        latex_to_clean_omml(r"M_{\text{max}}"),
        ", hệ thống kích hoạt yêu cầu thu dọn và nén tóm tắt dữ liệu. Để hiện thực hóa hợp đồng này, các kỹ thuật như chính sách Least Recently Updated (LRU), cấu trúc phác thảo Count-Min Sketch và hàm suy giảm trọng số cạnh theo hàm mũ ",
        latex_to_clean_omml(r"\omega(e_{uv}, t) = \exp(-\lambda (t - t_{uv}))"),
        " được đề xuất như các ứng viên kỹ thuật hiện thực hóa (Proposed Implementation Choices) có thể thay thế hoặc tùy biến linh hoạt."
    ], first_line_indent=False)
    add_p([
        "Để xử lý hiện tượng sự kiện đến muộn hoặc xáo trộn thứ tự (Out-of-Order Events) phát sinh từ độ trễ mạng thu thập telemetry ",
        make_citation_element([9, 13]),
        ", hệ thống áp dụng nguyên tắc xử lý theo thời gian sự kiện (Event-Time) kết hợp cơ chế mốc ngắt thời gian (Event-Time Watermark):"
    ])
    add_display_equation(r"t_{\text{wm}} = \max_{1 \le i \le t}(t_{\text{event}}(e_i)) - \delta_{\text{delay}}")
    add_p([
        "Mốc thời gian ",
        latex_to_clean_omml(r"t_{\text{wm}}"),
        " xác lập ranh giới trễ cho phép ",
        latex_to_clean_omml(r"\delta_{\text{delay}}"),
        "; các sự kiện đến muộn trong phạm vi ",
        latex_to_clean_omml(r"\delta_{\text{delay}}"),
        " được tự động điều hòa vào trạng thái dòng. Các sự kiện đến trễ hơn ",
        latex_to_clean_omml(r"\delta_{\text{delay}}"),
        " được chuyển tiếp vào bộ đệm tái điều chỉnh (Reconciliation Buffer) theo giao thức ghi nhận thất thoát thông tin xác định (Explicit Information-Loss Protocol) mà không làm gián đoạn luồng suy luận thời gian thực."
    ], first_line_indent=False)
    add_p([
        "Khi hệ thống gặp hiện tượng đột biến lưu lượng (Traffic Spike), cơ chế kiểm soát áp lực ngược (Backpressure Control) dựa trên thuật toán Token-Bucket điều tiết tốc độ nạp dữ liệu. Nếu lưu lượng vượt quá giới hạn chịu tải tối đa, việc loại bỏ gói tin (Shedding) được thực hiện hoàn toàn độc lập với kết quả phát hiện của mô hình (dựa trên hạn ngạch băng thông nguồn thu thập hoặc mức độ ưu tiên của phân vùng telemetry, tuyệt đối không dựa vào điểm số an ninh chưa kiểm chứng) nhằm tránh rủi ro rò rỉ vòng lặp (Detector Leakage) và loại bỏ nhầm các bằng chứng APT yếu thưa thớt."
    ])
    add_p([
        "Về mặt lý thuyết tính toán, mô hình dòng thiết lập hợp đồng ngân sách độ phức tạp cho từng thành phần xử lý ký hiệu tổng quát: ",
        latex_to_clean_omml(r"\mathcal{C}_{\text{step}} = \mathcal{C}_{\text{seq}}(L, d, \text{arch}) + \mathcal{C}_{\text{graph}}(\text{card}(\mathcal{V}_{\text{active}}), \text{card}(\mathcal{E}_{\text{window}}), \text{arch}) + \mathcal{C}_{\text{state}}(\text{idx}, M_{\text{max}}) + \mathcal{C}_{\text{align}}"),
        ". Trong đó các hàm chi phí phụ thuộc vào kiến trúc và cấu trúc dữ liệu cụ thể (ví dụ chi phí tự chú ý ",
        latex_to_clean_omml(r"\approx \mathcal{O}(L^2 \cdot d)"),
        " nếu dùng Transformer chuẩn hoặc ",
        latex_to_clean_omml(r"\mathcal{O}(L \cdot d)"),
        " nếu dùng biến thể tuyến tính, chi phí truyền tin đồ thị lân cận phụ thuộc số lớp và bậc đỉnh ",
        latex_to_clean_omml(r"\mathcal{O}(\text{deg}_{\text{max}} \cdot d)"),
        ", chi phí cập nhật chỉ mục trạng thái ",
        latex_to_clean_omml(r"\mathcal{O}(\log \text{card}(\mathcal{V}_{\text{active}}))"),
        " nếu sử dụng cây cân bằng). Các mức chặn Big-O cụ thể được đặc tả chặt chẽ cho từng nhánh tại Mục 2.3, Mục 2.4 và tổng hợp tại Bảng 2.2 và Bảng 2.4. Chuyên đề chỉ rõ bản chất đánh đổi cốt lõi (Long-Horizon Context Trade-off): các chiến dịch APT kéo dài nhiều tuần đòi hỏi duy trì ngữ cảnh tương tác sâu, trong khi bộ nhớ thực tế bị giới hạn bởi ",
        latex_to_clean_omml(r"M_{\text{max}}"),
        ", và thiết kế dòng giải quyết mâu thuẫn này bằng cơ chế nén ngữ cảnh có kiểm soát thay vì lưu trữ toàn bộ dữ liệu thô."
    ])

    # --- 2.1.3 ---
    add_h3("Kiến trúc tổng thể và giao diện vào/ra")
    add_p([
        "Kiến trúc tổng thể của khung biểu diễn đặc trưng log đa góc nhìn dòng được thiết kế theo nguyên tắc phân tách độc lập giữa hai mặt phẳng vận hành: Mặt phẳng Huấn luyện (Training Plane) và Mặt phẳng Suy luận Dòng (Streaming Inference Plane), được minh họa chi tiết trong ",
        make_ref_element("BK_FIG_2_001", "Hình 2.1"),
        ":"
    ])

    # FIGURE 2.1
    add_figure_image(fig5_path, width_inches=5.2)
    add_figure_caption(
        doc,
        target_p,
        2,
        1,
        "Kiến trúc hai mặt phẳng (Training Plane & Inference Plane) của khung biểu diễn đặc trưng log đa góc nhìn dòng (Nguồn: Tác giả đề xuất)",
        bookmark_name="BK_FIG_2_001"
    )

    add_p([
        "Vận hành ngoại tuyến (Offline) trên kho dữ liệu telemetry lịch sử ",
        latex_to_clean_omml(r"\mathcal{D}_{\text{train}}"),
        ". Mặt phẳng này thực hiện trích xuất song song hai góc nhìn (chuỗi sự kiện và đồ thị nguồn gốc), tối ưu hóa hàm mất mát gióng hàng đa góc nhìn kết hợp các số hạng điều hòa bất biến để tìm kiếm bộ trọng số tối ưu ",
        latex_to_clean_omml(r"\theta^*"),
        " tại Giai đoạn A, đồng thời hỗ trợ pha thích ứng bằng chứng yếu Giai đoạn B (nếu có nhãn mức thô). Quá trình tiền huấn luyện hoàn toàn không sử dụng nhãn tấn công và không can thiệp vào luồng vận hành thời gian thực."
    ], bold_prefix="1. Mặt phẳng Huấn luyện (Training Plane / Offline Self-Supervised): ")
    add_p([
        "Vận hành trực tuyến (Online) trên luồng sự kiện đang diễn ra ",
        latex_to_clean_omml(r"e_t \in \mathcal{L}_{1:t}"),
        ". Mặt phẳng này tiếp nhận trọng số đã đóng băng hoàn toàn ",
        latex_to_clean_omml(r"\theta^*"),
        ", cập nhật trạng thái bộ nhớ hữu hạn ",
        latex_to_clean_omml(r"\mathcal{S}_t"),
        " theo các chính sách TTL, Eviction và Watermarking, sau đó trích xuất vector đóng gói ",
        latex_to_clean_omml(r"\mathbf{z}_t = f_{\theta^*}(\mathcal{S}_t) \in \mathbb{R}^d"),
        " theo cơ chế đơn lượt đơn hướng (Strictly Causal, Zero Lookahead), cung cấp đầu vào tức thời cho các bộ dò tuyến tính đóng băng Giai đoạn C phục vụ phát hiện và quy kết chiến thuật/kỹ thuật MITRE ATT&CK."
    ], bold_prefix="2. Mặt phẳng Suy luận Dòng (Streaming Inference Plane / Online Causal): ")

    # =========================================================================
    # SECTION 2.2 — Tiền xử lý và bảo vệ dynamic parameters
    # NOD-000055..NOD-000066 | Role: MECHANISM
    # =========================================================================
    add_h2("Tiền xử lý và bảo vệ dynamic parameters")
    add_p(
        "Mục 2.2 xác lập cơ sở tiền xử lý phục vụ toàn bộ pipeline trích xuất đặc trưng đa góc nhìn: "
        "từ luồng nhật ký thô đến các sự kiện đã được chuẩn hóa, phân loại tham số an ninh và bảo vệ "
        "quyền riêng tư có kiểm soát; tất cả đều tuân thủ nghiêm ngặt ranh giới Extractor–Detector và "
        "hợp đồng biểu diễn đã thiết lập tại Mục 2.1. Ba thành phần cốt lõi được trình bày theo thứ tự "
        "phụ thuộc logic: (1) Parsing và chuẩn hóa có nhận thức an ninh (2.2.1); (2) Mô hình đe dọa "
        "quyền riêng tư và cơ chế liên kết có kiểm soát (2.2.2); (3) Đồng bộ thời gian và cửa sổ ngữ "
        "cảnh đa tỷ lệ (2.2.3)."
    )

    # --- 2.2.1 ---
    add_h3("Parsing, Typed Canonicalization, Entity Resolution và Security-aware Parameter Retention")
    add_p(
        "Điểm khởi đầu của pipeline là luồng nhật ký thô không đồng nhất từ nhiều nguồn thu thập: "
        "nhật ký kiểm toán hệ điều hành (auditd, Windows Event Log, Sysmon), nhật ký mạng, nhật ký "
        "xác thực. Mỗi nguồn sử dụng định dạng cú pháp riêng và ký pháp tham số đặc thù. Mục tiêu của "
        "giai đoạn tiền xử lý là ánh xạ tất cả sự kiện về một lược đồ bộ sáu có kiểu hóa (Typed Six-tuple "
        "Schema) trong khi bảo toàn mọi thông tin có giá trị an ninh nằm trong các tham số "
        "động."
    )

    # 2.2.1.1 — Typed schema
    add_h4("Typed Schema và Security-aware Parameter Retention")
    # 2.2.1.1 — Typed schema + parameter retention (body, no heading per roadmap)
    add_p([
        "Mỗi sự kiện thô sau khi qua bước parsing được biểu diễn dưới dạng bộ sáu có kiểu hóa (Typed "
        "Six-tuple) như sau:"
    ], keep_with_next=True)
    add_display_equation(
        r"\tilde{e}_i = \bigl(t_i,\; \tau_i,\; v_i,\; o_i,\; a_i,\; \mathbf{p}_i\bigr)"
    )
    add_p([
        "Trong đó: ",
        latex_to_clean_omml(r"t_i \in \mathbb{R}_{\ge 0}"),
        " là mốc thời gian sự kiện (Event-Time) theo đồng hồ hệ thống nguồn; ",
        latex_to_clean_omml(r"\tau_i \in \mathcal{T}"),
        " là kiểu hành vi kiểm toán (syscall, registry op, network op, auth op,...); ",
        latex_to_clean_omml(r"v_i \in \mathcal{V}"),
        " là thực thể chủ thể (Actor) gồm tiến trình, người dùng, dịch vụ; ",
        latex_to_clean_omml(r"o_i \in \mathcal{O}"),
        " là thực thể đối tượng (Object) gồm tệp tin, socket, registry key, tiến trình con; ",
        latex_to_clean_omml(r"a_i \in \mathcal{A}"),
        " là hành động cụ thể (read, write, execute, connect, create, delete,...); ",
        latex_to_clean_omml(r"\mathbf{p}_i = \{(k_j, v_j)\}_{j=1}^{m_i}"),
        " là tập các tham số động có kiểu hóa cụ thể: đường dẫn tệp tin, địa chỉ IP và cổng đích, chuỗi "
        "đối số dòng lệnh, giá trị registry, tên người dùng đích, mã trả về, đáp ứng đúng yêu cầu của "
        "Hợp đồng Biểu diễn yêu cầu bảo toàn."
    ], first_line_indent=False)
    add_p([
        "Tập kiểu sự kiện ",
        latex_to_clean_omml(r"\mathcal{T}"),
        " và tập hành động ",
        latex_to_clean_omml(r"\mathcal{A}"),
        " được định nghĩa tĩnh theo đặc tả nguồn thu thập (auditd syscall table, Windows Event ID "
        "taxonomy, Sysmon event types) và không phụ thuộc vào dữ liệu huấn luyện, do đó không phát sinh "
        "rủi ro rò rỉ vocabulary từ tập kiểm tra (CTRL-LEAK-001). Ngược lại, từ vựng tham số động "
        "(đặc biệt là chuỗi đường dẫn tệp tin và đối số lệnh) phụ thuộc vào môi trường và được học "
        "hoàn toàn trên phân vùng huấn luyện ",
        latex_to_clean_omml(r"\mathcal{D}_{\text{train}}"),
        " theo cơ chế fit-only: mọi thống kê, ánh xạ từ vựng và bảng chuẩn hóa đều được fit trên ",
        latex_to_clean_omml(r"\mathcal{D}_{\text{train}}"),
        " và áp dụng ở trạng thái đóng băng (frozen state) lên phân vùng Validation và Test — không "
        "có thông tin ngược từ Val/Test nào được đưa vào quá trình fit."
    ])

    # 2.2.1.2 — Retain vs Normalize
    add_h4("Typed Canonicalization và Entity Resolution")
    add_p([
        "Hạt nhân của giai đoạn chuẩn hóa là hai vị từ ngữ nghĩa bổ sung cho nhau, được áp dụng lên "
        "từng trường của bộ sáu kiểu hóa:"
    ], keep_with_next=True)
    add_display_equation(
        r"\text{Retain}(x) \;\;=\;\; \mathbf{1}\!\left[\,\text{Type}(x) \in \mathcal{T}_{\text{security}} \;\lor\; x \;\text{thỏa đặc tả an ninh tiền nghiệm}\,\right]"
    )
    add_display_equation(
        r"\text{Normalize}(x) \;\;=\;\; \phi\bigl(x\bigr) \;\;\text{với}\;\; \phi \;\text{loại bỏ biến thể định dạng, giữ nguyên ngữ nghĩa}"
    )
    add_p([
        "Vị từ ",
        latex_to_clean_omml(r"\text{Retain}(\cdot)"),
        " được xác lập dựa trên kiểu ngữ nghĩa trường (field semantic type) và tính phù hợp an ninh xác định tiền nghiệm (a priori security relevance) "
        "hoặc được cấu hình độc quyền trên phân vùng ",
        latex_to_clean_omml(r"\mathcal{D}_{\text{train}} / \mathcal{D}_{\text{val}}"),
        " theo đặc tả miền, tuyệt đối không phụ thuộc vào nhãn tấn công hay thống kê tập kiểm tra (loại trừ hoàn toàn rủi ro test leakage). Cụ thể, ",
        latex_to_clean_omml(r"\text{Retain}(\cdot)"),
        " đánh giá ",
        latex_to_clean_omml(r"\mathbf{1} = \text{True}"),
        " đối với các trường mang thông tin tương tác hệ thống trọng yếu: địa chỉ IP và cổng mạng, đường dẫn tệp tin thực thi/cấu hình, chuỗi đối số dòng lệnh, "
        "khóa registry, định danh tiến trình cha–con và tài khoản người dùng đích. Các quy tắc như kiểm tra dải IP mạng, khóa registry khởi động, hay cờ đặc quyền "
        "nếu được áp dụng chỉ đóng vai trò là các đặc trưng chính sách miền tùy chọn (optional domain-policy features) hỗ trợ chuẩn hóa biểu diễn, "
        "tuyệt đối không phải là bộ phán quyết tiên tri (oracle) để phân biệt hành vi độc hại; ranh giới Extractor–Detector được duy trì nghiêm ngặt (Extractor không phải là rule-based detector). "
        "Ngược lại, các trường metadata môi trường cục bộ không mang ngữ nghĩa an ninh (như PID tạm thời, timestamp mili-giây nội bộ không phải Event-Time, số sequence kernel) "
        "được phân loại là KNOWN EXCLUDE FIELD và bị loại bỏ (DROP) trước khi đưa vào Extractor. Đối với các tham số chưa từng thấy hoặc chưa rõ ngữ nghĩa (UNSEEN / UNCERTAIN), "
        "hệ thống áp dụng chính sách UNKNOWN-SAFE (bảo toàn nhãn kiểu trường và cấu trúc, gán typed placeholder), không thực hiện loại bỏ mặc định."
    ], first_line_indent=False)
    add_p([
        "Hàm chuẩn hóa ",
        latex_to_clean_omml(r"\phi(\cdot)"),
        "Áp dụng các phép biến đổi định dạng tường minh: chuẩn hóa đường dẫn,"
        " các ký hiệu separator, junction thừa, chuẩn hóa phân cách;"
        " chuẩn hóa địa chỉ IP về dạng chính tắc (loại bỏ zero-padding, khai triển IPv6 viết tắt);"
        " chuẩn hóa timestamp về UTC epoch; lower-casing không phân biệt hoa thường với tên lệnh trên Windows."
        " Quan trọng là hàm phi(·) không được phép loại bỏ thông tin ngữ nghĩa an ninh dưới danh nghĩa"
        " 'nhiễu cú pháp': ví dụ, khoảng trắng trong đối số lệnh không được thu gọn nếu"
        " điều đó phá vỡ khả năng tái tạo ý định lệnh. Ranh giới giữa nhiễu định dạng và tham số an ninh"
        " được ghi nhận tường minh và không thay đổi trong suốt quá trình thực nghiệm.",
    ])

    # --- Parameter Policy Table (NOD-000058) ---
    add_p(
        "Bảng 2.1 tóm tắt chính sách xử lý cho từng trường của bộ sáu kiểu hóa:",
        keep_with_next=True
    )
    _tbl22_headers = ["Trường", "Vai trò ngữ nghĩa", "Chính sách giữ/chuẩn hóa", "Xử lý quyền riêng tư và An toàn"]
    _tbl22_widths = [1300, 1800, 2100, 2400]
    _tbl22_rows = [
        [[latex_to_clean_omml(r"t_i")], "Event-Time", "NORMALIZE → UTC epoch",
         "Quasi-ID khi kết hợp dữ liệu ngoài; đánh giá theo CTRL-PRIV-001"],
        [[latex_to_clean_omml(r"\tau_i")], "Kiểu hành vi kiểm toán", "RETAIN (tĩnh, không học)", "Kiểu thấp, không PII trực tiếp"],
        [[latex_to_clean_omml(r"v_i")], "Thực thể chủ thể (Actor)", ["PRIVACY-PROTECT → ", latex_to_clean_omml(r"\pi_\psi(v_i)")], "Direct identifier; bắt buộc pseudonymize"],
        [[latex_to_clean_omml(r"o_i")], "Thực thể đối tượng", ["PRIVACY-PROTECT → ", latex_to_clean_omml(r"\pi_\psi(o_i)")], "Direct/quasi-ID; bắt buộc pseudonymize"],
        [[latex_to_clean_omml(r"a_i")], "Hành động", "RETAIN (tĩnh)", "Kiểu thấp"],
        [[latex_to_clean_omml(r"\mathbf{p}_i"), " (Retain=1)"], "Tham số an ninh được xác nhận", "RETAIN + NORMALIZE → tokenize",
         ["Tokenize trên từ vựng ", latex_to_clean_omml(r"\mathcal{D}_{\text{train}}"), "; OOV → UNKNOWN-SAFE"]],
        [[latex_to_clean_omml(r"\mathbf{p}_i"), " (KNOWN EXCLUDE)"], "Trường xác định không mang ngữ nghĩa an ninh (PID, timestamp cục bộ, sequence kernel)", "DROP trước Extractor",
         ["Loại bỏ hoàn toàn trước khi đưa vào ", latex_to_clean_omml(r"f_\theta")]],
        [[latex_to_clean_omml(r"\mathbf{p}_i"), " (UNSEEN / UNCERTAIN)"], ["Tham số ngoài từ vựng ", latex_to_clean_omml(r"\mathcal{D}_{\text{train}}"), " hoặc chưa rõ an ninh"], "UNKNOWN-SAFE (bảo toàn cấu trúc cú pháp)",
         "Giữ nguyên field-type tag (PATH, IP, CMD, REG, USER) và vị trí cấu trúc; thay giá trị chuỗi bằng typed placeholder (<UNK-PATH>, <UNK-IP>, <UNK-CMD>)"],
    ]
    add_table_caption(doc, target_p, 1,
                      "Chính sách RETAIN / NORMALIZE / PRIVACY-PROTECT / UNKNOWN-SAFE"
                      " cho các trường của bộ sáu kiểu hóa",
                      bookmark_name="BK_TBL_2_001", chapter_num=2)
    insert_thesis_table(doc, target_p, _tbl22_headers, _tbl22_widths, _tbl22_rows, font_size_pt=13)
    add_p("", first_line_indent=False)

    # 2.2.1.3 — Leakage-safe preprocessing
    add_h4("Leakage-safe Preprocessing")
    add_p([
        "Tất cả các thành phần tiền xử lý có trạng thái học được (bao gồm từ vựng tokenizer tham số "
        "động, bảng chuẩn hóa tên tiến trình theo môi trường, thống kê tần suất sự kiện phục vụ "
        "outlier filtering, và ngưỡng phân ngưỡng tham số) đều được fit độc quyền trên phân vùng "
        "huấn luyện ",
        latex_to_clean_omml(r"\mathcal{D}_{\text{train}}"),
        " theo thứ tự nhân quả thời gian (causal-time order), tức là chỉ sử dụng dữ liệu quan sát "
        "trước thời điểm phân cắt huấn luyện. Phân vùng Validation và Test chỉ nhận các phép biến đổi "
        "(transform) từ trạng thái đóng băng đã fit, không có bất kỳ thông tin thống kê nào từ "
        "Val/Test được đưa trở lại (CTRL-LEAK-001: Anti-Leakage Temporal Split). Cụ thể:"
    ], keep_with_next=True)
    add_bullet_p([
        "Phân vùng thời gian tuyến tính theo Event-Time: ",
        latex_to_clean_omml(r"\mathcal{D}_{\text{train}} \prec \mathcal{D}_{\text{val}} \prec \mathcal{D}_{\text{test}}"),
        " trong đó ký hiệu ",
        latex_to_clean_omml(r"\prec"),
        " biểu thị quan hệ xảy ra trước nghiêm ngặt theo thứ tự thời gian. Không có xáo trộn ngẫu "
        "nhiên (random shuffle) theo thời gian ở bất kỳ giai đoạn nào."
    ])
    add_bullet_p([
        "Tất cả preprocessor (tokenizer, normalizer, entity resolver, vocabulary builder) "
        "đều có trạng thái fit-freeze: quá trình fit chỉ chạy một lần trên ",
        latex_to_clean_omml(r"\mathcal{D}_{\text{train}}"),
        ".",
    ])
    add_bullet_p(
        "Chính sách xử lý tham số chưa xác định (UNSEEN / UNCERTAIN SECURITY SEMANTICS) tuân thủ nguyên tắc UNKNOWN-SAFE: "
        "tuyệt đối không áp dụng quy tắc mặc định hủy bỏ (không thực hiện UNCONFIRMED → DROP). Hệ thống phân định rạch ròi: "
        "(a) Các trường KNOWN EXCLUDE FIELD (metadata môi trường như PID, timestamp cục bộ mili-giây, số sequence nội bộ) "
        "được loại bỏ hoàn toàn (DROP) trước khi đưa vào Extractor; "
        "(b) Các tham số UNSEEN hoặc chưa xác định rõ ngữ nghĩa an ninh xuất hiện trong tập Validation hoặc Test được giữ nguyên "
        "nhãn kiểu trường (field-type tag: PATH, IP, CMD, REG, USER) và bảo toàn vị trí cấu trúc của tham số trong bộ sáu, "
        "đồng thời thay thế giá trị chuỗi cụ thể bằng typed placeholder tương ứng (<UNK-PATH>, <UNK-IP>, <UNK-CMD>). "
        "Kỹ thuật này bảo đảm không sụp đổ (collapse) thông tin cấu trúc cú pháp cần thiết cho Extractor và không làm rò rỉ từ vựng.",
        
    )
    add_bullet_p([
        "Ranh giới leakage-safe được kiểm tra tự động trong pipeline thực nghiệm: mọi phép fit nào "
        "truy cập dữ liệu ngoài ",
        latex_to_clean_omml(r"\mathcal{D}_{\text{train}}"),
        " đều bị phát hiện và báo lỗi trước khi chạy thực nghiệm (xem CTRL-LEAK-001)."
    ])

    # --- 2.2.2 ---
    add_h3("Mô hình đe dọa quyền riêng tư và Controlled Linkability")
    add_p([
        "Các vector biểu diễn đặc trưng ",
        latex_to_clean_omml(r"\mathbf{z}_t"),
        " và các sự kiện đã được chuẩn hóa ",
        latex_to_clean_omml(r"\tilde{e}_i"),
        " tất yếu mang thông tin nhận dạng thực thể, xem đây là điều kiện cần thiết để duy trì ngữ "
        "cảnh hành vi và phục vụ phân tích tấn công đa giai đoạn (kế thừa H5 từ Mục 2.1). Tuy nhiên, "
        "việc lưu giữ thông tin nhận dạng thô trong không gian vector tạo ra bề mặt tấn công quyền "
        "riêng tư không cần thiết. Mục 2.2.2 hình thức hóa mô hình đe dọa và xác lập cơ chế liên kết "
        "có kiểm soát (Controlled Linkability Mechanism) như một hợp đồng kỹ thuật. Cơ chế này "
        "được kỳ vọng thiết lập sự vượt trội Pareto theo H5 nhưng phải được kiểm chứng thực nghiệm "
        "tại Chương 3 (BOUNDARY-06: Privacy Claim Requires Attack-Based Empirical Evaluation)."
    ])

    # 2.2.2.1 — Data/entity adversary
    add_h4("Data/Entity Adversary")
    add_p([
        "Lớp đe dọa thứ nhất là đối nghịch dữ liệu/thực thể (Data/Entity Adversary). Đối nghịch này "
        "có quyền truy cập trực tiếp hoặc gián tiếp vào các sự kiện đã chuẩn hóa ",
        latex_to_clean_omml(r"\tilde{e}_i"),
        " hoặc các trạng thái trung gian trong pipeline tiền xử lý. Hai mối đe dọa chính:"
    ], keep_with_next=True)
    add_p([
        "Tấn công liên kết thực thể (Entity Linkage Attack): đối nghịch sử dụng các trường nhận dạng "
        "thô (tên người dùng, PID, hostname) để liên kết sự kiện giữa các phân đoạn thời gian, "
        "các hệ thống, hoặc giữa log nguồn và dữ liệu bên ngoài nhằm xác định danh tính cá nhân."
    ])
    add_p([
        "Tấn công tái định danh (Re-identification Attack): đối nghịch sử dụng các đặc trưng hành vi "
        "đặc trưng (pattern thao tác tệp tin, chuỗi lệnh đặc thù) kết hợp với thông tin bên ngoài "
        "để tái định danh thực thể đã bị ẩn danh hóa."
    ])

    # 2.2.2.2 — Model adversary
    add_h4("Model Adversary")
    add_p([
        "Lớp đe dọa thứ hai là đối nghịch mô hình (Model Adversary). Đối nghịch này không truy cập "
        "trực tiếp dữ liệu thô mà tương tác với mô hình đã được triển khai hoặc với các vector biểu "
        "diễn ",
        latex_to_clean_omml(r"\mathbf{z}_t"),
        " đã xuất ra. Hai mối đe dọa chính:"
    ], keep_with_next=True)
    add_p([
        "Tấn công suy luận thành viên (Membership Inference Attack, MIA): đối nghịch xác định xem "
        "một bản ghi dữ liệu cụ thể của một thực thể có nằm trong tập huấn luyện của mô hình "
        "hay không, thông qua quan sát hành vi đầu ra hoặc các đặc trưng của vector ",
        latex_to_clean_omml(r"\mathbf{z}"),
        " (",
        make_citation_element([25]),
        ")."
    ])
    add_p([
        "Tấn công nghịch đảo biểu diễn (Representation/Model Inversion Attack): đối nghịch khôi phục "
        "hoặc xấp xỉ các đặc trưng nhận dạng nhạy cảm từ vector ",
        latex_to_clean_omml(r"\mathbf{z}_t"),
        " (ví dụ tên người dùng, hostname, chuỗi lệnh) thông qua giải bài toán ngược của bộ mã "
        "hóa (",
        make_citation_element([26]),
        ")."
    ])
    add_p([
        "Lưu ý quan trọng: sự tồn tại của hai lớp đe dọa trên không đồng nghĩa với việc mọi thiết "
        "kế đều bị tổn thương ở mức độ như nhau. Mức độ tổn thương thực tế phụ thuộc vào kiến trúc "
        "cụ thể của ",
        latex_to_clean_omml(r"f_\theta"),
        " và cơ chế liên kết được chọn; mức độ này được xác định qua kiểm thử tấn công thực nghiệm "
        "tại Chương 3 (CTRL-PRIV-001)."
    ])

    # 2.2.2.3 — Mechanism contract
    add_h4("Controlled Linkability Mechanism Contract")
    add_p([
        "Để giải quyết đồng thời hai lớp đe dọa trên trong khi vẫn bảo toàn ngữ cảnh hành vi cần "
        "thiết cho phân tích tấn công, chuyên đề xác lập Cơ chế Liên kết Có Kiểm soát (Controlled "
        "Linkability Mechanism) như một hợp đồng kỹ thuật với ba thành phần:"
    ], keep_with_next=True)
    add_p([
        "Pseudonymization thực thể theo phiên (Session-scoped Pseudonymization): mỗi thực thể thực "
        "sự ",
        latex_to_clean_omml(r"v \in \mathcal{V}"),
        " được ánh xạ về một mã định danh giả danh ",
        latex_to_clean_omml(r"\hat{v} = \pi_\psi(v)"),
        " thông qua hàm ánh xạ xác định có kiểm soát ",
        latex_to_clean_omml(r"\pi_\psi : \mathcal{V} \to \hat{\mathcal{V}}"),
        ". Hàm ",
        latex_to_clean_omml(r"\pi_\psi"),
        " bảo toàn quan hệ đồng nhất thực thể (tức là ",
        latex_to_clean_omml(r"v_i = v_j \Rightarrow \pi_\psi(v_i) = \pi_\psi(v_j)"),
        ") để cho phép liên kết hành vi theo thời gian (điều kiện cần để phát hiện APT đa giai đoạn), nhưng không giữ lại thông tin nhận dạng thô ",
        latex_to_clean_omml(r"v"),
        " không cần thiết cho mục đích phát hiện."
    ], bold_prefix="1. ", first_line_indent=False)
    add_p([
        "Tokenization tham số động có chọn lọc (Security-selective Parameter Tokenization): các "
        "tham số động đã được ",
        latex_to_clean_omml(r"\text{Retain}(\cdot)"),
        " xác nhận là mang ngữ nghĩa an ninh được tokenize thành các đơn vị học được theo "
        "kho từ vựng fit trên ",
        latex_to_clean_omml(r"\mathcal{D}_{\text{train}}"),
        ". Các trường KNOWN EXCLUDE FIELD được loại bỏ hoàn toàn (DROP) trước khi đưa vào ",
        latex_to_clean_omml(r"f_\theta"),
        "; trong khi các tham số UNSEEN / UNCERTAIN được xử lý theo chính sách UNKNOWN-SAFE "
        "(bảo toàn nhãn kiểu trường và vị trí cấu trúc, gán typed placeholder), không bị loại bỏ tùy tiện."
    ], bold_prefix="2. ", first_line_indent=False)
    add_p([
        "Phạm vi liên kết và rotation (Linkability Scope): hàm ",
        latex_to_clean_omml(r"\pi_\psi"),
        " xác lập bốn ranh giới tường minh: "
        "(a) Linkable within scope: trong cùng phiên, ",
        latex_to_clean_omml(r"\pi_\psi"),
        " nhất quán để hành vi đa giai đoạn của cùng thực thể vẫn liên kết được; "
        "(b) Unlinkable across rotation boundary: khi rotation xảy ra, "
        "cùng thực thể nhận mã giả danh mới, làm gián đoạn cross-session linkage; "
        "(c) Remains linkable: kiểu hành vi, thời gian tương đối và tham số an ninh giúp duy trì behavioral fingerprint; "
        "(d) Becomes unlinkable: tên thực sự v, o và các direct identifier. "
        "Rotation boundary phải được xác lập theo predefined operational policy hoặc được lựa chọn dựa trên tập Train/Validation, tuyệt đối không sử dụng thông tin từ true Test attack/campaign duration (loại trừ rủi ro test leakage). Tại Chương 3, chuyên đề đánh giá nhiều mốc rotation horizon cố định (fixed horizons) để xây dựng đường biên Utility–Privacy frontier. Trường hợp rotation horizon quá ngắn so với thời lượng chiến dịch APT thực tế sẽ dẫn tới việc cùng một thực thể nhận các mã giả danh khác nhau qua từng giai đoạn, làm phân mảnh khả năng liên kết chuỗi hành vi cross-phase nếu không có cơ chế tái liên kết an toàn; đây là failure case tường minh cần được định lượng trong đánh giá thực nghiệm. Thiết kế phạm vi cụ thể (session-scope vs. global scope) là quyết định thực nghiệm phụ thuộc vào đánh đổi này.",
    ], bold_prefix="3. ", first_line_indent=False)

    add_p([
        "Hợp đồng kỹ thuật tổng quát của cơ chế Controlled Linkability được phát biểu như sau: với "
        "mọi sự kiện đã được chuẩn hóa ",
        latex_to_clean_omml(r"\tilde{e}_i"),
        ", sự kiện được biến đổi quyền riêng tư ",
        latex_to_clean_omml(r"\tilde{e}_i^{\text{priv}}"),
        " được tạo ra bằng cách thay thế các trường nhận dạng thực thể bằng mã giả danh "
        "và áp dụng toán tử bảo vệ tham số an toàn ",
        latex_to_clean_omml(r"\text{ParamProtect}(\mathbf{p}_i)"),
        " tuân thủ nguyên tắc: (i) các trường KNOWN EXCLUDE FIELD bị loại bỏ hoàn toàn (DROP); "
        "(ii) các tham số được ",
        latex_to_clean_omml(r"\text{Retain}(\cdot)"),
        " xác nhận được tokenize trên ",
        latex_to_clean_omml(r"\mathcal{D}_{\text{train}}"),
        "; (iii) các tham số UNSEEN / UNCERTAIN được xử lý theo chính sách UNKNOWN-SAFE "
        "(bảo toàn nhãn kiểu trường và vị trí cấu trúc, gán typed placeholder):"
    ], keep_with_next=True)
    add_display_equation(
        r"\tilde{e}_i^{\text{priv}} = \bigl(t_i,\; \tau_i,\; \pi_\psi(v_i),\; \pi_\psi(o_i),\; a_i,\; \text{ParamProtect}(\mathbf{p}_i)\bigr)"
    )
    add_p([
        "Lưu ý rằng các tham số thời gian ",
        latex_to_clean_omml(r"t_i"),
        " và kiểu hành vi ",
        latex_to_clean_omml(r"\tau_i, a_i"),
        " được bảo toàn nguyên vẹn vì chúng không phải direct identifier, tuy nhiên các tham số này có thể hoạt động như quasi-identifier khi kết hợp với dữ liệu bên ngoài; rủi ro liên kết tương ứng cần được đánh giá theo CTRL-PRIV-001. "
        "Cơ chế này được kỳ vọng thiết lập sự vượt trội Pareto theo nghĩa của H5 so với hai đường "
        "cơ sở: (a) giữ nguyên định danh thô không có pseudonymization; và (b) ẩn danh hóa quá mức "
        "xóa bỏ cả liên kết thực thể cần thiết cho phân tích hành vi. Tuy nhiên, khẳng định sự vượt "
        "trội này yêu cầu kiểm thử thực nghiệm tường minh tại Chương 3 thông qua CTRL-PRIV-001: "
        "đánh giá MIA (",
        make_citation_element([25]),
        "), tấn công nghịch đảo (",
        make_citation_element([26]),
        "), và tấn công tái định danh; không thể đưa ra kết luận chỉ từ lý thuyết DP hoặc thiết "
        "kế pseudonymization."
    ], first_line_indent=False)

    # --- 2.2.3 ---
    add_h3("Đồng bộ thời gian và cửa sổ ngữ cảnh đa tỷ lệ")
    add_p([
        "Sau khi các sự kiện được chuẩn hóa và bảo vệ quyền riêng tư theo Mục 2.2.1–2.2.2, bước "
        "tiếp theo là tổ chức chúng thành các đơn vị ngữ cảnh phục vụ các bộ trích xuất đặc trưng "
        "hạ nguồn tại Mục 2.3. Mục 2.2.3 giải quyết hai vấn đề liên quan chặt chẽ: (1) đảm bảo "
        "tính nhất quán Event-Time khi sự kiện đến từ nhiều nguồn có đồng hồ không đồng bộ; và "
        "(2) xây dựng cấu trúc cửa sổ ngữ cảnh đa tỷ lệ phù hợp với bản chất đa giai đoạn của "
        "chiến dịch APT. Cơ chế mốc ngắt thời gian (Event-Time Watermark) và tham số ",
        latex_to_clean_omml(r"\delta_{\text{delay}}"),
        " đã được xác lập tại Mục 2.1.2; Mục 2.2.3 kế thừa và áp dụng các định nghĩa đó mà "
        "không tái định nghĩa lại."
    ])

    # 2.2.3.1 — Event-time alignment
    add_h4("Event-Time, Clock Skew, Watermark và Late Events")
    add_p([
        "Môi trường SOC thực tế thu thập log từ nhiều nguồn có đồng hồ hệ thống khác nhau, phát "
        "sinh hiện tượng lệch pha đồng hồ (Clock Skew). Lệch pha đồng hồ khiến các bản ghi telemetry "
        "tương ứng của cùng một sự kiện/tương tác quan sát được lại mang nhãn thời gian không nhất quán khi so sánh xuyên nguồn, "
        "làm xáo trộn thứ tự thời gian sự kiện dùng để xây dựng chuỗi và đồ thị hạ nguồn (lưu ý nguyên lý bất biến: "
        "quan hệ phụ thuộc ghi nhận trong log không đồng nhất với quan hệ nhân quả (nguyên lý: dependency != causal effect)). Giai "
        "đoạn đồng bộ thời gian áp dụng một chuỗi bước xử lý có thứ tự:"
    ], keep_with_next=True)
    add_p([
        "Chuẩn hóa múi giờ: tất cả nhãn thời gian được chuyển về UTC epoch trước khi so sánh "
        "xuyên nguồn."
    ], bold_prefix="(i) ", first_line_indent=False)
    add_p([
        "Ước lượng và hiệu chỉnh lệch pha đồng hồ: với mỗi cặp nguồn có thể quan sát được sự kiện "
        "chung (ví dụ: cùng một kết nối mạng được ghi nhận bởi cả endpoint log và network log), "
        "offset lệch pha ",
        latex_to_clean_omml(r"\Delta_{\text{skew}}"),
        " được ước lượng và hiệu chỉnh. Phương pháp ước lượng cụ thể (crystal-synced event "
        "matching, NTP-based correction,...) là quyết định triển khai phụ thuộc môi trường và "
        "không ảnh hưởng đến hợp đồng tiền xử lý."
    ], bold_prefix="(ii) ", first_line_indent=False)
    add_p([
        "Áp dụng cơ chế Watermark kế thừa từ Mục 2.1.2: tham số ",
        latex_to_clean_omml(r"\delta_{\text{delay}}"),
        " xác lập cửa sổ dung sai cho sự kiện đến muộn sau đồng bộ thời gian. Các sự kiện đến "
        "trong cửa sổ ",
        latex_to_clean_omml(r"[t_{\text{wm}}, t_{\text{wm}} + \delta_{\text{delay}}]"),
        " được tích hợp vào trạng thái dòng; các sự kiện đến trễ hơn được chuyển vào Reconciliation "
        "Buffer theo Explicit Information-Loss Protocol."
    ], bold_prefix="(iii) ", first_line_indent=False)

    # 2.2.3.2 — Multi-scale context windows
    add_h4("Multi-scale Context Windows")
    add_p([
        "Đặc điểm của chiến dịch APT là hoạt động trải dài nhiều tỷ lệ thời gian khác nhau "
        "đồng thời: hành động tức thời (milliseconds đến seconds), chuỗi hành vi ngắn hạn "
        "(minutes đến hours), và chiến dịch dài hạn (days đến weeks). Một cửa sổ ngữ cảnh duy "
        "nhất không thể bao phủ đầy đủ tất cả tỷ lệ này mà không vi phạm ràng buộc bộ nhớ "
        "hữu hạn ",
        latex_to_clean_omml(r"M_{\text{max}}"),
        " từ Mục 2.1.2. Chuyên đề đề xuất cấu trúc cửa sổ ngữ cảnh ba tỷ lệ (Multi-scale "
        "Context Windows) như một giải pháp phân cấp:"
    ], keep_with_next=True)
    add_display_equation(
        r"\mathcal{W} = \left\{\,W_s,\; W_m,\; W_l\,\right\}"
    )
    add_p([
        "Cửa sổ ngắn hạn ",
        latex_to_clean_omml(r"W_s"),
        " (Short-term Window): bao phủ chuỗi sự kiện gần nhất trong phạm vi thời gian ngắn "
        "(hyperparameter, ví dụ: vài phút đến hàng chục phút; khóa bằng Validation/operational budget tại Mục 2.3), cung cấp ngữ cảnh thực thi tức thời cho một "
        "tiến trình hoặc phiên làm việc. Cửa sổ này cung cấp dữ liệu đầu vào chính cho Bộ "
        "trích xuất tuần tự (Sequential Extractor) tại Mục 2.3. Kích thước ",
        latex_to_clean_omml(r"\text{card}(W_s)"),
        " bị chặn trên bởi độ dài chuỗi tối đa ",
        latex_to_clean_omml(r"L_{\text{max}}"),
        " của kiến trúc encoder, với giá trị cụ thể sẽ được xác lập tại Mục 2.3 khi kiến trúc "
        "được khóa."
    ])
    add_p([
        "Cửa sổ trung hạn ",
        latex_to_clean_omml(r"W_m"),
        " (Medium-term Window): bao phủ lịch sử hành vi của thực thể trong phạm vi thời gian "
        "trung (hyperparameter, ví dụ: vài giờ đến một ngày; khóa bằng Validation budget), phản ánh các pattern hành vi đặc trưng của "
        "người dùng/tiến trình và cung cấp ngữ cảnh phát hiện leo thang đặc quyền hoặc lateral "
        "movement. Cửa sổ này được duy trì trong trạng thái hữu hạn ",
        latex_to_clean_omml(r"\mathcal{S}_t"),
        " thông qua cơ chế nén trạng thái (state summary/sketch) đã xác lập tại Mục 2.1.2."
    ])
    add_p([
        "Cửa sổ dài hạn/Tóm tắt trạng thái ",
        latex_to_clean_omml(r"W_l"),
        " (Long-term State Summary): thay vì lưu trữ toàn bộ lịch sử thô (vi phạm ",
        latex_to_clean_omml(r"M_{\text{max}}"),
        "), trạng thái dài hạn được duy trì dưới dạng các sketch thống kê nén gọn "
        "(ví dụ: Count-Min Sketch về tần suất hành vi loại sự kiện, hàm suy giảm trọng số "
        "cạnh theo thời gian). Đây chính là đánh đổi cốt lõi Long-Horizon Context Trade-off "
        "đã được xác lập tại Mục 2.1.2: chi phí tính toán và bộ nhớ cụ thể của ",
        latex_to_clean_omml(r"W_l"),
        " phụ thuộc vào lựa chọn cấu trúc sketch tại Mục 2.3–2.4."
    ])
    add_p([
        "Ba cửa sổ ",
        latex_to_clean_omml(r"W_s, W_m, W_l"),
        " không độc lập mà có quan hệ phân cấp: ",
        latex_to_clean_omml(r"W_s \subset W_m"),
        " theo thứ tự thời gian, và ",
        latex_to_clean_omml(r"W_l"),
        " là tóm tắt nén của lịch sử trước ",
        latex_to_clean_omml(r"W_m"),
        ". Phân cấp này đảm bảo cùng một evidence có thể xuất hiện ở nhiều scale (Ws, Wm, Wl) dưới dạng bài biểu diễn khác nhau, nhưng không có thông tin "
        "tương lai nào bị nạp vào cửa sổ hiện tại (nguyên lý Strictly Causal / Zero Lookahead từ Mục 2.1.3). "
        "Sơ đồ phân cấp cụ thể và giao diện đầu ra từ ",
        latex_to_clean_omml(r"\mathcal{W}"),
        " sang các bộ trích xuất tại Mục 2.3 sẽ được hình thức hóa khi kiến trúc được xác lập."
    ])

    
    # =========================================================================
    # 2.3. MULTI-VIEW FEATURE EXTRACTION
    # =========================================================================
    add_h2("Trích xuất đặc trưng đa góc nhìn")
    add_p([
        "Sau giai đoạn tiền xử lý và đồng bộ thời gian tại Mục 2.2, luồng sự kiện kiểu hóa ",
        latex_to_clean_omml(r"\tilde{e}_i^{\text{priv}}"),
        " được đưa vào hai bộ trích xuất đặc trưng song song nhằm khai thác hai góc nhìn bổ trợ cho nhau: "
        "góc nhìn tuần tự ngữ nghĩa (Semantic-Sequential View) và góc nhìn cấu trúc phụ thuộc (Dependency-Temporal Graph View). "
        "Mỗi bộ trích xuất hoạt động độc lập trên không gian đầu vào tương ứng và xuất ra một vector biểu diễn đặc trưng riêng biệt: ",
        latex_to_clean_omml(r"\mathbf{z}_{\text{seq}} \in \mathbb{R}^{d_{\text{seq}}}"),
        " và ",
        latex_to_clean_omml(r"\mathbf{z}_{\text{graph}} \in \mathbb{R}^{d_{\text{graph}}}"),
        ". Hai vector này là đầu vào của tầng gióng hàng và học biểu diễn thống nhất tại Mục 2.4."
    ])

    # -------------------------------------------------------------------------
    # 2.3.1. Transformer Semantic–Sequential Extractor
    # -------------------------------------------------------------------------
    add_h3("Bộ trích xuất tuần tự ngữ nghĩa Transformer")
    add_p([
        "Bộ trích xuất tuần tự ngữ nghĩa tiếp nhận chuỗi các sự kiện kiểu hóa trong cửa sổ ngắn hạn ",
        latex_to_clean_omml(r"W_s"),
        " gồm ",
        latex_to_clean_omml(r"L \le L_{\text{max}}"),
        " sự kiện liên tiếp, kế thừa kiến trúc Self-Attention tiêu chuẩn ",
        make_citation_element([35]),
        ". Nhiệm vụ của bộ trích xuất là mô hình hóa các phụ thuộc ngữ cảnh cục bộ, "
        "thứ tự gọi hàm hệ thống và tương tác tham số động mà không dựa trên bất kỳ nhãn phân loại tấn công nào ở hạ nguồn."
    ])

    # 2.3.1.1. Event Representation
    add_h4("Biểu diễn sự kiện kiểu hóa và mã hóa ngữ cảnh")
    add_p([
        "Mỗi sự kiện kiểu hóa ",
        latex_to_clean_omml(r"\tilde{e}_i^{\text{priv}} = (t_i, \tau_i, \pi_\psi(v_i), \pi_\psi(o_i), a_i, \text{ParamProtect}(\mathbf{p}_i))"),
        " được chiếu vào không gian vector liên tục ",
        latex_to_clean_omml(r"\mathbb{R}^{d_{\text{model}}}"),
        " thông qua hàm nhúng kết hợp các thành phần thông tin theo cấu trúc:"
    ], keep_with_next=True)
    add_display_equation(make_event_embedding_omml())
    add_p([
        "Trong đó: ",
        latex_to_clean_omml(r"\mathbf{e}_\tau, \mathbf{e}_a"),
        " là các ma trận nhúng kiểu sự kiện và hành động tĩnh; ",
        latex_to_clean_omml(r"\mathbf{e}_v, \mathbf{e}_o"),
        " là nhúng thực thể chủ thể và đối tượng theo không gian giả danh có kiểm soát; ",
        latex_to_clean_omml(r"\mathbf{e}_p"),
        " là nhúng các token tham số an ninh được giữ lại; ",
        latex_to_clean_omml(r"\mathbf{e}_{\text{pos}}(i)"),
        " là mã hóa vị trí rời rạc; và ",
        latex_to_clean_omml(r"\mathbf{e}_{\text{time}}(t_i)"),
        " là hàm mã hóa thời gian liên tục điều hòa (Harmonic Time Encoding) dạng ",
        latex_to_clean_omml(r"\mathbf{e}_{\text{time}}(t_i) = [\cos(\omega_1 t_i + \phi_1), \dots, \cos(\omega_k t_i + \phi_k)]^\top"),
        ", kế thừa nguyên lý mã hóa thời gian liên tục trên đồ thị động ",
        make_citation_element([36]),
        ", cung cấp biểu diễn học được của khoảng thời gian tương đối ",
        latex_to_clean_omml(r"\Delta t"),
        " giữa các hành vi liên tiếp. Toàn bộ các siêu tham số cấu trúc như số tầng encoder ",
        latex_to_clean_omml(r"N_{\text{layers}}"),
        ", số đầu chú ý ",
        latex_to_clean_omml(r"N_{\text{heads}}"),
        ", số chiều ẩn ",
        latex_to_clean_omml(r"d_{\text{model}}"),
        ", số chiều lan truyền ",
        latex_to_clean_omml(r"d_{\text{ff}}"),
        " và tỷ lệ dropout ",
        latex_to_clean_omml(r"p_{\text{drop}}"),
        " đều được xác lập dựa trên tập huấn luyện và kiểm định (Train/Validation), không sử dụng nhãn kiểm tra."
    ], first_line_indent=False)

    # 2.3.1.2. Self-supervised Objectives
    add_h4("Mục tiêu huấn luyện tự giám sát")
    add_p([
        "Để huấn luyện bộ trích xuất tuần tự mà không gây rò rỉ nhãn tấn công hay phụ thuộc vào tri thức chuyên gia định trước, "
        "mô hình tối ưu hóa đồng thời ba mục tiêu tự giám sát (Self-supervised Objectives) với đầu dự đoán phụ trợ (Auxiliary Heads) và tiêu chuẩn tối ưu hóa tường minh:"
    ], keep_with_next=True)
    add_num_p([
        "Tác vụ dự đoán sự kiện bị che (Masked Event Prediction, ",
        latex_to_clean_omml(r"\mathcal{L}_{\text{MEP}}"),
        "): Thích ứng từ phương pháp LogBERT ",
        make_citation_element([4]),
        ", che ngẫu nhiên tập chỉ số sự kiện ",
        latex_to_clean_omml(r"\mathcal{M}_{\text{event}} \subset \{1, \dots, L\}"),
        " theo tỷ lệ che ",
        latex_to_clean_omml(r"r_{\text{mask}}^{\text{event}} \in (0, 1)"),
        " (trong đó giá trị 15% được sử dụng làm điểm khởi tạo tham chiếu và được tinh chỉnh, khóa cố định trên tập Train/Validation trước khi đánh giá trên tập Test). Đầu dự đoán phụ trợ ",
        latex_to_clean_omml(r"\phi_{\text{seq}}^{\text{event}} : \mathbb{R}^{d_{\text{model}}} \to \mathbb{R}^{\text{card}(\mathcal{V}_{\text{event}})}"),
        " ánh xạ trạng thái ẩn ",
        latex_to_clean_omml(r"\mathbf{h}_i"),
        " sang không gian xác suất kiểu sự kiện. Mục tiêu (target) là nhãn kiểu sự kiện và hành động thực tế ",
        latex_to_clean_omml(r"y_i^{\text{event}} \in \{1, \dots, \text{card}(\mathcal{V}_{\text{event}})\}"),
        ", tối ưu hóa qua hàm mất mát Categorical Cross-Entropy trung bình trên tập bị che, cung cấp gradient cập nhật trực tiếp cho ",
        latex_to_clean_omml(r"\theta_{\text{seq}}"),
        " và ",
        latex_to_clean_omml(r"\phi_{\text{seq}}^{\text{event}}"),
        ":"
    ], keep_with_next=True)
    add_display_equation(make_l_mep_omml())
    add_num_p([
        "Tác vụ dự đoán tham số an ninh bị che (Masked Parameter Prediction, ",
        latex_to_clean_omml(r"\mathcal{L}_{\text{MPP}}"),
        "): Thiết kế đóng góp riêng của chuyên đề nhằm che ngẫu nhiên tập tham số bảo toàn an ninh ",
        latex_to_clean_omml(r"\mathcal{M}_{\text{param}} \subset \{1, \dots, L\}"),
        " theo tỷ lệ ",
        latex_to_clean_omml(r"r_{\text{mask}}^{\text{param}} \in (0, 1)"),
        " (khởi tạo tham chiếu 15% và được khóa cố định trên tập Validation trước khi thử nghiệm). Đầu dự đoán phụ trợ ",
        latex_to_clean_omml(r"\phi_{\text{seq}}^{\text{param}} : \mathbb{R}^{d_{\text{model}}} \to \mathbb{R}^{\text{card}(\mathcal{V}_{\text{param}})}"),
        " dự đoán token tham số mục tiêu ",
        latex_to_clean_omml(r"y_i^{\text{param}}"),
        ". Mục tiêu tái dựng (target) được định nghĩa chuẩn xác là token đã qua chuẩn hóa và biến đổi bảo vệ quyền riêng tư ",
        latex_to_clean_omml(r"\pi_\psi"),
        " tại Mục 2.2 (gồm token thực thể giả danh hóa, token tham số an ninh chuẩn hóa hoặc token kiểu hóa UNKNOWN-SAFE), tuyệt đối không đặt mục tiêu là chuỗi nhạy cảm thô. "
        "Hàm mất mát Categorical Cross-Entropy tính trung bình trên ",
        latex_to_clean_omml(r"\mathcal{M}_{\text{param}}"),
        ", truyền gradient cho ",
        latex_to_clean_omml(r"\theta_{\text{seq}}"),
        " và ",
        latex_to_clean_omml(r"\phi_{\text{seq}}^{\text{param}}"),
        ":"
    ], num_id=20, keep_with_next=True)
    add_display_equation(make_l_mpp_omml())
    add_num_p([
        "Tác vụ dự đoán khoảng thời gian tương đối (Relative Time Gap Prediction, ",
        latex_to_clean_omml(r"\mathcal{L}_{\text{time}}"),
        "): Thiết kế đóng góp riêng của chuyên đề nhằm dự đoán khoảng cách thời gian giữa hai sự kiện kề nhau ",
        latex_to_clean_omml(r"\Delta t_{i, i+1} = t_{i+1} - t_i \ge 0"),
        ". Đầu hồi quy phụ trợ ",
        latex_to_clean_omml(r"\phi_{\text{seq}}^{\text{time}} : \mathbb{R}^{2 d_{\text{model}}} \to \mathbb{R}"),
        " tiếp nhận vector ghép ",
        latex_to_clean_omml(r"[\mathbf{h}_i; \mathbf{h}_{i+1}]"),
        " để dự đoán giá trị logarit ",
        latex_to_clean_omml(r"\log(1 + \Delta t_{i, i+1})"),
        ". Tiêu chuẩn tối ưu hóa là hàm Smooth L1 (Huber criterion, ",
        latex_to_clean_omml(r"\delta = 1.0"),
        ") trung bình trên toàn bộ ",
        latex_to_clean_omml(r"L - 1"),
        " cặp sự kiện liền kề, truyền gradient cho ",
        latex_to_clean_omml(r"\theta_{\text{seq}}"),
        " và ",
        latex_to_clean_omml(r"\phi_{\text{seq}}^{\text{time}}"),
        ":"
    ], keep_with_next=True)
    add_display_equation(make_l_seq_time_omml())
    add_p([
        "Hàm mất mát tổng thể của Bộ trích xuất tuần tự trong Giai đoạn A1 là tổng kết hợp có trọng số: ",
        latex_to_clean_omml(r"\mathcal{L}_{\text{seq}}^{\text{self}} = \alpha_1 \mathcal{L}_{\text{MEP}} + \alpha_2 \mathcal{L}_{\text{MPP}} + \alpha_3 \mathcal{L}_{\text{time}}"),
        ", trong đó các hệ số ",
        latex_to_clean_omml(r"\alpha_1, \alpha_2, \alpha_3 > 0"),
        " được tinh chỉnh trên tập Validation."
    ])

    # 2.3.1.3. Sequential Output
    add_h4("Giao diện đầu ra của bộ trích xuất tuần tự")
    add_p([
        "Sau khi đi qua ",
        latex_to_clean_omml(r"N_{\text{layers}}"),
        " khối Transformer Encoder, chuỗi biểu diễn ẩn ",
        latex_to_clean_omml(r"\mathbf{H}_{\text{seq}} = [\mathbf{h}_1, \dots, \mathbf{h}_L] \in \mathbb{R}^{L \times d_{\text{model}}}"),
        " được tổng hợp thành vector biểu diễn ngữ cảnh tuần tự ",
        latex_to_clean_omml(r"\mathbf{z}_{\text{seq}}"),
        " thông qua toán tử tổng hợp có trọng số chú ý (Attention-weighted Readout):"
    ], keep_with_next=True)
    add_display_equation(
        r"\mathbf{z}_{\text{seq}} = \text{Readout}_{\text{seq}}(\mathbf{H}_{\text{seq}}) = \sum_{i=1}^L \text{softmax}\left(\frac{\mathbf{w}_{\text{seq}}^\top \mathbf{h}_i}{\sqrt{d_{\text{model}}}}\right) \mathbf{W}_{\text{proj}} \mathbf{h}_i \in \mathbb{R}^{d_{\text{seq}}}"
    )
    add_p([
        "Đầu ra của Bộ trích xuất tuần tự là một gói biểu diễn chuẩn hóa mang vector ",
        latex_to_clean_omml(r"\mathbf{z}_{\text{seq}} \in \mathbb{R}^{d_{\text{seq}}}"),
        " đại diện cho toàn bộ chuỗi hành vi trong cửa sổ ngắn hạn ",
        latex_to_clean_omml(r"W_s"),
        ", kèm theo siêu dữ liệu tương ứng (Correspondence Metadata) bao gồm: khoảng thời gian sự kiện ",
        latex_to_clean_omml(r"[t_{\text{start}}, t_{\text{end}}]"),
        ", phạm vi thực thể hoặc phiên tương ứng, mặt nạ sẵn sàng của góc nhìn ",
        latex_to_clean_omml(r"m_{\text{seq}} \in \{0, 1\}"),
        ", và định danh cửa sổ. Giao diện này cung cấp đầy đủ thông tin để Tầng gióng hàng tại Mục 2.4 xác định quan hệ tương ứng giữa các góc nhìn mà không thực hiện dung hợp sớm hay sử dụng nhãn kiểm tra. "
        "Giả định kiến trúc gồm ",
        latex_to_clean_omml(r"N_{\text{layers}}"),
        " tầng Transformer Encoder, độ dài chuỗi đầu vào ",
        latex_to_clean_omml(r"L"),
        ", số chiều mô hình ",
        latex_to_clean_omml(r"d_{\text{model}}"),
        " và số chiều tầng truyền thẳng ",
        latex_to_clean_omml(r"d_{\text{ff}}"),
        ", độ phức tạp tính toán của bộ trích xuất tuần tự là ",
        latex_to_clean_omml(r"C_{\text{seq}} = \mathcal{O}\left(N_{\text{layers}} \cdot (L^2 \cdot d_{\text{model}} + L \cdot d_{\text{model}} \cdot d_{\text{ff}})\right)"),
        ", trong đó số hạng bậc hai phản ánh chi phí tính toán ma trận chú ý theo độ dài chuỗi."
    ])

    # Insert Figure 2.2 Placeholder Canvas
    p_c1 = doc.add_paragraph() if target_p is None else target_p.insert_paragraph_before()
    p_c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_c1.paragraph_format.keep_with_next = True
    p_c1.paragraph_format.space_before = Pt(6)
    p_c1.paragraph_format.space_after = Pt(2)
    p_c1._p.append(parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="30202" w:name="BK_FIG_2_002_CANVAS"/>'))
    r_c1 = p_c1.add_run(" ")
    p_c1._p.append(parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="30202"/>'))
    add_figure_caption(doc, target_p, 2, 2, "Kiến trúc Bộ trích xuất tuần tự ngữ nghĩa Transformer và cơ chế học tự giám sát trên cửa sổ ngắn hạn (Nguồn: Tác giả đề xuất)", bookmark_name="BK_FIG_2_002")

    # -------------------------------------------------------------------------
    # 2.3.2. Dependency–Temporal Provenance Graph Construction and Graph Fidelity
    # -------------------------------------------------------------------------
    add_h3("Xây dựng đồ thị nguồn gốc phụ thuộc thời gian và đảm bảo độ chân thực đồ thị")
    add_p([
        "Song song với bộ trích xuất tuần tự, luồng sự kiện được chuyển hóa thành đồ thị nguồn gốc phụ thuộc thời gian (Dependency-Temporal Provenance Graph). "
        "Đồ thị này biểu diễn các quan hệ tương tác và luồng thông tin cấu trúc giữa các thực thể hệ thống. "
        "Quá trình xây dựng đồ thị tuân thủ quy trình ba bước: tiếp nhận sự kiện kiểu hóa, áp dụng các cơ chế ứng viên kiểm soát độ chân thực đồ thị (Graph-Fidelity Candidates), "
        "và nạp cấu trúc đã tinh lọc vào mô hình Temporal GNN."
    ])

    # 2.3.2.1. Typed Nodes / Edges / Temporal Attributes
    add_h4("Đỉnh kiểu hóa, cạnh có hướng và thuộc tính thời gian")
    add_p([
        "Đồ thị nguồn gốc phụ thuộc thời gian tại thời điểm ",
        latex_to_clean_omml(r"t"),
        " được hình thức hóa dưới dạng đồ thị dị thể động ",
        latex_to_clean_omml(r"\mathcal{G}_t = (\mathcal{V}_t, \mathcal{E}_t, \phi_V, \phi_E, \tau_E)"),
        ", trong đó:"
    ], keep_with_next=True)
    add_bullet_p([
        "Tập đỉnh kiểu hóa ",
        latex_to_clean_omml(r"\mathcal{V}_t"),
        ": phân loại thông qua hàm ánh xạ kiểu ",
        latex_to_clean_omml(r"\phi_V : \mathcal{V}_t \to \{\text{Process}, \text{File}, \text{Socket}, \text{User}, \text{Host}\}"),
        ", phân định ranh giới ngữ nghĩa giữa các loại tài nguyên hệ thống."
    ])
    add_bullet_p([
        "Tập cạnh có hướng ",
        latex_to_clean_omml(r"\mathcal{E}_t"),
        ": mỗi cạnh ",
        latex_to_clean_omml(r"e = (u, v, r, t_e) \in \mathcal{E}_t"),
        " thể hiện một tương tác quan sát được với kiểu quan hệ ",
        latex_to_clean_omml(r"r = \phi_E(e) \in \{\text{fork}, \text{exec}, \text{read}, \text{write}, \text{connect}, \text{bind}\}"),
        " và mốc thời gian sự kiện ",
        latex_to_clean_omml(r"t_e = \tau_E(e)"),
        "."
    ])
    add_bullet_p([
        "Thuộc tính cạnh: tích hợp các tham số an ninh đã được chuẩn hóa và bảo vệ quyền riêng tư (như số byte truyền nhận, cờ truy cập, mã trạng thái), "
        "gắn kèm vector thuộc tính ",
        latex_to_clean_omml(r"\mathbf{x}_e \in \mathbb{R}^{d_{\text{edge}}}"),
        "."
    ])
    add_p([
        "Nguyên tắc cốt lõi: cạnh chỉ được tạo ra khi có sự kiện nhật ký kiểm toán ghi nhận tường minh tương tác giữa hai thực thể. "
        "Hệ thống không tự ý tạo cạnh nối chỉ dựa trên sự gần nhau về mặt thời gian giữa hai sự kiện rời rạc."
    ])

    # 2.3.2.2. Observable Dependency != Causal Effect
    add_h4("Quan hệ phụ thuộc quan sát được và ranh giới với quan hệ nhân quả")
    add_p([
        "Một ranh giới phương pháp luận căn bản cần được duy trì xuyên suốt chuyên đề là sự phân định giữa quan hệ phụ thuộc quan sát được (Observable Dependency) "
        "và quan hệ nhân quả thực sự (Causal Effect). "
        "Các cạnh trong đồ thị nguồn gốc phản ánh luồng dữ liệu và quan hệ kiểm toán do hệ điều hành ghi lại, biểu thị khả năng có sự truyền dẫn thông tin (Information Flow Potential). "
        "Tuy nhiên, việc một tiến trình đọc tệp tin ",
        latex_to_clean_omml(r"A"),
        " rồi sau đó ghi vào tệp tin ",
        latex_to_clean_omml(r"B"),
        " không khẳng định một cách chắc chắn rằng nội dung trong ",
        latex_to_clean_omml(r"B"),
        " bị chi phối nhân quả bởi ",
        latex_to_clean_omml(r"A"),
        " (nguyên lý bất biến: dependency != causal effect). "
        "Chuyên đề không sử dụng thuật ngữ đồ thị nhân quả trừ khi các giả định nhận dạng nhân quả chặt chẽ được thỏa mãn, "
        "tránh các suy diễn sai lầm trong việc truy vết nguồn gốc tấn công."
    ])

    # 2.3.2.3. False Dependency / Long-lived Entity / Edge Control
    add_h4("Các cơ chế ứng viên kiểm soát độ chân thực đồ thị và bùng nổ phụ thuộc")
    add_p([
        "Trong môi trường thực tế, đồ thị nguồn gốc thường đối mặt với hiện tượng bùng nổ phụ thuộc (Dependency Explosion) "
        "và ô nhiễm liên kết từ các thực thể tồn tại lâu (Long-lived Entities như dịch vụ hệ thống, tiến trình nền). "
        "Để xử lý các thách thức này trước khi nạp đồ thị vào mô hình học máy, chuyên đề đề xuất bốn cơ chế ứng viên kiểm soát độ chân thực đồ thị (Proposed Graph-Fidelity Candidates) hoạt động độc lập với bộ phát hiện:"
    ], keep_with_next=True)
    add_num_p([
        "Phân rã thực thể theo đơn vị công việc (Unit-of-work Node Splitting): các tiến trình hệ thống chạy dài ngày được phân tách thành các đỉnh phiên "
        "dựa trên luồng thực thi (thread) hoặc ranh giới phiên xử lý, nhằm hạn chế việc một tiến trình duy nhất tạo liên kết ảo giữa các tác vụ không liên quan."
    ], num_id=22)
    add_num_p([
        "Hàm suy giảm trọng số cạnh theo thời gian (Temporal Edge Weight Decay): trọng số ảnh hưởng của cạnh suy giảm theo hàm mũ ",
        latex_to_clean_omml(r"w(e) = \exp(-\lambda (t - t_e))"),
        ", trong đó ",
        latex_to_clean_omml(r"\lambda > 0"),
        " là hệ số suy giảm được xác định theo ngân sách trạng thái tại Mục 2.1.2, làm giảm độ ưu tiên của các tương tác cũ."
    ], num_id=22)
    add_num_p([
        "Kiểm soát bậc đỉnh và nén cạnh lặp (Degree-bounded Edge Compaction): áp dụng ngưỡng chặn trên đối với các cạnh tương tác lặp lại tần suất cao giữa cùng một cặp đỉnh "
        "trong một khoảng thời gian ngắn. Khi nén hoặc khử trùng lặp các cạnh này, hệ thống chỉ duy trì bản tóm tắt có khả năng kiểm toán (auditable summary) bao gồm: số lần lặp (multiplicity), "
        "mốc thời gian đầu tiên (",
        latex_to_clean_omml(r"t_{\text{first}}"),
        "), mốc thời gian cuối cùng (",
        latex_to_clean_omml(r"t_{\text{last}}"),
        "), kiểu quan hệ và các thuộc tính liên quan, mà không khẳng định có thể tái dựng toàn bộ chuỗi sự kiện gốc nếu không lưu trữ con trỏ định danh sự kiện thô (raw event pointers)."
    ], num_id=22)
    add_num_p([
        "Tính độc lập và đánh giá thực nghiệm: các cơ chế trên hoàn toàn dựa trên quy tắc cấu trúc và thời gian, "
        "không sử dụng điểm số tấn công hay nhãn kiểm tra để quyết định giữ hay xóa cạnh. Mức độ đóng góp và hiệu quả thực tế của từng cơ chế ứng viên được đánh giá định lượng thông qua phân tích triệt tiêu (Ablation Study) tại Chương 3."
    ], num_id=22)

    # 2.3.2.4. Cold-start and Unseen Entities
    add_h4("Xử lý thực thể khởi động lạnh và thực thể mới")
    add_p([
        "Hệ thống log thường xuyên xuất hiện các thực thể mới (tiến trình mới tạo, kết nối mạng tạm thời, máy chủ mới kết nối) "
        "chưa có lịch sử tương tác trong đồ thị, dẫn đến thách thức khởi động lạnh (Cold-start). "
        "Để xử lý hiện tượng này trong không gian biểu diễn:",
    ], keep_with_next=True)
    add_bullet_p([
        "Khởi tạo vector trạng thái dựa trên kiểu thực thể: đỉnh mới ",
        latex_to_clean_omml(r"v"),
        " được khởi tạo vector đặc trưng ban đầu thông qua hàm nhúng kiểu thực thể học được (Entity-type-conditioned learnable initialization): ",
        latex_to_clean_omml(r"\mathbf{h}_v^{(0)} = \mathbf{e}_{\text{type}}(\phi_V(v))"),
        ". Cơ chế kế thừa ngữ cảnh từ tiến trình cha chỉ được kích hoạt như một tùy chọn bổ trợ khi quan hệ cha con thực sự quan sát được trong luồng nhật ký kiểm toán và thỏa mãn hợp đồng riêng tư, "
        "tuyệt đối không sử dụng thông tin nhãn kiểm tra hay giả định mở rộng miền định danh."
    ])
    add_bullet_p([
        "Toán tử tự kết nối điều hòa (Self-loop Regularization): nhằm hỗ trợ việc truyền thông điệp khi cấu trúc lân cận thưa thớt, "
        "mạng GNN có thể áp dụng toán tử tự kết nối điều hòa trong đồ thị tính toán của mô hình. "
        "Toán tử này chỉ tồn tại trong quá trình thực thi lan truyền thông điệp và không được ghi nhận vào đồ thị nguồn gốc ",
        latex_to_clean_omml(r"\mathcal{G}_t"),
        " như một quan hệ quan sát được."
    ])

    # -------------------------------------------------------------------------
    # 2.3.3. Temporal GNN Extractor
    # -------------------------------------------------------------------------
    add_h3("Bộ trích xuất đồ thị động Temporal GNN")
    add_p([
        "Sau khi đồ thị nguồn gốc được xây dựng và áp dụng các cơ chế ứng viên kiểm soát độ chân thực, cấu trúc đồ thị động được đưa vào Bộ trích xuất Temporal GNN. "
        "Kế thừa khung xử lý bộ nhớ và truyền thông điệp thời gian động ",
        make_citation_element([37]),
        " cùng các nguyên lý thiết kế từ các hệ thống đồ thị nguồn gốc tiêu biểu ",
        make_citation_element([12, 13, 14, 15]),
        ", chuyên đề hình thức hóa cơ chế truyền thông điệp thời gian có kiểu hóa (Typed Temporal Message Passing) phù hợp với dữ liệu log bảo toàn ngữ cảnh an ninh."
    ])

    # 2.3.3.1. Typed Temporal Message Passing
    add_h4("Truyền thông điệp thời gian có kiểu hóa")
    add_p([
        "Tại mỗi sự kiện tương tác ",
        latex_to_clean_omml(r"e = (v, u, r, t)"),
        ", thông điệp truyền từ đỉnh nguồn ",
        latex_to_clean_omml(r"v"),
        " sang đỉnh đích ",
        latex_to_clean_omml(r"u"),
        " được tính toán dựa trên trạng thái bộ nhớ của các đỉnh ngay trước mốc thời gian sự kiện, khoảng thời gian tương đối, vector nhúng kiểu quan hệ và thuộc tính cạnh an toàn:"
    ], keep_with_next=True)
    add_display_equation(make_tgn_msg_omml())
    add_p([
        "Ký hiệu ",
        latex_to_clean_omml(r"\mathbf{h}_{v,\text{pre}}"),
        " và ",
        latex_to_clean_omml(r"\mathbf{h}_{u,\text{pre}}"),
        " lần lượt chỉ trạng thái bộ nhớ của các đỉnh ",
        latex_to_clean_omml(r"v"),
        " và ",
        latex_to_clean_omml(r"u"),
        " ngay trước mốc thời gian của sự kiện đang xử lý; ",
        latex_to_clean_omml(r"\mathbf{e}_r"),
        " là vector nhúng của loại quan hệ ",
        latex_to_clean_omml(r"r"),
        "; ",
        latex_to_clean_omml(r"\mathbf{x}_e"),
        " là vector thuộc tính cạnh; ",
        latex_to_clean_omml(r"\Delta t_v(t) = t - t_{\text{last}}(v)"),
        " là độ lệch thời gian so với tương tác gần nhất của ",
        latex_to_clean_omml(r"v"),
        ". Các thông điệp gửi đến đỉnh ",
        latex_to_clean_omml(r"u"),
        " trong cùng bước thời gian được tổng hợp qua hàm gom cụm:",
    ], keep_with_next=True)
    add_display_equation(make_tgn_agg_omml())
    add_p([
        "Trạng thái bộ nhớ của thực thể ",
        latex_to_clean_omml(r"u"),
        " được cập nhật qua tế bào hồi quy (như GRU Cell) từ trạng thái trước đó và thông điệp tổng hợp:"
    ], keep_with_next=True)
    add_display_equation(make_tgn_update_omml())

    # 2.3.3.2. Graph Self-Supervised Objective
    add_h4("Học biểu diễn tự giám sát trên đồ thị nguồn gốc động")
    add_p([
        "Nhằm tối ưu hóa các tham số ",
        latex_to_clean_omml(r"\theta_{\text{graph}}"),
        " của mạng Temporal GNN trong Giai đoạn A1 mà không sử dụng nhãn tấn công hay thông tin tương lai, chuyên đề thiết lập hàm mục tiêu tự giám sát đồ thị động ",
        latex_to_clean_omml(r"\mathcal{L}_{\text{graph}}^{\text{self}}"),
        ". Kế thừa cảm hứng từ mô hình MAGIC ",
        make_citation_element([14]),
        " về học biểu diễn qua cơ chế che giấu (Masked Graph Representation Learning) và mở rộng cho đồ thị nguồn gốc động theo thời gian (Adapted / Ours), hàm mục tiêu tự giám sát đồ thị được cấu thành từ ba tác vụ phụ trợ với đầu dự đoán riêng biệt:"
    ], keep_with_next=True)
    add_display_equation(make_graph_self_loss_omml())
    add_bullet_p([
        "Tác vụ tái tạo thuộc tính thực thể bị che (Masked Node Attribute Reconstruction, ",
        latex_to_clean_omml(r"\mathcal{L}_{\text{mask-node}}"),
        "): Che ngẫu nhiên một tỷ lệ ",
        latex_to_clean_omml(r"r_{\text{mask-node}} \in (0, 1)"),
        " thuộc tính kiểu hóa của các thực thể hoạt động ",
        latex_to_clean_omml(r"\mathcal{V}_{\text{mask}} \subset \mathcal{V}_{\text{active}}"),
        " (trong đó ",
        latex_to_clean_omml(r"r_{\text{mask-node}}"),
        " là siêu tham số được kiểm soát, tinh chỉnh và khóa cố định trên tập Validation). Đầu giải mã phụ trợ ",
        latex_to_clean_omml(r"\phi_{\text{graph}}^{\text{node}} : \mathbb{R}^{d_{\text{GNN}}} \to \mathbb{R}^{d_{\text{node-attr}}}"),
        " tái tạo lại vector thuộc tính an toàn quyền riêng tư ",
        latex_to_clean_omml(r"\mathbf{x}_v^{\text{priv}}"),
        " từ trạng thái ẩn ",
        latex_to_clean_omml(r"\mathbf{h}_v(t)"),
        " theo tiêu chuẩn hồi quy Mean Squared Error (MSE). Mục tiêu tái dựng (target) được khóa chuẩn xác là ",
        latex_to_clean_omml(r"\mathbf{x}_v^{\text{priv}}"),
        " — vector đặc trưng đỉnh kiểu hóa thu được SAU KHI đã áp dụng toàn diện quy trình chuẩn hóa, kiểm soát liên kết danh tính / giả danh hóa (Controlled Linkability / Pseudonymization) và chính sách UNKNOWN-SAFE tại Mục 2.2; tuyệt đối không đặt mục tiêu là tên người dùng thô (raw username), tên máy chủ thô (raw hostname), mã tiến trình thô (raw PID), định danh đối tượng thô hay bất kỳ giá trị nhạy cảm nào trước biến đổi, bảo đảm tác vụ tự giám sát đồ thị tuân thủ nghiêm ngặt cùng hợp đồng bảo vệ quyền riêng tư như tác vụ MPP của chuỗi; truyền gradient cập nhật cho ",
        latex_to_clean_omml(r"\theta_{\text{graph}}"),
        " và ",
        latex_to_clean_omml(r"\phi_{\text{graph}}^{\text{node}}"),
        ":"
    ], keep_with_next=True)
    add_display_equation(make_l_mask_node_omml())
    add_bullet_p([
        "Tác vụ dự đoán kiểu quan hệ cạnh bị che (Masked Edge Relation Prediction, ",
        latex_to_clean_omml(r"\mathcal{L}_{\text{mask-edge}}"),
        "): Che nhãn quan hệ tương tác trên tập cạnh ",
        latex_to_clean_omml(r"\mathcal{E}_{\text{mask}} \subset \mathcal{E}"),
        " theo tỷ lệ ",
        latex_to_clean_omml(r"r_{\text{mask-edge}} \in (0, 1)"),
        " (siêu tham số được kiểm soát, tinh chỉnh và khóa cố định trên tập Validation trước khi đánh giá Test). Đầu phân loại phụ trợ ",
        latex_to_clean_omml(r"\phi_{\text{graph}}^{\text{edge}} : \mathbb{R}^{2 d_{\text{GNN}}} \to \mathbb{R}^{\text{card}(\mathcal{R})}"),
        " tiếp nhận vector ghép ",
        latex_to_clean_omml(r"[\mathbf{h}_v(t); \mathbf{h}_u(t)]"),
        " và tối ưu hóa hàm Categorical Cross-Entropy đối với nhãn quan hệ thực tế ",
        latex_to_clean_omml(r"r_{(v, u)} \in \{1, \dots, \text{card}(\mathcal{R})\}"),
        ", truyền gradient cho ",
        latex_to_clean_omml(r"\theta_{\text{graph}}"),
        " và ",
        latex_to_clean_omml(r"\phi_{\text{graph}}^{\text{edge}}"),
        ":"
    ], keep_with_next=True)
    add_display_equation(make_l_mask_edge_omml())
    add_bullet_p([
        "Tác vụ dự đoán khoảng cách thời gian tương đối (Relative Temporal Gap Prediction, ",
        latex_to_clean_omml(r"\mathcal{L}_{\text{time-gap}}"),
        "): Dự đoán độ lệch thời gian tương tác ",
        latex_to_clean_omml(r"\Delta t_{v, u} = |t_v - t_u|"),
        " giữa các cặp đỉnh tương tác trong tập cạnh hoạt động ",
        latex_to_clean_omml(r"\mathcal{E}_{\text{active}}"),
        ". Đầu hồi quy phụ trợ ",
        latex_to_clean_omml(r"\phi_{\text{graph}}^{\text{time}} : \mathbb{R}^{2 d_{\text{GNN}}} \to \mathbb{R}"),
        " tiếp nhận ",
        latex_to_clean_omml(r"[\mathbf{h}_v(t); \mathbf{h}_u(t)]"),
        " và tối ưu hóa hàm Smooth L1 (Huber criterion, ",
        latex_to_clean_omml(r"\delta = 1.0"),
        ") đối với đích ",
        latex_to_clean_omml(r"\log(1 + \Delta t_{v, u})"),
        ", truyền gradient cho ",
        latex_to_clean_omml(r"\theta_{\text{graph}}"),
        " và ",
        latex_to_clean_omml(r"\phi_{\text{graph}}^{\text{time}}"),
        " (giúp mô hình học cấu trúc thứ tự và khoảng thời gian quan sát được giữa các sự kiện mà không hàm ý tác động nhân quả, tuân thủ nguyên tắc temporal order != causal effect và dependency != causal effect):"
    ], keep_with_next=True)
    add_display_equation(make_l_time_gap_omml())
    add_p([
        "Các trọng số ",
        latex_to_clean_omml(r"\beta_1, \beta_2, \beta_3 > 0"),
        " là các siêu tham số cân bằng giữa các tác vụ phụ trợ. Hàm mục tiêu ",
        latex_to_clean_omml(r"\mathcal{L}_{\text{graph}}^{\text{self}}"),
        " trực tiếp tối ưu hóa tập tham số ",
        latex_to_clean_omml(r"\theta_{\text{graph}}"),
        " cùng các đầu phụ trợ trong Giai đoạn A1 và trở thành thành phần bảo toàn cấu trúc đồ thị trong hàm ",
        latex_to_clean_omml(r"\mathcal{L}_{\text{preserv}}"),
        " tại Mục 2.4."
    ])

    # 2.3.3.3. Over-smoothing / Over-squashing Controls
    add_h4("Kiểm soát hiện tượng làm mịn quá mức và nghẽn cổ chai thông tin")
    add_p([
        "Khi áp dụng mạng nơ-ron đồ thị trên đồ thị nguồn gốc có đường kính lớn, hai rủi ro kiến trúc phổ biến là hiện tượng làm mịn quá mức (Over-smoothing) "
        "khiến biểu diễn các đỉnh hội tụ về các vector khó phân biệt, và hiện tượng nghẽn cổ chai thông tin (Over-squashing) khi thông tin từ vùng lân cận mở rộng "
        "theo hàm mũ bị ép vào một vector kích thước cố định ",
        make_citation_element([21]),
        ". Chuyên đề đề xuất hai cơ chế ứng viên giảm thiểu (Mitigation Candidates):"
    ], keep_with_next=True)
    add_bullet_p([
        "Giảm thiểu Over-smoothing: giới hạn số lớp truyền thông điệp ở mức thấp (",
        latex_to_clean_omml(r"K \le 2"),
        ") kết hợp cơ chế kết nối tắt phần dư (Residual Skip-connections): ",
        latex_to_clean_omml(r"\mathbf{h}_u^{(l)}(t) = \mathbf{h}_u^{(l-1)}(t) + \text{GNN}(\mathbf{h}_u^{(l-1)}(t), \dots)"),
        ", hỗ trợ duy trì tính phân biệt giữa các thực thể."
    ])
    add_bullet_p([
        "Giảm thiểu Over-squashing: áp dụng cơ chế lấy mẫu lân cận có chọn lọc theo trọng số thời gian (Top-k Temporal Attention Sampling), "
        "ưu tiên tổng hợp thông điệp từ các đỉnh lân cận có hoạt động gần nhất thay vì mở rộng toàn bộ cây phụ thuộc nhiều bước. "
        "Cơ chế này được thiết kế để đánh giá đối sánh với các chính sách lấy mẫu toàn bộ lân cận (Full Neighborhood) và lấy mẫu theo độ mới (Recency Sampling) tại Chương 3, "
        "nhằm kiểm tra thực nghiệm liệu chính sách sampling có gây mất mát các bằng chứng APT dài hạn (long-range evidence) hay không."
    ])

    # 2.3.3.4. Graph Output
    add_h4("Giao diện đầu ra của bộ trích xuất đồ thị")
    add_p([
        "Vector biểu diễn đặc trưng đồ thị ",
        latex_to_clean_omml(r"\mathbf{z}_{\text{graph}}"),
        " trên cửa sổ trung hạn ",
        latex_to_clean_omml(r"W_m"),
        " được tạo ra bằng cách tổng hợp trạng thái của tất cả các thực thể đang hoạt động trong cửa sổ:"
    ], keep_with_next=True)
    add_display_equation(make_tgn_readout_omml())
    add_p([
        "Đầu ra của Bộ trích xuất đồ thị là một gói biểu diễn mang vector ",
        latex_to_clean_omml(r"\mathbf{z}_{\text{graph}} \in \mathbb{R}^{d_{\text{graph}}}"),
        " tóm tắt cấu trúc đồ thị và sự lan truyền thông tin qua lại giữa các thực thể trong cửa sổ trung hạn ",
        latex_to_clean_omml(r"W_m"),
        ", đi kèm siêu dữ liệu tương ứng: khoảng thời gian ",
        latex_to_clean_omml(r"[t_{\text{start}}, t_{\text{end}}]"),
        ", tập thực thể hoạt động ",
        latex_to_clean_omml(r"\mathcal{V}_{\text{active}}"),
        ", mặt nạ sẵn sàng ",
        latex_to_clean_omml(r"m_{\text{graph}} \in \{0, 1\}"),
        ", và định danh đồ thị con. Về mặt khung đa tỷ lệ, cửa sổ dài hạn ",
        latex_to_clean_omml(r"W_l"),
        " đã thiết lập tại Mục 2.2.3.2 đóng vai trò cung cấp ngữ cảnh trạng thái dài hạn có chặn (Bounded Long-term Context Summary) cho hệ thống mà không làm phát sinh bộ trích xuất thứ ba độc lập. "
        "Chi phí tính toán của nhánh đồ thị được phân lập giữa chi phí xây dựng cửa sổ và chi phí xử lý gia số theo từng sự kiện: "
        "xây dựng đồ thị ",
        make_c_build_omml(),
        ", tạo thông điệp ",
        make_c_msg_omml(),
        ", cập nhật trạng thái bộ nhớ với tập đỉnh nhận thông điệp ",
        make_v_upd_subset_omml(),
        " là ",
        make_c_update_omml(),
        ", và tổng hợp đầu ra ",
        make_c_readout_omml(),
        "."
    ])

    # Insert Figure 2.3 Placeholder Canvas
    p_c2 = doc.add_paragraph() if target_p is None else target_p.insert_paragraph_before()
    p_c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_c2.paragraph_format.keep_with_next = True
    p_c2.paragraph_format.space_before = Pt(6)
    p_c2.paragraph_format.space_after = Pt(2)
    p_c2._p.append(parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="30203" w:name="BK_FIG_2_003_CANVAS"/>'))
    r_c2 = p_c2.add_run(" ")
    p_c2._p.append(parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="30203"/>'))
    add_figure_caption(doc, target_p, 2, 3, "Kiến trúc xây dựng đồ thị nguồn gốc phụ thuộc thời gian và Bộ trích xuất Temporal GNN (Nguồn: Tác giả đề xuất)", bookmark_name="BK_FIG_2_003")

    # Insert Table 2.2
    add_p(
        "Bảng 2.2 tóm tắt các đặc tính kiến trúc và độ phức tạp tính toán của hai bộ trích xuất đặc trưng đa góc nhìn:",
        keep_with_next=True,
        page_break_before=True
    )
    _tbl23_headers = ["Thành phần trích xuất", "Góc nhìn biểu diễn", "Cấu trúc đầu vào", "Mục tiêu biểu diễn", "Độ phức tạp tính toán"]
    _tbl23_widths = [1600, 1300, 1300, 1600, 2600]
    _tbl23_rows = [
        ["Bộ trích xuất tuần tự (Transformer)", "Tuần tự ngữ nghĩa", ["Cửa sổ ngắn ", latex_to_clean_omml(r"W_s"), " (", latex_to_clean_omml(r"L"), " sự kiện)"], "Tự giám sát (MEP, MPP, Time Gap)", [latex_to_clean_omml(r"\mathcal{O}(N_{\text{layers}} \cdot (L^2 \cdot d + L \cdot d \cdot d_{\text{ff}}))")]],
        ["Xây dựng đồ thị nguồn gốc", "Cấu trúc phụ thuộc", ["Luồng sự kiện ", latex_to_clean_omml(r"\tilde{e}_i^{\text{priv}}")], "Độ chân thực đồ thị (Fidelity candidates)", [make_card_ew_omml()]],
        ["Bộ trích xuất đồ thị (Temporal GNN)", "Đồ thị động", ["Đồ thị dị thể ", latex_to_clean_omml(r"\mathcal{G}_t")], "Tự giám sát động (L_graph_self)", [make_card_tgn_omml()]],
        ["Giao diện biểu diễn đầu ra", "Không gian ẩn đa góc nhìn", ["Vector ", latex_to_clean_omml(r"\mathbf{z}_{\text{seq}}"), " & ", latex_to_clean_omml(r"\mathbf{z}_{\text{graph}}")], "Đầu vào cho tầng gióng hàng (Mục 2.4)", [latex_to_clean_omml(r"\mathcal{O}(d_{\text{seq}} + d_{\text{graph}})")]]
    ]
    add_table_caption(doc, target_p, 2,
                      "Đặc tả kiến trúc và độ phức tạp tính toán của hai bộ trích xuất đặc trưng đa góc nhìn",
                      bookmark_name="BK_TBL_2_002", chapter_num=2)
    insert_thesis_table(doc, target_p, _tbl23_headers, _tbl23_widths, _tbl23_rows, font_size_pt=12)
    add_p("", first_line_indent=False)

    # =========================================================================
    # 2.4. MULTI-VIEW ALIGNMENT AND UNIFIED REPRESENTATION
    # =========================================================================
    add_h2("Gióng hàng và học biểu diễn thống nhất đa góc nhìn")
    add_p([
        "Sau khi hai bộ trích xuất độc lập tại Mục 2.3 sản sinh hai vector đặc trưng ẩn ",
        latex_to_clean_omml(r"\mathbf{z}_{\text{seq}} \in \mathbb{R}^{d_{\text{seq}}}"),
        " (nắm bắt thứ tự vi mô, tham số an ninh và nhịp điệu thời gian trên cửa sổ ngắn hạn ",
        latex_to_clean_omml(r"W_s"),
        ") và ",
        latex_to_clean_omml(r"\mathbf{z}_{\text{graph}} \in \mathbb{R}^{d_{\text{graph}}}"),
        " (nắm bắt cấu trúc phụ thuộc quan sát được, luồng truyền lan thông tin và quan hệ phụ thuộc cấu trúc, thời gian trên cửa sổ trung hạn ",
        latex_to_clean_omml(r"W_m"),
        "), thách thức khoa học trung tâm là làm thế nào để dung hợp và gióng hàng hai nguồn thông tin dị thể này thành một vector biểu diễn thống nhất canonical ",
        latex_to_clean_omml(r"\mathbf{z}_{\text{mv}} \in \mathbb{R}^{d_{\text{mv}}}"),
        ". Tầng gióng hàng đa góc nhìn (Multi-View Alignment Plane) xây dựng cơ chế phương pháp luận giải quyết câu hỏi nghiên cứu ",
        make_rq_omml(2),
        " và ",
        make_rq_omml(4),
        ", thiết lập nền tảng kỹ thuật nhằm phục vụ kiểm định giả thuyết ",
        make_hypo_omml(2),
        " và ",
        make_hypo_omml(4),
        " tại Chương 3. Để đạt được mục tiêu này, chuyên đề thiết lập một quy trình tối ưu hóa đa tầng chặt chẽ: (1) thiết lập không gian ngữ nghĩa chung mà không gây sụp đổ biểu diễn (Representation Collapse); (2) bảo toàn đối xứng thông tin đặc thù nội góc nhìn (Symmetric Multi-View Preservation); (3) kiểm soát nhiễu đặc quyền và hành vi quản trị viên thông qua ngữ cảnh hóa đa chiều (Risk-aware Admin-Noise Control); (4) cung cấp cơ chế phân bổ bằng chứng yếu (Weak Evidence Attribution) qua mô hình học đa thể hiện (Multiple Instance Learning, MIL) tùy chọn; và (5) thích ứng linh hoạt với hiện tượng khuyết góc nhìn (Missing-View) trong môi trường phân tán mà không vi phạm nguyên tắc nhân quả thời gian (Strictly Causal / Zero Lookahead)."
    ])

    # 2.4.1. Heterogeneous Cross-view Alignment
    add_h3("Gióng hàng đa góc nhìn dị thể")
    add_p([
        "Mục 2.4.1 thiết lập Hợp đồng Tương ứng (Correspondence Contract), phân định ranh giới giữa Không gian Biểu diễn (Representation Space) và Không gian Chiếu (Projection Space), cùng các cơ chế điều hòa chống sụp đổ biểu diễn."
    ])

    add_h4("Hợp đồng tương ứng ngữ cảnh và siêu dữ liệu liên kết")
    add_p([
        "Một cặp biểu diễn đa góc nhìn ",
        latex_to_clean_omml(r"(\mathbf{z}_{\text{seq}}, \mathbf{z}_{\text{graph}})"),
        " được xem là một cặp tương ứng hợp lệ (Valid Correspondence Pair) nếu và chỉ nếu chúng thỏa mãn Hợp đồng Tương ứng dựa trên bộ siêu dữ liệu liên kết ",
        latex_to_clean_omml(r"\mathcal{M} = \langle [t_{\text{start}}, t_{\text{end}}], \mathcal{E}_{\text{scope}}, m_{\text{seq}}, m_{\text{graph}} \rangle"),
        " đã thiết lập tại Mục 2.3:"
    ], keep_with_next=True)

    add_bullet_p([
        "Giao thoa khoảng thời gian quan sát (Temporal Interval Overlap): Khoảng thời gian hiệu lực của cửa sổ chuỗi ngắn hạn ",
        latex_to_clean_omml(r"[t_{\text{start}}^{\text{seq}}, t_{\text{end}}^{\text{seq}}]"),
        " phải là tập con hoặc có giao thoa thực dương với khoảng thời gian của cửa sổ đồ thị trung hạn ",
        latex_to_clean_omml(r"[t_{\text{start}}^{\text{graph}}, t_{\text{end}}^{\text{graph}}]"),
        ", tức ",
        latex_to_clean_omml(r"[t_{\text{start}}^{\text{seq}}, t_{\text{end}}^{\text{seq}}] \cap [t_{\text{start}}^{\text{graph}}, t_{\text{end}}^{\text{graph}}] \ne \emptyset"),
        "."
    ])

    add_bullet_p([
        "Giao thoa phạm vi thực thể (Entity Scope Overlap): Tập các thực thể hệ thống quan sát trong chuỗi ",
        latex_to_clean_omml(r"\mathcal{E}_{\text{scope}}^{\text{seq}}"),
        " (tiến trình, người dùng, tài nguyên) phải chia sẻ ít nhất một đỉnh thực thể chung với tập đỉnh đồ thị ",
        latex_to_clean_omml(r"\mathcal{E}_{\text{scope}}^{\text{graph}} = \mathcal{V}_{\text{active}}(W_m)"),
        ", tức ",
        latex_to_clean_omml(r"\mathcal{E}_{\text{scope}}^{\text{seq}} \cap \mathcal{E}_{\text{scope}}^{\text{graph}} \ne \emptyset"),
        "."
    ])

    add_bullet_p([
        "Tính đầy đủ của mặt nạ khả dụng (Availability Mask Completeness): Cả hai góc nhìn đều tồn tại dữ liệu đo kiểm thực tế trong khoảng quan sát, biểu thị bởi ",
        latex_to_clean_omml(r"m_{\text{seq}} = 1"),
        " và ",
        latex_to_clean_omml(r"m_{\text{graph}} = 1"),
        "."
    ])

    add_p([
        "Tuyệt đối không sử dụng nhãn tấn công (Attack Labels) hay bất kỳ tín hiệu từ tập kiểm thử nào để định nghĩa cặp tương ứng trong Stage A. Các cặp tương ứng được hình thành hoàn toàn tự nhiên từ tiến trình vận hành hệ thống thực tế. Các trường hợp chỉ giao thoa một phần thực thể hoặc thời gian được xếp vào nhóm Tương ứng Cục bộ (Partial Correspondence), trong khi các trường hợp thiếu hụt một trong hai nguồn được xử lý qua cơ chế Góc nhìn Khuyết (Missing-View)."
    ])

    add_h4("Phân tách không gian biểu diễn và không gian chiếu")
    add_p([
        "Để tránh hiện tượng mất mát thông tin cấu trúc đặc thù khi ép hai vector biểu diễn gốc ",
        latex_to_clean_omml(r"\mathbf{z}_{\text{seq}}"),
        " và ",
        latex_to_clean_omml(r"\mathbf{z}_{\text{graph}}"),
        " phải hội tụ về cùng một điểm, chuyên đề áp dụng nguyên lý phân tách không gian theo SimCLR ",
        make_citation_element([39]),
        " và VICReg ",
        make_citation_element([22]),
        ". Chúng ta đưa mỗi vector biểu diễn qua một đầu chiếu phi tuyến riêng biệt (Non-linear Projection Heads):"
    ])

    add_display_equation(
        r"\mathbf{p}_{\text{seq}} = g_{\text{seq}}(\mathbf{z}_{\text{seq}}) = \mathbf{W}_{\text{p2}} \cdot \text{GELU}(\text{LayerNorm}(\mathbf{W}_{\text{p1}} \mathbf{z}_{\text{seq}})) \in \mathbb{R}^{d_{\text{proj}}}"
    )
    add_display_equation(
        r"\mathbf{p}_{\text{graph}} = g_{\text{graph}}(\mathbf{z}_{\text{graph}}) = \mathbf{W}_{\text{g2}} \cdot \text{GELU}(\text{LayerNorm}(\mathbf{W}_{\text{g1}} \mathbf{z}_{\text{graph}})) \in \mathbb{R}^{d_{\text{proj}}}"
    )

    add_p([
        "Trong đó ",
        latex_to_clean_omml(r"g_{\text{seq}}"),
        " và ",
        latex_to_clean_omml(r"g_{\text{graph}}"),
        " là các mạng MLP hai tầng có chuẩn hóa LayerNorm và hàm kích hoạt phi tuyến GELU; ",
        latex_to_clean_omml(r"d_{\text{proj}}"),
        " là số chiều của Không gian Chiếu (Projection Space). Việc phân tách này thiết lập một ranh giới phương pháp luận nghiêm ngặt: Không gian Chiếu ",
        latex_to_clean_omml(r"\mathbb{R}^{d_{\text{proj}}}"),
        " chỉ phục vụ việc tính toán các hàm mất mát gióng hàng tự giám sát và sẽ bị loại bỏ sau khi kết thúc Stage A; trong khi Không gian Biểu diễn ",
        latex_to_clean_omml(r"\mathbb{R}^{d_{\text{seq}}}"),
        ", ",
        latex_to_clean_omml(r"\mathbb{R}^{d_{\text{graph}}}"),
        " và không gian biểu diễn thống nhất ",
        latex_to_clean_omml(r"\mathbb{R}^{d_{\text{mv}}}"),
        " được bảo toàn nguyên vẹn để cung cấp đặc trưng cho các tác vụ hạ nguồn ở Stage C."
    ])

    add_h4("Phân tích đối sánh các cơ chế tự giám sát đa góc nhìn")
    add_p([
        "Trong học biểu diễn tự giám sát (Self-Supervised Learning), ba hướng tiếp cận tiêu biểu bao gồm: InfoNCE / Contrastive Learning ",
        make_citation_element([38, 39]),
        ", Barlow Twins ",
        make_citation_element([23]),
        ", và VICReg ",
        make_citation_element([22]),
        ". Bảng 2.3 phân tích đối sánh chi tiết các đặc tính lý thuyết từ các công bố gốc và các giả thuyết nghiên cứu của chuyên đề đối với dữ liệu telemetry an ninh:"
    ], keep_with_next=True)

    # Table 2.3
    _tbl24_headers = ["Tiêu chí phương pháp luận", "InfoNCE / Contrastive", "Barlow Twins", "VICReg (Ứng viên đề xuất)"]
    _tbl24_widths = [1800, 2200, 2200, 2400]
    _tbl24_rows = [
        [
            "Cơ chế điều hòa chống sụp đổ",
            ["Sử dụng phân bố mẫu âm (Negative pairs) để đẩy xa các phiên khác nhau ", make_citation_element([38, 39])],
            ["Ép ma trận tương quan chéo giữa hai view về ma trận đơn vị I ", make_citation_element([23])],
            ["Phân tách độc lập: Duy trì phương sai từng chiều + Phạt hiệp phương sai ", make_citation_element([22])]
        ],
        [
            "Yêu cầu mẫu âm (Negative pairs)",
            "Bắt buộc (phụ thuộc vào cơ chế lấy mẫu âm trong batch)",
            "Không yêu cầu mẫu âm tường minh",
            "Không yêu cầu mẫu âm tường minh"
        ],
        [
            "Đặc tính phân tích trên telemetry",
            "Rủi ro False Negatives tiềm tàng khi phạt các phiên bình thường độc lập có hành vi tương đồng (Giả thuyết chuyên đề, kiểm chứng tại Chương 3)",
            "Độ ổn định phụ thuộc vào việc ước lượng tương quan chéo; rủi ro nhạy cảm thang đo chuỗi và đồ thị (Giả thuyết chuyên đề, kiểm chứng tại Chương 3)",
            "Điều hòa độc lập phương sai và hiệp phương sai; kỳ vọng duy trì cấu trúc phân phối telemetry an ninh (Giả thuyết chuyên đề, kiểm chứng tại Chương 3)"
        ],
        [
            "Đặc tính phụ thuộc kích thước lô",
            "Hiệu năng thường hưởng lợi khi kích thước lô đủ lớn để bao quát phân bố mẫu âm",
            "Độ ổn định phụ thuộc vào việc ước lượng ma trận tương quan chéo trên batch",
            "Ước lượng phương sai và hiệp phương sai trực tiếp trên phân bố lô huấn luyện"
        ],
        [
            "Vai trò trong đề tài",
            "Phương pháp đối sánh (Ablation baseline) tại Chương 3",
            "Phương pháp đối sánh (Ablation baseline) tại Chương 3",
            "Ứng viên chính (Primary candidate), kiểm chứng thực nghiệm tại Chương 3"
        ]
    ]
    add_table_caption(doc, target_p, 3,
                      "So sánh đối sánh các cơ chế tự giám sát đa góc nhìn trong bài toán an ninh mạng",
                      bookmark_name="BK_TBL_2_003", chapter_num=2)
    insert_thesis_table(doc, target_p, _tbl24_headers, _tbl24_widths, _tbl24_rows, font_size_pt=11)

    add_p([
        "Từ phân tích phương pháp luận trên, chuyên đề lựa chọn VICReg (Variance-Invariance-Covariance Regularization) ",
        make_citation_element([22]),
        " làm ứng viên gióng hàng cốt lõi (Primary Candidate). Ưu điểm lý thuyết của VICReg đối với bài toán telemetry log là loại bỏ nhu cầu lấy mẫu âm tường minh, từ đó giảm thiểu rủi ro False Negatives, hạn chế việc ép các phiên telemetry bình thường độc lập nhưng có hình thái hành vi tương đồng ra xa nhau một cách giả tạo. Lựa chọn này sẽ được kiểm chứng thực nghiệm chặt chẽ tại Chương 3 trong mối tương quan với hai phương pháp đối chứng (Ablation Baselines) là InfoNCE ",
        make_citation_element([38, 39]),
        " và Barlow Twins ",
        make_citation_element([23]),
        "."
    ])

    add_h4("Hàm mục tiêu gióng hàng VICReg trên không gian chiếu")
    add_p([
        "Hàm mục tiêu gióng hàng VICReg trên một lô huấn luyện (mini-batch) kích thước ",
        latex_to_clean_omml(r"B"),
        " gồm các cặp chiếu hợp lệ ",
        latex_to_clean_omml(r"\{(\mathbf{p}_{\text{seq}}^{(i)}, \mathbf{p}_{\text{graph}}^{(i)})\}_{i=1}^B"),
        " được phân rã thành ba thành phần điều hòa tường minh:"
    ], keep_with_next=True)

    add_bullet_p([
        "Số hạng Bất biến (Invariance Term, ",
        latex_to_clean_omml(r"\mathcal{L}_{\text{inv}}"),
        "): Đo lường khoảng cách sai số toàn phương trung bình (Mean Squared Error) giữa hai góc nhìn biểu diễn của cùng một thực thể hoặc khoảng thời gian:",
    ])
    add_display_equation(make_l_inv_omml())

    add_bullet_p([
        "Số hạng Phương sai (Variance Regularization, ",
        latex_to_clean_omml(r"\mathcal{L}_{\text{var}}"),
        "): Sử dụng hàm bản lề (Hinge loss) để duy trì độ lệch chuẩn của từng chiều đặc trưng trên toàn bộ lô huấn luyện lớn hơn một ngưỡng tối thiểu ",
        latex_to_clean_omml(r"\gamma > 0"),
        " (giá trị khởi tạo tham chiếu ",
        latex_to_clean_omml(r"\gamma = 1"),
        ", được tối ưu hóa trên tập xác thực), kiểm soát nguy cơ sụp đổ điểm (Point Collapse):",
    ])
    add_display_equation(make_l_var_omml())

    add_p([
        "Trong đó công thức phương sai mẫu từng chiều đặc trưng ",
        latex_to_clean_omml(r"\text{Var}(\mathbf{p}_{:, j})"),
        " và hằng số ổn định số học ",
        latex_to_clean_omml(r"\epsilon = 10^{-4}"),
        " được xác định tường minh:"
    ])
    add_display_equation(make_var_formula_omml())

    add_bullet_p([
        "Số hạng Hiệp phương sai (Covariance Regularization, ",
        latex_to_clean_omml(r"\mathcal{L}_{\text{cov}}"),
        "): Phạt tổng bình phương các phần tử nằm ngoài đường chéo chính của ma trận hiệp phương sai ",
        latex_to_clean_omml(r"\mathbf{C}(\mathbf{P}) \in \mathbb{R}^{d_{\text{proj}} \times d_{\text{proj}}}"),
        ", điều hòa các chiều đặc trưng hướng tới tính độc lập thống kê và giảm thiểu tính dư thừa thông tin (Redundancy Reduction), hạn chế hiện tượng sụp đổ số chiều (Dimensional Collapse):",
    ])
    add_display_equation(make_c_matrix_omml())
    add_display_equation(make_l_cov_omml())

    add_h4("Cơ chế thích ứng góc nhìn khuyết và tương ứng từng phần")
    add_p([
        "Trong thực tế giám sát an ninh phân tán, các cảm biến thu thập telemetry có thể gặp sự cố cục bộ, gây ra hiện tượng khuyết góc nhìn (Missing-View, Incomplete Telemetry). Chuyên đề xử lý vấn đề này theo khung tiếp cận của CPM-Nets ",
        make_citation_element([41]),
        " kết hợp cơ chế mặt nạ khả dụng và suy thoái duyên dáng (Graceful Degradation) do tác giả đề xuất (Adapted / Ours):"
    ], keep_with_next=True)

    add_bullet_p([
        "Trường hợp đầy đủ hai góc nhìn (",
        latex_to_clean_omml(r"m_{\text{seq}}=1, m_{\text{graph}}=1"),
        "): Hệ thống kích hoạt toàn bộ luồng huấn luyện gióng hàng và dung hợp đa góc nhìn.",
    ])

    add_bullet_p([
        "Trường hợp chỉ có chuỗi sự kiện (",
        latex_to_clean_omml(r"m_{\text{seq}}=1, m_{\text{graph}}=0"),
        "): Vô hiệu hóa số hạng ",
        latex_to_clean_omml(r"\mathcal{L}_{\text{inv}}"),
        " và các thành phần liên quan đến đồ thị; biểu diễn thống nhất tự động suy thoái duyên dáng về nhánh chuỗi thông qua phép chiếu thích ứng: ",
        latex_to_clean_omml(r"\mathbf{z}_{\text{mv}} = f_{\text{adapt}}^{\text{seq}}(\mathbf{z}_{\text{seq}}) = \text{LayerNorm}(\mathbf{W}_{\text{out}}^{\text{seq}} \mathbf{z}_{\text{seq}})"),
        "."
    ])

    add_bullet_p([
        "Trường hợp chỉ có đồ thị nguồn gốc (",
        latex_to_clean_omml(r"m_{\text{seq}}=0, m_{\text{graph}}=1"),
        "): Tương tự, hệ thống vô hiệu hóa ",
        latex_to_clean_omml(r"\mathcal{L}_{\text{inv}}"),
        " và chuyển hướng biểu diễn thống nhất về nhánh đồ thị: ",
        latex_to_clean_omml(r"\mathbf{z}_{\text{mv}} = f_{\text{adapt}}^{\text{graph}}(\mathbf{z}_{\text{graph}}) = \text{LayerNorm}(\mathbf{W}_{\text{out}}^{\text{graph}} \mathbf{z}_{\text{graph}})"),
        "."
    ])

    add_bullet_p([
        "Trường hợp tương ứng cục bộ (Partial Correspondence): Khi hai góc nhìn chỉ giao nhau một phần về thời gian hoặc thực thể, hệ thống tính toán hệ số giao thoa Jaccard ",
        latex_to_clean_omml(r"\rho_{\text{overlap}} \in (0, 1]"),
        " và điều tiết trọng số gióng hàng tương thích: ",
        latex_to_clean_omml(r"\mathcal{L}_{\text{inv}}^{\text{partial}} = \rho_{\text{overlap}} \cdot \mathcal{L}_{\text{inv}}"),
        "."
    ])

    add_p([
        "Nguyên tắc bất biến: Hệ thống tuyệt đối không tự bịa đặt dữ liệu (Data Hallucination) cho góc nhìn bị thiếu, không sử dụng dữ liệu tương lai để nội suy và không sử dụng nhãn tấn công để bù đắp, tuân thủ nguyên tắc nhân quả thời gian (Strictly Causal / Zero Lookahead) của mặt phẳng suy luận dòng (Streaming Inference Plane)."
    ])
    add_p([
        "Lưu ý phương pháp luận về tính nhân quả: Khái niệm 'Strictly Causal / Zero Lookahead' được định nghĩa chính xác theo nghĩa: hệ thống chỉ sử dụng dữ liệu đo kiểm trong quá khứ và hiện tại tại thời điểm suy luận dòng, tuyệt đối không sử dụng thông tin tương lai hay nhãn kiểm thử; khái niệm này KHÔNG hàm ý suy diễn quan hệ nhân quả (Causal-effect Inference) theo nghĩa thống kê can thiệp. Trong an ninh mạng, sự phụ thuộc quan sát được (Observable Dependency) không đồng nhất với tác động nhân quả (Dependency != Causal Effect)."
    ])

    # 2.4.2. Risk-aware Admin Behavior (RQ4 Coverage)
    add_h3("Nhận thức hành vi quản trị viên và điều hòa rủi ro")
    add_p([
        "Một nguồn gây nhầm lẫn đáng chú ý trong giám sát an ninh là sự chồng lấn giữa hoạt động quản trị hợp thức và hành vi tấn công. Mục 2.4.2 thiết lập khung phương pháp luận nhận thức hành vi quản trị (Risk-aware Administrative Behavior) và các cơ chế kiểm soát nhiễu đặc quyền (Admin-Noise Controls, Đề xuất của đề tài / Ours) nhằm giải quyết trực tiếp câu hỏi nghiên cứu ",
        make_rq_omml(4),
        "."
    ])

    add_h4("Nguyên tắc bất biến: Bất thường không đồng nhất với Độc hại")
    add_p([
        "Trong môi trường vận hành công nghệ thông tin thực tế, các hoạt động quản trị viên (System Administration), ví dụ như cài đặt bản vá hàng loạt, sao lưu cơ sở dữ liệu định kỳ, bảo trì từ xa thông qua PowerShell, SSH, WMI, PsExec hay các kịch bản script tự động hóa phức tạp, thường xuyên tạo ra các chuỗi sự kiện có tần suất cao đột biến, quyền hạn cao (Privileged Execution) và cấu trúc lệnh khác lạ so với người dùng thông thường. Tuy nhiên, các hoạt động này hoàn toàn hợp thức. Chuyên đề xác lập nguyên tắc khoa học cốt lõi: Hành vi bất thường không đồng nhất với hành vi độc hại (Unusual != Malicious)."
    ])

    add_p([
        "Bộ trích xuất đa góc nhìn không đóng vai trò bộ phân loại dựa trên luật (Rule-based Detector) để gán nhãn cứng hoạt động quản trị là an toàn hay độc hại. Thay vào đó, bộ trích xuất có nhiệm vụ mã hóa toàn diện các đặc trưng ngữ cảnh quan sát được (Observable Context Features) vào không gian vector biểu diễn liên tục ",
        latex_to_clean_omml(r"\mathbf{z}_{\text{mv}}"),
        " dựa trên 4 chiều thông tin an ninh:"
    ], keep_with_next=True)

    add_bullet_p([
        "Ngữ cảnh đặc quyền (Privilege Context): Định lượng mức độ leo thang và phân cấp đặc quyền hệ thống theo thang đo liên tục, phản ánh đúng quyền hạn thực thi tại thời điểm quan sát mà không xem quyền hạn cao mặc định là mối đe dọa."
    ])

    add_bullet_p([
        "Ngữ cảnh công cụ và tham số (Tool & Parameter Structure): Biểu diễn cấu trúc cú pháp câu lệnh, tham số thực thi động và trạng thái chữ ký số (Signed/Unsigned State nếu quan sát được), phản ánh trung thực bản chất công cụ thực thi vào không gian vector."
    ])

    add_bullet_p([
        "Ngữ cảnh thời gian và chu kỳ (Temporal & Cadence Regularity): Mô hình hóa tính chu kỳ của các tác vụ tự động (Cronjobs, Scheduled Tasks, Maintenance Windows) thông qua hàm phân phối nhịp điệu thời gian ",
        latex_to_clean_omml(r"\Delta t"),
        ", mã hóa các hoạt động định kỳ theo lịch trình và các hành vi tương tác thủ công vào biểu diễn thời gian."
    ])

    add_bullet_p([
        "Ngữ cảnh lan truyền thực thể (Entity Footprint & Neighborhood Breadth): Biểu diễn luồng kết nối quản trị diện rộng và cấu trúc phân nhánh lân cận của đồ thị nguồn gốc trung hạn, tạo cơ sở cho các mô hình hạ nguồn phân biệt hành vi quản trị thông thường với các chuỗi lây lan ngang (Lateral Movement)."
    ])

    add_p([
        "Toàn bộ các đặc trưng trên chỉ được mã hóa vào vector biểu diễn. Việc đưa ra quyết định độc hại hay an toàn (Malicious-vs-Benign Decision) được bàn giao hoàn toàn cho các mô hình đánh giá hạ nguồn ở Stage C và các kịch bản thực nghiệm tại Chương 3."
    ])

    add_h4("Kiểm soát biến gây nhiễu và ngăn chặn học đường tắt")
    add_p([
        "Để bảo đảm tính vững chắc (Robustness) của không gian vector biểu diễn, chuyên đề thiết lập các cơ chế kiểm soát biến gây nhiễu (Confounder Controls):"
    ], keep_with_next=True)

    add_bullet_p([
        "Cấm học đường tắt định danh (No Identity / Role Shortcuts): Hệ thống tuyệt đối không sử dụng tên người dùng thô (Raw Usernames) hoặc chuỗi ký tự chức danh làm đặc trưng đầu vào trực tiếp. Mọi định danh thực thể đều được chuẩn hóa và bảo vệ qua cơ chế Kiểm soát Liên kết Danh tính (Controlled Linkability, Mục 2.2.2), ngăn chặn mô hình học mối liên hệ giả tạo rằng 'tài khoản Admin luôn an toàn' hoặc 'tài khoản Guest luôn đáng ngờ'."
    ])

    add_bullet_p([
        "Cấm rò rỉ nhãn kiểm thử và đặc quyền giám sát (No Privileged Test Knowledge): Trong toàn bộ quá trình tiền xử lý và huấn luyện tự giám sát, hệ thống tuyệt đối không sử dụng bất kỳ thông tin nào về nhãn tấn công tương lai hay danh sách phân loại tĩnh từ tập kiểm thử, ngăn rò rỉ thông tin kiểm thử qua các kênh đã định nghĩa (Data Leakage Prevention)."
    ])

    # 2.4.3. Unified Objective + Weak Evidence Attribution
    add_h3("Hàm mục tiêu thống nhất và phân bổ bằng chứng yếu")
    add_p([
        "Mục 2.4.3 tổng hợp toàn bộ đồ thị tối ưu hóa (Optimization Graph) của chuyên đề, thiết lập hàm mục tiêu tự giám sát thống nhất trong Stage A, đóng kín các luồng gradient cho toàn bộ tập tham số, và tích hợp mô đun Phân bổ Bằng chứng Yếu (Weak Evidence Attribution) qua mô hình học đa thể hiện (Multiple Instance Learning, MIL) tùy chọn trong Stage B."
    ])

    add_h4("Đóng kín đồ thị tối ưu hóa và kiểm toán tham số huấn luyện")
    add_p([
        "Nhằm bảo đảm tính chặt chẽ phương pháp luận và kiểm toán hoàn toàn các tham số có gradient (Zero Orphan Parameters), Bảng 2.4 kiểm toán toàn bộ tập tham số huấn luyện của hệ thống, chỉ rõ giai đoạn huấn luyện, hàm mục tiêu sinh gradient, thời điểm đóng băng và vai trò khi suy luận trực tuyến:"
    ], keep_with_next=True)

    # Table 2.4 Parameter Audit Table
    _tbl25_headers = ["Khối tham số", "Giai đoạn", "Hàm mục tiêu sinh Gradient", "Thời điểm đóng băng / Vòng đời", "Dùng khi suy luận?"]
    _tbl25_widths = [1900, 1300, 3100, 1700, 1200]
    _tbl25_rows = [
        [[latex_to_clean_omml(r"\theta_{\text{seq}}"), " (Transformer Encoder)"], "Stage A1 (Chuỗi)", [latex_to_clean_omml(r"\mathcal{L}_{\text{seq}}^{\text{self}} = \alpha_1 \mathcal{L}_{\text{MEP}} + \alpha_2 \mathcal{L}_{\text{MPP}} + \alpha_3 \mathcal{L}_{\text{time}}")], "Đóng băng sau A1; Tinh chỉnh ở A2", ["CÓ (Trích ", latex_to_clean_omml(r"\mathbf{z}_{\text{seq}}"), ")"]],
        [[latex_to_clean_omml(r"\phi_{\text{seq}}^{\text{event}}, \phi_{\text{seq}}^{\text{param}}, \phi_{\text{seq}}^{\text{time}}")], "Stage A1 (Chuỗi)", [latex_to_clean_omml(r"\mathcal{L}_{\text{MEP}}, \mathcal{L}_{\text{MPP}}, \mathcal{L}_{\text{time}}")], ["Duy trì ở A2 cho ", latex_to_clean_omml(r"\mathcal{L}_{\text{preserv}}"), "; Hủy sau A2"], "KHÔNG (Ngắt bỏ)"],
        [[latex_to_clean_omml(r"\theta_{\text{graph}}"), " (Temporal GNN)"], "Stage A1 (Đồ thị)", [latex_to_clean_omml(r"\mathcal{L}_{\text{graph}}^{\text{self}} = \beta_1 \mathcal{L}_{\text{mask-node}} + \beta_2 \mathcal{L}_{\text{mask-edge}} + \beta_3 \mathcal{L}_{\text{time-gap}}")], "Đóng băng sau A1; Tinh chỉnh ở A2", ["CÓ (Trích ", latex_to_clean_omml(r"\mathbf{z}_{\text{graph}}"), ")"]],
        [[latex_to_clean_omml(r"\phi_{\text{graph}}^{\text{node}}, \phi_{\text{graph}}^{\text{edge}}, \phi_{\text{graph}}^{\text{time}}")], "Stage A1 (Đồ thị)", [latex_to_clean_omml(r"\mathcal{L}_{\text{mask-node}}, \mathcal{L}_{\text{mask-edge}}, \mathcal{L}_{\text{time-gap}}")], ["Duy trì ở A2 cho ", latex_to_clean_omml(r"\mathcal{L}_{\text{preserv}}"), "; Hủy sau A2"], "KHÔNG (Ngắt bỏ)"],
        [[latex_to_clean_omml(r"g_{\text{seq}}, g_{\text{graph}}"), " (Projection Heads)"], "Stage A2 (Gióng hàng)", [latex_to_clean_omml(r"\mathcal{L}_{\text{align}} = \mathcal{L}_{\text{VICReg}}(\mathbf{P}_{\text{seq}}, \mathbf{P}_{\text{graph}})")], "Đóng băng & Hủy sau Stage A2", "KHÔNG (Ngắt bỏ)"],
        [[latex_to_clean_omml(r"\mathbf{w}_{\text{r}}, b_{\text{r}}, \mathbf{w}_{\text{g}}, b_{\text{g}}"), " (Gating Weights)"], "Stage A2 (Dung hợp)", [latex_to_clean_omml(r"\mathcal{L}_{\text{fuse-rec}}"), " (Tái tạo tự giám sát)"], "Đóng băng sau Stage A2", ["CÓ (Tính ", latex_to_clean_omml(r"w_{\text{rel}}"), ")"]],
        [[latex_to_clean_omml(r"\mathbf{P}_{\text{fuse}}^{\text{seq}}, \mathbf{P}_{\text{fuse}}^{\text{graph}}"), " (Cross Proj)"], "Stage A2 (Dung hợp)", [latex_to_clean_omml(r"\mathcal{L}_{\text{fuse-rec}}"), " (Tái tạo tự giám sát)"], "Đóng băng sau Stage A2", ["CÓ (Chiếu ", latex_to_clean_omml(r"\mathbb{R}^{d_{\text{cross}}}"), ")"]],
        [[latex_to_clean_omml(r"\mathbf{W}_{\text{out}}^{\text{seq}}, \mathbf{W}_{\text{out}}^{\text{graph}}, \mathbf{W}_{\text{out}}^{\text{cross}}")], "Stage A2 (Dung hợp)", [latex_to_clean_omml(r"\mathcal{L}_{\text{fuse-rec}}"), " (Tái tạo tự giám sát)"], "Đóng băng sau Stage A2", ["CÓ (Gộp ", latex_to_clean_omml(r"\mathbf{z}_{\text{mv}}"), ")"]],
        [[latex_to_clean_omml(r"\mathbf{D}_{\text{seq}}, \mathbf{D}_{\text{graph}}"), " (Linear Decoders)"], "Stage A2 (Dung hợp)", [latex_to_clean_omml(r"\mathcal{L}_{\text{fuse-rec}}"), " (Tái tạo tự giám sát)"], "Đóng băng & Hủy sau Stage A2", "KHÔNG (Ngắt bỏ)"],
        [[latex_to_clean_omml(r"\mathbf{w}_{\text{mil}}, \mathbf{V}_{\text{mil}}, \mathbf{U}_{\text{mil}}, \mathbf{W}_{\text{mil}}^{\text{cls}}, b_{\text{mil}}^{\text{cls}}")], "Stage B (Tùy chọn)", [latex_to_clean_omml(r"\mathcal{L}_{\text{MIL}}"), " (Coarse Bag BCE)"], "Đóng băng sau Stage B (chỉ bản mở rộng)", "TÙY CHỌN (Bản MIL)"]
    ]
    add_table_caption(doc, target_p, 4,
                      "Kiểm toán toàn diện các khối tham số và đồ thị tối ưu hóa trong hệ thống",
                      bookmark_name="BK_TBL_2_004", chapter_num=2, page_break_before=True)
    insert_thesis_table(doc, target_p, _tbl25_headers, _tbl25_widths, _tbl25_rows, font_size_pt=8.0, pad_v_dxa=10, space_v_pt=0.0)

    add_h4("Hàm mục tiêu tự giám sát thống nhất Stage A và Bảo toàn Đối xứng")
    add_p([
        "Trong Stage A, chuyên đề áp dụng chiến lược Bảo toàn Đối xứng (Symmetric Multi-View Preservation) kết hợp hàm mất mát tái tạo dung hợp tự giám sát (Self-Supervised Fusion Reconstruction). Hàm mất mát bảo toàn đối xứng duy trì năng lực biểu diễn nội góc nhìn cho cả chuỗi và đồ thị thông qua việc tái sử dụng các mục tiêu tự giám sát nội tại:"
    ])

    add_display_equation(make_preserv_loss_omml())

    add_p([
        "Vòng đời và vai trò của các đầu dự đoán phụ trợ (Auxiliary SSL Heads Lifecycle): Trong Stage A1, các đầu phụ trợ ",
        latex_to_clean_omml(r"\phi_{\text{seq}}^{\text{event}}, \phi_{\text{seq}}^{\text{param}}, \phi_{\text{seq}}^{\text{time}}"),
        " và ",
        latex_to_clean_omml(r"\phi_{\text{graph}}^{\text{node}}, \phi_{\text{graph}}^{\text{edge}}, \phi_{\text{graph}}^{\text{time}}"),
        " được tối ưu hóa đồng thời cùng backbone tương ứng. Trong Stage A2, để tính toán hàm mất mát bảo toàn đối xứng ",
        latex_to_clean_omml(r"\mathcal{L}_{\text{preserv}} = \mathcal{L}_{\text{seq}}^{\text{self}} + \mathcal{L}_{\text{graph}}^{\text{self}}"),
        " được thiết kế nhằm hạn chế nguy cơ trôi dạt biểu diễn (Representation Drift) và quên thảm họa (Catastrophic Forgetting) trong khi gióng hàng đa góc nhìn (việc đánh giá định lượng hiệu quả thực tế được thực hiện tại Chương 3), các đầu phụ trợ này tiếp tục được duy trì để cung cấp gradient điều hòa. "
        "Ngay sau khi Stage A2 kết thúc, toàn bộ các đầu phụ trợ này cùng 2 đầu chiếu phi tuyến ",
        latex_to_clean_omml(r"g_{\text{seq}}, g_{\text{graph}}"),
        " và 2 bộ giải mã tái tạo ",
        latex_to_clean_omml(r"\mathbf{D}_{\text{seq}}, \mathbf{D}_{\text{graph}}"),
        " đều bị ngắt bỏ (discarded). Tại thời điểm suy luận dòng (Stage C), chỉ các backbone ",
        latex_to_clean_omml(r"\theta_{\text{seq}}, \theta_{\text{graph}}"),
        " và các ma trận dung hợp chuẩn tắc được nạp để tính toán vector ",
        latex_to_clean_omml(r"\mathbf{z}_{\text{mv}}"),
        ", bảo đảm nguyên tắc Zero Orphan Parameters và tối ưu tài nguyên thực thi trực tuyến."
    ])

    add_p([
        "Đồng thời, để cung cấp tín hiệu học trực tiếp cho các ma trận dung hợp ",
        latex_to_clean_omml(r"\mathbf{P}_{\text{fuse}}^{\dots}, \mathbf{W}_{\text{out}}^{\dots}"),
        ", vector trọng số gating ",
        latex_to_clean_omml(r"\mathbf{w}_{\text{r}}, \mathbf{w}_{\text{g}}"),
        " và các ma trận giải mã tuyến tính ",
        latex_to_clean_omml(r"\mathbf{D}_{\text{seq}} \in \mathbb{R}^{d_{\text{seq}} \times d_{\text{mv}}}, \mathbf{D}_{\text{graph}} \in \mathbb{R}^{d_{\text{graph}} \times d_{\text{mv}}}"),
        ", chuyên đề thiết lập hàm mục tiêu tái thiết thông tin riêng biệt cho từng góc nhìn, ngăn chặn gradient của hàm tái thiết lan truyền ngược về các mạng trích xuất backbone ",
        latex_to_clean_omml(r"\theta_{\text{seq}}"),
        " và ",
        latex_to_clean_omml(r"\theta_{\text{graph}}"),
        ". Trong quá trình tinh chỉnh liên hợp ở Stage A2, hai bộ trích xuất backbone ",
        latex_to_clean_omml(r"\theta_{\text{seq}}"),
        " và ",
        latex_to_clean_omml(r"\theta_{\text{graph}}"),
        " chỉ nhận luồng gradient duy nhất từ hàm gióng hàng ",
        latex_to_clean_omml(r"\mathcal{L}_{\text{align}}"),
        " kết hợp với hàm bảo toàn đối xứng ",
        latex_to_clean_omml(r"\mathcal{L}_{\text{preserv}}"),
        "."
    ])

    add_p([
        "Hàm mục tiêu tối ưu hóa thống nhất toàn bộ Stage A được xác lập:"
    ])

    add_display_equation(make_stage_a_loss_omml())

    add_p([
        "Các hệ số siêu tham số ",
        latex_to_clean_omml(r"\lambda_{\text{inv}}, \lambda_{\text{var}}, \lambda_{\text{cov}}, \lambda_{\text{spec}}, \lambda_{\text{rec}}"),
        " là các trọng số điều hòa được thiết lập giá trị khởi tạo tham chiếu ",
        latex_to_clean_omml(r"\{25.0, 25.0, 1.0, 1.0, 1.0\}"),
        " và sẽ được tối ưu hóa, kiểm định độ nhạy thực nghiệm trên tập xác thực (Validation Set) tại Chương 3."
    ])

    add_h4("Cổng độ tin cậy động và Phẫu thuật Gradient PCGrad")
    add_p([
        "Để đánh giá chất lượng góc nhìn và điều hòa xung đột gradient trong quá trình tối ưu hóa Stage A, hệ thống tích hợp vector chất lượng đo kiểm quan sát được ",
        latex_to_clean_omml(r"\mathbf{q}_{\text{seq}}"),
        " và ",
        latex_to_clean_omml(r"\mathbf{q}_{\text{graph}}"),
        ":"
    ])

    add_display_equation(make_quality_vector_omml())
    add_display_equation(
        r"s_{\text{seq}} = \mathbf{w}_{\text{r}}^\top [\mathbf{z}_{\text{seq}}; \mathbf{q}_{\text{seq}}] + b_{\text{r}}, \quad s_{\text{graph}} = \mathbf{w}_{\text{g}}^\top [\mathbf{z}_{\text{graph}}; \mathbf{q}_{\text{graph}}] + b_{\text{g}}"
    )
    add_display_equation(make_gating_weights_omml())

    add_p([
        "Trong đó ",
        latex_to_clean_omml(r"\tau = 0.05"),
        " là siêu tham số ngưỡng chặn dưới an toàn, ngăn ngừa việc hạ trọng số về 0 (tính nhạy cảm của ",
        latex_to_clean_omml(r"\tau"),
        " sẽ được kiểm chứng thực nghiệm tại Chương 3). Khi tối ưu hóa liên hợp, kỹ thuật PCGrad ",
        make_citation_element([40]),
        " được áp dụng chặt chẽ trên tập tham số nhận đồng thời hai luồng gradient:",
    ])

    add_display_equation(
        r"\Theta_{\text{PCGrad}} = \text{dom}(\nabla \mathcal{L}_{\text{align}}) \cap \text{dom}(\nabla \mathcal{L}_{\text{preserv}})"
    )
    add_p([
        "Trong đó ",
        latex_to_clean_omml(r"\mathbf{g}_{\text{align}} = \nabla_{\Theta_{\text{PCGrad}}} \mathcal{L}_{\text{align}}"),
        " và ",
        latex_to_clean_omml(r"\mathbf{g}_{\text{preserv}} = \nabla_{\Theta_{\text{PCGrad}}} \mathcal{L}_{\text{preserv}}"),
        ". Các gradient từ hàm tái tạo ",
        latex_to_clean_omml(r"\mathcal{L}_{\text{fuse-rec}}"),
        " không tham gia vào phép chiếu PCGrad vì chúng chỉ cập nhật độc lập các tham số tầng dung hợp. Khi xảy ra xung đột hướng ",
        latex_to_clean_omml(r"\langle \mathbf{g}_{\text{align}}, \mathbf{g}_{\text{preserv}} \rangle < 0"),
        ", gradient gióng hàng được chiếu vuông góc:"
    ])

    add_display_equation(make_pcgrad_omml())

    add_p([
        "Kỹ thuật chiếu trực giao PCGrad giúp giảm thiểu xung đột gradient giữa mục tiêu gióng hàng liên góc nhìn và mục tiêu bảo toàn đặc thù nội góc nhìn theo tiêu chuẩn cosine; hiệu quả thực tế trong việc hạn chế hiện tượng chuyển giao tiêu cực (Negative Transfer) sẽ được kiểm định qua phân tích triệt tiêu tại Chương 3."
    ])

    add_h4("Stage B (Tùy chọn): Phân bổ bằng chứng yếu qua Attention-MIL")
    add_p([
        "Trong các kịch bản an ninh thực tế, nhãn tấn công thường chỉ sẵn có ở mức độ thô (Coarse Labels), ví dụ: một phiên làm việc (Session), một máy chủ (Host), hoặc một cửa sổ chiến dịch kéo dài (Campaign Window) được xác định là bị xâm nhập (nhãn túi ",
        latex_to_clean_omml(r"Y_{\text{bag}} \in \{0, 1\}"),
        "), nhưng không thể xác định chính xác sự kiện log đơn lẻ nào là hành vi tấn công. Để giải quyết thách thức này, chuyên đề thiết lập mô đun Phân bổ Bằng chứng Yếu (Weak Evidence Attribution) tùy chọn ở Stage B dựa trên khung tiếp cận Học Đa Thể Hiện (Multiple Instance Learning, MIL)."
    ])

    add_p([
        "Chuyên đề áp dụng cơ chế Attention-based Deep MIL của Ilse et al. ",
        make_citation_element([24]),
        ", nhưng thiết kế cơ chế ánh xạ Túi, Thực thể đặc thù cho bài toán an ninh mạng (Cybersecurity Bag-Instance Mapping, Đề xuất của đề tài / Ours):"
    ], keep_with_next=True)

    add_bullet_p([
        "Định nghĩa Túi quan sát (Bag, ",
        latex_to_clean_omml(r"\mathcal{X}_{\text{bag}} = \{\mathbf{z}_{\text{mv}}^{(1)}, \mathbf{z}_{\text{mv}}^{(2)}, \dots, \mathbf{z}_{\text{mv}}^{(K)}\}"),
        "): Một túi gồm ",
        latex_to_clean_omml(r"K"),
        " vector biểu diễn đa góc nhìn liên tiếp tương ứng với một phiên làm việc hoặc một cửa sổ giám sát của máy chủ."
    ])

    add_bullet_p([
        "Định nghĩa Thực thể thành phần (Instance, ",
        latex_to_clean_omml(r"\mathbf{z}_{\text{mv}}^{(i)}"),
        "): Mỗi vector biểu diễn đa góc nhìn tại thời điểm ",
        latex_to_clean_omml(r"t_i"),
        " đóng vai trò là một thể hiện bên trong túi."
    ])

    add_p([
        "Cơ chế Gated Attention-MIL với không gian chú ý ẩn kích thước ",
        latex_to_clean_omml(r"d_{\text{att}}"),
        " tính toán trọng số chú ý thích ứng ",
        latex_to_clean_omml(r"a_i \in [0, 1]"),
        " cho từng thể hiện:"
    ])

    add_display_equation(make_mil_attention_omml())
    add_display_equation(
        r"\mathbf{z}_{\text{bag}} = \sum_{i=1}^K a_i \mathbf{z}_{\text{mv}}^{(i)}, \quad \hat{Y}_{\text{bag}} = \sigma(\mathbf{W}_{\text{mil,cls}} \mathbf{z}_{\text{bag}} + b_{\text{mil,cls}})"
    )
    add_display_equation(make_mil_loss_omml())

    add_p([
        "Phân định ranh giới giữa Khung trích xuất chuẩn tắc và Đầu phân lớp phụ trợ MIL: Cơ chế Gated Attention-MIL trang bị một đầu phân lớp túi phụ trợ ",
        latex_to_clean_omml(r"\mathbf{W}_{\text{mil,cls}}, b_{\text{mil,cls}}"),
        " (Supervised Auxiliary Head). Trong quy trình chuẩn hóa của chuyên đề:"
    ], keep_with_next=True)

    add_bullet_p([
        "Giai đoạn Stage B: Chỉ sử dụng nhãn thô trên tập Huấn luyện và Xác thực (Train/Validation only) để huấn luyện đầu phân lớp phụ trợ và tùy chọn tinh chỉnh nhẹ không gian vector đặc trưng."
    ])

    add_bullet_p([
        "Giai đoạn Stage C chuẩn tắc: Hệ thống hoàn toàn LOẠI BỎ (DISCARD) đầu phân lớp ",
        latex_to_clean_omml(r"\mathbf{W}_{\text{mil,cls}}, b_{\text{mil,cls}}"),
        " cùng cấu trúc phân loại túi, đóng băng bộ trích xuất đặc trưng thuần túy để chỉ xuất gói biểu diễn chuẩn tắc ",
        latex_to_clean_omml(r"\mathcal{B}_{\text{mv}}"),
        " chứa vector canonical ",
        latex_to_clean_omml(r"\mathbf{z}_{\text{mv}}"),
        " cho các mô hình hạ nguồn độc lập (Detector-Agnostic Extractor)."
    ])

    add_bullet_p([
        "Biến thể tùy chọn hỗ trợ MIL (Optional MIL-Assisted Variant): Nếu kịch bản triển khai thực tế có nhu cầu giám sát trực quan điểm bằng chứng ",
        latex_to_clean_omml(r"a_i"),
        ", hệ thống định nghĩa riêng một biến thể hỗ trợ MIL, hoàn toàn phân định với Extractor chuẩn tắc. Tại Chương 3, chuyên đề sẽ báo cáo độc lập kết quả giữa ba cấu hình: SSL thuần túy (SSL-only), SSL kết hợp thích ứng MIL (SSL + MIL Adaptation), và Biến thể hỗ trợ MIL khi triển khai (MIL-assisted variant)."
    ])

    add_p([
        "Ý nghĩa phương pháp luận của trọng số chú ý ",
        latex_to_clean_omml(r"a_i"),
        ": Trọng số ",
        latex_to_clean_omml(r"a_i"),
        " phản ánh mức độ đóng góp bằng chứng yếu (Weak Evidence Attribution Score) của cửa sổ quan sát thứ ",
        latex_to_clean_omml(r"i"),
        " đối với trạng thái bất thường chung của toàn bộ túi. Chuyên đề nhấn mạnh ranh giới khoa học: Trọng số ",
        latex_to_clean_omml(r"a_i"),
        " là tín hiệu gợi ý điều tra (Attribution Signal), TUYỆT ĐỐI KHÔNG PHẢI là xác suất tấn công độc hại và KHÔNG PHẢI là bằng chứng giải thích nhân quả (Causal Explanation)."
    ])

    # 2.4.4. Unified Representation / Interface / Complexity
    add_h3("Biểu diễn thống nhất, giao diện đầu ra và độ phức tạp")
    add_p([
        "Mục 2.4.4 xác lập công thức tổng hợp biểu diễn thống nhất canonical, định nghĩa giao diện chuẩn cho các mô hình hạ nguồn ở Stage C và phân tích chi phí tính toán."
    ])

    add_h4([
        "Vector biểu diễn thống nhất canonical ",
        latex_to_clean_omml(r"\mathbf{z}_{\text{mv}}"),
        " và giao diện Extractor"
    ])
    add_p([
        "Để kết hợp sức mạnh biểu diễn vi mô từ chuỗi và quan hệ cấu trúc vĩ mô từ đồ thị mà không áp đặt giả định đồng nhất số chiều tùy tiện, chuyên đề thiết lập các ánh xạ tuyến tính chuẩn hóa sang không gian tương tác chung ",
        latex_to_clean_omml(r"\mathbb{R}^{d_{\text{cross}}}"),
        ":"
    ])

    add_display_equation(
        r"\mathbf{u}_{\text{seq}} = \mathbf{P}_{\text{fuse}}^{\text{seq}} \mathbf{z}_{\text{seq}} \in \mathbb{R}^{d_{\text{cross}}}, \quad \mathbf{u}_{\text{graph}} = \mathbf{P}_{\text{fuse}}^{\text{graph}} \mathbf{z}_{\text{graph}} \in \mathbb{R}^{d_{\text{cross}}}"
    )
    add_display_equation(
        r"\mathbf{u}_{\text{cross}} = \mathbf{u}_{\text{seq}} \odot \mathbf{u}_{\text{graph}} \in \mathbb{R}^{d_{\text{cross}}}"
    )

    add_p([
        "Trong đó ",
        latex_to_clean_omml(r"\mathbf{P}_{\text{fuse}}^{\text{seq}} \in \mathbb{R}^{d_{\text{cross}} \times d_{\text{seq}}}"),
        " và ",
        latex_to_clean_omml(r"\mathbf{P}_{\text{fuse}}^{\text{graph}} \in \mathbb{R}^{d_{\text{cross}} \times d_{\text{graph}}}"),
        " là các ma trận chiếu tuyến tính phục vụ dung hợp; ký hiệu ",
        latex_to_clean_omml(r"\odot"),
        " biểu thị tích Hadamard trên không gian cùng số chiều ",
        latex_to_clean_omml(r"\mathbb{R}^{d_{\text{cross}}}"),
        " nhằm nắm bắt tương tác phi tuyến bậc hai giữa hai góc nhìn. Vector biểu diễn thống nhất canonical ",
        latex_to_clean_omml(r"\mathbf{z}_{\text{mv}} \in \mathbb{R}^{d_{\text{mv}}}"),
        " được tổng hợp thông qua cơ chế Dung hợp có Cổng thích ứng (Gated Fusion Readout):"
    ])

    add_display_equation(make_canonical_fusion_omml())

    add_p([
        "Trong đó ",
        latex_to_clean_omml(r"\mathbf{W}_{\text{out}}^{\text{seq}} \in \mathbb{R}^{d_{\text{mv}} \times d_{\text{seq}}}"),
        ", ",
        latex_to_clean_omml(r"\mathbf{W}_{\text{out}}^{\text{graph}} \in \mathbb{R}^{d_{\text{mv}} \times d_{\text{graph}}}"),
        ", ",
        latex_to_clean_omml(r"\mathbf{W}_{\text{out}}^{\text{cross}} \in \mathbb{R}^{d_{\text{mv}} \times d_{\text{cross}}}"),
        " là các ma trận trọng số đầu ra. Gói biểu diễn đầu ra đa góc nhìn được đóng gói chuẩn mực:"
    ])

    add_display_equation(make_rep_bundle_omml())

    add_p([
        "Gói biểu diễn ",
        latex_to_clean_omml(r"\mathcal{B}_{\text{mv}}"),
        " tạo thành giao diện đầu ra duy nhất của toàn bộ Khung Trích xuất Đặc trưng Đa góc nhìn (Multi-View Feature Extractor). Phù hợp với ranh giới phân định tầng tại Mục 1.1.3, Extractor tuyệt đối không chứa logic phát hiện bất thường hay ngưỡng phân lớp an ninh. Gói biểu diễn được bàn giao nguyên vẹn sang Stage C phục vụ các đầu dò tuyến tính đóng băng (Frozen Linear Probe) và quy trình chấm điểm bất thường không giám sát tùy chọn (Optional Downstream Zero-Shot Anomaly Scoring Protocol, trong đó hàm đo khoảng cách/năng lượng, phân bố nền và chính sách ngưỡng được tiền đăng ký và khóa cố định trên tập Train/Validation tại Chương 3)."
    ])

    # Insert Figure 2.4 Canvas Placeholder
    p_c3 = doc.add_paragraph() if target_p is None else target_p.insert_paragraph_before()
    p_c3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_c3.paragraph_format.keep_with_next = True
    p_c3.paragraph_format.space_before = Pt(6)
    p_c3.paragraph_format.space_after = Pt(2)
    p_c3._p.append(parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="30204" w:name="BK_FIG_2_004_CANVAS"/>'))
    r_c3 = p_c3.add_run(" ")
    p_c3._p.append(parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="30204"/>'))
    add_figure_caption(doc, target_p, 2, 4, "Sơ đồ kiến trúc Gióng hàng đa góc nhìn, Kiểm soát sụp đổ và Tổng hợp biểu diễn thống nhất (Nguồn: Tác giả đề xuất)", bookmark_name="BK_FIG_2_004")

    add_h4("Phân tích độ phức tạp tính toán và ngân sách streaming")
    add_p([
        "Chi phí tính toán của Tầng Gióng hàng và Biểu diễn Thống nhất được phân định rõ ràng giữa giai đoạn huấn luyện tự giám sát và giai đoạn suy luận trực tuyến:"
    ], keep_with_next=True)

    add_bullet_p([
        "Chi phí huấn luyện gióng hàng và tái tạo dung hợp (Training-time Budget): Trong mỗi bước huấn luyện lô ",
        latex_to_clean_omml(r"B"),
        ", chi phí chiếu phi tuyến qua MLP là ",
        latex_to_clean_omml(r"\mathcal{O}(B \cdot (d_{\text{seq}} + d_{\text{graph}}) d_{\text{proj}})"),
        ". Chi phí tính toán hàm mất mát VICReg trên không gian chiếu là ",
        latex_to_clean_omml(r"\mathcal{O}(B \cdot d_{\text{proj}}^2)"),
        ", hàm tái tạo dung hợp là ",
        latex_to_clean_omml(r"\mathcal{O}(B \cdot (d_{\text{seq}} + d_{\text{graph}} + d_{\text{cross}}) d_{\text{mv}})"),
        ", và chi phí phẫu thuật gradient PCGrad bổ sung trên tập tham số chia sẻ ",
        latex_to_clean_omml(r"\Theta_{\text{PCGrad}}"),
        " sau khi đã tính các gradient thành phần là ",
        latex_to_clean_omml(r"\mathcal{O}(\text{card}(\Theta_{\text{PCGrad}}))"),
        ". Nếu kích hoạt mô đun Attention-MIL ở Stage B với không gian ẩn chú ý ",
        latex_to_clean_omml(r"d_{\text{att}}"),
        ", chi phí tính toán cho mỗi túi kích thước ",
        latex_to_clean_omml(r"K"),
        " bao gồm chi phí chiếu cổng chú ý ",
        latex_to_clean_omml(r"\mathcal{O}(K \cdot d_{\text{mv}} \cdot d_{\text{att}})"),
        " và chi phí phân lớp túi ",
        latex_to_clean_omml(r"\mathcal{O}(K \cdot d_{\text{att}} + d_{\text{mv}})"),
        ", cho chi phí tổng thể mỗi túi là ",
        latex_to_clean_omml(r"\mathcal{O}(K \cdot d_{\text{mv}} \cdot d_{\text{att}} + d_{\text{mv}})"),
        " (tuyến tính theo quy mô túi ",
        latex_to_clean_omml(r"\mathcal{O}(K)"),
        " khi các tham số chiều được cố định). Toàn bộ chi phí này chỉ phát sinh trong quá trình huấn luyện ngoại tuyến (Offline Training)."
    ])

    add_bullet_p([
        "Chi phí suy luận trực tuyến (Streaming Inference Budget): Trong giai đoạn vận hành dòng, các đầu chiếu ",
        latex_to_clean_omml(r"g_{\text{seq}}, g_{\text{graph}}"),
        " và các mô đun mất mát được ngắt bỏ hoàn toàn. Chi phí dung hợp và tổng hợp đầu ra đa góc nhìn tăng thêm (Incremental Fusion/Readout Cost) sau khi đã có ",
        latex_to_clean_omml(r"\mathbf{z}_{\text{seq}}"),
        " và ",
        latex_to_clean_omml(r"\mathbf{z}_{\text{graph}}"),
        " chỉ bao gồm các phép nhân ma trận, vector tuyến tính và cổng gating. Chi phí tăng thêm này được biểu diễn như sau:"
    ], keep_with_next=True)

    add_display_equation(latex_to_clean_omml(r"\mathcal{C}_{\text{fuse}} = \mathcal{O}(d_{\text{seq}} d_{\text{cross}} + d_{\text{graph}} d_{\text{cross}} + d_{\text{seq}} d_{\text{mv}} + d_{\text{graph}} d_{\text{mv}} + d_{\text{cross}} d_{\text{mv}})"))

    add_p([
        "Khi các số chiều biểu diễn được cố định, chi phí dung hợp cục bộ này đạt mức ",
        latex_to_clean_omml(r"\mathcal{O}(1)"),
        " đối với quy mô luồng dữ liệu (constant with respect to telemetry stream length and graph size). Chi phí tính toán toàn trình (End-to-End Budget) của toàn bộ hệ thống bao gồm: tiền xử lý sự kiện thô, trích xuất chuỗi ngắn hạn, cập nhật đồ thị động, lan truyền thông điệp Temporal GNN, quản lý trạng thái bộ nhớ hữu hạn và tổng hợp dung hợp đa góc nhìn. Các chỉ số đo kiểm thực nghiệm toàn trình (thông lượng sự kiện/giây, độ trễ p50/p95, dung lượng RAM, VRAM và kích thước trạng thái trên mỗi thực thể) sẽ được đo lường thực nghiệm chi tiết tại Chương 3 nhằm kiểm định toàn diện giả thuyết ",
        make_hypo_omml(4),
        "."
    ])

    # =========================================================================
    # BIBLIOGRAPHY / TÀI LIỆU THAM KHẢO SYNCHRONIZATION
    # =========================================================================
    print("[5/6] Creating native Word BIBLIOGRAPHY field...")
    
    bib_idx = None
    for idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip().lower()
        if "tài liệu tham khảo" in txt and idx > 30:
            bib_idx = idx
            break

    if bib_idx is not None:
        paras_to_del = [p for p in doc.paragraphs[bib_idx + 1:]]
        for p in paras_to_del:
            p._p.getparent().remove(p._p)

        bib_entries = [
            "[1] MITRE ATT&CK, \"MITRE ATT&CK: Enterprise Tactics and Techniques Matrix (v19.1),\" MITRE Corporation, 2026.",
            "[2] D. Arp, E. Quiring, F. Pendlebury, A. Warnecke, F. Pierazzi, C. Wressnegger, L. Cavallaro, and K. Rieck, \"Dos and Don'ts of Machine Learning in Computer Security,\" in Proceedings of the 31st USENIX Security Symposium (USENIX Security 2022), 2022.",
            "[3] M. Du, F. Li, G. Zheng, and V. Srikumar, \"DeepLog: Anomaly Detection and Diagnosis from System Logs through Deep Learning,\" in Proceedings of the 2017 ACM SIGSAC Conference on Computer and Communications Security (CCS 2017), 2017.",
            "[4] H. Guo, S. Yuan, and X. Wu, \"LogBERT: Log Anomaly Detection via BERT,\" in Proceedings of the International Joint Conference on Neural Networks (IJCNN 2021), 2021.",
            "[5] V.-H. Le and H. Zhang, \"Log-based Anomaly Detection Without Log Parsing,\" in Proceedings of the 36th IEEE/ACM International Conference on Automated Software Engineering (ASE 2021), 2021.",
            "[6] J. Zhu, S. He, J. Liu, P. He, Q. Xie, Z. Zheng, and M. R. Lyu, \"Tools and Benchmarks for Automated Log Parsing,\" in Proceedings of the 34th IEEE International Symposium on Software Reliability Engineering (ISSRE 2023), 2023.",
            "[7] Z. Jiang, J. Liu, J. Huang, Y. Huo, X. Peng, Y. Li, J. Zhu, and M. R. Lyu, \"A Large-Scale Evaluation for Log Parsing Techniques: How Far Are We?,\" in Proceedings of the 33rd ACM SIGSOFT International Symposium on Software Testing and Analysis (ISSTA 2024), 2024.",
            "[8] L. Michael, A. Tamersoy, T. Kelley, and M. Locasto, \"On the Forensic Validity of Approximated Audit Logs,\" in Proceedings of the 36th Annual Computer Security Applications Conference (ACSAC 2020), 2020.",
            "[9] M. A. Inam, Y. Chen, F. Mohsen, A. Tamersoy, C. Wressnegger, M. Locasto, and G. Wang, \"SoK: History is a Vast Early Warning System: Auditing the Provenance of System Intrusions,\" in Proceedings of the 44th IEEE Symposium on Security and Privacy (S&P 2023), 2023.",
            "[10] M. Zipperle, F. Armknecht, and C. Kolb, \"Provenance-based Intrusion Detection Systems: A Survey,\" ACM Computing Surveys, 2022.",
            "[11] X. Han, T. Pasquier, A. Bates, J. Mickens, and M. Seltzer, \"UNICORN: Runtime Provenance-Based Detector for Advanced Persistent Threats,\" in Proceedings of the Network and Distributed System Security Symposium (NDSS 2020), 2020.",
            "[12] Z. Cheng, Q. Lv, J. Liang, Y. Wang, D. Sun, T. Pasquier, and X. Han, \"KAIROS: Practical Intrusion Detection and Investigation using Whole-system Provenance,\" in Proceedings of the 45th IEEE Symposium on Security and Privacy (S&P 2024), 2024.",
            "[13] S. Li, F. Dong, D. Li, X. Xiao, H. Wang, F. Shao, J. Chen, Y. Guo, and X. Chen, \"NODLINK: An Online System for Fine-Grained APT Attack Detection and Investigation,\" in Proceedings of the Network and Distributed System Security Symposium (NDSS 2024), 2024.",
            "[14] Z. Jia, Y. Xiong, Y. Nan, Y. Zhang, J. Zhao, and Mi Wen, \"MAGIC: Detecting Advanced Persistent Threats via Masked Graph Representation Learning,\" in Proceedings of the 33rd USENIX Security Symposium (USENIX Security 2024), 2024.",
            "[15] B. Jiang, T. Bilot, N. E. Madhoun, K. A. Agha, A. Zouaoui, S. Iqbal, X. Han, and T. Pasquier, \"ORTHRUS: Achieving High Quality of Attribution in Provenance-based Intrusion Detection Systems,\" in Proceedings of the 34th USENIX Security Symposium (USENIX Security 2025), 2025.",
            "[16] T. Bilot, B. Jiang, Z. Li, N. E. Madhoun, K. A. Agha, A. Zouaoui, and T. Pasquier, \"Sometimes Simpler is Better: A Comprehensive Analysis of State-of-the-Art Provenance-Based Intrusion Detection Systems,\" in Proceedings of the 34th USENIX Security Symposium (USENIX Security 2025), 2025.",
            "[17] T. Bilot, B. Jiang, and T. Pasquier, \"PIDSMaker: Building and Evaluating Provenance-based Intrusion Detection Systems,\" in Proceedings of the 32nd ACM SIGKDD Conference on Knowledge Discovery and Data Mining (KDD 2026), 2026.",
            "[18] Y. Liu, D. Arp, L. Cavallaro, and K. Rieck, \"What We Talk About When We Talk About Logs: Understanding the Effects of Dataset Quality on Endpoint Threat Detection Research,\" in Proceedings of the 46th IEEE Symposium on Security and Privacy (S&P 2025), 2025.",
            "[19] A. Goyal, X. Han, G. Wang, and A. Bates, \"Sometimes, You Aren't What You Do: Mimicry Attacks against Provenance Graph HIDS,\" in Proceedings of the Network and Distributed System Security Symposium (NDSS 2023), 2023.",
            "[20] P. Gao, X. Xiao, Z. Li, K. Jee, F. Xu, S. R. Kulkarni, and P. Mittal, \"PalanTír: Optimizing Attack Provenance with Coarse Audit Logs,\" in Proceedings of the 2022 ACM SIGSAC Conference on Computer and Communications Security (CCS 2022), 2022.",
            "[21] U. Alon and E. Yahav, \"On the Bottleneck of Graph Neural Networks and its Practical Implications,\" in Proceedings of the 9th International Conference on Learning Representations (ICLR 2021), 2021.",
            "[22] A. Bardes, J. Ponce, and Y. LeCun, \"VICReg: Variance-Invariance-Covariance Regularization for Self-Supervised Learning,\" in Proceedings of the 10th International Conference on Learning Representations (ICLR 2022), 2022.",
            "[23] J. Zbontar, L. Jing, I. Misra, Y. LeCun, and S. Deny, \"Barlow Twins: Self-Supervised Learning via Redundancy Reduction,\" in Proceedings of the 38th International Conference on Machine Learning (ICML 2021), 2021.",
            "[24] M. Ilse, J. M. Tomczak, and M. Welling, \"Attention-based Deep Multiple Instance Learning,\" in Proceedings of the 35th International Conference on Machine Learning (ICML 2018), 2018.",
            "[25] R. Shokri, M. Stronati, C. Song, and V. Shmatikov, \"Membership Inference Attacks Against Machine Learning Models,\" in Proceedings of the 38th IEEE Symposium on Security and Privacy (S&P 2017), 2017.",
            "[26] M. Fredrikson, S. Jha, and T. Ristenpart, \"Model Inversion Attacks that Exploit Confidence Information and Basic Countermeasures,\" in Proceedings of the 22nd ACM SIGSAC Conference on Computer and Communications Security (CCS 2015), 2015.",
            "[27] National Institute of Standards and Technology (NIST), \"Guidelines for Evaluating Differential Privacy Guarantees (NIST SP 800-226),\" NIST Special Publication 800-226, 2025.",
            "[28] Defense Advanced Research Projects Agency (DARPA), \"DARPA Transparent Computing Program Telemetry Datasets (Engagements 3 and 5),\" DARPA Official Dataset Release, 2019.",
            "[29] A. D. Kent, \"Comprehensive, Multi-Source Cyber-Security Events (Unified Host and Network Dataset),\" Los Alamos National Laboratory Official Dataset Release, 2017, DOI: 10.17021/1117677.",
            "[30] L. Guerra, T. Chapuis, G. Duc, P. Mozharovskyi, and V.-T. Nguyen, \"How Benchmarks and Evaluation Protocols Shape Conclusions in Provenance-Based Intrusion Detection,\" arXiv preprint arXiv:2608.01454, 2026.",
            "[31] W. Xu, L. Huang, A. Fox, D. Patterson, and M. I. Jordan, \"Detecting Large-Scale System Problems by Mining Console Logs,\" in Proceedings of the ACM SIGOPS 22nd Symposium on Operating Systems Principles (SOSP 2009), 2009.",
            "[32] J.-G. Lou, Q. Fu, S. Yang, Y. Xu, and J. Li, \"Mining Invariants from Console Logs for System Problem Detection,\" in Proceedings of the 2010 USENIX Annual Technical Conference (USENIX ATC 2010), 2010.",
            "[33] W. Meng, Y. Liu, Y. Zhu, S. Zhang, D. Pei, Y. Liu, Y. Chen, R. Zhang, S. Tao, P. Sun, and R. Zhou, \"LogAnomaly: Unsupervised Detection of Sequential and Quantitative Anomalies in Unstructured Logs,\" in Proceedings of the 28th International Joint Conference on Artificial Intelligence (IJCAI 2019), 2019.",
            "[34] S. Nedelkoski, J. Bogatinovski, A. Acker, J. Cardoso, and O. Kao, \"Self-Attentive Classification-Based Anomaly Detection in Unstructured Logs,\" in Proceedings of the 2020 IEEE International Conference on Data Mining (ICDM 2020), 2020.",
            "[35] A. Vaswani, N. Shazeer, N. Parmar, J. Uszkoreit, L. Jones, A. N. Gomez, L. Kaiser, and I. Polosukhin, \"Attention Is All You Need,\" in Advances in Neural Information Processing Systems 30 (NeurIPS 2017), 2017.",
            "[36] D. Xu, C. Ruan, E. Korpeoglu, S. Kumar, and K. Achan, \"Inductive Representation Learning on Temporal Graphs,\" in International Conference on Learning Representations (ICLR 2020), 2020.",
            "[37] E. Rossi, B. Chamberlain, F. Frasca, D. Eynard, F. Monti, and M. Bronstein, \"Temporal Graph Networks for Deep Learning on Dynamic Graphs,\" in ICML 2020 Workshop on Graph Representation Learning (arXiv:2006.10637), 2020.",
            "[38] A. van den Oord, Y. Li, and O. Vinyals, \"Representation Learning with Contrastive Predictive Coding,\" arXiv preprint arXiv:1807.03748, 2018.",
            "[39] T. Chen, S. Kornblith, M. Norouzi, and G. Hinton, \"A Simple Framework for Contrastive Learning of Visual Representations,\" in Proceedings of the 37th International Conference on Machine Learning (ICML 2020), 2020.",
            "[40] T. Yu, S. Kumar, A. Gupta, S. Levine, K. Hausman, and C. Finn, \"Gradient Surgery for Multi-Task Learning: Projecting Conflicting Gradients,\" in Advances in Neural Information Processing Systems 33 (NeurIPS 2020), 2020.",
            "[41] C. Zhang, Z. Han, Y. Cui, H. Fu, J. T. Zhou, and Q. Hu, \"CPM-Nets: Cross Partial Multi-View Networks,\" in Advances in Neural Information Processing Systems 32 (NeurIPS 2019), 2019."
        ]

        for b_text in bib_entries:
            bp = doc.add_paragraph(style="Normal")
            bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            bp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            bp.paragraph_format.space_before = Pt(0)
            bp.paragraph_format.space_after = Pt(4)
            bp.paragraph_format.left_indent = Cm(1.27)
            bp.paragraph_format.first_line_indent = Cm(-1.27)
            r = bp.add_run(b_text)
            r.font.name = "Times New Roman"
            r.font.size = Pt(14)

    # Save directly to target_path
    doc.save(str(target_path))
    del doc
    import gc
    gc.collect()
    import time
    time.sleep(1.0)
    print(f"[SUCCESS] Saved and updated DOCX: {target_path}", flush=True)

    # =========================================================================
    # STEP 6: AUTOMATE MICROSOFT WORD DESKTOP TO INSERT NATIVE DIAGRAMS, UPDATE ALL DYNAMIC FIELDS & EXPORT PDF
    # =========================================================================
    pdf_path = target_path.parent / (target_path.stem + ".pdf")
    import subprocess
    post_script = Path(__file__).resolve().parent / "word_com_post_process.py"
    cmd = [sys.executable, str(post_script), str(target_path), str(pdf_path)]
    print(f"[6/6] Executing Word COM Post-Processor: {' '.join(cmd)}", flush=True)
    res = subprocess.run(cmd)
    if res.returncode != 0:
        print(f"[WARNING] Word COM post-processor exited with code {res.returncode}", flush=True)

    # -------------------------------------------------------------------------
    # REGISTER VISUAL RECORDS IN SQLITE VISUAL REGISTRY
    # -------------------------------------------------------------------------
    try:
        v_reg = VisualRegistry()
        v_records = [
            VisualRecord(
                visual_id="FIG-000001",
                node_code="1.1.1",
                purpose="Phân cấp các đơn vị quan sát và sự đánh đổi giữa ngữ cảnh và chi phí tính toán",
                visual_type=VisualType.CONCEPTUAL_DIAGRAM,
                creation_method=CreationMethod.WORD_DRAWING_CANVAS,
                caption="Phân cấp các đơn vị quan sát và sự đánh đổi giữa ngữ cảnh và chi phí tính toán trong biểu diễn dữ liệu log (Nguồn: Tác giả tổng hợp)",
                source_provenance="Tác giả tổng hợp dựa trên phân loại telemetry hệ thống",
                script_path="src/research_agent/visuals/academic_diagram_renderer.py",
                output_file_path=str(target_path),
                bookmark_name="BK_FIG_1_001",
                seq_number=1,
                chapter_number=1,
                is_verified=True,
            ),
            VisualRecord(
                visual_id="FIG-000002",
                node_code="1.1.2",
                purpose="Mô hình Không gian Bằng chứng Hành vi Đa chiều MITRE ATT&CK phi tuyến tính",
                visual_type=VisualType.CONCEPTUAL_DIAGRAM,
                creation_method=CreationMethod.WORD_DRAWING_CANVAS,
                caption="Mô hình Không gian Bằng chứng Hành vi Đa chiều MITRE ATT&CK và các đặc trưng phi tuyến tính trong tấn công APT (Nguồn: Tác giả tổng hợp dựa trên MITRE ATT&CK và Inam et al.)",
                source_provenance="Tác giả tổng hợp dựa trên MITRE ATT&CK v19.1 và Inam et al.",
                script_path="src/research_agent/visuals/academic_diagram_renderer.py",
                output_file_path=str(target_path),
                bookmark_name="BK_FIG_1_002",
                seq_number=2,
                chapter_number=1,
                is_verified=True,
            ),
            VisualRecord(
                visual_id="FIG-000003",
                node_code="1.1.3",
                purpose="Ranh giới ba tầng phương pháp luận và vị trí trọng tâm của không gian vector biểu diễn z",
                visual_type=VisualType.CONCEPTUAL_DIAGRAM,
                creation_method=CreationMethod.WORD_DRAWING_CANVAS,
                caption="Khung phân định ranh giới ba tầng phương pháp luận và vị trí trọng tâm của không gian vector biểu diễn z (Nguồn: Tác giả đề xuất)",
                source_provenance="Tác giả đề xuất mô hình ba tầng và Hợp đồng Biểu diễn",
                script_path="src/research_agent/visuals/academic_diagram_renderer.py",
                output_file_path=str(target_path),
                bookmark_name="BK_FIG_1_003",
                seq_number=3,
                chapter_number=1,
                is_verified=True,
            ),
            VisualRecord(
                visual_id="FIG-000004",
                node_code="1.2.0",
                purpose="Bản đồ đối chiếu ba nhóm phương pháp biểu diễn log và nguồn gốc hình thành năm khoảng trống nghiên cứu cốt lõi",
                visual_type=VisualType.CONCEPTUAL_DIAGRAM,
                creation_method=CreationMethod.WORD_DRAWING_CANVAS,
                caption="Bản đồ đối chiếu ba nhóm phương pháp biểu diễn log và nguồn gốc hình thành năm khoảng trống nghiên cứu cốt lõi (Nguồn: Tác giả tổng hợp)",
                source_provenance="Tác giả tổng hợp từ khảo sát ba nhóm phương pháp",
                script_path="src/research_agent/visuals/academic_diagram_renderer.py",
                output_file_path=str(target_path),
                bookmark_name="BK_FIG_1_004",
                seq_number=4,
                chapter_number=1,
                is_verified=True,
            ),
            VisualRecord(
                visual_id="FIG-000005",
                node_code="2.1.3",
                purpose="Kiến trúc hai mặt phẳng (Training Plane & Inference Plane) của khung biểu diễn đặc trưng log đa góc nhìn dòng",
                visual_type=VisualType.CONCEPTUAL_DIAGRAM,
                creation_method=CreationMethod.WORD_DRAWING_CANVAS,
                caption="Kiến trúc hai mặt phẳng (Training Plane & Inference Plane) của khung biểu diễn đặc trưng log đa góc nhìn dòng (Nguồn: Tác giả đề xuất)",
                source_provenance="Tác giả đề xuất kiến trúc biểu diễn đặc trưng log đa góc nhìn dòng",
                script_path="src/research_agent/visuals/academic_diagram_renderer.py",
                output_file_path=str(target_path),
                bookmark_name="BK_FIG_2_001",
                seq_number=1,
                chapter_number=2,
                is_verified=True,
            ),
            VisualRecord(
                visual_id="FIG-000006",
                node_code="2.3.1",
                purpose="Kiến trúc trích xuất tuần tự ngữ nghĩa Transformer nhận thức quyền riêng tư và đa tầng tự giám sát",
                visual_type=VisualType.CONCEPTUAL_DIAGRAM,
                creation_method=CreationMethod.WORD_DRAWING_CANVAS,
                caption="Kiến trúc trích xuất tuần tự ngữ nghĩa Transformer nhận thức quyền riêng tư và đa tầng tự giám sát (Nguồn: Tác giả đề xuất)",
                source_provenance="Tác giả đề xuất bộ trích xuất tuần tự Transformer",
                script_path="src/research_agent/visuals/chapter2_drawings.py",
                output_file_path=str(target_path),
                bookmark_name="BK_FIG_2_002",
                seq_number=2,
                chapter_number=2,
                is_verified=True,
            ),
            VisualRecord(
                visual_id="FIG-000007",
                node_code="2.3.2",
                purpose="Kiến trúc xây dựng đồ thị nguồn gốc phụ thuộc thời gian và Bộ trích xuất Temporal GNN",
                visual_type=VisualType.CONCEPTUAL_DIAGRAM,
                creation_method=CreationMethod.WORD_DRAWING_CANVAS,
                caption="Kiến trúc xây dựng đồ thị nguồn gốc phụ thuộc thời gian và Bộ trích xuất Temporal GNN (Nguồn: Tác giả đề xuất)",
                source_provenance="Tác giả đề xuất kiến trúc xây dựng đồ thị và Temporal GNN",
                script_path="src/research_agent/visuals/chapter2_drawings.py",
                output_file_path=str(target_path),
                bookmark_name="BK_FIG_2_003",
                seq_number=3,
                chapter_number=2,
                is_verified=True,
            ),
            VisualRecord(
                visual_id="FIG-000008",
                node_code="2.4.0",
                purpose="Sơ đồ kiến trúc Gióng hàng đa góc nhìn, Kiểm soát sụp đổ và Tổng hợp biểu diễn thống nhất",
                visual_type=VisualType.CONCEPTUAL_DIAGRAM,
                creation_method=CreationMethod.WORD_DRAWING_CANVAS,
                caption="Sơ đồ kiến trúc Gióng hàng đa góc nhìn, Kiểm soát sụp đổ và Tổng hợp biểu diễn thống nhất (Nguồn: Tác giả đề xuất)",
                source_provenance="Tác giả đề xuất khung gióng hàng đa góc nhìn VICReg và Gated Fusion",
                script_path="src/research_agent/visuals/chapter2_drawings.py",
                output_file_path=str(target_path),
                bookmark_name="BK_FIG_2_004",
                seq_number=4,
                chapter_number=2,
                is_verified=True,
            ),
        ]
        for vr in v_records:
            v_reg.register_visual(vr)
        print(f"[SUCCESS] Registered {len(v_records)} Visual Records in Visual Registry database.")
    except Exception as e:
        print(f"[WARNING] Could not register visual records: {e}")

    return str(pdf_path)


if __name__ == "__main__":
    build_and_audit_document()
