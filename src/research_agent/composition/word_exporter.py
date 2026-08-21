"""
Module to directly generate and append fully formatted Microsoft Word (.docx) thesis documents.
Standard: Times New Roman 14pt, 1.5 line spacing, standard thesis margins (Top 2.5cm, Bottom 2.5cm, Left 3.5cm, Right 2.0cm).
"""

import os
from pathlib import Path
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def create_or_load_thesis_docx(file_path: str = r"D:\Research\Luan_An_Tien_Si.docx") -> docx.Document:
    p = Path(file_path)
    if p.exists():
        doc = docx.Document(str(p))
    else:
        doc = docx.Document()
        # Set page margins
        for s in doc.sections:
            s.top_margin = Cm(2.5)
            s.bottom_margin = Cm(2.5)
            s.left_margin = Cm(3.5)
            s.right_margin = Cm(2.0)

        # Set default Normal style
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(14)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return doc


def set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    xml_borders = (
        f'<w:tcBorders {nsdecls("w")}>\n'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '</w:tcBorders>'
    )
    tcBorders = parse_xml(xml_borders)
    tcPr.append(tcBorders)


def add_heading_1(doc: docx.Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)
    run.bold = True


def add_heading_2(doc: docx.Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = True


def add_heading_3(doc: docx.Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = True
    run.italic = True


def add_body_paragraph(doc: docx.Document, text: str, bold_prefix: str = None, first_line_indent: bool = True):
    p = doc.add_paragraph()
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Times New Roman"
        r_pre.font.size = Pt(14)
        r_pre.bold = True

    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)


def add_table_data(doc: docx.Document, headers: list, rows_data: list):
    cols_count = len(headers)
    tbl = doc.add_table(rows=len(rows_data) + 1, cols=cols_count)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for c_idx, h in enumerate(headers):
        cell = tbl.cell(0, c_idx)
        cell.text = h
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(11)
            r.bold = True

    # Body rows
    for r_idx, row in enumerate(rows_data):
        for c_idx, val in enumerate(row):
            cell = tbl.cell(r_idx + 1, c_idx)
            cell.text = val
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(6)
    p_after.paragraph_format.space_after = Pt(6)


def export_section_1_1(output_file: str = r"D:\Research\Luan_An_Tien_Si.docx"):
    doc = create_or_load_thesis_docx(output_file)

    add_heading_1(doc, "CHƯƠNG 1: TỔNG QUAN VỀ PHƯƠNG PHÁP TRÍCH XUẤT ĐẶC TRƯNG LOG VÀ THÁCH THỨC BẢO TOÀN NGỮ CẢNH AN TOÀN")
    add_heading_2(doc, "1.1. Bài toán biểu diễn log trong phát hiện tấn công đa giai đoạn")
    add_body_paragraph(
        doc,
        "Trong các Trung tâm Điều hành An ninh mạng (Security Operations Center - SOC) hiện đại, dữ liệu nhật ký hệ thống (system logs) và nhật ký kiểm toán (audit logs) đóng vai trò là nguồn bằng chứng trung tâm phục vụ phát hiện, điều tra và ứng phó các chiến dịch tấn công có chủ đích (Advanced Persistent Threats - APT). Khác với văn bản ngôn ngữ tự nhiên thông thường hoặc tín hiệu thị giác máy tính, dữ liệu log sở hữu cấu trúc bán hình thức (semi-structured), mang tính phụ thuộc thời gian nghiêm ngặt, phản ánh các quan hệ phụ thuộc thực thi giữa các tiến trình và tài nguyên hệ điều hành, đồng thời chứa đựng các ngữ nghĩa an ninh đặc thù. Mục 1.1 tập trung hình thức hóa bài toán biểu diễn đặc trưng log phục vụ phát hiện tấn công đa giai đoạn, phân tích bản chất không gian dữ liệu doanh nghiệp, xác lập mô hình hành vi phi tuyến tính trên ma trận MITRE ATT&CK và thiết lập khung Hợp đồng Biểu diễn (Representation Contract) nhằm bảo toàn các bất biến an toàn cho không gian vector đặc trưng z."
    )

    add_heading_3(doc, "1.1.1. Không gian dữ liệu log doanh nghiệp: tốc độ cao, mất cân bằng cực đoan và phân phối biến đổi")
    add_body_paragraph(
        doc,
        "Không gian dữ liệu nhật ký trong môi trường mạng doanh nghiệp được tổng hợp từ ba nhóm nguồn telemetry chính với cấu trúc và định dạng không đồng nhất. Nhóm thứ nhất là nhật ký kiểm toán máy chủ (Host Audit Logs), bao gồm Linux Auditd, Windows Event Log / Sysmon và Linux eBPF (Extended Berkeley Packet Filter). Nguồn dữ liệu này ghi nhận trực tiếp các sự kiện ở mức nhân hệ điều hành thông qua việc chặn bắt các lời gọi hệ thống (syscalls), bao gồm: khởi tạo tiến trình (execve, CreateProcess - Sysmon Event ID 1), nạp thư viện động (ImageLoaded - Sysmon Event ID 7), thao tác tệp tin (open, unlink, FileCreate - Sysmon Event ID 11), sửa đổi cấu hình registry (RegSetValue - Sysmon Event ID 13), cùng các thao tác mở và kết nối socket mạng (connect, accept - Sysmon Event ID 3). Nhóm thứ hai là nhật ký luồng mạng (Network Flow & Protocol Logs), được thu thập từ Zeek, Suricata hoặc NetFlow/IPFIX, cung cấp siêu dữ liệu kết nối giữa các nút mạng, giao dịch DNS, chứng chỉ TLS/SSL và thông lượng gói tin. Nhóm thứ ba là nhật ký ứng dụng và dịch vụ (Application & Service Logs), phát sinh từ máy chủ web (Nginx, Apache), cơ sở dữ liệu, dịch vụ phân tán (HDFS) cùng hệ thống điều phối container (Kubernetes Audit Logs)."
    )
    add_body_paragraph(
        doc,
        "Tính dị thể sâu sắc của dữ liệu đặt ra bài toán khoa học về việc lựa chọn đơn vị quan sát (Unit of Observation) phù hợp cho mô hình học biểu diễn. Việc phân cấp đơn vị quan sát quyết định trực tiếp đến mức độ bảo toàn thông tin và độ phức tạp tính toán:"
    )

    headers_1 = ["Mức độ hạt biểu diễn", "Đơn vị quan sát", "Dữ liệu đại diện", "Ưu điểm cốt lõi", "Rào cản và Thách thức ngữ nghĩa"]
    rows_1 = [
        ["1. Mức từ tố (Token)", "Chuỗi con, từ khóa rời rạc", "Từ khóa tĩnh, địa chỉ IP, mã lỗi hex", "Dễ dàng xử lý bằng kỹ thuật nhúng từ vựng", "Mất hoàn toàn cấu trúc cú pháp và quan hệ liên kết trường"],
        ["2. Mức sự kiện (Event)", "Một dòng log đơn lẻ", "Bản ghi telemetry tại thời điểm t", "Bảo toàn đầy đủ thuộc tính cục bộ tại thời điểm t", "Thiếu ngữ cảnh chuỗi tuần tự và lịch sử tương tác trước đó"],
        ["3. Mức chuỗi / Phiên (Sequence / Session)", "Cửa sổ trượt hoặc phiên tiến trình", "Chuỗi sự kiện [e_{t-k+1}, ..., e_t] theo thời gian", "Nắm bắt quan hệ phụ thuộc thứ tự thời gian cục bộ", "Nhạy cảm với nhiễu xen kẽ (interleaving) từ các luồng chạy song song"],
        ["4. Mức thực thể (Entity)", "Định danh tác nhân (Host, User, IP, Process Instance)", "Lịch sử tương tác gom cụm theo thực thể", "Phân lập rõ ràng ranh giới hành vi của từng chủ thể", "Khó phát hiện các hành vi tấn công phối hợp vượt ranh giới thực thể"],
        ["5. Mức đồ thị (Graph)", "Đồ thị nguồn gốc (Provenance Graph)", "Đồ thị luồng phụ thuộc dị thể G", "Mô hình hóa toàn diện quan hệ phụ thuộc đa thực thể", "Bùng nổ kích thước đồ thị (dependency explosion) và chi phí tính toán"]
    ]
    add_table_data(doc, headers_1, rows_1)

    add_body_paragraph(
        doc,
        "Bên cạnh tính dị thể, dữ liệu log doanh nghiệp chịu áp lực vận hành khắc nghiệt về mặt thông lượng và tỷ lệ phân bố nhãn. Trong các môi trường mạng lớn, hệ thống tiếp nhận luồng sự kiện liên tục với thông lượng cao, đòi hỏi thuật toán biểu diễn đặc trưng phải xử lý theo cơ chế dòng (streaming) với độ phức tạp tính toán tuyến tính O(N) mà không đòi hỏi lưu giữ toàn bộ lịch sử đồ thị trong bộ nhớ truy cập ngẫu nhiên. Đồng thời, tỷ lệ các dòng log liên quan đến hành vi tấn công trong thực tế thường ở mức rất thấp, tạo ra sự mất cân bằng nhãn cực đoan khiến các hàm mất mát học máy thông thường có xu hướng xem nhẹ nhóm thiểu số. Trong khi đó, việc suy thoái không gian vector (sụp đổ chiều biểu diễn) nảy sinh khi các mô hình tự giám sát không áp dụng các cơ chế điều hòa phương sai - hiệp phương sai phù hợp để duy trì tính đa dạng của các chiều đặc trưng."
    )
    add_body_paragraph(
        doc,
        "Dưới góc độ phân phối thời gian, dữ liệu log liên tục biến đổi trong môi trường vận hành dài hạn. Để tránh các sai lệch phương pháp luận khi đánh giá mô hình, luận án phân định tường minh bốn cơ chế trôi dạt dữ liệu (Drift Taxonomy):"
    )
    add_body_paragraph(doc, "Bản chất hành vi và mục đích tấn công thay đổi theo thời gian mặc dù cấu trúc định dạng log không đổi: P_t(Y | X) != P_{t+1}(Y | X).", bold_prefix="• Concept Drift: ")
    add_body_paragraph(doc, "Việc nâng cấp phần mềm, cập nhật bản vá hoặc thay đổi cấu hình làm biến đổi cấu trúc chuỗi mẫu log: P_t(X_template) != P_{t+1}(X_template).", bold_prefix="• Template Drift: ")
    add_body_paragraph(doc, "Lưu lượng người dùng, tần suất giao tác nghiệp vụ hoặc cơ cấu dịch vụ hệ thống biến động theo chu kỳ: P_t(X) != P_{t+1}(X).", bold_prefix="• Population Drift: ")
    add_body_paragraph(doc, "Không gian vector tiềm ẩn z bị suy giảm năng lực phân tách do dữ liệu đầu vào trôi dạt khỏi vùng phân phối huấn luyện ban đầu: P_t(z | X) != P_{t+1}(z | X).", bold_prefix="• Representation Drift: ")

    add_heading_3(doc, "1.1.2. Hành vi tấn công đa giai đoạn và ánh xạ đa nhãn MITRE ATT&CK")
    add_body_paragraph(
        doc,
        "Trong phân tích an ninh mạng, một số công trình áp dụng mô hình trạng thái Markov tuyến tính tuần tự để xâu chuỗi các giai đoạn tấn công từ Thâm nhập ban đầu (Initial Access), Thực thi (Execution), Duy trì (Persistence), Leo thang đặc quyền (Privilege Escalation) cho đến Đánh cắp dữ liệu (Exfiltration)."
    )
    add_body_paragraph(
        doc,
        "Tuy nhiên, các quan sát thực nghiệm trên các chiến dịch tấn công cho thấy hành vi của kẻ tấn công mang bản chất phi tuyến tính: (1) Nhảy cóc giai đoạn (Step Skipping): Kẻ tấn công có thể khai thác trực tiếp lỗ hổng thực thi mã từ xa để trích xuất dữ liệu ra ngoài mà không cần thiết lập cơ chế duy trì hay di chuyển ngang; (2) Lặp vòng kỹ thuật (Tactic Looping & Interleaving): Kỹ thuật thu thập thông tin nội bộ (Discovery) thường được lặp lại nhiều lần xen kẽ giữa các bước leo thang đặc quyền và chiếm đoạt thông tin xác thực; (3) Phân nhánh tiến trình song song (Parallel Branching): Kẻ tấn công có thể khởi tạo đồng thời nhiều luồng tiến trình con độc lập trên các tiến trình hợp lệ khác nhau nhằm phân tán sự theo dõi của hệ thống phòng thủ."
    )
    add_body_paragraph(
        doc,
        "Do đó, luận án xác lập nguyên tắc: Ma trận MITRE ATT&CK được mô hình hóa thành một Không gian Bằng chứng Hành vi Đa chiều (Multi-label Behavioral Evidence Space) Y subseteq {0, 1}^{|T|}, trong đó một chuỗi sự kiện hoặc cây tiến trình có thể đồng thời kích hoạt nhiều nhãn chiến thuật (Tactics) và kỹ thuật (Techniques) tại cùng một thời điểm quan sát."
    )
    add_body_paragraph(
        doc,
        "Về mặt dữ liệu thực nghiệm, việc mô hình hóa hành vi tấn công đòi hỏi phải phân định chính xác đặc tính gán nhãn và mức độ hạt (Label Granularity) của từng bộ dữ liệu chuẩn: (1) DARPA Transparent Computing (TC E3/E5) cung cấp telemetry kiểm toán mức nhân hệ điều hành với nhãn mặt đất được ghi nhận ở mức kịch bản tấn công tổng thể kèm theo danh sách các nút và cạnh liên quan trong đồ thị nguồn gốc; (2) LANL CyberEvents (2017) kết hợp xác thực máy tính, tiến trình và luồng mạng quy mô lớn với nhãn độc hại gắn theo mốc thời gian và tài khoản/máy tính bị đội đỏ xâm nhập; (3) HDFS Benchmark ghi nhận thao tác khối dữ liệu Hadoop với nhãn bất thường phản ánh sự cố kỹ thuật hoặc lỗi khối (Block Anomaly), không phải tấn công APT; (4) BGL Benchmark ghi nhận nhật ký máy siêu điện toán BlueGene/L với nhãn cảnh báo lỗi phần cứng và hệ thống phản ánh độ tin cậy vận hành."
    )
    add_body_paragraph(
        doc,
        "Đặc biệt, sự xuất hiện của nhiễu từ hành vi quản trị viên (Admin-Noise) là một thách thức thực tế lớn. Quản trị viên hệ thống thường xuyên sử dụng các công cụ dòng lệnh hợp lệ (Living-off-the-Land Binaries - LOLBins) như PowerShell, WMI (wmic.exe), SSH, certutil.exe hay vssadmin.exe cho mục đích bảo trì, sao lưu hoặc kiểm tra mạng. Các hành vi này có cấu trúc cú pháp tương tự như kỹ thuật của kẻ tấn công. Nếu mô hình chỉ dựa vào sự xuất hiện của từ khóa, hệ thống sẽ phát sinh nhiều cảnh báo sai (False Positives). Do đó, bài toán biểu diễn đòi hỏi cơ chế phân bổ bằng chứng yếu (Weak Evidence Attribution) thông qua khung Học Đa Thể hiện (Multiple Instance Learning - MIL)."
    )

    add_heading_3(doc, "1.1.3. Các mức biểu diễn dữ liệu và Hợp đồng Biểu diễn (Representation Contract)")
    add_body_paragraph(
        doc,
        "Nhằm định hình rõ ràng các yêu cầu đối với không gian vector đặc trưng z in R^d và ngăn ngừa rủi ro học đường tắt (Shortcut Learning), mô hình biểu diễn bắt buộc phải tuân thủ Hợp đồng Biểu diễn (Representation Contract) được định nghĩa dưới dạng bộ ba hình thức C_representation = < P_preserve, I_invariant, E_exclude >:"
    )

    headers_2 = ["Nhóm quy tắc", "Ý nghĩa phương pháp luận", "Danh mục thuộc tính telemetry tương ứng"]
    rows_2 = [
        ["PRESERVE (Bảo toàn)", "Bảo toàn các thuộc tính mang ngữ nghĩa an ninh và quan hệ phụ thuộc thực thi trong không gian vector z.", "• Thứ tự thời gian cục bộ giữa các sự kiện liền kề.\n• Quan hệ tiến trình cha-con (parent_process_guid -> process_guid).\n• Cấu trúc liên kết đa thực thể (Tiến trình - Tệp tin - Socket - Registry).\n• Tham số an ninh động: IP đích, cổng, đường dẫn tệp, command_line, return_code, ImageHash."],
        ["INVARIANT (Bất biến)", "Duy trì tính bất biến của vector z trước các biến đổi hình thức không làm thay đổi bản chất hành vi.", "• Biến đổi cú pháp vô hại: khoảng trắng thừa, dấu xuống dòng, chữ hoa/thường trong câu lệnh không phân biệt hoa-thường.\n• Định danh tiến trình tạm thời ngẫu nhiên (PID, Thread ID biến động qua mỗi lần khởi chạy).\n• Hoán vị thứ tự giữa các trường thuộc tính độc lập không có quan hệ phụ thuộc tuần tự."],
        ["EXCLUDE (Triệt tiêu)", "Loại bỏ các thuộc tính gây học vẹt, tạo đường tắt giả định hoặc vi phạm quyền riêng tư khỏi vector z.", "• Định danh máy chủ cố định gây thiên lệch (Host UUID, tên máy trạm thử nghiệm cụ thể, địa chỉ MAC tĩnh).\n• Dấu vết phân vùng nhân tạo (Synthetic Split Artifacts) phát sinh trong quá trình chia tập train/test.\n• Mọi đặc trưng học đường tắt (Shortcut Features) có tương quan ngẫu nhiên với nhãn trong tập huấn luyện nhưng không có giá trị tổng quát hóa."]
    ]
    add_table_data(doc, headers_2, rows_2)

    add_body_paragraph(
        doc,
        "Cần nhấn mạnh rằng, nguyên tắc PRESERVE tập trung vào việc bảo toàn ngữ nghĩa an ninh động chứ không đồng nghĩa với việc giữ lại toàn bộ định danh thô của người dùng hay máy chủ, nhằm đảm bảo khả năng liên kết có kiểm soát (controlled linkability) và tương thích với các tiêu chuẩn đánh giá quyền riêng tư. Tương tự, nguyên tắc EXCLUDE thiết lập các ràng buộc phủ định đối với các giả định đường tắt đã biết và các biến số gây rò rỉ phân vùng, thay vì giả định rằng mọi đường tắt đều có thể nhận biết trước khi thực nghiệm."
    )
    add_body_paragraph(doc, "Đi đôi với Hợp đồng Biểu diễn, luận án thiết lập nguyên tắc phân định ranh giới phương pháp luận ba tầng độc lập:")
    add_body_paragraph(
        doc,
        "Đảm nhiệm tiền xử lý dữ liệu thô, phân tích cú pháp sơ bộ, chuẩn hóa kiểu dữ liệu trường và làm sạch dữ liệu. Tầng này không gánh vác nhiệm vụ học biểu diễn ngữ cảnh an ninh sâu.",
        bold_prefix="1. Tầng 1 — Trích xuất đặc trưng cơ sở (Feature Extraction): "
    )
    add_body_paragraph(
        doc,
        "Thiết lập ánh xạ f_theta: X -> z in R^d từ cấu trúc chuỗi sự kiện và đồ thị nguồn gốc sang không gian vector tiềm ẩn. Toàn bộ năng lực bảo toàn ngữ nghĩa an ninh và tính bất biến được đóng gói trọn vẹn bên trong vector z.",
        bold_prefix="2. Tầng 2 — Học không gian biểu diễn (Representation Learning — Trọng tâm Luận án): "
    )
    add_body_paragraph(
        doc,
        "Đánh giá chất lượng của vector biểu diễn z thông qua các bộ thăm dò tuyến tính đóng băng tham số (Frozen Linear Probing): y_hat = sigma(W^T z + b). Trong đó tham số theta của bộ trích xuất đặc trưng Tầng 2 được giữ cố định hoàn toàn trong suốt quá trình đánh giá ở Tầng 3. Quy tắc này bảo đảm bộ phân loại hạ nguồn không làm thay nhiệm vụ trích xuất đặc trưng của Tầng 2.",
        bold_prefix="3. Tầng 3 — Phát hiện và phân loại hạ nguồn (Downstream Detection): "
    )

    doc.save(output_file)
    print(f"Document successfully written to: {output_file}")


if __name__ == "__main__":
    export_section_1_1()
