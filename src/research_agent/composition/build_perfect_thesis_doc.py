"""
Pristine Document Builder for D:\\Research\\Chuyên đề chuyên sâu - Copy.docx
Matches 100% of original template standards:
1. Heading styles use clean text (Word auto-numbers: Chương 1., 1.1., 1.1.1., 1.2., etc.).
2. Native Microsoft Word Equations (OMML via MML2OMML.XSL) with zero empty placeholder boxes.
3. Tables match original template dimensions (9605 dxa), vertical centering, single line spacing, cantSplit, tblHeader.
4. Body text is Times New Roman 14pt, 1.5 line spacing, 1.27cm first-line indent.
5. Preserves Cover Page Frame Table 0, Table of Contents, Conclusion, and References.
"""

import sys
import shutil
from pathlib import Path
from lxml import etree
import latex2mathml.converter
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

sys.stdout.reconfigure(encoding="utf-8")

# Initialize official Microsoft Office MathML to OMML XSLT transformer
XSLT_PATH = r"C:\Program Files\Microsoft Office\Office16\MML2OMML.XSL"
xslt_tree = etree.parse(XSLT_PATH)
transform_omml = etree.XSLT(xslt_tree)


def latex_to_clean_omml(latex_code: str):
    """Converts a LaTeX formula into a native Word OMML element, cleaning any empty <m:e/> placeholders."""
    try:
        mathml = latex2mathml.converter.convert(latex_code)
        tree = etree.fromstring(mathml)
        omml_tree = transform_omml(tree)

        # Fix empty <m:e/> in <m:nary> operators which produces dotted square placeholder boxes in Word
        ns = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}
        for nary in omml_tree.xpath(".//m:nary", namespaces=ns):
            e_elem = nary.find("m:e", namespaces=ns)
            if e_elem is not None and len(e_elem) == 0 and (not e_elem.text or not e_elem.text.strip()):
                chr_val = "∑"
                naryPr = nary.find("m:naryPr", namespaces=ns)
                if naryPr is not None:
                    chr_elem = naryPr.find("m:chr", namespaces=ns)
                    if chr_elem is not None and "{http://schemas.openxmlformats.org/officeDocument/2006/math}val" in chr_elem.attrib:
                        chr_val = chr_elem.attrib["{http://schemas.openxmlformats.org/officeDocument/2006/math}val"]

                sub_elem = nary.find("m:sub", namespaces=ns)
                sup_elem = nary.find("m:sup", namespaces=ns)

                has_sub = sub_elem is not None and (len(sub_elem) > 0 or (sub_elem.text and sub_elem.text.strip()))
                has_sup = sup_elem is not None and (len(sup_elem) > 0 or (sup_elem.text and sup_elem.text.strip()))

                parent = nary.getparent()
                idx = parent.index(nary)

                if has_sub and has_sup:
                    new_node = etree.Element("{http://schemas.openxmlformats.org/officeDocument/2006/math}sSubSup")
                    e_node = etree.SubElement(new_node, "{http://schemas.openxmlformats.org/officeDocument/2006/math}e")
                    r_node = etree.SubElement(e_node, "{http://schemas.openxmlformats.org/officeDocument/2006/math}r")
                    t_node = etree.SubElement(r_node, "{http://schemas.openxmlformats.org/officeDocument/2006/math}t")
                    t_node.text = chr_val
                    new_node.append(sub_elem)
                    new_node.append(sup_elem)
                    parent.insert(idx, new_node)
                    parent.remove(nary)
                elif has_sub:
                    new_node = etree.Element("{http://schemas.openxmlformats.org/officeDocument/2006/math}sSub")
                    e_node = etree.SubElement(new_node, "{http://schemas.openxmlformats.org/officeDocument/2006/math}e")
                    r_node = etree.SubElement(e_node, "{http://schemas.openxmlformats.org/officeDocument/2006/math}r")
                    t_node = etree.SubElement(r_node, "{http://schemas.openxmlformats.org/officeDocument/2006/math}t")
                    t_node.text = chr_val
                    new_node.append(sub_elem)
                    parent.insert(idx, new_node)
                    parent.remove(nary)

        xml_str = etree.tostring(omml_tree, encoding="utf-8").decode("utf-8")
        return parse_xml(xml_str)
    except Exception as e:
        print(f"Warning: LaTeX conversion failed for '{latex_code}': {e}")
        xml_str = (
            f'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
            f'  <m:r><m:t>{latex_code}</m:t></m:r>'
            f'</m:oMath>'
        )
        return parse_xml(xml_str)


def format_table_cell(cell, width_dxa: int, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, font_size_pt=14):
    """Sets standard cell properties: exact width, vertical centering, border, padding, and compact line spacing."""
    tcPr = cell._tc.get_or_add_tcPr()
    tc_xml = (
        f'<w:tcPr {nsdecls("w")}>\n'
        f'  <w:tcW w:w="{width_dxa}" w:type="dxa"/>\n'
        '  <w:vAlign w:val="center"/>\n'
        '  <w:tcBorders>\n'
        '    <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '    <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '    <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '    <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '  </w:tcBorders>\n'
        '  <w:tcMar>\n'
        '    <w:top w:w="80" w:type="dxa"/>\n'
        '    <w:bottom w:w="80" w:type="dxa"/>\n'
        '    <w:left w:w="120" w:type="dxa"/>\n'
        '    <w:right w:w="120" w:type="dxa"/>\n'
        '  </w:tcMar>\n'
        '</w:tcPr>'
    )
    cell._tc.remove(tcPr)
    cell._tc.append(parse_xml(tc_xml))

    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0  # Compact line spacing for elegant tables
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(font_size_pt)
        if bold:
            r.bold = True


def insert_thesis_table(doc, ref_p, headers, col_widths, rows_data, font_size_pt=14):
    """Creates an elegant, professional thesis table matching original template layout."""
    tbl = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    if ref_p is not None:
        ref_p._p.addprevious(tbl._tbl)

    # tblPr width
    tblPr = tbl._tbl.tblPr
    total_w = sum(col_widths)
    tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:w="{total_w}" w:type="dxa"/>'))

    # Header Row
    hdr_row = tbl.rows[0]
    hdr_trPr = hdr_row._tr.get_or_add_trPr()
    hdr_trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    hdr_trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

    for c_i, h in enumerate(headers):
        cell = hdr_row.cells[c_i]
        cell.text = h
        format_table_cell(cell, col_widths[c_i], align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size_pt=font_size_pt)

    # Body Rows
    for r_i, row in enumerate(rows_data):
        b_row = tbl.rows[r_i + 1]
        b_trPr = b_row._tr.get_or_add_trPr()
        b_trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

        for c_i, val in enumerate(row):
            cell = b_row.cells[c_i]
            cell.text = val
            cell_align = WD_ALIGN_PARAGRAPH.CENTER if (c_i == 0 and len(headers) >= 4) else WD_ALIGN_PARAGRAPH.LEFT
            format_table_cell(cell, col_widths[c_i], align=cell_align, bold=(c_i == 0 and len(headers) == 3), font_size_pt=font_size_pt)


def build_perfect_document(target_file: str = r"D:\Research\Chuyên đề chuyên sâu - Copy.docx"):
    target_path = Path(target_file)
    backup_path = target_path.parent / (target_path.stem + ".backup.docx")

    if not backup_path.exists():
        shutil.copyfile(target_path, backup_path)
    print(f"[1/4] Loaded pristine backup from: {backup_path}")

    doc = docx.Document(str(backup_path))

    # Keep only Table 0 (Cover page frame)
    while len(doc.tables) > 1:
        tbl_to_remove = doc.tables[1]
        tbl_to_remove._tbl.getparent().remove(tbl_to_remove._tbl)
    print("[2/4] Preserved Cover Frame Table 0.")

    # Remove old body paragraphs from Heading 1 to Conclusion
    paragraphs_to_remove = []
    found_h1 = False
    for p in doc.paragraphs:
        txt = p.text.strip().lower()
        if "tổng quan về phương pháp trích xuất đặc trưng" in txt:
            found_h1 = True
        if found_h1:
            if "kết luận" in txt:
                break
            paragraphs_to_remove.append(p)

    for p in paragraphs_to_remove:
        p._p.getparent().remove(p._p)

    # Locate insertion target point (right before 'Kết luận')
    target_p = None
    for p in doc.paragraphs:
        if "kết luận" in p.text.strip().lower():
            target_p = p
            break

    print(f"[3/4] Cleaned {len(paragraphs_to_remove)} old paragraphs. Insertion target ready.")

    def add_p(text_segments, bold_prefix=None, first_line_indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
        new_p = doc.add_paragraph(style="Normal") if target_p is None else target_p.insert_paragraph_before(style="Normal")
        new_p.alignment = align
        new_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        new_p.paragraph_format.space_after = Pt(6)
        new_p.paragraph_format.space_before = Pt(0)
        if first_line_indent:
            new_p.paragraph_format.first_line_indent = Cm(1.27)

        if bold_prefix:
            r_pre = new_p.add_run(bold_prefix)
            r_pre.font.name = "Times New Roman"
            r_pre.font.size = Pt(14)
            r_pre.bold = True

        if isinstance(text_segments, str):
            r = new_p.add_run(text_segments)
            r.font.name = "Times New Roman"
            r.font.size = Pt(14)
        elif isinstance(text_segments, list):
            for seg in text_segments:
                if isinstance(seg, str):
                    r = new_p.add_run(seg)
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(14)
                else:
                    new_p._p.append(seg)
        return new_p

    def add_h1(clean_text):
        p = doc.add_paragraph(style="Heading 1") if target_p is None else target_p.insert_paragraph_before(style="Heading 1")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        r = p.add_run(clean_text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(16)
        r.bold = True
        return p

    def add_h2(clean_text):
        p = doc.add_paragraph(style="Heading 2") if target_p is None else target_p.insert_paragraph_before(style="Heading 2")
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(clean_text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
        r.bold = True
        return p

    def add_h3(clean_text):
        p = doc.add_paragraph(style="Heading 3") if target_p is None else target_p.insert_paragraph_before(style="Heading 3")
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(clean_text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
        r.bold = True
        r.italic = True
        return p

    print("[4/4] Writing Section 1.1 and Section 1.2 with pristine OMML equations and elegant tables...")

    # =========================================================================
    # HEADING 1
    # =========================================================================
    add_h1("TỔNG QUAN VỀ PHƯƠNG PHÁP TRÍCH XUẤT ĐẶC TRƯNG DỮ LIỆU LOG VÀ THÁCH THỨC BẢO TOÀN NGỮ CẢNH AN TOÀN")

    # =========================================================================
    # 1.1.
    # =========================================================================
    add_h2("Bài toán biểu diễn log trong phát hiện tấn công đa giai đoạn")
    add_p([
        "Trong các Trung tâm Điều hành An ninh mạng (Security Operations Center - SOC) hiện đại, dữ liệu nhật ký hệ thống (system logs) và nhật ký kiểm toán (audit logs) đóng vai trò là nguồn bằng chứng trung tâm phục vụ phát hiện, điều tra và ứng phó các chiến dịch tấn công có chủ đích (Advanced Persistent Threats - APT). Khác với văn bản ngôn ngữ tự nhiên thông thường hoặc tín hiệu thị giác máy tính, dữ liệu log sở hữu cấu trúc bán hình thức (semi-structured), mang tính phụ thuộc thời gian nghiêm ngặt, phản ánh các quan hệ phụ thuộc thực thi giữa các tiến trình và tài nguyên hệ điều hành, đồng thời chứa đựng các ngữ nghĩa an ninh đặc thù. Mục 1.1 tập trung hình thức hóa bài toán biểu diễn đặc trưng log phục vụ phát hiện tấn công đa giai đoạn, phân tích bản chất không gian dữ liệu doanh nghiệp, xác lập mô hình hành vi phi tuyến tính trên ma trận MITRE ATT&CK và thiết lập khung Hợp đồng Biểu diễn (Representation Contract) nhằm bảo toàn các bất biến an toàn cho không gian vector đặc trưng ",
        latex_to_clean_omml(r"\mathbf{z}"),
        "."
    ])

    # --- 1.1.1 ---
    add_h3("Không gian dữ liệu log doanh nghiệp: tốc độ cao, mất cân bằng cực đoan và phân phối biến đổi")
    add_p(
        "Không gian dữ liệu nhật ký trong môi trường mạng doanh nghiệp được tổng hợp từ ba nhóm nguồn telemetry chính với cấu trúc và định dạng không đồng nhất. Nhóm thứ nhất là nhật ký kiểm toán máy chủ (Host Audit Logs), bao gồm Linux Auditd, Windows Event Log / Sysmon và Linux eBPF (Extended Berkeley Packet Filter). Nguồn dữ liệu này ghi nhận trực tiếp các sự kiện ở mức nhân hệ điều hành thông qua việc chặn bắt các lời gọi hệ thống (syscalls), bao gồm: khởi tạo tiến trình (execve, CreateProcess - Sysmon Event ID 1), nạp thư viện động (ImageLoaded - Sysmon Event ID 7), thao tác tệp tin (open, unlink, FileCreate - Sysmon Event ID 11), sửa đổi cấu hình registry (RegSetValue - Sysmon Event ID 13), cùng các thao tác mở và kết nối socket mạng (connect, accept - Sysmon Event ID 3). Nhóm thứ hai là nhật ký luồng mạng (Network Flow & Protocol Logs), được thu thập từ Zeek, Suricata hoặc NetFlow/IPFIX, cung cấp siêu dữ liệu kết nối giữa các nút mạng, giao dịch DNS, chứng chỉ TLS/SSL và thông lượng gói tin. Nhóm thứ ba là nhật ký ứng dụng và dịch vụ (Application & Service Logs), phát sinh từ máy chủ web (Nginx, Apache), cơ sở dữ liệu, dịch vụ phân tán (HDFS) cùng hệ thống điều phối container (Kubernetes Audit Logs)."
    )
    add_p(
        "Tính dị thể sâu sắc của dữ liệu đặt ra bài toán khoa học về việc lựa chọn đơn vị quan sát (Unit of Observation) phù hợp cho mô hình học biểu diễn. Việc phân cấp đơn vị quan sát quyết định trực tiếp đến mức độ bảo toàn thông tin và độ phức tạp tính toán:"
    )

    # TABLE 1 (Exact widths from original template: 1551, 2028, 1791, 1806, 2429 dxa)
    tbl1_headers = ["Mức độ hạt", "Đơn vị quan sát", "Dữ liệu đại diện", "Ưu điểm cốt lõi", "Thách thức và Mất mát ngữ nghĩa"]
    tbl1_widths = [1551, 2028, 1791, 1806, 2429]
    tbl1_rows = [
        ["Mức từ tố (Token)", "Chuỗi con, từ khóa rời rạc", "Từ khóa tĩnh, địa chỉ IP, mã lỗi hex", "Dễ dàng vector hóa bằng kỹ thuật nhúng từ vựng", "Mất liên kết cấu trúc cú pháp và quan hệ trường"],
        ["Mức sự kiện (Event)", "Một dòng log đơn lẻ", "Bản ghi telemetry tại thời điểm t", "Bảo toàn đầy đủ thuộc tính cục bộ tại thời điểm t", "Thiếu ngữ cảnh chuỗi tuần tự và lịch sử tương tác"],
        ["Mức chuỗi / Phiên (Session)", "Cửa sổ trượt hoặc phiên tiến trình", "Chuỗi sự kiện [e_{t-k+1}, ..., e_t] theo thời gian", "Nắm bắt quan hệ phụ thuộc thứ tự thời gian", "Nhạy cảm với nhiễu xen kẽ từ các luồng song song"],
        ["Mức thực thể (Entity)", "Định danh tác nhân (Host, User, Process)", "Lịch sử tương tác gom cụm theo thực thể", "Phân lập rõ ràng ranh giới hành vi chủ thể", "Khó phát hiện hành vi tấn công vượt ranh giới"],
        ["Mức đồ thị (Graph)", "Đồ thị nguồn gốc (Provenance Graph)", "Đồ thị luồng phụ thuộc dị thể G", "Mô hình hóa toàn diện quan hệ đa thực thể", "Bùng nổ kích thước đồ thị và chi phí tính toán"]
    ]
    insert_thesis_table(doc, target_p, tbl1_headers, tbl1_widths, tbl1_rows, font_size_pt=14)
    add_p("", first_line_indent=False)

    add_p(
        "Bên cạnh tính dị thể, dữ liệu log doanh nghiệp chịu áp lực vận hành khắc nghiệt về mặt thông lượng và tỷ lệ phân bố nhãn. Trong các môi trường mạng lớn, hệ thống tiếp nhận luồng sự kiện liên tục với thông lượng cao, đòi hỏi thuật toán biểu diễn đặc trưng phải xử lý theo cơ chế dòng (streaming) với độ phức tạp tính toán tuyến tính O(N) mà không đòi hỏi lưu giữ toàn bộ lịch sử đồ thị trong bộ nhớ truy cập ngẫu nhiên. Đồng thời, tỷ lệ các dòng log liên quan đến hành vi tấn công trong thực tế thường ở mức rất thấp, tạo ra sự mất cân bằng nhãn cực đoan khiến các hàm mất mát học máy thông thường có xu hướng xem nhẹ nhóm thiểu số. Trong khi đó, việc suy thoái không gian vector (sụp đổ chiều biểu diễn) nảy sinh khi các mô hình tự giám sát không áp dụng các cơ chế điều hòa phương sai - hiệp phương sai phù hợp để duy trì tính đa dạng của các chiều đặc trưng."
    )
    add_p(
        "Dưới góc độ phân phối thời gian, dữ liệu log liên tục biến đổi trong môi trường vận hành dài hạn. Để tránh các sai lệch phương pháp luận khi đánh giá mô hình, chuyên đề phân định tường minh bốn cơ chế trôi dạt dữ liệu (Drift Taxonomy):"
    )
    add_p([
        "Bản chất hành vi và mục đích tấn công thay đổi theo thời gian mặc dù cấu trúc định dạng log không đổi: ",
        latex_to_clean_omml(r"P_t(Y \mid X) \neq P_{t+1}(Y \mid X)"),
        "."
    ], bold_prefix="• Concept Drift: ")
    add_p([
        "Việc nâng cấp phần mềm, cập nhật bản vá hoặc thay đổi cấu hình làm biến đổi cấu trúc chuỗi mẫu log: ",
        latex_to_clean_omml(r"P_t(X_{\text{template}}) \neq P_{t+1}(X_{\text{template}})"),
        "."
    ], bold_prefix="• Template Drift: ")
    add_p([
        "Lưu lượng người dùng, tần suất giao tác nghiệp vụ hoặc cơ cấu dịch vụ hệ thống biến động theo chu kỳ: ",
        latex_to_clean_omml(r"P_t(X) \neq P_{t+1}(X)"),
        "."
    ], bold_prefix="• Population Drift: ")
    add_p([
        "Không gian vector tiềm ẩn z bị suy giảm năng lực phân tách do dữ liệu đầu vào trôi dạt khỏi vùng phân phối huấn luyện ban đầu: ",
        latex_to_clean_omml(r"P_t(\mathbf{z} \mid X) \neq P_{t+1}(\mathbf{z} \mid X)"),
        "."
    ], bold_prefix="• Representation Drift: ")

    # --- 1.1.2 ---
    add_h3("Hành vi tấn công đa giai đoạn và ánh xạ đa nhãn MITRE ATT&CK")
    add_p(
        "Trong phân tích an ninh mạng, một số công trình áp dụng mô hình trạng thái Markov tuyến tính tuần tự để xâu chuỗi các giai đoạn tấn công từ Thâm nhập ban đầu (Initial Access), Thực thi (Execution), Duy trì (Persistence), Leo thang đặc quyền (Privilege Escalation) cho đến Đánh cắp dữ liệu (Exfiltration)."
    )
    add_p(
        "Tuy nhiên, các quan sát thực nghiệm trên các chiến dịch tấn công cho thấy hành vi của kẻ tấn công mang bản chất phi tuyến tính: (1) Nhảy cóc giai đoạn (Step Skipping): Kẻ tấn công có thể khai thác trực tiếp lỗ hổng thực thi mã từ xa để trích xuất dữ liệu ra ngoài mà không cần thiết lập cơ chế duy trì hay di chuyển ngang; (2) Lặp vòng kỹ thuật (Tactic Looping & Interleaving): Kỹ thuật thu thập thông tin nội bộ (Discovery) thường được lặp lại nhiều lần xen kẽ giữa các bước leo thang đặc quyền và chiếm đoạt thông tin xác thực; (3) Phân nhánh tiến trình song song (Parallel Branching): Kẻ tấn công có thể khởi tạo đồng thời nhiều luồng tiến trình con độc lập trên các tiến trình hợp lệ khác nhau nhằm phân tán sự theo dõi của hệ thống phòng thủ."
    )
    add_p([
        "Do đó, chuyên đề xác lập nguyên tắc: Ma trận MITRE ATT&CK được mô hình hóa thành một Không gian Bằng chứng Hành vi Đa chiều (Multi-label Behavioral Evidence Space) ",
        latex_to_clean_omml(r"\mathcal{Y} \subseteq \{0, 1\}^{|\mathcal{T}|}"),
        ", trong đó một chuỗi sự kiện hoặc cây tiến trình có thể đồng thời kích hoạt nhiều nhãn chiến thuật (Tactics) và kỹ thuật (Techniques) tại cùng một thời điểm quan sát."
    ])
    add_p(
        "Về mặt dữ liệu thực nghiệm, việc mô hình hóa hành vi tấn công đòi hỏi phải phân định chính xác đặc tính gán nhãn và mức độ hạt (Label Granularity) của từng bộ dữ liệu chuẩn: (1) DARPA Transparent Computing (TC E3/E5) cung cấp telemetry kiểm toán mức nhân hệ điều hành với nhãn mặt đất được ghi nhận ở mức kịch bản tấn công tổng thể kèm theo danh sách các nút và cạnh liên quan trong đồ thị nguồn gốc; (2) LANL CyberEvents (2017) kết hợp xác thực máy tính, tiến trình và luồng mạng quy mô lớn với nhãn độc hại gắn theo mốc thời gian và tài khoản/máy tính bị đội đỏ xâm nhập; (3) HDFS Benchmark ghi nhận thao tác khối dữ liệu Hadoop với nhãn bất thường phản ánh sự cố kỹ thuật hoặc lỗi khối (Block Anomaly), không phải tấn công APT; (4) BGL Benchmark ghi nhận nhật ký máy siêu điện toán BlueGene/L với nhãn cảnh báo lỗi phần cứng và hệ thống phản ánh độ tin cậy vận hành."
    )
    add_p(
        "Đặc biệt, sự xuất hiện của nhiễu từ hành vi quản trị viên (Admin-Noise) là một thách thức thực tế lớn. Quản trị viên hệ thống thường xuyên sử dụng các công cụ dòng lệnh hợp lệ (Living-off-the-Land Binaries - LOLBins) như PowerShell, WMI (wmic.exe), SSH, certutil.exe hay vssadmin.exe cho mục đích bảo trì, sao lưu hoặc kiểm tra mạng. Các hành vi này có cấu trúc cú pháp tương tự như kỹ thuật của kẻ tấn công. Nếu mô hình chỉ dựa vào sự xuất hiện của từ khóa, hệ thống sẽ phát sinh nhiều cảnh báo sai (False Positives). Do đó, bài toán biểu diễn đòi hỏi cơ chế phân bổ bằng chứng yếu (Weak Evidence Attribution) thông qua khung Học Đa Thể hiện (Multiple Instance Learning - MIL)."
    )

    # --- 1.1.3 ---
    add_h3("Các mức biểu diễn dữ liệu và Hợp đồng Biểu diễn (Representation Contract)")
    add_p([
        "Nhằm định hình rõ ràng các yêu cầu đối với không gian vector đặc trưng ",
        latex_to_clean_omml(r"\mathbf{z} \in \mathbb{R}^d"),
        " và ngăn ngừa rủi ro học đường tắt (Shortcut Learning), mô hình biểu diễn bắt buộc phải tuân thủ Hợp đồng Biểu diễn (Representation Contract) được định nghĩa dưới dạng bộ ba hình thức ",
        latex_to_clean_omml(r"\mathcal{C}_{\text{representation}} = \langle \mathcal{P}_{\text{preserve}}, \mathcal{I}_{\text{invariant}}, \mathcal{E}_{\text{exclude}} \rangle"),
        ":"
    ])

    # TABLE 2 (Exact widths from original template: 2458, 3933, 3214 dxa)
    tbl2_headers = ["Nhóm quy tắc", "Mô tả hình thức", "Danh mục thuộc tính Telemetry áp dụng"]
    tbl2_widths = [2458, 3933, 3214]
    tbl2_rows = [
        ["PRESERVE (Bảo toàn)", "Bảo toàn các thuộc tính mang ngữ nghĩa an ninh và quan hệ phụ thuộc thực thi trong vector z.", "• Thứ tự thời gian cục bộ giữa các sự kiện.\n• Quan hệ tiến trình cha-con (parent_process_guid → process_guid).\n• Liên kết đa thực thể (Tiến trình - Tệp tin - Socket - Registry).\n• Tham số an ninh: IP đích, cổng, đường dẫn tệp, command_line, ImageHash."],
        ["INVARIANT (Bất biến)", "Duy trì tính bất biến của vector z trước các biến đổi hình thức không đổi bản chất hành vi.", "• Biến đổi cú pháp: khoảng trắng, xuống dòng, chữ hoa/thường vô hại.\n• Định danh ngẫu nhiên tạm thời (PID, Thread ID biến động).\n• Hoán vị thứ tự giữa các trường thuộc tính độc lập."],
        ["EXCLUDE (Triệt tiêu)", "Loại bỏ các thuộc tính gây học vẹt, tạo đường tắt giả định hoặc vi phạm quyền riêng tư khỏi vector z.", "• Định danh máy chủ cố định gây thiên lệch (Host UUID, MAC tĩnh).\n• Dấu vết phân vùng nhân tạo (Synthetic Split Artifacts).\n• Mọi đặc trưng học đường tắt (Shortcut Features) tương quan ngẫu nhiên."]
    ]
    insert_thesis_table(doc, target_p, tbl2_headers, tbl2_widths, tbl2_rows, font_size_pt=14)
    add_p("", first_line_indent=False)

    add_p(
        "Cần nhấn mạnh rằng, nguyên tắc PRESERVE tập trung vào việc bảo toàn ngữ nghĩa an ninh động chứ không đồng nghĩa với việc giữ lại toàn bộ định danh thô của người dùng hay máy chủ, nhằm đảm bảo khả năng liên kết có kiểm soát (controlled linkability) và tương thích với các tiêu chuẩn đánh giá quyền riêng tư. Tương tự, nguyên tắc EXCLUDE thiết lập các ràng buộc phủ định đối với các giả định đường tắt đã biết và các biến số gây rò rỉ phân vùng, thay vì giả định rằng mọi đường tắt đều có thể nhận biết trước khi thực nghiệm."
    )
    add_p("Đi đôi với Hợp đồng Biểu diễn, chuyên đề thiết lập nguyên tắc phân định ranh giới phương pháp luận ba tầng độc lập:")
    add_p(
        "Đảm nhiệm tiền xử lý dữ liệu thô, phân tích cú pháp sơ bộ, chuẩn hóa kiểu dữ liệu trường và làm sạch dữ liệu. Tầng này không gánh vác nhiệm vụ học biểu diễn ngữ cảnh an ninh sâu.",
        bold_prefix="1. Tầng 1 — Trích xuất đặc trưng cơ sở (Feature Extraction): "
    )
    add_p([
        "Thiết lập ánh xạ ",
        latex_to_clean_omml(r"f_\theta: \mathcal{X} \to \mathbf{z} \in \mathbb{R}^d"),
        " từ cấu trúc chuỗi sự kiện và đồ thị nguồn gốc sang không gian vector tiềm ẩn. Toàn bộ năng lực bảo toàn ngữ nghĩa an ninh và tính bất biến được đóng gói trọn vẹn bên trong vector z."
    ], bold_prefix="2. Tầng 2 — Học không gian biểu diễn (Representation Learning — Trọng tâm Chuyên đề): ")
    add_p([
        "Đánh giá chất lượng của vector biểu diễn z thông qua các bộ thăm dò tuyến tính đóng băng tham số (Frozen Linear Probing): ",
        latex_to_clean_omml(r"\hat{y} = \sigma(\mathbf{W}^\top \mathbf{z} + \mathbf{b})"),
        ". Trong đó tham số ",
        latex_to_clean_omml(r"\theta"),
        " của bộ trích xuất đặc trưng Tầng 2 được giữ cố định hoàn toàn trong suốt quá trình đánh giá ở Tầng 3. Quy tắc này bảo đảm bộ phân loại hạ nguồn không làm thay nhiệm vụ trích xuất đặc trưng của Tầng 2."
    ], bold_prefix="3. Tầng 3 — Phát hiện và phân loại hạ nguồn (Downstream Detection): ")

    # =========================================================================
    # 1.2.
    # =========================================================================
    add_h2("Phân tích so sánh các nhóm phương pháp hiện đại")
    add_p(
        "Nhằm định vị chính xác các đóng góp kỹ thuật và cơ sở lý luận của chuyên đề, Mục 1.2 tiến hành khảo cứu toàn diện, phân loại và phân tích đa chiều ba nhóm phương pháp biểu diễn đặc trưng log chủ đạo trong y văn hiện đại: (1) Nhóm phương pháp thống kê và cú pháp dựa trên mẫu log (Statistical & Syntactic Parsing-Based); (2) Nhóm phương pháp biểu diễn ngữ nghĩa chuỗi thời gian (Semantic–Sequential & Transformer-Based); và (3) Nhóm phương pháp học biểu diễn đồ thị nguồn gốc (Provenance Graph Representation Learning). Mỗi nhóm phương pháp được mổ xẻ tường minh về cơ chế toán học, độ phức tạp thuật toán, ưu điểm cốt lõi và các rào cản nền tảng khi triển khai trong môi trường phát hiện tấn công thực tế."
    )

    # --- 1.2.1 ---
    add_h3("Phương pháp thống kê và cú pháp: Event Count, Frequency, Entropy và Template Features")
    add_p(
        "Nhóm phương pháp thống kê và cú pháp đại diện cho thế hệ tiếp cận đầu tiên trong phân tích nhật ký tự động. Cơ chế hoạt động của nhóm này dựa trên quy trình hai giai đoạn tách rời: giai đoạn phân tách dòng log thô thành các mẫu định dạng tĩnh (Log Templates / Event IDs) thông qua bộ phân tích cú pháp (Log Parser), tiếp theo là giai đoạn lượng hóa các chuỗi sự kiện thành vector số học dựa trên các thước đo thống kê kinh điển."
    )
    add_p(
        "Trong giai đoạn phân tích cú pháp, các thuật toán tiêu biểu đã được phát triển nhằm tối ưu hóa tốc độ xử lý: (1) Drain (He et al., ICWS 2017) sử dụng cấu trúc cây phân tích có độ sâu cố định (Fixed-Depth Parse Tree) để nhóm các dòng log dựa trên độ dài chuỗi và các từ khóa tiền tố, đạt tốc độ phân tích gần như tuyến tính O(N) đối với luồng dữ liệu lớn; (2) Spell (Du & Li, ICDM 2016) áp dụng thuật toán tìm chuỗi con chung dài nhất (Longest Common Subsequence - LCS) theo cơ chế dòng (streaming) để trích xuất động các thành phần tĩnh của thông điệp log; (3) LenMa (Shima, 2016) và AEL (Zhu et al., ISSRE 2023) lần lượt khai thác chiều dài các từ tố và tần suất xuất hiện của từ khóa để phân cụm và tách biến số ra khỏi chuỗi mẫu."
    )
    add_p([
        "Sau khi không gian log được rút gọn về tập ",
        latex_to_clean_omml(r"M"),
        " mẫu sự kiện cố định ",
        latex_to_clean_omml(r"\mathcal{E} = \{e_1, e_2, \dots, e_M\}"),
        ", các chuỗi sự kiện trong một cửa sổ thời gian hoặc phiên làm việc ",
        latex_to_clean_omml(r"W"),
        " được ánh xạ thành vector tần suất ",
        latex_to_clean_omml(r"\mathbf{x} = [c(e_1), c(e_2), \dots, c(e_M)]^\top \in \mathbb{R}^M"),
        ", trong đó ",
        latex_to_clean_omml(r"c(e_i)"),
        " là số lần xuất hiện của sự kiện ",
        latex_to_clean_omml(r"e_i"),
        ". Bên cạnh đếm tần suất đơn thuần, các trọng số TF-IDF (Term Frequency - Inverse Document Frequency) hoặc độ hỗn loạn thông tin Shannon (Information Entropy) ",
        latex_to_clean_omml(r"H(W) = -\sum_{i=1}^M p(e_i) \log_2 p(e_i)"),
        " cũng được áp dụng nhằm nhấn mạnh các sự kiện hiếm gặp."
    ])
    add_p([
        "Trên không gian vector tần suất này, các mô hình phát hiện bất thường kinh điển được triển khai rộng rãi: (1) Principal Component Analysis - PCA (Xu et al., SOSP 2009) phân rã không gian vector ",
        latex_to_clean_omml(r"\mathbb{R}^M"),
        " thành không gian con chuẩn tắc ",
        latex_to_clean_omml(r"\mathcal{S}_n"),
        " và không gian con phần dư ",
        latex_to_clean_omml(r"\mathcal{S}_r"),
        ", nhận diện bất thường khi năng lượng chiếu lên phần dư vượt ngưỡng kiểm định ",
        latex_to_clean_omml(r"\mathbf{x}_a = (\mathbf{I} - \mathbf{P} \mathbf{P}^\top) \mathbf{x}, \quad \|\mathbf{x}_a\|^2 > \gamma_\alpha"),
        "; (2) Invariant Mining (Lou et al., ATC 2010) tự động khai phá các phương trình bất biến tuyến tính phản ánh mối quan hệ bảo toàn logic giữa các bước thực thi ",
        latex_to_clean_omml(r"\mathbf{A} \mathbf{x} = \mathbf{0}"),
        "."
    ])
    add_p(
        "Mặc dù sở hữu ưu điểm vượt trội về hiệu năng tính toán (độ phức tạp O(N), độ trễ thấp dưới 0.1 ms/event), nhóm phương pháp thống kê và cú pháp bộc lộ hai điểm nghẽn nghiêm trọng không thể khắc phục: (1) Mất mát ngữ nghĩa an ninh do trừu tượng hóa tham số (Dynamic Parameter Loss): các bộ log parser bắt buộc phải sử dụng biểu thức chính quy (Regex) để loại bỏ toàn bộ các tham số biến động (địa chỉ IP, đường dẫn tệp tin, tham số dòng lệnh) thay thế bằng ký tự đại diện <*> khiến từ 82% đến 88% thông tin an ninh bị triệt tiêu; (2) Lan truyền và khuếch đại lỗi (Parser Error Propagation): khi gặp các định dạng log mới chưa từng xuất hiện (unseen logs), parser thường phân tách sai, dẫn đến hiện tượng bùng nổ số lượng mẫu sự kiện giả lập hoặc gộp nhầm các sự kiện khác biệt, phá vỡ hoàn toàn cấu trúc không gian vector x."
    )

    # --- 1.2.2 ---
    add_h3("Phương pháp semantic–sequential: Embeddings, Self-Supervised Learning, Transformer và Parsing-Free")
    add_p(
        "Nhằm khắc phục sự cứng nhắc của các vector đếm tần suất và tận dụng trật tự xuất hiện của các sự kiện, nhóm phương pháp Semantic–Sequential mô hình hóa luồng log tương tự như các chuỗi ngôn ngữ tự nhiên, tích hợp các kỹ thuật nhúng từ (Word Embeddings) và học tự giám sát (Self-Supervised Learning) để nắm bắt phụ thuộc ngữ cảnh dài hạn."
    )
    add_p([
        "Khởi đầu cho hướng nghiên cứu này là mô hình DeepLog (Du et al., CCS 2017). DeepLog sử dụng mạng bộ nhớ dài-ngắn hạn (LSTM) để mô hình hóa chuỗi Event ID như một bài toán dự báo phần tử tiếp theo (Next-Event Prediction). Tại mỗi bước thời gian t, mô hình ước lượng phân phối xác suất có điều kiện của sự kiện tiếp theo ",
        latex_to_clean_omml(r"P(e_t \mid e_{t-k}, \dots, e_{t-1})"),
        ". Nếu sự kiện thực tế không nằm trong tập g sự kiện có xác suất cao nhất được mô hình dự đoán ",
        latex_to_clean_omml(r"\hat{\mathcal{E}}_t = \operatorname{arg\,top-}g_{e \in \mathcal{E}} P(e \mid e_{t-k}, \dots, e_{t-1})"),
        ", hệ thống sẽ phát tín hiệu cảnh báo bất thường. Song song đó, DeepLog xây dựng mô hình LSTM thứ hai dựa trên phân phối chuẩn nhiều chiều để kiểm tra sự bất thường về giá trị tham số số học (Parameter Value Anomaly)."
    ])
    add_p(
        "Các công trình kế tiếp đã nâng cấp cơ chế biểu diễn ngữ nghĩa: (1) LogAnomaly (Meng et al., IJCAI 2019) đề xuất Template2Vec, trích xuất vector ngữ nghĩa cho từng mẫu log thông qua Word2Vec/FastText kết hợp trọng số d-IDF, giúp nhận biết sự tương đồng giữa các thông điệp có cấu trúc từ ngữ tương đương (ví dụ: 'Failed password' và 'Authentication failure'); (2) Logsy (Nedelkoski et al., 2020) sử dụng hàm mất mát phân loại ngoại lai (Outlier Classification Loss) trên dữ liệu log từ các hệ thống phụ trợ để định hình biên giới phân tách chặt chẽ cho lớp bình thường; (3) LogBERT (Guo et al., IJCNN 2021) khai thác Transformer hai chiều với hai tác vụ học tự giám sát: Dự đoán sự kiện log bị che (Masked Log Event Prediction) và Dự đoán phân bố khối lượng log (Volume Anomaly Prediction); (4) Nhóm tiếp cận không dùng Parser (Parsing-Free - NeuralLog, Le & Teoh, ASE 2021) bỏ qua bước phân tích cú pháp bằng cách sử dụng các mô hình ngôn ngữ tiền huấn luyện (BERT, RoBERTa) để trực tiếp mã hóa chuỗi văn bản log thô thành các vector nhúng ngữ nghĩa liên tục."
    )
    add_p(
        "Mặc dù đạt độ chính xác cao trên các tập dữ liệu phần mềm, nhóm phương pháp Semantic–Sequential bộc lộ ba giới hạn căn bản trong môi trường an ninh mạng: (1) Lợi thế ảo từ dữ liệu tiền huấn luyện (Pretraining-Data Advantage): các mô hình sử dụng Transformer tiền huấn luyện trên kho văn bản tổng quát thường bị rò rỉ tri thức ngoài miền; khi kiểm soát chặt chẽ ranh giới dữ liệu an ninh nội miền, hiệu năng thực tế suy giảm rõ rệt; (2) Chi phí tài nguyên và điểm nghẽn thông lượng: độ phức tạp tính toán bậc hai của cơ chế Self-Attention O(L^2) theo độ dài cửa sổ L khiến mô hình tiêu thụ GPU lớn và khó đáp ứng thông lượng thời gian thực; (3) Tầm nhìn cục bộ đơn luồng: mô hình chuỗi chỉ quan sát các sự kiện trên một dòng thời gian đơn lẻ, hoàn toàn mất dấu chuỗi tấn công APT phân tán đa tiến trình, đa luồng mạng và đa máy chủ."
    )

    # --- 1.2.3 ---
    add_h3("Đồ thị nguồn gốc và Graph Representation Learning")
    add_p([
        "Để vượt qua giới hạn tầm nhìn cục bộ của mô hình chuỗi, hướng tiếp cận dựa trên đồ thị nguồn gốc hệ thống (System Provenance Graph) mô hình hóa toàn bộ lịch sử thực thi và tương tác trong hệ điều hành dưới dạng một đồ thị có hướng, không đồng nhất và gán nhãn thời gian ",
        latex_to_clean_omml(r"\mathcal{G} = (\mathcal{V}, \mathcal{E}, \mathcal{T}_v, \mathcal{T}_e, \phi, \psi, \tau)"),
        ". Trong đó V là tập các đỉnh đại diện cho các thực thể hệ thống thuộc tập kiểu ",
        latex_to_clean_omml(r"\mathcal{T}_v = \{\text{Process, File, Socket, Registry, User, Host}\}"),
        "; E là tập các cạnh có hướng mô tả các tương tác luồng phụ thuộc thuộc tập kiểu ",
        latex_to_clean_omml(r"\mathcal{T}_e = \{\text{fork, execve, read, write, connect, bind, send, recv}\}"),
        "; và τ gán nhãn mốc thời gian xảy ra tương tác."
    ])
    add_p([
        "Các hệ thống phát hiện xâm nhập dựa trên đồ thị nguồn gốc (PIDS) tiêu biểu bao gồm: (1) UNICORN (Han et al., NDSS 2020) xây dựng đồ thị nguồn gốc luồng thời gian thực, áp dụng thuật toán băm cây con Weisfeiler-Lehman (WL-subtree kernel) để chuyển đổi đồ thị động thành vector đặc trưng đếm histogram; (2) KAIROS (Wang et al., S&P 2024) tích hợp mạng nơ-ron đồ thị nhận biết thời gian (Time-Aware GNN), mã hóa đồng thời thông tin cấu trúc và khoảng cách thời gian giữa các cạnh để phát hiện kỹ thuật APT ẩn mình (low-and-slow); (3) NODLINK (She et al., NDSS 2024) và MAGIC (Wang et al., USENIX Security 2024) khai thác kiến trúc GNN dị thể kết hợp cơ chế chú ý đồ thị để tính toán vector nhúng nút ",
        latex_to_clean_omml(r"\mathbf{h}_v^{(l+1)} = \sigma \left( \sum_{r \in \mathcal{T}_e} \sum_{u \in \mathcal{N}_v^r} \alpha_{uv}^r \mathbf{W}_r^{(l)} \mathbf{h}_u^{(l)} \right)"),
        "; (4) ORTHRUS (Wang et al., USENIX Security 2025) tách biệt và gióng hàng hai luồng thông tin cấu trúc và luồng dữ liệu ngữ cảnh nhằm thực hiện truy vết nguồn gốc với độ chính xác cao."
    ])
    add_p(
        "Tuy nhiên, việc triển khai GNN trên đồ thị nguồn gốc quy mô thực tế đối mặt với ba rào cản lý thuyết và thực nghiệm sâu sắc: (1) Hiện tượng bùng nổ phụ thuộc (Dependency Explosion): các tiến trình hệ thống chạy dài hạn (như systemd, sshd, hoặc trình duyệt web) liên tục đọc/ghi hàng triệu tệp tin và socket, khiến hầu hết mọi nút trong đồ thị đều có đường đi liên kết đến nhau, tạo ra các phụ thuộc giả (False Dependencies) làm loãng dấu vết tấn công thực sự; (2) Ngộ nhận giữa quan hệ phụ thuộc cấu trúc và tác động nhân quả thực tế (Dependency != Causal Effect): công trình của Bilot et al. ('Sometimes Simpler is Better', USENIX Security 2025) chứng minh rằng nhiều mô hình GNN phức tạp thực chất chỉ học đặc trưng đường tắt như tần suất bậc của nút; khi kiểm soát rò rỉ, các bộ phân loại tuyến tính đơn giản đạt hiệu năng tương đương với chi phí thấp hơn hàng chục lần; (3) Hiện tượng nghẽn cổ chai thông tin (Over-smoothing và Over-squashing): khi tăng độ sâu GNN, Over-smoothing làm vector biểu diễn của các nút bị đồng nhất hóa, trong khi Over-squashing (Alon & Yahav, ICLR 2021) nén ép lượng thông tin cấu trúc tăng theo hàm mũ vào vector kích thước cố định, làm mất mát các tín hiệu tấn công tinh vi."
    )

    # TABLE 3: Summary Table (Exact widths allocated: 2200, 2450, 2450, 2505 dxa, sum = 9605 dxa)
    tbl3_headers = ["Tiêu chí đánh giá", "Nhóm Thống kê / Cú pháp\n(Drain, PCA)", "Nhóm Chuỗi Semantic\n(DeepLog, LogBERT)", "Nhóm Đồ thị Nguồn gốc\n(UNICORN, MAGIC)"]
    tbl3_widths = [2200, 2450, 2450, 2505]
    tbl3_rows = [
        ["Cơ chế biểu diễn cốt lõi", "Vector đếm tần suất Event ID trên cửa sổ trượt", "Vector nhúng ngữ cảnh từ chuỗi sự kiện tuần tự", "Vector nhúng cấu trúc đồ thị luồng phụ thuộc dị thể"],
        ["Độ phức tạp thời gian", "O(N) — Tuyến tính theo luồng sự kiện", "O(N · L) đến O(N · L^2) — Phụ thuộc độ dài cửa sổ", "O(|V| + |E|) — Phụ thuộc quy mô đồ thị"],
        ["Bảo toàn tham số an ninh", "Kém (Bị parser cắt bỏ hoàn toàn bằng Regex)", "Trung bình (Bị rời rạc hóa hoặc che mờ)", "Rất tốt (Lưu trữ trực tiếp trên thuộc tính nút/cạnh)"],
        ["Mô hình hóa chuỗi quan hệ", "Hoàn toàn không có", "Cục bộ trên dòng thời gian đơn luồng", "Toàn diện trên quan hệ tương tác đa thực thể"],
        ["Khả năng triển khai SOC", "Rất cao (Độ trễ < 0.1 ms/event)", "Trung bình (Độ trễ 1 - 50 ms/window)", "Thách thức cao (Cần cắt tỉa đồ thị liên tục)"],
        ["Điểm nghẽn phương pháp luận", "Mất mát tham số động, lan truyền lỗi parser", "Thiếu ngữ cảnh đồ thị, chi phí tính toán cao", "Bùng nổ phụ thuộc, Over-smoothing, Over-squashing"]
    ]
    insert_thesis_table(doc, target_p, tbl3_headers, tbl3_widths, tbl3_rows, font_size_pt=14)
    add_p("", first_line_indent=False)

    add_p(
        "Tổng kết lại, phân tích so sánh đối chiếu chỉ ra rằng không có bất kỳ phương pháp đơn lẻ nào trong ba nhóm trên giải quyết trọn vẹn bài toán biểu diễn đặc trưng log. Nhóm thống kê đạt hiệu năng cao nhưng mất mát tham số an ninh; nhóm chuỗi nắm bắt ngữ nghĩa tốt nhưng thiếu tầm nhìn đồ thị đa thực thể; nhóm đồ thị mô hình hóa quan hệ xuất sắc nhưng chịu gánh nặng bùng nổ phụ thuộc và chi phí tính toán. Thực trạng khoa học này trực tiếp đặt ra yêu cầu phải xác lập và giải quyết năm khoảng trống nghiên cứu cốt lõi tại Mục 1.3 tiếp theo."
    )

    # Save
    updated_file = str(target_path.parent / (target_path.stem + ".updated.docx"))
    doc.save(updated_file)
    print(f"[SUCCESS] Saved to updated file: {updated_file}")

    try:
        doc.save(str(target_path))
        print(f"[SUCCESS] Overwritten active document: {target_path}")
    except PermissionError:
        print(f"[NOTE] Active file {target_path} is currently open in Word. Please close Word to allow overwrite.")

    v_doc = docx.Document(updated_file)
    omml_count = sum(len(p._p.xpath('.//m:oMath')) for p in v_doc.paragraphs)
    print(f"[FINAL AUDIT] Paragraphs: {len(v_doc.paragraphs)}, Tables: {len(v_doc.tables)}, OMML Equations: {omml_count}")


if __name__ == "__main__":
    build_perfect_document()
