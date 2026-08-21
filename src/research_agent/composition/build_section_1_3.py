"""
Module to write Section 1.3 (1.3.1 -> RQ1, 1.3.2 -> RQ2, 1.3.3 -> RQ3, 1.3.4 -> RQ4, 1.3.5 -> RQ5)
directly into D:\\Research\\Chuyên đề chuyên sâu - Copy.docx.
Preserves Section 1.1, Section 1.2, Title Page Table 0, TOC, Conclusion, and Bibliography.
"""

import sys
import shutil
from pathlib import Path
from lxml import etree
import latex2mathml.converter
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

sys.stdout.reconfigure(encoding="utf-8")

# Initialize official Microsoft Office MathML to OMML XSLT transformer
XSLT_PATH = r"C:\Program Files\Microsoft Office\Office16\MML2OMML.XSL"
xslt_tree = etree.parse(XSLT_PATH)
transform_omml = etree.XSLT(xslt_tree)


def latex_to_clean_omml(latex_code: str):
    """Converts a LaTeX formula into a native Word OMML element, cleaning any empty <m:e/> placeholders."""
    try:
        mathml = latex2mathml.converter.convert(latex_code)
        tree = etree.fromstring(mathml)
        omml_tree = transform_omml(tree)

        ns = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}
        for nary in omml_tree.xpath(".//m:nary", namespaces=ns):
            e_elem = nary.find("m:e", namespaces=ns)
            if e_elem is not None and len(e_elem) == 0 and (not e_elem.text or not e_elem.text.strip()):
                chr_val = "∑"
                naryPr = nary.find("m:naryPr", namespaces=ns)
                if naryPr is not None:
                    chr_elem = naryPr.find("m:chr", namespaces=ns)
                    if chr_elem is not None and "{http://schemas.openxmlformats.org/officeDocument/2006/math}val" in chr_elem.attrib:
                        chr_val = chr_elem.attrib["{http://schemas.openxmlformats.org/officeDocument/2006/math}val"]

                sub_elem = nary.find("m:sub", namespaces=ns)
                sup_elem = nary.find("m:sup", namespaces=ns)

                has_sub = sub_elem is not None and (len(sub_elem) > 0 or (sub_elem.text and sub_elem.text.strip()))
                has_sup = sup_elem is not None and (len(sup_elem) > 0 or (sup_elem.text and sup_elem.text.strip()))

                parent = nary.getparent()
                idx = parent.index(nary)

                if has_sub and has_sup:
                    new_node = etree.Element("{http://schemas.openxmlformats.org/officeDocument/2006/math}sSubSup")
                    e_node = etree.SubElement(new_node, "{http://schemas.openxmlformats.org/officeDocument/2006/math}e")
                    r_node = etree.SubElement(e_node, "{http://schemas.openxmlformats.org/officeDocument/2006/math}r")
                    t_node = etree.SubElement(r_node, "{http://schemas.openxmlformats.org/officeDocument/2006/math}t")
                    t_node.text = chr_val
                    new_node.append(sub_elem)
                    new_node.append(sup_elem)
                    parent.insert(idx, new_node)
                    parent.remove(nary)
                elif has_sub:
                    new_node = etree.Element("{http://schemas.openxmlformats.org/officeDocument/2006/math}sSub")
                    e_node = etree.SubElement(new_node, "{http://schemas.openxmlformats.org/officeDocument/2006/math}e")
                    r_node = etree.SubElement(e_node, "{http://schemas.openxmlformats.org/officeDocument/2006/math}r")
                    t_node = etree.SubElement(r_node, "{http://schemas.openxmlformats.org/officeDocument/2006/math}t")
                    t_node.text = chr_val
                    new_node.append(sub_elem)
                    parent.insert(idx, new_node)
                    parent.remove(nary)

        xml_str = etree.tostring(omml_tree, encoding="utf-8").decode("utf-8")
        return parse_xml(xml_str)
    except Exception as e:
        print(f"Warning: LaTeX conversion failed for '{latex_code}': {e}")
        xml_str = (
            f'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
            f'  <m:r><m:t>{latex_code}</m:t></m:r>'
            f'</m:oMath>'
        )
        return parse_xml(xml_str)


def append_section_1_3(target_file: str = r"D:\Research\Chuyên đề chuyên sâu - Copy.docx"):
    target_path = Path(target_file)
    backup_path = target_path.parent / (target_path.stem + ".backup.docx")

    doc = docx.Document(str(target_path))

    # Find the insertion point: right before 'Kết luận'
    target_p = None
    for idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip().lower()
        if "kết luận" in txt and idx > 30:
            target_p = p
            break

    if target_p is None:
        raise ValueError("Could not find insertion point before 'Kết luận'")

    print(f"[1/3] Located insertion point before 'Kết luận' at paragraph index.")

    def add_p(text_segments, bold_prefix=None, first_line_indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
        new_p = target_p.insert_paragraph_before(style="Normal")
        new_p.alignment = align
        new_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        new_p.paragraph_format.space_after = Pt(6)
        new_p.paragraph_format.space_before = Pt(0)
        if first_line_indent:
            new_p.paragraph_format.first_line_indent = Cm(1.27)

        if bold_prefix:
            r_pre = new_p.add_run(bold_prefix)
            r_pre.font.name = "Times New Roman"
            r_pre.font.size = Pt(14)
            r_pre.bold = True

        if isinstance(text_segments, str):
            r = new_p.add_run(text_segments)
            r.font.name = "Times New Roman"
            r.font.size = Pt(14)
        elif isinstance(text_segments, list):
            for seg in text_segments:
                if isinstance(seg, str):
                    r = new_p.add_run(seg)
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(14)
                else:
                    new_p._p.append(seg)
        return new_p

    def add_h2(clean_text):
        p = target_p.insert_paragraph_before(style="Heading 2")
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(clean_text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
        r.bold = True
        return p

    def add_h3(clean_text):
        p = target_p.insert_paragraph_before(style="Heading 3")
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(clean_text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
        r.bold = True
        r.italic = True
        return p

    print("[2/3] Writing Section 1.3 (1.3.1 -> RQ1, 1.3.2 -> RQ2, 1.3.3 -> RQ3, 1.3.4 -> RQ4, 1.3.5 -> RQ5)...")

    # =========================================================================
    # 1.3. CÁC KHOẢNG TRỐNG NGHIÊN CỨU CỐT LÕI
    # =========================================================================
    add_h2("Các khoảng trống nghiên cứu cốt lõi")
    add_p(
        "Từ kết quả khảo sát và phân tích đối chiếu ba nhóm phương pháp biểu diễn đặc trưng log tại Mục 1.2, có thể nhận thấy rằng mặc dù các kỹ thuật thống kê, mô hình chuỗi ngữ nghĩa và học biểu diễn đồ thị nguồn gốc đã đạt được nhiều bước tiến quan trọng, việc ứng dụng chúng vào môi trường phát hiện tấn công mạng doanh nghiệp thực tế vẫn đối mặt với những rào cản nền tảng chưa được giải quyết thấu đáo [2, 6, 9, 18]. Nhằm thiết lập cơ sở khoa học vững chắc và định hình phạm vi nghiên cứu, luận án tổng kết năm khoảng trống nghiên cứu cốt lõi (Research Gaps) tương ứng với năm câu hỏi nghiên cứu (Research Questions - RQ) định hướng cho toàn bộ các đề xuất phương pháp luận tiếp theo."
    )

    # -------------------------------------------------------------------------
    # 1.3.1. KHOẢNG TRỐNG 1 -> RQ1
    # -------------------------------------------------------------------------
    add_h3("Khoảng trống 1: Mất mát ngữ nghĩa an ninh trong quá trình trừu tượng hóa tham số")
    add_p(
        "Dữ liệu thực nghiệm từ các nghiên cứu tiền nhiệm chỉ ra rằng các bộ phân tích cú pháp (Log Parsers) đóng vai trò then chốt trong việc giảm chiều không gian văn bản log thành các mẫu định dạng tĩnh [6, 7]. Tuy nhiên, các kỹ thuật này áp dụng quy tắc biểu thức chính quy để loại bỏ toàn bộ các chuỗi biến động (như địa chỉ IP, tên tài khoản, đường dẫn tệp tin, cổng mạng, tham số dòng lệnh) và thay thế bằng các ký tự đại diện <*> nhằm tạo ra các mẫu log dùng chung."
    )
    add_p(
        "Giới hạn quan trọng chưa được giải quyết là: Sự tương đương về mặt cú pháp mẫu log (Template Equivalence) không đồng nghĩa với sự tương đương về mặt ngữ nghĩa an ninh (Security Semantic Equivalence) [8, 18]. Ví dụ, cùng một mẫu thông điệp xác thực thất bại nhưng xuất phát từ một tài khoản quản trị viên đặc quyền root trong mạng nội bộ mang mức độ nghiêm trọng an ninh hoàn toàn khác so với từ một tài khoản người dùng thông thường; hoặc hai dòng lệnh powershell.exe thực thi lệnh quản trị hợp lệ và tải mã độc mã hóa base64 đều bị ánh xạ về cùng một Event ID tiến trình. Khi bộ trích xuất đặc trưng triệt tiêu toàn bộ các tham số động này, không gian biểu diễn bị mất đi phần lớn các tín hiệu ngữ cảnh quan trọng để phân biệt hành vi độc hại với hành vi thông thường."
    )
    add_p(
        "Khoảng trống nghiên cứu này đặt ra thách thức: Làm thế nào để thiết lập một cơ chế trích xuất đặc trưng có khả năng lọc bỏ các nhiễu cú pháp vô hại nhưng vẫn bảo toàn nguyên vẹn các tham số động mang ngữ nghĩa an ninh trọng yếu trong không gian vector biểu diễn? Từ đó, câu hỏi nghiên cứu thứ nhất được xác lập:"
    )
    add_p(
        "Có thể xây dựng representation loại bỏ nhiễu cú pháp nhưng vẫn bảo toàn các dynamic parameters có ý nghĩa an toàn quan trọng hay không?",
        bold_prefix="• Câu hỏi nghiên cứu 1 (RQ1 — Representation Fidelity): "
    )

    # -------------------------------------------------------------------------
    # 1.3.2. KHOẢNG TRỐNG 2 -> RQ2
    # -------------------------------------------------------------------------
    add_h3("Khoảng trống 2: Bất đồng bộ và suy thoái trong gióng hàng biểu diễn đa góc nhìn")
    add_p([
        "Các công trình nghiên cứu hiện đại đã chỉ ra rằng dữ liệu log sở hữu tính đa góc nhìn nội tại: góc nhìn thống kê tần suất phản ánh mật độ hoạt động [6], góc nhìn chuỗi sự kiện phản ánh trật tự thời gian cục bộ [3, 4], và góc nhìn đồ thị nguồn gốc phản ánh luồng phụ thuộc cấu trúc đa thực thể ",
        latex_to_clean_omml(r"\mathcal{G}"),
        " [9, 10, 11]. Việc kết hợp các góc nhìn này được kỳ vọng sẽ cung cấp bức tranh toàn cảnh về chiến dịch tấn công."
    ])
    add_p(
        "Tuy nhiên, khi tiến hành hợp nhất hoặc gióng hàng (Cross-view Alignment) các không gian biểu diễn dị thể này trong không gian vector tiềm ẩn, các mô hình học biểu diễn thường gặp phải ba rào cản lý thuyết nghiêm trọng [22, 23]: (1) Hiện tượng sụp đổ chiều biểu diễn (Representation Collapse), trong đó các nhánh biểu diễn bị suy thoái thành các vector hằng số hoặc mất tính phân tách; (2) Hiện tượng chuyển giao tiêu cực (Negative Transfer), khi sự kết hợp giữa các góc nhìn làm suy giảm độ chính xác tổng thể do nhiễu từ một góc nhìn lấn át tín hiệu hữu ích của góc nhìn khác; và (3) Vấn đề góc nhìn bị khuyết hoặc tương ứng từng phần (Missing-view & Partial Correspondence), nảy sinh khi một số luồng telemetry bị chậm trễ, thất thoát gói tin hoặc không có liên kết 1-1 đồng thời giữa chuỗi sự kiện và đồ thị nguồn gốc."
    )
    add_p(
        "Khoảng trống nghiên cứu cốt lõi ở đây là: Thiếu một khung toán học chặt chẽ để gióng hàng đồng thời các góc nhìn dị thể trong không gian tiềm ẩn mà không gây sụp đổ biểu diễn, đồng thời vẫn bảo tồn được các thông tin đặc thù hữu ích của từng góc nhìn riêng biệt. Do đó, câu hỏi nghiên cứu thứ hai được hình thành:"
    )
    add_p(
        "Có thể căn chỉnh các view dị thể mà không gây representation collapse, negative transfer, đồng thời vẫn bảo toàn thông tin hữu ích đặc thù của từng view hay không?",
        bold_prefix="• Câu hỏi nghiên cứu 2 (RQ2 — Cross-View Alignment): "
    )

    # -------------------------------------------------------------------------
    # 1.3.3. KHOẢNG TRỐNG 3 -> RQ3
    # -------------------------------------------------------------------------
    add_h3("Khoảng trống 3: Rò rỉ thông tin quy trình, học đường tắt và trôi dạt biểu diễn")
    add_p(
        "Công trình tổng kết của Arp et al. (USENIX Security 2022) [2] và Liu et al. (IEEE S&P 2025) [18] đã cảnh báo sâu sắc về tình trạng sai lệch phương pháp luận trong nghiên cứu học máy áp dụng cho an ninh mạng. Trong bài toán trích xuất đặc trưng log, các mô hình rất dễ học phải các đặc trưng đường tắt giả định (Dataset Shortcuts) có tương quan ngẫu nhiên với nhãn độc hại trong tập huấn luyện nhưng không phản ánh bản chất tấn công thực tế."
    )
    add_p(
        "Phân tích hệ thống chỉ ra sáu kênh rò rỉ thông tin tiềm tàng (Leakage Pathways) thường xuất hiện trong quy trình thực nghiệm: (1) Rò rỉ từ vựng/bộ phân tích cú pháp (Parser/Vocabulary Leakage), khi bộ parser được học trên toàn bộ dữ liệu gồm cả tập kiểm thử; (2) Rò rỉ chuẩn hóa thống kê (Normalization/Statistics Leakage), khi các tham số trung bình/phương sai hoặc trọng số TF-IDF được tính toán xuyên qua ranh giới tập train và test; (3) Rò rỉ thực thể/máy chủ thử nghiệm (Host/Entity/Campaign Holdout Leakage), khi dữ liệu tấn công trên cùng một máy chủ xuất hiện ở cả hai tập; (4) Rò rỉ điều chỉnh ngưỡng và siêu tham số (Threshold/Hyperparameter Leakage); (5) Rò rỉ dữ liệu tiền huấn luyện ngoài miền (Pretraining Data Leakage); và (6) Rò rỉ thời gian tương lai (Future Temporal Information Leakage), khi mô hình vô tình sử dụng sự kiện ở thời điểm tương lai để trích xuất biểu diễn cho hiện tại."
    )
    add_p(
        "Khoảng trống nghiên cứu quan trọng là: Làm thế nào để phân định tường minh giữa tín hiệu an ninh hợp lệ (Legitimate Security Signal) và đặc trưng đường tắt nhân tạo (Dataset Shortcut Artifacts), đồng thời đảm bảo không gian biểu diễn duy trì được năng lực phân tách khi dữ liệu vận hành bị trôi dạt phân phối theo thời gian? Luận án xác lập câu hỏi nghiên cứu thứ ba:"
    )
    add_p(
        "Representation có còn hữu ích sau khi loại bỏ shortcut của dataset và khi phân phối dữ liệu thay đổi hay không?",
        bold_prefix="• Câu hỏi nghiên cứu 3 (RQ3 — Validity Without Shortcuts): "
    )

    # -------------------------------------------------------------------------
    # 1.3.4. KHOẢNG TRỐNG 4 -> RQ4
    # -------------------------------------------------------------------------
    add_h3("Khoảng trống 4: Gán nhãn mức thô, phân bổ bằng chứng yếu và nhiễu quản trị viên")
    add_p(
        "Trong thực tế giám sát an ninh mạng doanh nghiệp, dữ liệu nhãn mặt đất (Ground Truth) thường chỉ được ghi nhận ở mức độ hạt rất thô (Coarse Labels), chẳng hạn như gán nhãn một phiên làm việc kéo dài nhiều giờ hoặc toàn bộ một máy chủ bị xâm nhập trong một khoảng thời gian [18, 28, 29]. Bên trong phiên làm việc đó, phần lớn các dòng log vẫn là các thao tác hệ thống bình thường, và chỉ có một số lượng rất nhỏ sự kiện thực sự phản ánh kỹ thuật tấn công."
    )
    add_p(
        "Thách thức này càng trở nên phức tạp do sự xuất hiện liên tục của hành vi quản trị viên hệ thống (Admin-Noise / LOLBins) [9, 18]. Các kỹ sư quản trị thường xuyên thực thi các lệnh bảo trì, chẩn đoán mạng hoặc sao lưu dữ liệu bằng các công cụ dòng lệnh hợp lệ (như PowerShell, WMI, SSH, vssadmin). Các hành vi này sở hữu cấu trúc cú pháp và quyền hạn thực thi tương đồng với kỹ thuật của kẻ tấn công, nhưng hoàn toàn mang mục đích hợp pháp (Benign-but-risky). Một hành vi bất thường hoặc hiếm gặp không tự thân đồng nghĩa với hành vi tấn công độc hại (Unusual ≠ Malicious, Anomaly ≠ Attack)."
    )
    add_p(
        "Khoảng trống nghiên cứu ở đây là: Làm thế nào để phân bổ chính xác bằng chứng tấn công yếu (Weak Evidence Attribution) từ các nhãn mức thô mà không học nhầm các hành vi quản trị hợp pháp thành độc hại? Đây là cơ sở để xác lập câu hỏi nghiên cứu thứ tư:"
    )
    add_p(
        "Có thể gán đúng attack evidence dưới coarse labels mà không học nhầm các hành vi quản trị hợp pháp thành malicious hay không?",
        bold_prefix="• Câu hỏi nghiên cứu 4 (RQ4 — Weak Evidence Attribution): "
    )

    # -------------------------------------------------------------------------
    # 1.3.5. KHOẢNG TRỐNG 5 -> RQ5
    # -------------------------------------------------------------------------
    add_h3("Khoảng trống 5: Đánh đổi giữa bảo toàn liên kết an ninh và rủi ro quyền riêng tư")
    add_p(
        "Để phát hiện và điều tra các cuộc tấn công APT kéo dài, mô hình học biểu diễn đòi hỏi phải duy trì khả năng liên kết có kiểm soát (Controlled Linkability) giữa các chuỗi hành vi của cùng một thực thể (người dùng, máy chủ, tiến trình) qua nhiều mốc thời gian khác nhau [9, 10]."
    )
    add_p(
        "Tuy nhiên, việc lưu giữ khả năng liên kết này trong không gian vector biểu diễn trực tiếp làm nảy sinh các nguy cơ nghiêm trọng về quyền riêng tư và an toàn thông tin [25, 26, 27]. Các vector biểu diễn đặc trưng tiềm ẩn có nguy cơ bị kẻ tấn công khai thác thông qua các kỹ thuật tấn công suy luận thành viên (Membership Inference Attacks - MIA [25]) để xác định xem dữ liệu của một thực thể có nằm trong tập huấn luyện hay không, hoặc tấn công nghịch đảo biểu diễn (Representation / Model Inversion Attacks [26]) nhằm khôi phục lại các định danh nhạy cảm của người dùng và cấu hình mạng nội bộ. Cần nhấn mạnh rằng, một mô hình được thiết kế có nhận thức về quyền riêng tư (Privacy-Aware) không tự động đồng nghĩa với việc đã đạt được khả năng bảo vệ quyền riêng tư vững chắc (Privacy-Preserving) nếu chưa trải qua các kiểm thử thực nghiệm tấn công nghiêm ngặt [27]."
    )
    add_p(
        "Khoảng trống nghiên cứu này đòi hỏi phải phân tích tường minh mối quan hệ đánh đổi: Đâu là điểm cân bằng giữa năng lực duy trì liên kết thực thể phục vụ phân tích an ninh và mức độ rò rỉ thông tin riêng tư của không gian vector đặc trưng? Từ đó, câu hỏi nghiên cứu thứ năm được xác lập:"
    )
    add_p(
        "Đâu là sự cân bằng chấp nhận được giữa entity continuity và privacy leakage để representation vẫn hữu ích cho phân tích an toàn?",
        bold_prefix="• Câu hỏi nghiên cứu 5 (RQ5 — Privacy-Security Trade-Off): "
    )

    # Save
    updated_file = str(target_path.parent / (target_path.stem + ".updated.docx"))
    doc.save(updated_file)
    print(f"[3/3] Saved updated file with Section 1.3 to: {updated_file}")

    try:
        doc.save(str(target_path))
        print(f"[SUCCESS] Overwritten active document: {target_path}")
    except PermissionError:
        print(f"[NOTE] Active file {target_path} is currently open in Word. Please close Word to allow overwrite.")

    v_doc = docx.Document(updated_file)
    omml_count = sum(len(p._p.xpath('.//m:oMath')) for p in v_doc.paragraphs)
    print(f"[FINAL AUDIT] Total Paragraphs: {len(v_doc.paragraphs)}, Total Tables: {len(v_doc.tables)}, OMML Equations: {omml_count}")


if __name__ == "__main__":
    append_section_1_3()
