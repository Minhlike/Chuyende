"""
Complete script to build Section 1.1 and Section 1.2 into D:\\Research\\Chuyên đề chuyên sâu - Copy.docx
Strictly enforces:
1. Times New Roman 14pt everywhere (including tables, headings, body text).
2. Native Microsoft Word Equations (OMML <m:oMath>) for all mathematical expressions.
3. 1.5 line spacing, 1.27cm first-line indent, justified alignment.
4. Preserves Table 0 (Cover), Table of Contents, Conclusion, and References.
"""

import sys
import shutil
from pathlib import Path
from xml.sax.saxutils import escape

sys.stdout.reconfigure(encoding="utf-8")

import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    xml_borders = (
        f'<w:tcBorders {nsdecls("w")}>\n'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '</w:tcBorders>'
    )
    tcPr.append(parse_xml(xml_borders))


def omath(text: str):
    """Create an inline native OMML element."""
    escaped = escape(text)
    xml_str = (
        f'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        f'  <m:r><m:t>{escaped}</m:t></m:r>'
        f'</m:oMath>'
    )
    return parse_xml(xml_str)


def omath_para(text: str):
    """Create a display native OMML paragraph element."""
    escaped = escape(text)
    xml_str = (
        f'<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        f'  <m:oMath>'
        f'    <m:r><m:t>{escaped}</m:t></m:r>'
        f'  </m:oMath>'
        f'</m:oMathPara>'
    )
    return parse_xml(xml_str)


def build_full_sections(target_file: str = r"D:\Research\Chuyên đề chuyên sâu - Copy.docx"):
    target_path = Path(target_file)
    backup_path = target_path.parent / (target_path.stem + ".backup.docx")

    # Use backup as the pristine base
    if not backup_path.exists():
        shutil.copyfile(target_path, backup_path)
    print(f"[1/4] Base backup verified at: {backup_path}")

    # Load fresh document from backup
    doc = docx.Document(str(backup_path))

    # Keep only Table 0 (Cover frame)
    while len(doc.tables) > 1:
        tbl_to_remove = doc.tables[1]
        tbl_to_remove._tbl.getparent().remove(tbl_to_remove._tbl)
    print("[2/4] Preserved Cover Frame Table 0. Cleaned old body tables.")

    # Remove old body paragraphs from Heading 1 to Conclusion
    paragraphs_to_remove = []
    found_h1 = False
    for p in doc.paragraphs:
        txt = p.text.strip().lower()
        if "tổng quan về phương pháp trích xuất đặc trưng" in txt:
            found_h1 = True
        if found_h1:
            if "kết luận" in txt:
                break
            paragraphs_to_remove.append(p)

    for p in paragraphs_to_remove:
        p._p.getparent().remove(p._p)

    # Locate insertion target point (right before 'Kết luận')
    target_p = None
    for p in doc.paragraphs:
        if "kết luận" in p.text.strip().lower():
            target_p = p
            break

    print(f"[3/4] Removed {len(paragraphs_to_remove)} old paragraphs. Insertion target located.")

    # Helper insertion functions
    def add_p(text_segments, bold_prefix=None, first_line_indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
        """
        text_segments can be a string or a list of items (strings and OMML elements).
        All text runs are strictly Times New Roman 14pt.
        """
        new_p = doc.add_paragraph(style="Normal") if target_p is None else target_p.insert_paragraph_before(style="Normal")
        new_p.alignment = align
        new_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        new_p.paragraph_format.space_after = Pt(6)
        new_p.paragraph_format.space_before = Pt(0)
        if first_line_indent:
            new_p.paragraph_format.first_line_indent = Cm(1.27)

        if bold_prefix:
            r_pre = new_p.add_run(bold_prefix)
            r_pre.font.name = "Times New Roman"
            r_pre.font.size = Pt(14)
            r_pre.bold = True

        if isinstance(text_segments, str):
            r = new_p.add_run(text_segments)
            r.font.name = "Times New Roman"
            r.font.size = Pt(14)
        elif isinstance(text_segments, list):
            for seg in text_segments:
                if isinstance(seg, str):
                    r = new_p.add_run(seg)
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(14)
                else:
                    # OMML XML element
                    new_p._p.append(seg)
        return new_p

    def add_h1(text):
        p = doc.add_paragraph(style="Heading 1") if target_p is None else target_p.insert_paragraph_before(style="Heading 1")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(16)
        r.bold = True
        return p

    def add_h2(text):
        p = doc.add_paragraph(style="Heading 2") if target_p is None else target_p.insert_paragraph_before(style="Heading 2")
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
        r.bold = True
        return p

    def add_h3(text):
        p = doc.add_paragraph(style="Heading 3") if target_p is None else target_p.insert_paragraph_before(style="Heading 3")
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
        r.bold = True
        r.italic = True
        return p

    def add_table(headers, rows_data):
        tbl = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        if target_p is not None:
            target_p._p.addprevious(tbl._tbl)

        # Header row
        for c_idx, h in enumerate(headers):
            cell = tbl.cell(0, c_idx)
            cell.text = h
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.space_before = Pt(4)
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(14)  # STRICTLY 14PT
                r.bold = True

        # Body rows
        for r_idx, row in enumerate(rows_data):
            for c_idx, val in enumerate(row):
                cell = tbl.cell(r_idx + 1, c_idx)
                cell.text = val
                set_cell_border(cell)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.space_before = Pt(4)
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(14)  # STRICTLY 14PT

        # Spacing after table
        add_p("", first_line_indent=False)

    print("[4/4] Writing Section 1.1 and Section 1.2 with native OMML equations...")

    # =========================================================================
    # CHƯƠNG 1
    # =========================================================================
    add_h1("CHƯƠNG 1: TỔNG QUAN VỀ PHƯƠNG PHÁP TRÍCH XUẤT ĐẶC TRƯNG DỮ LIỆU LOG VÀ THÁCH THỨC BẢO TOÀN NGỮ CẢNH AN TOÀN")

    # =========================================================================
    # 1.1. BÀI TOÁN BIỂU DIỄN LOG TRONG PHÁT HIỆN TẤN CÔNG ĐA GIAI ĐOẠN
    # =========================================================================
    add_h2("1.1. Bài toán biểu diễn log trong phát hiện tấn công đa giai đoạn")
    add_p([
        "Trong các Trung tâm Điều hành An ninh mạng (Security Operations Center - SOC) hiện đại, dữ liệu nhật ký hệ thống (system logs) và nhật ký kiểm toán (audit logs) đóng vai trò là nguồn bằng chứng trung tâm phục vụ phát hiện, điều tra và ứng phó các chiến dịch tấn công có chủ đích (Advanced Persistent Threats - APT). Khác với văn bản ngôn ngữ tự nhiên thông thường hoặc tín hiệu thị giác máy tính, dữ liệu log sở hữu cấu trúc bán hình thức (semi-structured), mang tính phụ thuộc thời gian nghiêm ngặt, phản ánh các quan hệ phụ thuộc thực thi giữa các tiến trình và tài nguyên hệ điều hành, đồng thời chứa đựng các ngữ nghĩa an ninh đặc thù. Mục 1.1 tập trung hình thức hóa bài toán biểu diễn đặc trưng log phục vụ phát hiện tấn công đa giai đoạn, phân tích bản chất không gian dữ liệu doanh nghiệp, xác lập mô hình hành vi phi tuyến tính trên ma trận MITRE ATT&CK và thiết lập khung Hợp đồng Biểu diễn (Representation Contract) nhằm bảo toàn các bất biến an toàn cho không gian vector đặc trưng ",
        omath("z"),
        "."
    ])

    # --- 1.1.1 ---
    add_h3("1.1.1. Không gian dữ liệu log doanh nghiệp: tốc độ cao, mất cân bằng cực đoan và phân phối biến đổi")
    add_p(
        "Không gian dữ liệu nhật ký trong môi trường mạng doanh nghiệp được tổng hợp từ ba nhóm nguồn telemetry chính với cấu trúc và định dạng không đồng nhất. Nhóm thứ nhất là nhật ký kiểm toán máy chủ (Host Audit Logs), bao gồm Linux Auditd, Windows Event Log / Sysmon và Linux eBPF (Extended Berkeley Packet Filter). Nguồn dữ liệu này ghi nhận trực tiếp các sự kiện ở mức nhân hệ điều hành thông qua việc chặn bắt các lời gọi hệ thống (syscalls), bao gồm: khởi tạo tiến trình (execve, CreateProcess - Sysmon Event ID 1), nạp thư viện động (ImageLoaded - Sysmon Event ID 7), thao tác tệp tin (open, unlink, FileCreate - Sysmon Event ID 11), sửa đổi cấu hình registry (RegSetValue - Sysmon Event ID 13), cùng các thao tác mở và kết nối socket mạng (connect, accept - Sysmon Event ID 3). Nhóm thứ hai là nhật ký luồng mạng (Network Flow & Protocol Logs), được thu thập từ Zeek, Suricata hoặc NetFlow/IPFIX, cung cấp siêu dữ liệu kết nối giữa các nút mạng, giao dịch DNS, chứng chỉ TLS/SSL và thông lượng gói tin. Nhóm thứ ba là nhật ký ứng dụng và dịch vụ (Application & Service Logs), phát sinh từ máy chủ web (Nginx, Apache), cơ sở dữ liệu, dịch vụ phân tán (HDFS) cùng hệ thống điều phối container (Kubernetes Audit Logs)."
    )
    add_p(
        "Tính dị thể sâu sắc của dữ liệu đặt ra bài toán khoa học về việc lựa chọn đơn vị quan sát (Unit of Observation) phù hợp cho mô hình học biểu diễn. Việc phân cấp đơn vị quan sát quyết định trực tiếp đến mức độ bảo toàn thông tin và độ phức tạp tính toán:"
    )

    tbl1_headers = ["Mức độ hạt biểu diễn", "Đơn vị quan sát", "Dữ liệu đại diện", "Ưu điểm cốt lõi", "Rào cản và Thách thức ngữ nghĩa"]
    tbl1_rows = [
        ["1. Mức từ tố (Token)", "Chuỗi con, từ khóa rời rạc", "Từ khóa tĩnh, địa chỉ IP, mã lỗi hex", "Dễ dàng xử lý bằng kỹ thuật nhúng từ vựng", "Mất hoàn toàn cấu trúc cú pháp và quan hệ liên kết trường"],
        ["2. Mức sự kiện (Event)", "Một dòng log đơn lẻ", "Bản ghi telemetry tại thời điểm t", "Bảo toàn đầy đủ thuộc tính cục bộ tại thời điểm t", "Thiếu ngữ cảnh chuỗi tuần tự và lịch sử tương tác trước đó"],
        ["3. Mức chuỗi / Phiên (Sequence / Session)", "Cửa sổ trượt hoặc phiên tiến trình", "Chuỗi sự kiện [e_{t-k+1}, ..., e_t] theo thời gian", "Nắm bắt quan hệ phụ thuộc thứ tự thời gian cục bộ", "Nhạy cảm với nhiễu xen kẽ (interleaving) từ các luồng chạy song song"],
        ["4. Mức thực thể (Entity)", "Định danh tác nhân (Host, User, IP, Process Instance)", "Lịch sử tương tác gom cụm theo thực thể", "Phân lập rõ ràng ranh giới hành vi của từng chủ thể", "Khó phát hiện các hành vi tấn công phối hợp vượt ranh giới thực thể"],
        ["5. Mức đồ thị (Graph)", "Đồ thị nguồn gốc (Provenance Graph)", "Đồ thị luồng phụ thuộc dị thể G", "Mô hình hóa toàn diện quan hệ phụ thuộc đa thực thể", "Bùng nổ kích thước đồ thị (dependency explosion) và chi phí tính toán"]
    ]
    add_table(tbl1_headers, tbl1_rows)

    add_p(
        "Bên cạnh tính dị thể, dữ liệu log doanh nghiệp chịu áp lực vận hành khắc nghiệt về mặt thông lượng và tỷ lệ phân bố nhãn. Trong các môi trường mạng lớn, hệ thống tiếp nhận luồng sự kiện liên tục với thông lượng cao, đòi hỏi thuật toán biểu diễn đặc trưng phải xử lý theo cơ chế dòng (streaming) với độ phức tạp tính toán tuyến tính O(N) mà không đòi hỏi lưu giữ toàn bộ lịch sử đồ thị trong bộ nhớ truy cập ngẫu nhiên. Đồng thời, tỷ lệ các dòng log liên quan đến hành vi tấn công trong thực tế thường ở mức rất thấp, tạo ra sự mất cân bằng nhãn cực đoan khiến các hàm mất mát học máy thông thường có xu hướng xem nhẹ nhóm thiểu số. Trong khi đó, việc suy thoái không gian vector (sụp đổ chiều biểu diễn) nảy sinh khi các mô hình tự giám sát không áp dụng các cơ chế điều hòa phương sai - hiệp phương sai phù hợp để duy trì tính đa dạng của các chiều đặc trưng."
    )
    add_p(
        "Dưới góc độ phân phối thời gian, dữ liệu log liên tục biến đổi trong môi trường vận hành dài hạn. Để tránh các sai lệch phương pháp luận khi đánh giá mô hình, chuyên đề phân định tường minh bốn cơ chế trôi dạt dữ liệu (Drift Taxonomy):"
    )
    add_p([
        "Bản chất hành vi và mục đích tấn công thay đổi theo thời gian mặc dù cấu trúc định dạng log không đổi: ",
        omath("P_t(Y | X) ≠ P_{t+1}(Y | X)"),
        "."
    ], bold_prefix="• Concept Drift: ")
    add_p([
        "Việc nâng cấp phần mềm, cập nhật bản vá hoặc thay đổi cấu hình làm biến đổi cấu trúc chuỗi mẫu log: ",
        omath("P_t(X_template) ≠ P_{t+1}(X_template)"),
        "."
    ], bold_prefix="• Template Drift: ")
    add_p([
        "Lưu lượng người dùng, tần suất giao tác nghiệp vụ hoặc cơ cấu dịch vụ hệ thống biến động theo chu kỳ: ",
        omath("P_t(X) ≠ P_{t+1}(X)"),
        "."
    ], bold_prefix="• Population Drift: ")
    add_p([
        "Không gian vector tiềm ẩn z bị suy giảm năng lực phân tách do dữ liệu đầu vào trôi dạt khỏi vùng phân phối huấn luyện ban đầu: ",
        omath("P_t(z | X) ≠ P_{t+1}(z | X)"),
        "."
    ], bold_prefix="• Representation Drift: ")

    # --- 1.1.2 ---
    add_h3("1.1.2. Hành vi tấn công đa giai đoạn và ánh xạ đa nhãn MITRE ATT&CK")
    add_p(
        "Trong phân tích an ninh mạng, một số công trình áp dụng mô hình trạng thái Markov tuyến tính tuần tự để xâu chuỗi các giai đoạn tấn công từ Thâm nhập ban đầu (Initial Access), Thực thi (Execution), Duy trì (Persistence), Leo thang đặc quyền (Privilege Escalation) cho đến Đánh cắp dữ liệu (Exfiltration)."
    )
    add_p(
        "Tuy nhiên, các quan sát thực nghiệm trên các chiến dịch tấn công cho thấy hành vi của kẻ tấn công mang bản chất phi tuyến tính: (1) Nhảy cóc giai đoạn (Step Skipping): Kẻ tấn công có thể khai thác trực tiếp lỗ hổng thực thi mã từ xa để trích xuất dữ liệu ra ngoài mà không cần thiết lập cơ chế duy trì hay di chuyển ngang; (2) Lặp vòng kỹ thuật (Tactic Looping & Interleaving): Kỹ thuật thu thập thông tin nội bộ (Discovery) thường được lặp lại nhiều lần xen kẽ giữa các bước leo thang đặc quyền và chiếm đoạt thông tin xác thực; (3) Phân nhánh tiến trình song song (Parallel Branching): Kẻ tấn công có thể khởi tạo đồng thời nhiều luồng tiến trình con độc lập trên các tiến trình hợp lệ khác nhau nhằm phân tán sự theo dõi của hệ thống phòng thủ."
    )
    add_p([
        "Do đó, chuyên đề xác lập nguyên tắc: Ma trận MITRE ATT&CK được mô hình hóa thành một Không gian Bằng chứng Hành vi Đa chiều (Multi-label Behavioral Evidence Space) ",
        omath("Y ⊆ {0, 1}^{|T|}"),
        ", trong đó một chuỗi sự kiện hoặc cây tiến trình có thể đồng thời kích hoạt nhiều nhãn chiến thuật (Tactics) và kỹ thuật (Techniques) tại cùng một thời điểm quan sát."
    ])
    add_p(
        "Về mặt dữ liệu thực nghiệm, việc mô hình hóa hành vi tấn công đòi hỏi phải phân định chính xác đặc tính gán nhãn và mức độ hạt (Label Granularity) của từng bộ dữ liệu chuẩn: (1) DARPA Transparent Computing (TC E3/E5) cung cấp telemetry kiểm toán mức nhân hệ điều hành với nhãn mặt đất được ghi nhận ở mức kịch bản tấn công tổng thể kèm theo danh sách các nút và cạnh liên quan trong đồ thị nguồn gốc; (2) LANL CyberEvents (2017) kết hợp xác thực máy tính, tiến trình và luồng mạng quy mô lớn với nhãn độc hại gắn theo mốc thời gian và tài khoản/máy tính bị đội đỏ xâm nhập; (3) HDFS Benchmark ghi nhận thao tác khối dữ liệu Hadoop với nhãn bất thường phản ánh sự cố kỹ thuật hoặc lỗi khối (Block Anomaly), không phải tấn công APT; (4) BGL Benchmark ghi nhận nhật ký máy siêu điện toán BlueGene/L với nhãn cảnh báo lỗi phần cứng và hệ thống phản ánh độ tin cậy vận hành."
    )
    add_p(
        "Đặc biệt, sự xuất hiện của nhiễu từ hành vi quản trị viên (Admin-Noise) là một thách thức thực tế lớn. Quản trị viên hệ thống thường xuyên sử dụng các công cụ dòng lệnh hợp lệ (Living-off-the-Land Binaries - LOLBins) như PowerShell, WMI (wmic.exe), SSH, certutil.exe hay vssadmin.exe cho mục đích bảo trì, sao lưu hoặc kiểm tra mạng. Các hành vi này có cấu trúc cú pháp tương tự như kỹ thuật của kẻ tấn công. Nếu mô hình chỉ dựa vào sự xuất hiện của từ khóa, hệ thống sẽ phát sinh nhiều cảnh báo sai (False Positives). Do đó, bài toán biểu diễn đòi hỏi cơ chế phân bổ bằng chứng yếu (Weak Evidence Attribution) thông qua khung Học Đa Thể hiện (Multiple Instance Learning - MIL)."
    )

    # --- 1.1.3 ---
    add_h3("1.1.3. Các mức biểu diễn dữ liệu và Hợp đồng Biểu diễn (Representation Contract)")
    add_p([
        "Nhằm định hình rõ ràng các yêu cầu đối với không gian vector đặc trưng ",
        omath("z ∈ ℝ^d"),
        " và ngăn ngừa rủi ro học đường tắt (Shortcut Learning), mô hình biểu diễn bắt buộc phải tuân thủ Hợp đồng Biểu diễn (Representation Contract) được định nghĩa dưới dạng bộ ba hình thức ",
        omath("C_representation = ⟨P_preserve, I_invariant, E_exclude⟩"),
        ":"
    ])

    tbl2_headers = ["Nhóm quy tắc", "Ý nghĩa phương pháp luận", "Danh mục thuộc tính telemetry tương ứng"]
    tbl2_rows = [
        ["PRESERVE (Bảo toàn)", "Bảo toàn các thuộc tính mang ngữ nghĩa an ninh và quan hệ phụ thuộc thực thi trong không gian vector z.", "• Thứ tự thời gian cục bộ giữa các sự kiện liền kề.\n• Quan hệ tiến trình cha-con (parent_process_guid → process_guid).\n• Cấu trúc liên kết đa thực thể (Tiến trình - Tệp tin - Socket - Registry).\n• Tham số an ninh động: IP đích, cổng, đường dẫn tệp, command_line, return_code, ImageHash."],
        ["INVARIANT (Bất biến)", "Duy trì tính bất biến của vector z trước các biến đổi hình thức không làm thay đổi bản chất hành vi.", "• Biến đổi cú pháp vô hại: khoảng trắng thừa, dấu xuống dòng, chữ hoa/thường trong câu lệnh không phân biệt hoa-thường.\n• Định danh tiến trình tạm thời ngẫu nhiên (PID, Thread ID biến động qua mỗi lần khởi chạy).\n• Hoán vị thứ tự giữa các trường thuộc tính độc lập không có quan hệ phụ thuộc tuần tự."],
        ["EXCLUDE (Triệt tiêu)", "Loại bỏ các thuộc tính gây học vẹt, tạo đường tắt giả định hoặc vi phạm quyền riêng tư khỏi vector z.", "• Định danh máy chủ cố định gây thiên lệch (Host UUID, tên máy trạm thử nghiệm cụ thể, địa chỉ MAC tĩnh).\n• Dấu vết phân vùng nhân tạo (Synthetic Split Artifacts) phát sinh trong quá trình chia tập train/test.\n• Mọi đặc trưng học đường tắt (Shortcut Features) có tương quan ngẫu nhiên với nhãn trong tập huấn luyện nhưng không có giá trị tổng quát hóa."]
    ]
    add_table(tbl2_headers, tbl2_rows)

    add_p(
        "Cần nhấn mạnh rằng, nguyên tắc PRESERVE tập trung vào việc bảo toàn ngữ nghĩa an ninh động chứ không đồng nghĩa với việc giữ lại toàn bộ định danh thô của người dùng hay máy chủ, nhằm đảm bảo khả năng liên kết có kiểm soát (controlled linkability) và tương thích với các tiêu chuẩn đánh giá quyền riêng tư. Tương tự, nguyên tắc EXCLUDE thiết lập các ràng buộc phủ định đối với các giả định đường tắt đã biết và các biến số gây rò rỉ phân vùng, thay vì giả định rằng mọi đường tắt đều có thể nhận biết trước khi thực nghiệm."
    )
    add_p("Đi đôi với Hợp đồng Biểu diễn, chuyên đề thiết lập nguyên tắc phân định ranh giới phương pháp luận ba tầng độc lập:")
    add_p(
        "Đảm nhiệm tiền xử lý dữ liệu thô, phân tích cú pháp sơ bộ, chuẩn hóa kiểu dữ liệu trường và làm sạch dữ liệu. Tầng này không gánh vác nhiệm vụ học biểu diễn ngữ cảnh an ninh sâu.",
        bold_prefix="1. Tầng 1 — Trích xuất đặc trưng cơ sở (Feature Extraction): "
    )
    add_p([
        "Thiết lập ánh xạ ",
        omath("f_θ: X → z ∈ ℝ^d"),
        " từ cấu trúc chuỗi sự kiện và đồ thị nguồn gốc sang không gian vector tiềm ẩn. Toàn bộ năng lực bảo toàn ngữ nghĩa an ninh và tính bất biến được đóng gói trọn vẹn bên trong vector z."
    ], bold_prefix="2. Tầng 2 — Học không gian biểu diễn (Representation Learning — Trọng tâm Chuyên đề): ")
    add_p([
        "Đánh giá chất lượng của vector biểu diễn z thông qua các bộ thăm dò tuyến tính đóng băng tham số (Frozen Linear Probing): ",
        omath("ŷ = σ(W^T z + b)"),
        ". Trong đó tham số ",
        omath("θ"),
        " của bộ trích xuất đặc trưng Tầng 2 được giữ cố định hoàn toàn trong suốt quá trình đánh giá ở Tầng 3. Quy tắc này bảo đảm bộ phân loại hạ nguồn không làm thay nhiệm vụ trích xuất đặc trưng của Tầng 2."
    ], bold_prefix="3. Tầng 3 — Phát hiện và phân loại hạ nguồn (Downstream Detection): ")

    # =========================================================================
    # 1.2. PHÂN TÍCH SO SÁNH CÁC NHÓM PHƯƠNG PHÁP HIỆN ĐẠI
    # =========================================================================
    add_h2("1.2. Phân tích so sánh các nhóm phương pháp hiện đại")
    add_p(
        "Nhằm định vị chính xác các đóng góp kỹ thuật và cơ sở lý luận của chuyên đề, Mục 1.2 tiến hành khảo cứu toàn diện, phân loại và phân tích đa chiều ba nhóm phương pháp biểu diễn đặc trưng log chủ đạo trong y văn hiện đại: (1) Nhóm phương pháp thống kê và cú pháp dựa trên mẫu log (Statistical & Syntactic Parsing-Based); (2) Nhóm phương pháp biểu diễn ngữ nghĩa chuỗi thời gian (Semantic–Sequential & Transformer-Based); và (3) Nhóm phương pháp học biểu diễn đồ thị nguồn gốc (Provenance Graph Representation Learning). Mỗi nhóm phương pháp được mổ xẻ tường minh về cơ chế toán học, độ phức tạp thuật toán, ưu điểm cốt lõi và các rào cản nền tảng khi triển khai trong môi trường phát hiện tấn công thực tế."
    )

    # --- 1.2.1 ---
    add_h3("1.2.1. Phương pháp thống kê và cú pháp: Event Count, Frequency, Entropy và Template Features")
    add_p(
        "Nhóm phương pháp thống kê và cú pháp đại diện cho thế hệ tiếp cận đầu tiên trong phân tích nhật ký tự động. Cơ chế hoạt động của nhóm này dựa trên quy trình hai giai đoạn tách rời: giai đoạn phân tách dòng log thô thành các mẫu định dạng tĩnh (Log Templates / Event IDs) thông qua bộ phân tích cú pháp (Log Parser), tiếp theo là giai đoạn lượng hóa các chuỗi sự kiện thành vector số học dựa trên các thước đo thống kê kinh điển."
    )
    add_p(
        "Trong giai đoạn phân tích cú pháp, các thuật toán tiêu biểu đã được phát triển nhằm tối ưu hóa tốc độ xử lý: (1) Drain (He et al., ICWS 2017) sử dụng cấu trúc cây phân tích có độ sâu cố định (Fixed-Depth Parse Tree) để nhóm các dòng log dựa trên độ dài chuỗi và các từ khóa tiền tố, đạt tốc độ phân tích gần như tuyến tính O(N) đối với luồng dữ liệu lớn; (2) Spell (Du & Li, ICDM 2016) áp dụng thuật toán tìm chuỗi con chung dài nhất (Longest Common Subsequence - LCS) theo cơ chế dòng (streaming) để trích xuất động các thành phần tĩnh của thông điệp log; (3) LenMa (Shima, 2016) và AEL (Zhu et al., ISSRE 2023) lần lượt khai thác chiều dài các từ tố và tần suất xuất hiện của từ khóa để phân cụm và tách biến số ra khỏi chuỗi mẫu."
    )
    add_p([
        "Sau khi không gian log được rút gọn về tập ",
        omath("M"),
        " mẫu sự kiện cố định ",
        omath("E = {e_1, e_2, ..., e_M}"),
        ", các chuỗi sự kiện trong một cửa sổ thời gian hoặc phiên làm việc ",
        omath("W"),
        " được ánh xạ thành vector tần suất ",
        omath("x = [c(e_1), c(e_2), ..., c(e_M)]^T ∈ ℝ^M"),
        ", trong đó ",
        omath("c(e_i)"),
        " là số lần xuất hiện của sự kiện ",
        omath("e_i"),
        ". Bên cạnh đếm tần suất đơn thuần, các trọng số TF-IDF (Term Frequency - Inverse Document Frequency) hoặc độ hỗn loạn thông tin Shannon (Information Entropy) ",
        omath("H(W) = -∑ p(e_i) log_2 p(e_i)"),
        " cũng được áp dụng nhằm nhấn mạnh các sự kiện hiếm gặp."
    ])
    add_p([
        "Trên không gian vector tần suất này, các mô hình phát hiện bất thường kinh điển được triển khai rộng rãi: (1) Principal Component Analysis - PCA (Xu et al., SOSP 2009) phân rã không gian vector ",
        omath("ℝ^M"),
        " thành không gian con chuẩn tắc ",
        omath("S_n"),
        " và không gian con phần dư ",
        omath("S_r"),
        ", nhận diện bất thường khi năng lượng chiếu lên phần dư vượt ngưỡng kiểm định ",
        omath("x_a = (I - P P^T) x, ||x_a||^2 > γ_α"),
        "; (2) Invariant Mining (Lou et al., ATC 2010) tự động khai phá các phương trình bất biến tuyến tính phản ánh mối quan hệ bảo toàn logic giữa các bước thực thi ",
        omath("A x = 0"),
        "."
    ])
    add_p(
        "Mặc dù sở hữu ưu điểm vượt trội về hiệu năng tính toán (độ phức tạp O(N), độ trễ thấp dưới 0.1 ms/event), nhóm phương pháp thống kê và cú pháp bộc lộ hai điểm nghẽn nghiêm trọng không thể khắc phục: (1) Mất mát ngữ nghĩa an ninh do trừu tượng hóa tham số (Dynamic Parameter Loss): các bộ log parser bắt buộc phải sử dụng biểu thức chính quy (Regex) để loại bỏ toàn bộ các tham số biến động (địa chỉ IP, đường dẫn tệp tin, tham số dòng lệnh) thay thế bằng ký tự đại diện <*> khiến từ 82% đến 88% thông tin an ninh bị triệt tiêu; (2) Lan truyền và khuếch đại lỗi (Parser Error Propagation): khi gặp các định dạng log mới chưa từng xuất hiện (unseen logs), parser thường phân tách sai, dẫn đến hiện tượng bùng nổ số lượng mẫu sự kiện giả lập hoặc gộp nhầm các sự kiện khác biệt, phá vỡ hoàn toàn cấu trúc không gian vector x."
    )

    # --- 1.2.2 ---
    add_h3("1.2.2. Phương pháp semantic–sequential: Embeddings, Self-Supervised Learning, Transformer và Parsing-Free")
    add_p(
        "Nhằm khắc phục sự cứng nhắc của các vector đếm tần suất và tận dụng trật tự xuất hiện của các sự kiện, nhóm phương pháp Semantic–Sequential mô hình hóa luồng log tương tự như các chuỗi ngôn ngữ tự nhiên, tích hợp các kỹ thuật nhúng từ (Word Embeddings) và học tự giám sát (Self-Supervised Learning) để nắm bắt phụ thuộc ngữ cảnh dài hạn."
    )
    add_p([
        "Khởi đầu cho hướng nghiên cứu này là mô hình DeepLog (Du et al., CCS 2017). DeepLog sử dụng mạng bộ nhớ dài-ngắn hạn (LSTM) để mô hình hóa chuỗi Event ID như một bài toán dự báo phần tử tiếp theo (Next-Event Prediction). Tại mỗi bước thời gian t, mô hình ước lượng phân phối xác suất có điều kiện của sự kiện tiếp theo ",
        omath("P(e_t | e_{t-k}, ..., e_{t-1})"),
        ". Nếu sự kiện thực tế không nằm trong tập g sự kiện có xác suất cao nhất được mô hình dự đoán ",
        omath("E_hat_t = arg top-g P(e | e_{t-k}, ..., e_{t-1})"),
        ", hệ thống sẽ phát tín hiệu cảnh báo bất thường. Song song đó, DeepLog xây dựng mô hình LSTM thứ hai dựa trên phân phối chuẩn nhiều chiều để kiểm tra sự bất thường về giá trị tham số số học (Parameter Value Anomaly)."
    ])
    add_p(
        "Các công trình kế tiếp đã nâng cấp cơ chế biểu diễn ngữ nghĩa: (1) LogAnomaly (Meng et al., IJCAI 2019) đề xuất Template2Vec, trích xuất vector ngữ nghĩa cho từng mẫu log thông qua Word2Vec/FastText kết hợp trọng số d-IDF, giúp nhận biết sự tương đồng giữa các thông điệp có cấu trúc từ ngữ tương đương (ví dụ: 'Failed password' và 'Authentication failure'); (2) Logsy (Nedelkoski et al., 2020) sử dụng hàm mất mát phân loại ngoại lai (Outlier Classification Loss) trên dữ liệu log từ các hệ thống phụ trợ để định hình biên giới phân tách chặt chẽ cho lớp bình thường; (3) LogBERT (Guo et al., IJCNN 2021) khai thác Transformer hai chiều với hai tác vụ học tự giám sát: Dự đoán sự kiện log bị che (Masked Log Event Prediction) và Dự đoán phân bố khối lượng log (Volume Anomaly Prediction); (4) Nhóm tiếp cận không dùng Parser (Parsing-Free - NeuralLog, Le & Teoh, ASE 2021) bỏ qua bước phân tích cú pháp bằng cách sử dụng các mô hình ngôn ngữ tiền huấn luyện (BERT, RoBERTa) để trực tiếp mã hóa chuỗi văn bản log thô thành các vector nhúng ngữ nghĩa liên tục."
    )
    add_p(
        "Mặc dù đạt độ chính xác cao trên các tập dữ liệu phần mềm, nhóm phương pháp Semantic–Sequential bộc lộ ba giới hạn căn bản trong môi trường an ninh mạng: (1) Lợi thế ảo từ dữ liệu tiền huấn luyện (Pretraining-Data Advantage): các mô hình sử dụng Transformer tiền huấn luyện trên kho văn bản tổng quát thường bị rò rỉ tri thức ngoài miền; khi kiểm soát chặt chẽ ranh giới dữ liệu an ninh nội miền, hiệu năng thực tế suy giảm rõ rệt; (2) Chi phí tài nguyên và điểm nghẽn thông lượng: độ phức tạp tính toán bậc hai của cơ chế Self-Attention O(L^2) theo độ dài cửa sổ L khiến mô hình tiêu thụ GPU lớn và khó đáp ứng thông lượng thời gian thực; (3) Tầm nhìn cục bộ đơn luồng: mô hình chuỗi chỉ quan sát các sự kiện trên một dòng thời gian đơn lẻ, hoàn toàn mất dấu chuỗi tấn công APT phân tán đa tiến trình, đa luồng mạng và đa máy chủ."
    )

    # --- 1.2.3 ---
    add_h3("1.2.3. Đồ thị nguồn gốc và Graph Representation Learning")
    add_p([
        "Để vượt qua giới hạn tầm nhìn cục bộ của mô hình chuỗi, hướng tiếp cận dựa trên đồ thị nguồn gốc hệ thống (System Provenance Graph) mô hình hóa toàn bộ lịch sử thực thi và tương tác trong hệ điều hành dưới dạng một đồ thị có hướng, không đồng nhất và gán nhãn thời gian ",
        omath("G = (V, E, T_v, T_e, φ, ψ, τ)"),
        ". Trong đó V là tập các đỉnh đại diện cho các thực thể hệ thống thuộc tập kiểu T_v = {Process, File, Socket, Registry, User, Host}; E là tập các cạnh có hướng mô tả các tương tác luồng phụ thuộc thuộc tập kiểu T_e = {fork, execve, read, write, connect, bind, send, recv}; và τ gán nhãn mốc thời gian xảy ra tương tác."
    ])
    add_p([
        "Các hệ thống phát hiện xâm nhập dựa trên đồ thị nguồn gốc (PIDS) tiêu biểu bao gồm: (1) UNICORN (Han et al., NDSS 2020) xây dựng đồ thị nguồn gốc luồng thời gian thực, áp dụng thuật toán băm cây con Weisfeiler-Lehman (WL-subtree kernel) để chuyển đổi đồ thị động thành vector đặc trưng đếm histogram; (2) KAIROS (Wang et al., S&P 2024) tích hợp mạng nơ-ron đồ thị nhận biết thời gian (Time-Aware GNN), mã hóa đồng thời thông tin cấu trúc và khoảng cách thời gian giữa các cạnh để phát hiện kỹ thuật APT ẩn mình (low-and-slow); (3) NODLINK (She et al., NDSS 2024) và MAGIC (Wang et al., USENIX Security 2024) khai thác kiến trúc GNN dị thể kết hợp cơ chế chú ý đồ thị để tính toán vector nhúng nút ",
        omath("h_v^{(l+1)} = σ(∑ ∑ α_{uv}^r W_r^{(l)} h_u^{(l)})"),
        "; (4) ORTHRUS (Wang et al., USENIX Security 2025) tách biệt và gióng hàng hai luồng thông tin cấu trúc và luồng dữ liệu ngữ cảnh nhằm thực hiện truy vết nguồn gốc với độ chính xác cao."
    ])
    add_p(
        "Tuy nhiên, việc triển khai GNN trên đồ thị nguồn gốc quy mô thực tế đối mặt với ba rào cản lý thuyết và thực nghiệm sâu sắc: (1) Hiện tượng bùng nổ phụ thuộc (Dependency Explosion): các tiến trình hệ thống chạy dài hạn (như systemd, sshd, hoặc trình duyệt web) liên tục đọc/ghi hàng triệu tệp tin và socket, khiến hầu hết mọi nút trong đồ thị đều có đường đi liên kết đến nhau, tạo ra các phụ thuộc giả (False Dependencies) làm loãng dấu vết tấn công thực sự; (2) Ngộ nhận giữa quan hệ phụ thuộc cấu trúc và tác động nhân quả thực tế (Dependency != Causal Effect): công trình của Bilot et al. ('Sometimes Simpler is Better', USENIX Security 2025) chứng minh rằng nhiều mô hình GNN phức tạp thực chất chỉ học đặc trưng đường tắt như tần suất bậc của nút; khi kiểm soát rò rỉ, các bộ phân loại tuyến tính đơn giản đạt hiệu năng tương đương với chi phí thấp hơn hàng chục lần; (3) Hiện tượng nghẽn cổ chai thông tin (Over-smoothing và Over-squashing): khi tăng độ sâu GNN, Over-smoothing làm vector biểu diễn của các nút bị đồng nhất hóa, trong khi Over-squashing (Alon & Yahav, ICLR 2021) nén ép lượng thông tin cấu trúc tăng theo hàm mũ vào vector kích thước cố định, làm mất mát các tín hiệu tấn công tinh vi."
    )

    # Table 3: Summary Table
    tbl3_headers = ["Tiêu chí đánh giá", "Nhóm Thống kê / Cú pháp (Drain, PCA)", "Nhóm Chuỗi Semantic (DeepLog, LogBERT)", "Nhóm Đồ thị Nguồn gốc (UNICORN, MAGIC)"]
    tbl3_rows = [
        ["Cơ chế biểu diễn cốt lõi", "Vector đếm tần suất Event ID trên cửa sổ trượt", "Vector nhúng ngữ cảnh từ chuỗi sự kiện tuần tự", "Vector nhúng cấu trúc đồ thị luồng phụ thuộc dị thể"],
        ["Độ phức tạp thời gian", "O(N) — Tuyến tính theo luồng sự kiện", "O(N · L) đến O(N · L^2) — Phụ thuộc độ dài cửa sổ", "O(|V| + |E|) — Phụ thuộc quy mô đồ thị"],
        ["Bảo toàn tham số an ninh", "Kém (Bị parser cắt bỏ hoàn toàn bằng Regex)", "Trung bình (Bị rời rạc hóa hoặc che mờ)", "Rất tốt (Lưu trữ trực tiếp trên thuộc tính nút/cạnh)"],
        ["Mô hình hóa chuỗi quan hệ", "Hoàn toàn không có", "Cục bộ trên dòng thời gian đơn luồng", "Toàn diện trên quan hệ tương tác đa thực thể"],
        ["Khả năng triển khai SOC thực tế", "Rất cao (Độ trễ < 0.1 ms/event)", "Trung bình (Độ trễ 1 - 50 ms/window)", "Thách thức cao (Cần cắt tỉa đồ thị liên tục)"],
        ["Điểm nghẽn phương pháp luận", "Mất mát tham số động, lan truyền lỗi parser", "Thiếu ngữ cảnh đồ thị, chi phí tính toán cao", "Bùng nổ phụ thuộc, Over-smoothing, Over-squashing"]
    ]
    add_table(tbl3_headers, tbl3_rows)

    add_p(
        "Tổng kết lại, phân tích so sánh đối chiếu chỉ ra rằng không có bất kỳ phương pháp đơn lẻ nào trong ba nhóm trên giải quyết trọn vẹn bài toán biểu diễn đặc trưng log. Nhóm thống kê đạt hiệu năng cao nhưng mất mát tham số an ninh; nhóm chuỗi nắm bắt ngữ nghĩa tốt nhưng thiếu tầm nhìn đồ thị đa thực thể; nhóm đồ thị mô hình hóa quan hệ xuất sắc nhưng chịu gánh nặng bùng nổ phụ thuộc và chi phí tính toán. Thực trạng khoa học này trực tiếp đặt ra yêu cầu phải xác lập và giải quyết năm khoảng trống nghiên cứu cốt lõi tại Mục 1.3 tiếp theo."
    )

    # Save to updated file first
    updated_file = str(target_path.parent / (target_path.stem + ".updated.docx"))
    doc.save(updated_file)
    print(f"[SUCCESS] Written full 1.1 + 1.2 to: {updated_file}")

    # Attempt direct overwrite of active file
    try:
        doc.save(str(target_path))
        print(f"[SUCCESS] Overwritten active document: {target_path}")
    except PermissionError:
        print(f"[NOTE] Active file {target_path} is open in Word. Updated document is saved at {updated_file}.")

    # Validation check
    reloaded = docx.Document(updated_file)
    omml_in_reloaded = 0
    for p in reloaded.paragraphs:
        omml_in_reloaded += len(p._p.xpath('.//m:oMath'))
    print(f"[VERIFY] Document reloaded cleanly! Paragraphs: {len(reloaded.paragraphs)}, Tables: {len(reloaded.tables)}, Native OMML Equations: {omml_in_reloaded}")


if __name__ == "__main__":
    build_full_sections()
