"""
Script to safely patch D:\\Research\\Chuyên đề chuyên sâu - Copy.docx
Preserves Title Page (Table 0), Table of Contents, Conclusion, and References.
Rebuilds Section 1.1 (1.1.1, 1.1.2, 1.1.3) through 1.2.1 with verified academic prose,
5-level representation table, Representation Contract table, and clean mathematical equations.
"""

import sys
import shutil
from pathlib import Path
sys.stdout.reconfigure(encoding="utf-8")
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    xml_borders = (
        f'<w:tcBorders {nsdecls("w")}>\n'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '</w:tcBorders>'
    )
    tcPr.append(parse_xml(xml_borders))


def patch_document(target_file: str = r"D:\Research\Chuyên đề chuyên sâu - Copy.docx"):
    target_path = Path(target_file)
    if not target_path.exists():
        raise FileNotFoundError(f"Target file not found: {target_file}")

    # 1. Create safe backup
    backup_path = target_path.parent / (target_path.stem + ".backup.docx")
    shutil.copyfile(target_path, backup_path)
    print(f"[1/4] Created safe backup at: {backup_path}")

    # 2. Open original document
    doc = docx.Document(str(target_path))

    # Identify boundary paragraphs
    # Paragraph 0..23: Front matter (Title, TOC, Intro)
    # Paragraph 24: Heading 1 (TỔNG QUAN VỀ PHƯƠNG PHÁP...)
    # Paragraph 73: Conclusion (KẾT LUẬN)
    
    # We will build a clean new document that copies the Front matter, appends rebuilt 1.1 -> 1.2.1, and appends Conclusion & References
    new_doc = docx.Document(str(backup_path))
    
    # Let's inspect paragraphs in new_doc
    # We want to replace paragraphs from index 24 to the paragraph before 'Kết luận'
    # Find exact indices
    p_h1_idx = None
    p_conclusion_idx = None
    
    for idx, p in enumerate(new_doc.paragraphs):
        txt = p.text.strip().lower()
        if "tổng quan về phương pháp trích xuất đặc trưng" in txt and p_h1_idx is None:
            p_h1_idx = idx
        if txt == "kết luận" or "kết luận" in txt and p_conclusion_idx is None and idx > 30:
            p_conclusion_idx = idx

    print(f"[2/4] Identified boundaries: Heading 1 at [{p_h1_idx}], Conclusion at [{p_conclusion_idx}]")

    # Let's create a fresh document using the original document's styles and sections
    # To avoid XML corruption, we will construct a clean output document based on backup_path
    clean_doc = docx.Document(str(backup_path))
    
    # Remove old paragraphs from p_h1_idx to p_conclusion_idx - 1
    # Also remove the 2 old content tables (Table 1 and Table 2 in doc.tables)
    # Table 0 is the Title page box (preserve it!)
    # Let's safely remove old tables (Table 1 and Table 2)
    while len(clean_doc.tables) > 1:
        tbl_to_remove = clean_doc.tables[1]
        tbl_to_remove._tbl.getparent().remove(tbl_to_remove._tbl)
    print(f"[3/4] Preserved Title Page Table. Cleaned up old content tables.")

    # Remove the old paragraphs from index 24 to the conclusion
    # In python-docx, deleting paragraphs from bottom to top prevents index shifting:
    paragraphs_to_remove = []
    found_h1 = False
    for p in clean_doc.paragraphs:
        txt = p.text.strip().lower()
        if "tổng quan về phương pháp trích xuất đặc trưng" in txt:
            found_h1 = True
        if found_h1:
            if "kết luận" in txt:
                break
            paragraphs_to_remove.append(p)

    for p in paragraphs_to_remove:
        p._p.getparent().remove(p._p)

    # Now let's locate the paragraph right before 'Kết luận' to insert our clean sections
    # Find insertion point
    target_p = None
    for p in clean_doc.paragraphs:
        if "kết luận" in p.text.strip().lower():
            target_p = p
            break

    def insert_para_before(ref_p, text, style_name="Normal", bold_prefix=None, first_line_indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
        new_p = clean_doc.add_paragraph(style=style_name) if ref_p is None else ref_p.insert_paragraph_before(style=style_name)
        new_p.alignment = align
        new_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        new_p.paragraph_format.space_after = Pt(6)
        new_p.paragraph_format.space_before = Pt(0)
        if first_line_indent:
            new_p.paragraph_format.first_line_indent = Cm(1.27)
        
        if bold_prefix:
            r_pre = new_p.add_run(bold_prefix)
            r_pre.font.name = "Times New Roman"
            r_pre.font.size = Pt(14)
            r_pre.bold = True

        r = new_p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
        return new_p

    def insert_heading_1(ref_p, text):
        p = clean_doc.add_paragraph(style="Heading 1") if ref_p is None else ref_p.insert_paragraph_before(style="Heading 1")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(16)
        r.bold = True
        return p

    def insert_heading_2(ref_p, text):
        p = clean_doc.add_paragraph(style="Heading 2") if ref_p is None else ref_p.insert_paragraph_before(style="Heading 2")
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
        r.bold = True
        return p

    def insert_heading_3(ref_p, text):
        p = clean_doc.add_paragraph(style="Heading 3") if ref_p is None else ref_p.insert_paragraph_before(style="Heading 3")
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
        r.bold = True
        r.italic = True
        return p

    def insert_table(ref_p, headers, rows_data):
        tbl = clean_doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Move table before ref_p in XML
        if ref_p is not None:
            ref_p._p.addprevious(tbl._tbl)

        # Headers
        for c_idx, h in enumerate(headers):
            cell = tbl.cell(0, c_idx)
            cell.text = h
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
                r.bold = True

        # Rows
        for r_idx, row in enumerate(rows_data):
            for c_idx, val in enumerate(row):
                cell = tbl.cell(r_idx + 1, c_idx)
                cell.text = val
                set_cell_border(cell)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)

        # Space after table
        insert_para_before(ref_p, "", first_line_indent=False)

    # NOW INSERT ALL AUDITED SECTIONS FROM 1.1 TO 1.2.1
    print("[4/4] Writing audited Sections 1.1 -> 1.2.1 into Document...")

    insert_heading_1(target_p, "CHƯƠNG 1: TỔNG QUAN VỀ PHƯƠNG PHÁP TRÍCH XUẤT ĐẶC TRƯNG DỮ LIỆU LOG VÀ THÁCH THỨC BẢO TOÀN NGỮ CẢNH AN TOÀN")
    insert_heading_2(target_p, "1.1. Bài toán biểu diễn log trong phát hiện tấn công đa giai đoạn")
    insert_para_before(
        target_p,
        "Trong các Trung tâm Điều hành An ninh mạng (Security Operations Center - SOC) hiện đại, dữ liệu nhật ký hệ thống (system logs) và nhật ký kiểm toán (audit logs) đóng vai trò là nguồn bằng chứng trung tâm phục vụ phát hiện, điều tra và ứng phó các chiến dịch tấn công có chủ đích (Advanced Persistent Threats - APT). Khác với văn bản ngôn ngữ tự nhiên thông thường hoặc tín hiệu thị giác máy tính, dữ liệu log sở hữu cấu trúc bán hình thức (semi-structured), mang tính phụ thuộc thời gian nghiêm ngặt, phản ánh các quan hệ phụ thuộc thực thi giữa các tiến trình và tài nguyên hệ điều hành, đồng thời chứa đựng các ngữ nghĩa an ninh đặc thù. Mục 1.1 tập trung hình thức hóa bài toán biểu diễn đặc trưng log phục vụ phát hiện tấn công đa giai đoạn, phân tích bản chất không gian dữ liệu doanh nghiệp, xác lập mô hình hành vi phi tuyến tính trên ma trận MITRE ATT&CK và thiết lập khung Hợp đồng Biểu diễn (Representation Contract) nhằm bảo toàn các bất biến an toàn cho không gian vector đặc trưng z."
    )

    insert_heading_3(target_p, "1.1.1. Không gian dữ liệu log doanh nghiệp: tốc độ cao, mất cân bằng cực đoan và phân phối biến đổi")
    insert_para_before(
        target_p,
        "Không gian dữ liệu nhật ký trong môi trường mạng doanh nghiệp được tổng hợp từ ba nhóm nguồn telemetry chính với cấu trúc và định dạng không đồng nhất. Nhóm thứ nhất là nhật ký kiểm toán máy chủ (Host Audit Logs), bao gồm Linux Auditd, Windows Event Log / Sysmon và Linux eBPF (Extended Berkeley Packet Filter). Nguồn dữ liệu này ghi nhận trực tiếp các sự kiện ở mức nhân hệ điều hành thông qua việc chặn bắt các lời gọi hệ thống (syscalls), bao gồm: khởi tạo tiến trình (execve, CreateProcess - Sysmon Event ID 1), nạp thư viện động (ImageLoaded - Sysmon Event ID 7), thao tác tệp tin (open, unlink, FileCreate - Sysmon Event ID 11), sửa đổi cấu hình registry (RegSetValue - Sysmon Event ID 13), cùng các thao tác mở và kết nối socket mạng (connect, accept - Sysmon Event ID 3). Nhóm thứ hai là nhật ký luồng mạng (Network Flow & Protocol Logs), được thu thập từ Zeek, Suricata hoặc NetFlow/IPFIX, cung cấp siêu dữ liệu kết nối giữa các nút mạng, giao dịch DNS, chứng chỉ TLS/SSL và thông lượng gói tin. Nhóm thứ ba là nhật ký ứng dụng và dịch vụ (Application & Service Logs), phát sinh từ máy chủ web (Nginx, Apache), cơ sở dữ liệu, dịch vụ phân tán (HDFS) cùng hệ thống điều phối container (Kubernetes Audit Logs)."
    )
    insert_para_before(
        target_p,
        "Tính dị thể sâu sắc của dữ liệu đặt ra bài toán khoa học về việc lựa chọn đơn vị quan sát (Unit of Observation) phù hợp cho mô hình học biểu diễn. Việc phân cấp đơn vị quan sát quyết định trực tiếp đến mức độ bảo toàn thông tin và độ phức tạp tính toán:"
    )

    headers_tbl1 = ["Mức độ hạt biểu diễn", "Đơn vị quan sát", "Dữ liệu đại diện", "Ưu điểm cốt lõi", "Rào cản và Thách thức ngữ nghĩa"]
    rows_tbl1 = [
        ["1. Mức từ tố (Token)", "Chuỗi con, từ khóa rời rạc", "Từ khóa tĩnh, địa chỉ IP, mã lỗi hex", "Dễ dàng xử lý bằng kỹ thuật nhúng từ vựng", "Mất hoàn toàn cấu trúc cú pháp và quan hệ liên kết trường"],
        ["2. Mức sự kiện (Event)", "Một dòng log đơn lẻ", "Bản ghi telemetry tại thời điểm t", "Bảo toàn đầy đủ thuộc tính cục bộ tại thời điểm t", "Thiếu ngữ cảnh chuỗi tuần tự và lịch sử tương tác trước đó"],
        ["3. Mức chuỗi / Phiên (Sequence / Session)", "Cửa sổ trượt hoặc phiên tiến trình", "Chuỗi sự kiện [e_{t-k+1}, ..., e_t] theo thời gian", "Nắm bắt quan hệ phụ thuộc thứ tự thời gian cục bộ", "Nhạy cảm với nhiễu xen kẽ (interleaving) từ các luồng chạy song song"],
        ["4. Mức thực thể (Entity)", "Định danh tác nhân (Host, User, IP, Process Instance)", "Lịch sử tương tác gom cụm theo thực thể", "Phân lập rõ ràng ranh giới hành vi của từng chủ thể", "Khó phát hiện các hành vi tấn công phối hợp vượt ranh giới thực thể"],
        ["5. Mức đồ thị (Graph)", "Đồ thị nguồn gốc (Provenance Graph)", "Đồ thị luồng phụ thuộc dị thể G", "Mô hình hóa toàn diện quan hệ phụ thuộc đa thực thể", "Bùng nổ kích thước đồ thị (dependency explosion) và chi phí tính toán"]
    ]
    insert_table(target_p, headers_tbl1, rows_tbl1)

    insert_para_before(
        target_p,
        "Bên cạnh tính dị thể, dữ liệu log doanh nghiệp chịu áp lực vận hành khắc nghiệt về mặt thông lượng và tỷ lệ phân bố nhãn. Trong các môi trường mạng lớn, hệ thống tiếp nhận luồng sự kiện liên tục với thông lượng cao, đòi hỏi thuật toán biểu diễn đặc trưng phải xử lý theo cơ chế dòng (streaming) với độ phức tạp tính toán tuyến tính O(N) mà không đòi hỏi lưu giữ toàn bộ lịch sử đồ thị trong bộ nhớ truy cập ngẫu nhiên. Đồng thời, tỷ lệ các dòng log liên quan đến hành vi tấn công trong thực tế thường ở mức rất thấp, tạo ra sự mất cân bằng nhãn cực đoan khiến các hàm mất mát học máy thông thường có xu hướng xem nhẹ nhóm thiểu số. Trong khi đó, việc suy thoái không gian vector (sụp đổ chiều biểu diễn) nảy sinh khi các mô hình tự giám sát không áp dụng các cơ chế điều hòa phương sai - hiệp phương sai phù hợp để duy trì tính đa dạng của các chiều đặc trưng."
    )
    insert_para_before(
        target_p,
        "Dưới góc độ phân phối thời gian, dữ liệu log liên tục biến đổi trong môi trường vận hành dài hạn. Để tránh các sai lệch phương pháp luận khi đánh giá mô hình, luận án phân định tường minh bốn cơ chế trôi dạt dữ liệu (Drift Taxonomy):"
    )
    insert_para_before(target_p, "Bản chất hành vi và mục đích tấn công thay đổi theo thời gian mặc dù cấu trúc định dạng log không đổi: P_t(Y | X) ≠ P_{t+1}(Y | X).", bold_prefix="• Concept Drift: ")
    insert_para_before(target_p, "Việc nâng cấp phần mềm, cập nhật bản vá hoặc thay đổi cấu hình làm biến đổi cấu trúc chuỗi mẫu log: P_t(X_template) ≠ P_{t+1}(X_template).", bold_prefix="• Template Drift: ")
    insert_para_before(target_p, "Lưu lượng người dùng, tần suất giao tác nghiệp vụ hoặc cơ cấu dịch vụ hệ thống biến động theo chu kỳ: P_t(X) ≠ P_{t+1}(X).", bold_prefix="• Population Drift: ")
    insert_para_before(target_p, "Không gian vector tiềm ẩn z bị suy giảm năng lực phân tách do dữ liệu đầu vào trôi dạt khỏi vùng phân phối huấn luyện ban đầu: P_t(z | X) ≠ P_{t+1}(z | X).", bold_prefix="• Representation Drift: ")

    insert_heading_3(target_p, "1.1.2. Hành vi tấn công đa giai đoạn và ánh xạ đa nhãn MITRE ATT&CK")
    insert_para_before(
        target_p,
        "Trong phân tích an ninh mạng, một số công trình áp dụng mô hình trạng thái Markov tuyến tính tuần tự để xâu chuỗi các giai đoạn tấn công từ Thâm nhập ban đầu (Initial Access), Thực thi (Execution), Duy trì (Persistence), Leo thang đặc quyền (Privilege Escalation) cho đến Đánh cắp dữ liệu (Exfiltration)."
    )
    insert_para_before(
        target_p,
        "Tuy nhiên, các quan sát thực nghiệm trên các chiến dịch tấn công cho thấy hành vi của kẻ tấn công mang bản chất phi tuyến tính: (1) Nhảy cóc giai đoạn (Step Skipping): Kẻ tấn công có thể khai thác trực tiếp lỗ hổng thực thi mã từ xa để trích xuất dữ liệu ra ngoài mà không cần thiết lập cơ chế duy trì hay di chuyển ngang; (2) Lặp vòng kỹ thuật (Tactic Looping & Interleaving): Kỹ thuật thu thập thông tin nội bộ (Discovery) thường được lặp lại nhiều lần xen kẽ giữa các bước leo thang đặc quyền và chiếm đoạt thông tin xác thực; (3) Phân nhánh tiến trình song song (Parallel Branching): Kẻ tấn công có thể khởi tạo đồng thời nhiều luồng tiến trình con độc lập trên các tiến trình hợp lệ khác nhau nhằm phân tán sự theo dõi của hệ thống phòng thủ."
    )
    insert_para_before(
        target_p,
        "Do đó, luận án xác lập nguyên tắc: Ma trận MITRE ATT&CK được mô hình hóa thành một Không gian Bằng chứng Hành vi Đa chiều (Multi-label Behavioral Evidence Space) Y ⊆ {0, 1}^{|T|}, trong đó một chuỗi sự kiện hoặc cây tiến trình có thể đồng thời kích hoạt nhiều nhãn chiến thuật (Tactics) và kỹ thuật (Techniques) tại cùng một thời điểm quan sát."
    )
    insert_para_before(
        target_p,
        "Về mặt dữ liệu thực nghiệm, việc mô hình hóa hành vi tấn công đòi hỏi phải phân định chính xác đặc tính gán nhãn và mức độ hạt (Label Granularity) của từng bộ dữ liệu chuẩn: (1) DARPA Transparent Computing (TC E3/E5) cung cấp telemetry kiểm toán mức nhân hệ điều hành với nhãn mặt đất được ghi nhận ở mức kịch bản tấn công tổng thể kèm theo danh sách các nút và cạnh liên quan trong đồ thị nguồn gốc; (2) LANL CyberEvents (2017) kết hợp xác thực máy tính, tiến trình và luồng mạng quy mô lớn với nhãn độc hại gắn theo mốc thời gian và tài khoản/máy tính bị đội đỏ xâm nhập; (3) HDFS Benchmark ghi nhận thao tác khối dữ liệu Hadoop với nhãn bất thường phản ánh sự cố kỹ thuật hoặc lỗi khối (Block Anomaly), không phải tấn công APT; (4) BGL Benchmark ghi nhận nhật ký máy siêu điện toán BlueGene/L với nhãn cảnh báo lỗi phần cứng và hệ thống phản ánh độ tin cậy vận hành."
    )
    insert_para_before(
        target_p,
        "Đặc biệt, sự xuất hiện của nhiễu từ hành vi quản trị viên (Admin-Noise) là một thách thức thực tế lớn. Quản trị viên hệ thống thường xuyên sử dụng các công cụ dòng lệnh hợp lệ (Living-off-the-Land Binaries - LOLBins) như PowerShell, WMI (wmic.exe), SSH, certutil.exe hay vssadmin.exe cho mục đích bảo trì, sao lưu hoặc kiểm tra mạng. Các hành vi này có cấu trúc cú pháp tương tự như kỹ thuật của kẻ tấn công. Nếu mô hình chỉ dựa vào sự xuất hiện của từ khóa, hệ thống sẽ phát sinh nhiều cảnh báo sai (False Positives). Do đó, bài toán biểu diễn đòi hỏi cơ chế phân bổ bằng chứng yếu (Weak Evidence Attribution) thông qua khung Học Đa Thể hiện (Multiple Instance Learning - MIL)."
    )

    insert_heading_3(target_p, "1.1.3. Các mức biểu diễn dữ liệu và Hợp đồng Biểu diễn (Representation Contract)")
    insert_para_before(
        target_p,
        "Nhằm định hình rõ ràng các yêu cầu đối với không gian vector đặc trưng z ∈ ℝ^d và ngăn ngừa rủi ro học đường tắt (Shortcut Learning), mô hình biểu diễn bắt buộc phải tuân thủ Hợp đồng Biểu diễn (Representation Contract) được định nghĩa dưới dạng bộ ba hình thức C_representation = ⟨P_preserve, I_invariant, E_exclude⟩:"
    )

    headers_tbl2 = ["Nhóm quy tắc", "Ý nghĩa phương pháp luận", "Danh mục thuộc tính telemetry tương ứng"]
    rows_tbl2 = [
        ["PRESERVE (Bảo toàn)", "Bảo toàn các thuộc tính mang ngữ nghĩa an ninh và quan hệ phụ thuộc thực thi trong không gian vector z.", "• Thứ tự thời gian cục bộ giữa các sự kiện liền kề.\n• Quan hệ tiến trình cha-con (parent_process_guid → process_guid).\n• Cấu trúc liên kết đa thực thể (Tiến trình - Tệp tin - Socket - Registry).\n• Tham số an ninh động: IP đích, cổng, đường dẫn tệp, command_line, return_code, ImageHash."],
        ["INVARIANT (Bất biến)", "Duy trì tính bất biến của vector z trước các biến đổi hình thức không làm thay đổi bản chất hành vi.", "• Biến đổi cú pháp vô hại: khoảng trắng thừa, dấu xuống dòng, chữ hoa/thường trong câu lệnh không phân biệt hoa-thường.\n• Định danh tiến trình tạm thời ngẫu nhiên (PID, Thread ID biến động qua mỗi lần khởi chạy).\n• Hoán vị thứ tự giữa các trường thuộc tính độc lập không có quan hệ phụ thuộc tuần tự."],
        ["EXCLUDE (Triệt tiêu)", "Loại bỏ các thuộc tính gây học vẹt, tạo đường tắt giả định hoặc vi phạm quyền riêng tư khỏi vector z.", "• Định danh máy chủ cố định gây thiên lệch (Host UUID, tên máy trạm thử nghiệm cụ thể, địa chỉ MAC tĩnh).\n• Dấu vết phân vùng nhân tạo (Synthetic Split Artifacts) phát sinh trong quá trình chia tập train/test.\n• Mọi đặc trưng học đường tắt (Shortcut Features) có tương quan ngẫu nhiên với nhãn trong tập huấn luyện nhưng không có giá trị tổng quát hóa."]
    ]
    insert_table(target_p, headers_tbl2, rows_tbl2)

    insert_para_before(
        target_p,
        "Cần nhấn mạnh rằng, nguyên tắc PRESERVE tập trung vào việc bảo toàn ngữ nghĩa an ninh động chứ không đồng nghĩa với việc giữ lại toàn bộ định danh thô của người dùng hay máy chủ, nhằm đảm bảo khả năng liên kết có kiểm soát (controlled linkability) và tương thích với các tiêu chuẩn đánh giá quyền riêng tư. Tương tự, nguyên tắc EXCLUDE thiết lập các ràng buộc phủ định đối với các giả định đường tắt đã biết và các biến số gây rò rỉ phân vùng, thay vì giả định rằng mọi đường tắt đều có thể nhận biết trước khi thực nghiệm."
    )
    insert_para_before(target_p, "Đi đôi với Hợp đồng Biểu diễn, luận án thiết lập nguyên tắc phân định ranh giới phương pháp luận ba tầng độc lập:")
    insert_para_before(
        target_p,
        "Đảm nhiệm tiền xử lý dữ liệu thô, phân tích cú pháp sơ bộ, chuẩn hóa kiểu dữ liệu trường và làm sạch dữ liệu. Tầng này không gánh vác nhiệm vụ học biểu diễn ngữ cảnh an ninh sâu.",
        bold_prefix="1. Tầng 1 — Trích xuất đặc trưng cơ sở (Feature Extraction): "
    )
    insert_para_before(
        target_p,
        "Thiết lập ánh xạ f_θ: X → z ∈ ℝ^d từ cấu trúc chuỗi sự kiện và đồ thị nguồn gốc sang không gian vector tiềm ẩn. Toàn bộ năng lực bảo toàn ngữ nghĩa an ninh và tính bất biến được đóng gói trọn vẹn bên trong vector z.",
        bold_prefix="2. Tầng 2 — Học không gian biểu diễn (Representation Learning — Trọng tâm Luận án): "
    )
    insert_para_before(
        target_p,
        "Đánh giá chất lượng của vector biểu diễn z thông qua các bộ thăm dò tuyến tính đóng băng tham số (Frozen Linear Probing): ŷ = σ(W^T z + b). Trong đó tham số θ của bộ trích xuất đặc trưng Tầng 2 được giữ cố định hoàn toàn trong suốt quá trình đánh giá ở Tầng 3. Quy tắc này bảo đảm bộ phân loại hạ nguồn không làm thay nhiệm vụ trích xuất đặc trưng của Tầng 2.",
        bold_prefix="3. Tầng 3 — Phát hiện và phân loại hạ nguồn (Downstream Detection): "
    )

    insert_heading_2(target_p, "1.2. Phân tích so sánh các nhóm phương pháp hiện đại")
    insert_para_before(
        target_p,
        "Nhằm định vị chính xác các đóng góp kỹ thuật và cơ sở lý luận của chuyên đề, Mục 1.2 tiến hành khảo cứu toàn diện, phân loại và phân tích đa chiều ba nhóm phương pháp biểu diễn đặc trưng log chủ đạo trong y văn hiện đại: (1) Nhóm phương pháp thống kê và cú pháp dựa trên mẫu log (Statistical & Syntactic Parsing-Based); (2) Nhóm phương pháp biểu diễn ngữ nghĩa chuỗi thời gian (Semantic–Sequential & Transformer-Based); và (3) Nhóm phương pháp học biểu diễn đồ thị nguồn gốc (Provenance Graph Representation Learning). Mỗi nhóm phương pháp được mổ xẻ tường minh về cơ chế toán học, độ phức tạp thuật toán, ưu điểm cốt lõi và các rào cản nền tảng khi triển khai trong môi trường phát hiện tấn công thực tế."
    )

    insert_heading_3(target_p, "1.2.1. Phương pháp thống kê và cú pháp: Event Count, Frequency, Entropy và Template Features")
    insert_para_before(
        target_p,
        "Nhóm phương pháp thống kê và cú pháp đại diện cho thế hệ tiếp cận đầu tiên trong phân tích nhật ký tự động. Cơ chế hoạt động của nhóm này dựa trên quy trình hai giai đoạn tách rời: giai đoạn phân tách dòng log thô thành các mẫu định dạng tĩnh (Log Templates / Event IDs) thông qua bộ phân tích cú pháp (Log Parser), tiếp theo là giai đoạn lượng hóa các chuỗi sự kiện thành vector số học dựa trên các thước đo thống kê kinh điển."
    )
    insert_para_before(
        target_p,
        "Trong giai đoạn phân tích cú pháp, các thuật toán tiêu biểu đã được phát triển nhằm tối ưu hóa tốc độ xử lý: (1) Drain (He et al., ICWS 2017) sử dụng cấu trúc cây phân tích có độ sâu cố định (Fixed-Depth Parse Tree) để nhóm các dòng log dựa trên độ dài chuỗi và các từ khóa tiền tố, đạt tốc độ phân tích gần như tuyến tính O(N) đối với luồng dữ liệu lớn; (2) Spell (Du & Li, ICDM 2016) áp dụng thuật toán tìm chuỗi con chung dài nhất (Longest Common Subsequence - LCS) theo cơ chế dòng (streaming) để trích xuất động các thành phần tĩnh của thông điệp log; (3) LenMa (Shima, 2016) và AEL (Zhu et al., ISSRE 2023) lần lượt khai thác chiều dài các từ tố và tần suất xuất hiện của từ khóa để phân cụm và tách biến số ra khỏi chuỗi mẫu."
    )
    insert_para_before(
        target_p,
        "Sau khi không gian log được rút gọn về tập M mẫu sự kiện cố định Ɛ = {e_1, e_2, ..., e_M}, các chuỗi sự kiện trong một cửa sổ thời gian hoặc phiên làm việc W được ánh xạ thành vector tần suất x = [c(e_1), c(e_2), ..., c(e_M)]^T ∈ ℝ^M, trong đó c(e_i) là số lần xuất hiện của sự kiện e_i. Bên cạnh đếm tần suất đơn thuần, các trọng số TF-IDF (Term Frequency - Inverse Document Frequency) hoặc độ hỗn loạn thông tin Shannon (Information Entropy) H(W) = -∑ p(e_i) log_2 p(e_i) cũng được áp dụng nhằm nhấn mạnh các sự kiện hiếm gặp."
    )
    insert_para_before(
        target_p,
        "Trên không gian vector tần suất này, các mô hình phát hiện bất thường kinh điển được triển khai rộng rãi: (1) Principal Component Analysis - PCA (Xu et al., SOSP 2009) phân rã không gian vector ℝ^M thành không gian con chuẩn tắc S_n và không gian con phần dư S_r, nhận diện bất thường khi năng lượng chiếu lên phần dư vượt ngưỡng x_a = (I - P P^T) x, ||x_a||^2 > γ_α; (2) Invariant Mining (Lou et al., ATC 2010) tự động khai phá các phương trình bất biến tuyến tính phản ánh mối quan hệ bảo toàn logic giữa các bước thực thi A x = 0."
    )
    insert_para_before(
        target_p,
        "Mặc dù sở hữu ưu điểm vượt trội về hiệu năng tính toán (độ phức tạp O(N), độ trễ thấp), nhóm phương pháp thống kê và cú pháp bộc lộ hai điểm nghẽn nghiêm trọng không thể khắc phục: (1) Mất mát ngữ nghĩa an ninh do trừu tượng hóa tham số (Dynamic Parameter Loss): các bộ log parser bắt buộc phải sử dụng biểu thức chính quy (Regex) để loại bỏ toàn bộ các tham số biến động (địa chỉ IP, đường dẫn tệp tin, tham số dòng lệnh) thay thế bằng ký tự đại diện <*> khiến thông tin an ninh bị triệt tiêu; (2) Lan truyền và khuếch đại lỗi (Parser Error Propagation): khi gặp các định dạng log mới chưa từng xuất hiện (unseen logs), parser thường phân tách sai, dẫn đến hiện tượng bùng nổ số lượng mẫu sự kiện giả lập hoặc gộp nhầm các sự kiện khác biệt, phá vỡ hoàn toàn cấu trúc không gian vector x."
    )

    # 4. Save and Validate
    updated_file = str(target_path.parent / (target_path.stem + ".updated.docx"))
    clean_doc.save(updated_file)
    print(f"[SUCCESS] Safely generated updated document at: {updated_file}")

    try:
        clean_doc.save(target_file)
        print(f"[SUCCESS] Directly overwritten active file: {target_file}")
    except PermissionError:
        print(f"[NOTE] File {target_file} is currently open in Microsoft Word. Content is saved in {updated_file}.")

    # Reload to verify zero corruption
    verified_doc = docx.Document(updated_file)
    print(f"[VERIFY] Document reloaded cleanly! Total paragraphs: {len(verified_doc.paragraphs)}, Total tables: {len(verified_doc.tables)}")


if __name__ == "__main__":
    patch_document()
