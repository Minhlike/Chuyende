"""
Master Academic Word Document Builder - Chapter 1 & Chapter 2 Section 2.1
Compiles pristine scientific thesis document directly into Microsoft Word 2016 (.docx)
with Native Word Citations, OMML Equations, Vector Visuals, Captions, and COM Automation.
"""

import os
import shutil
import zipfile
import uuid
import win32com.client as win32
import pythoncom
from datetime import datetime, timezone
from pathlib import Path
from xml.sax.saxutils import escape

import docx
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn




def latex_to_clean_omml(latex_code: str) -> OxmlElement:
    """Converts a LaTeX expression to a clean Office Math (w:oMath) element."""
    clean_code = (
        latex_code.replace(r"\mathbf", "")
        .replace(r"\mathcal", "")
        .replace(r"\mathbb", "")
        .replace(r"\text", "")
        .replace(r"\quad", " ")
        .replace(r"\qquad", "  ")
        .replace(r"\top", "T")
        .replace(r"\mid", "|")
        .replace(r"\_", "_")
        .replace("{", "")
        .replace("}", "")
    )
    omml_xml = (
        f'<m:oMath {nsdecls("m")}>\n'
        '  <m:r>\n'
        '    <m:rPr>\n'
        '      <m:scr m:val="roman"/>\n'
        '      <m:sty m:val="p"/>\n'
        '    </m:rPr>\n'
        f'    <m:t>{escape(clean_code)}</m:t>\n'
        '  </m:r>\n'
        '</m:oMath>'
    )
    return parse_xml(omml_xml)


def make_citation_element(items):
    """Creates native Word Citation field elements for one or more source keys/indices."""
    if isinstance(items, (int, str)):
        items = [items]

    elems = []
    for i, item in enumerate(items):
        if isinstance(item, int):
            tag = f"SRC{item:06d}"
            num_str = str(item)
        else:
            tag = item.replace("-", "")
            num_str = str(int(item.split("-")[1])) if "-" in item else item

        if i > 0:
            sep_xml = (
                f'<w:r {nsdecls("w")}>\n'
                '  <w:rPr>\n'
                '    <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>\n'
                '    <w:sz w:val="28"/>\n'
                '  </w:rPr>\n'
                '  <w:t xml:space="preserve">, </w:t>\n'
                '</w:r>'
            )
            elems.append(parse_xml(sep_xml))

        fld_xml = (
            f'<w:fldSimple {nsdecls("w")} w:instr="CITATION {tag} \\l 1033 ">\n'
            '  <w:r>\n'
            '    <w:rPr>\n'
            '      <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>\n'
            '      <w:sz w:val="28"/>\n'
            '    </w:rPr>\n'
            f'    <w:t>[{num_str}]</w:t>\n'
            '  </w:r>\n'
            '</w:fldSimple>'
        )
        elems.append(parse_xml(fld_xml))
    return elems


def make_ref_element(bookmark_name: str, fallback_text: str, font_size_pt: float = 14.0):
    """Creates a native Word REF dynamic cross-reference field element."""
    sz_val = int(font_size_pt * 2)
    ref_xml = (
        f'<w:fldSimple {nsdecls("w")} w:instr="REF {bookmark_name} \\h ">\n'
        '  <w:r>\n'
        '    <w:rPr>\n'
        '      <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>\n'
        f'      <w:sz w:val="{sz_val}"/>\n'
        '    </w:rPr>\n'
        f'    <w:t>{fallback_text}</w:t>\n'
        '  </w:r>\n'
        '</w:fldSimple>'
    )
    return parse_xml(ref_xml)


def add_table_caption(doc, target_p, seq_num: int, title_content, bookmark_name: str = None):
    """Inserts a native Word Caption paragraph with SEQ Bảng field and keepWithNext."""
    cap_p = doc.add_paragraph(style="Caption") if target_p is None else target_p.insert_paragraph_before(style="Caption")
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_before = Pt(12)
    cap_p.paragraph_format.space_after = Pt(4)
    cap_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    cap_p.paragraph_format.keep_with_next = True

    cap_xml = (
        f'<w:pPr {nsdecls("w")}>\n'
        '  <w:pStyle w:val="Caption"/>\n'
        '  <w:jc w:val="center"/>\n'
        '  <w:spacing w:before="240" w:after="80" w:line="360" w:lineRule="auto"/>\n'
        '  <w:keepNext/>\n'
        '</w:pPr>'
    )
    cap_p._p.remove(cap_p._p.pPr)
    cap_p._p.append(parse_xml(cap_xml))

    if bookmark_name:
        cap_p._p.append(parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="20{seq_num}" w:name="{bookmark_name}"/>'))

    r1 = cap_p.add_run(f"Bảng {seq_num//10 if seq_num >= 10 else 1}.")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(14)
    r1.bold = True

    seq_xml = (
        f'<w:fldSimple {nsdecls("w")} w:instr="SEQ Bảng \\* ARABIC ">\n'
        '  <w:r>\n'
        '    <w:rPr>\n'
        '      <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>\n'
        '      <w:b/>\n'
        '      <w:sz w:val="28"/>\n'
        '    </w:rPr>\n'
        f'    <w:t>{seq_num % 10 if seq_num >= 10 else seq_num}</w:t>\n'
        '  </w:r>\n'
        '</w:fldSimple>'
    )
    cap_p._p.append(parse_xml(seq_xml))

    if bookmark_name:
        cap_p._p.append(parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="20{seq_num}"/>'))

    colon_r = cap_p.add_run(": ")
    colon_r.font.name = "Times New Roman"
    colon_r.font.size = Pt(14)
    colon_r.bold = True

    items = [title_content] if isinstance(title_content, str) else title_content
    for item in items:
        if isinstance(item, str):
            r = cap_p.add_run(item)
            r.font.name = "Times New Roman"
            r.font.size = Pt(14)
            r.bold = True
        elif isinstance(item, list):
            for sub in item:
                cap_p._p.append(sub)
        else:
            cap_p._p.append(item)

    return cap_p


def add_figure_caption(doc, target_p, seq_label: str, title_content, bookmark_name: str = None, seq_id: int = 1):
    """Inserts a native Word Caption paragraph with SEQ Hình field placed below the figure."""
    cap_p = doc.add_paragraph(style="Caption") if target_p is None else target_p.insert_paragraph_before(style="Caption")
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_before = Pt(6)
    cap_p.paragraph_format.space_after = Pt(12)
    cap_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    cap_xml = (
        f'<w:pPr {nsdecls("w")}>\n'
        '  <w:pStyle w:val="Caption"/>\n'
        '  <w:jc w:val="center"/>\n'
        '  <w:spacing w:before="120" w:after="240" w:line="360" w:lineRule="auto"/>\n'
        '</w:pPr>'
    )
    cap_p._p.remove(cap_p._p.pPr)
    cap_p._p.append(parse_xml(cap_xml))

    if bookmark_name:
        cap_p._p.append(parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="10{seq_id}" w:name="{bookmark_name}"/>'))

    r1 = cap_p.add_run(f"Hình {seq_label.split('.')[0]}.")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(14)
    r1.bold = True

    seq_num = seq_label.split('.')[1] if '.' in seq_label else str(seq_id)
    seq_xml = (
        f'<w:fldSimple {nsdecls("w")} w:instr="SEQ Hình \\* ARABIC ">\n'
        '  <w:r>\n'
        '    <w:rPr>\n'
        '      <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>\n'
        '      <w:b/>\n'
        '      <w:sz w:val="28"/>\n'
        '    </w:rPr>\n'
        f'    <w:t>{seq_num}</w:t>\n'
        '  </w:r>\n'
        '</w:fldSimple>'
    )
    cap_p._p.append(parse_xml(seq_xml))

    if bookmark_name:
        cap_p._p.append(parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="10{seq_id}"/>'))

    colon_r = cap_p.add_run(": ")
    colon_r.font.name = "Times New Roman"
    colon_r.font.size = Pt(14)
    colon_r.bold = True

    items = [title_content] if isinstance(title_content, str) else title_content
    for item in items:
        if isinstance(item, str):
            r = cap_p.add_run(item)
            r.font.name = "Times New Roman"
            r.font.size = Pt(14)
            r.bold = True
        elif isinstance(item, list):
            for sub in item:
                cap_p._p.append(sub)
        else:
            cap_p._p.append(item)

    return cap_p


def format_table_cell(cell, width_dxa: int, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, font_size_pt=14):
    """Sets standard cell properties: exact width, vertical centering, border, padding, and compact line spacing."""
    tcPr = cell._tc.get_or_add_tcPr()
    tc_xml = (
        f'<w:tcPr {nsdecls("w")}>\n'
        f'  <w:tcW w:w="{width_dxa}" w:type="dxa"/>\n'
        '  <w:vAlign w:val="center"/>\n'
        '  <w:tcBorders>\n'
        '    <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '    <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '    <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '    <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '  </w:tcBorders>\n'
        '  <w:tcMar>\n'
        '    <w:top w:w="80" w:type="dxa"/>\n'
        '    <w:bottom w:w="80" w:type="dxa"/>\n'
        '    <w:left w:w="120" w:type="dxa"/>\n'
        '    <w:right w:w="120" w:type="dxa"/>\n'
        '  </w:tcMar>\n'
        '</w:tcPr>'
    )
    cell._tc.remove(tcPr)
    cell._tc.append(parse_xml(tc_xml))

    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(font_size_pt)
        if bold:
            r.bold = True


def insert_clean_table(doc, target_p, headers, rows_data, col_widths, font_size_pt=13):
    """Creates a cleanly bordered table inserted before target_p."""
    tbl = doc.add_table(rows=len(rows_data) + 1, cols=len(headers)) if target_p is None else target_p.insert_paragraph_before()._p.addprevious(parse_xml(f'<w:tbl {nsdecls("w")}/>'))
    if target_p is not None:
        tbl = docx.table.Table(tbl, doc)

    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = tbl._tbl.tblPr
    tbl_pr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:w="0" w:type="auto"/>'))
    tbl_pr.append(parse_xml(f'<w:tblBorders {nsdecls("w")}>\n  <w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>\n  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>\n  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n</w:tblBorders>'))

    hdr_row = tbl.rows[0]
    hdr_trPr = hdr_row._tr.get_or_add_trPr()
    hdr_trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    hdr_trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

    for c_i, h in enumerate(headers):
        cell = hdr_row.cells[c_i]
        cell.text = h
        format_table_cell(cell, col_widths[c_i], align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size_pt=font_size_pt)

    for r_i, row in enumerate(rows_data):
        b_row = tbl.rows[r_i + 1]
        b_trPr = b_row._tr.get_or_add_trPr()
        b_trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

        for c_i, val in enumerate(row):
            cell = b_row.cells[c_i]
            cell.text = val
            cell_align = WD_ALIGN_PARAGRAPH.CENTER if (c_i == 0 and len(headers) >= 4) else WD_ALIGN_PARAGRAPH.LEFT
            format_table_cell(cell, col_widths[c_i], align=cell_align, bold=(c_i == 0 and len(headers) == 3), font_size_pt=font_size_pt)


def generate_perfect_sources_xml(sources):
    """Generates valid Microsoft Word Bibliography Sources CustomXML with Corporate Authors & IEEE style."""
    lines = ['<?xml version="1.0" encoding="UTF-8" standalone="no"?>']
    lines.append('<b:Sources SelectedStyle="\\IEEE.XSL" StyleName="IEEE" xmlns:b="http://schemas.openxmlformats.org/officeDocument/2006/bibliography" xmlns="http://schemas.openxmlformats.org/officeDocument/2006/bibliography">')

    corporate_map = {
        "SRC-000001": "MITRE ATT&CK",
        "SRC-000027": "National Institute of Standards and Technology (NIST)",
        "SRC-000028": "Defense Advanced Research Projects Agency (DARPA)"
    }

    for s in sources:
        tag = s.source_id.replace("-", "")
        venue = s.venue or ""
        if any(w in venue for w in ["Proceedings", "Conference", "Symposium", "NDSS", "S&P", "CCS", "ICLR", "ICML", "ACSAC", "ISSTA", "ISSRE", "ASE", "IJCNN", "KDD", "IJCAI", "ICDM", "SOSP", "ATC", "USENIX"]):
            stype = "ConferenceProceedings"
        elif any(w in venue for w in ["Journal", "Surveys", "IEEE Transactions", "ACM"]):
            stype = "ArticleInAPeriodical"
        elif "Standard" in venue or "NIST" in venue or "DARPA" in venue or "LANL" in venue:
            stype = "Report"
        elif "arxiv" in venue.lower():
            stype = "InternetSite"
        else:
            stype = "Report"

        guid = str(uuid.uuid5(uuid.NAMESPACE_DNS, f"thesis.sources.{s.source_id}")).upper()

        lines.append("  <b:Source>")
        lines.append(f"    <b:Tag>{tag}</b:Tag>")
        lines.append(f"    <b:SourceType>{stype}</b:SourceType>")
        lines.append(f"    <b:Guid>{{{guid}}}</b:Guid>")
        lines.append(f"    <b:Title>{escape(s.title)}</b:Title>")
        lines.append(f"    <b:Year>{s.year}</b:Year>")
        if stype == "ConferenceProceedings":
            lines.append(f"    <b:ConferenceName>{escape(venue)}</b:ConferenceName>")
        elif stype == "ArticleInAPeriodical":
            lines.append(f"    <b:JournalName>{escape(venue)}</b:JournalName>")
        elif stype == "InternetSite":
            lines.append(f"    <b:InternetSiteTitle>{escape(venue)}</b:InternetSiteTitle>")
        else:
            lines.append(f"    <b:Institution>{escape(venue)}</b:Institution>")

        if s.source_id in corporate_map:
            lines.append("    <b:Author>")
            lines.append("      <b:Author>")
            lines.append(f"        <b:Corporate>{escape(corporate_map[s.source_id])}</b:Corporate>")
            lines.append("      </b:Author>")
            lines.append("    </b:Author>")
        elif s.source_id == "SRC-000029":
            lines.append("    <b:Author>")
            lines.append("      <b:Author>")
            lines.append("        <b:NameList>")
            lines.append("          <b:Person><b:Last>Kent</b:Last><b:First>Alexander D.</b:First></b:Person>")
            lines.append("        </b:NameList>")
            lines.append("      </b:Author>")
            lines.append("    </b:Author>")
            lines.append("    <b:Institution>Los Alamos National Laboratory</b:Institution>")
        else:
            lines.append("    <b:Author>")
            lines.append("      <b:Author>")
            lines.append("        <b:NameList>")
            for author_name in s.authors:
                parts = author_name.strip().split()
                if len(parts) > 1:
                    first = " ".join(parts[:-1])
                    last = parts[-1]
                    lines.append(f"          <b:Person><b:Last>{escape(last)}</b:Last><b:First>{escape(first)}</b:First></b:Person>")
                else:
                    lines.append(f"          <b:Person><b:Last>{escape(author_name)}</b:Last></b:Person>")
            lines.append("        </b:NameList>")
            lines.append("      </b:Author>")
            lines.append("    </b:Author>")

        lines.append("  </b:Source>")

    lines.append("</b:Sources>")
    return "\n".join(lines)
