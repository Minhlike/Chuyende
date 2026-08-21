"""
Master Document Builder - Chapter 1 & Chapter 2 Section 2.1
Assembles the complete thesis into Word 2016 (.docx) with Word Desktop COM automation,
generating updated TOC, List of Figures, List of Tables, and Bibliography.
"""

import os
import shutil
import zipfile
import win32com.client as win32
import pythoncom
from datetime import datetime, timezone
from pathlib import Path
from xml.sax.saxutils import escape

import docx
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

from research_agent.storage.db import DatabaseManager
from research_agent.storage.repository import ResearchRepository
from research_agent.visuals.academic_diagram_renderer import generate_all_figures
from research_agent.composition.academic_builder_base import (
    latex_to_clean_omml,
    make_citation_element,
    make_ref_element,
    add_table_caption,
    add_figure_caption,
    insert_clean_table,
    generate_perfect_sources_xml,
)


def build_master_thesis_document(target_file: str = r"D:\Research\Chuyên đề chuyên sâu - Copy.docx"):
    repo = ResearchRepository(DatabaseManager())
    sources = repo.list_sources()

    # Sync Master Sources.xml in AppData
    try:
        master_xml_str = generate_perfect_sources_xml(sources)
        master_xml_path = os.path.expandvars(r"%APPDATA%\Microsoft\Bibliography\Sources.xml")
        with open(master_xml_path, "w", encoding="utf-8") as f:
            f.write(master_xml_str)
        print(f"[0/6] Synchronized %APPDATA%\\Microsoft\\Bibliography\\Sources.xml with {len(sources)} verified sources.")
    except Exception as e:
        print(f"[WARNING] Could not sync master Sources.xml: {e}")

    # Generate high-resolution figures
    fig_paths = generate_all_figures()
    fig1_path, fig2_path, fig3_path, fig4_path, fig5_path = fig_paths

    target_path = Path(target_file)
    template_path = Path(r"D:\Research\Chuyên đề chuyên sâu - Copy.backup.docx")

    print(f"[1/6] Loading template from: {template_path}")
    doc = docx.Document(str(template_path))

    # Identify insertion point (Cover frame and TOC preserved)
    insertion_p_idx = None
    for idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip().lower()
        if "lời nói đầu" in txt:
            insertion_p_idx = idx
            break

    if insertion_p_idx is None:
        target_p = doc.paragraphs[-1]
    else:
        target_p = doc.paragraphs[insertion_p_idx]

    # Clean old body paragraphs
    cleaned_count = 0
    p_curr = target_p
    while p_curr is not None:
        p_next = p_curr._p.getnext()
        p_curr._p.getparent().remove(p_curr._p)
        cleaned_count += 1
        p_curr = docx.text.paragraph.Paragraph(p_next, doc) if p_next is not None and p_next.tag.endswith('p') else None
    print(f"[2/6] Cleaned {cleaned_count} old body paragraphs. Insertion target ready.")

    target_p = None  # Append to end

    # Helper paragraph builders
    def add_h1(text):
        p = doc.add_paragraph(style="Heading 1")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(16)
        r.bold = True
        return p

    def add_h2(text):
        p = doc.add_paragraph(style="Heading 2")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
        r.bold = True
        return p

    def add_h3(text):
        p = doc.add_paragraph(style="Heading 3")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
        r.bold = True
        r.italic = True
        return p

    def add_p(runs, bold_prefix=None, first_line_indent=True):
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Cm(1.27) if first_line_indent else Cm(0)

        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.font.name = "Times New Roman"
            r_pre.font.size = Pt(14)
            r_pre.bold = True

        if isinstance(runs, str):
            runs = [runs]

        for item in runs:
            if isinstance(item, str):
                r = p.add_run(item)
                r.font.name = "Times New Roman"
                r.font.size = Pt(14)
            elif isinstance(item, list):
                for sub in item:
                    p._p.append(sub)
            else:
                p._p.append(item)
        return p

    def add_display_equation(omml_or_latex):
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.keep_with_next = True
        if isinstance(omml_or_latex, str):
            elem = latex_to_clean_omml(omml_or_latex)
        else:
            elem = omml_or_latex
        p._p.append(elem)
        return p

    def add_figure_image(image_path, width_inches=6.2):
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.keep_with_next = True
        r = p.add_run()
        r.add_picture(str(image_path), width=Inches(width_inches))
        return p

    # =========================================================================
    # PREAMBLE: LỜI NÓI ĐẦU
    # =========================================================================
    add_h1("Lời nói đầu")
    add_p(
        "Trong bối cảnh hạ tầng công nghệ thông tin doanh nghiệp phát triển mạnh mẽ và ngày càng chuyển dịch sâu rộng sang kiến trúc phân tán, điện toán đám mây và microservices, các cuộc tấn công mạng có chủ đích và có tổ chức (Advanced Persistent Threats - APT) đã trở thành mối đe dọa thường trực, phức tạp và nguy hiểm nhất đối với an ninh quốc gia và sự toàn vẹn của các tổ chức kinh tế. Các tác nhân đe dọa APT hiện đại không chỉ sử dụng mã độc tinh vi mà còn khai thác triệt để các công cụ quản trị hệ thống hợp lệ (Living-off-the-Land - LotL), che giấu hành vi xâm nhập trong các luồng hoạt động thông thường và kéo dài thời gian hoạt động qua nhiều tuần, nhiều tháng. Trước thực tế đó, dữ liệu nhật ký hệ thống (System Audit Logs) và telemetry giám sát điểm cuối đóng vai trò là nguồn bằng chứng lịch sử toàn diện và trung thực nhất giúp các hệ thống phòng thủ phát hiện sớm các chiến dịch tấn công đa giai đoạn."
    )
    add_p(
        "Chuyên đề nghiên cứu chuyên sâu này tập trung giải quyết bài toán cốt lõi: Thiết lập cơ sở phương pháp luận và xây dựng khung biểu diễn đặc trưng dữ liệu log đa góc nhìn, bảo toàn ngữ cảnh an toàn và tối ưu hóa cho môi trường xử lý dòng thời gian thực. Bằng cách phân định ranh giới phương pháp luận ba tầng và xác lập Hợp đồng Biểu diễn (Representation Contract), nghiên cứu mở đường cho việc xây dựng các biểu diễn vector đặc trưng chất lượng cao, bền vững trước nhiễu quản trị và trôi dạt phân phối, phục vụ đắc lực cho công tác phát hiện và điều tra nguồn gốc an ninh mạng."
    )

    # =========================================================================
    # CHAPTER 1: TỔNG QUAN VỀ PHƯƠNG PHÁP TRÍCH XUẤT ĐẶC TRƯNG DỮ LIỆU LOG... (FROZEN)
    # =========================================================================
    add_h1("Chương 1. TỔNG QUAN VỀ PHƯƠNG PHÁP TRÍCH XUẤT ĐẶC TRƯNG DỮ LIỆU LOG VÀ THÁCH THỨC BẢO TOÀN NGỮ CẢNH AN TOÀN")

    # --- 1.1 ---
    add_h2("1.1. Bài toán biểu diễn log trong phát hiện tấn công đa giai đoạn")

    # --- 1.1.1 ---
    add_h3("Không gian dữ liệu log doanh nghiệp: tốc độ cao, mất cân bằng cực đoan và phân phối biến đổi")
    add_p([
        "Trong các Trung tâm Điều hành An ninh mạng (Security Operations Center - SOC) hiện đại, khối lượng dữ liệu nhật ký kiểm toán và sự kiện giám sát điểm cuối thu thập từ hạ tầng doanh nghiệp tăng trưởng theo cấp số nhân, đạt quy mô hàng chục triệu đến hàng tỷ bản ghi mỗi ngày ",
        make_citation_element([9, 8]),
        ". Không gian dữ liệu này sở hữu ba đặc tính kỹ thuật khắc nghiệt ",
        make_citation_element([18, 1]),
        ": (1) Tốc độ sinh dữ liệu cao (High Velocity) đòi hỏi các thuật toán xử lý phải vận hành trực tuyến với độ trễ thấp và tài nguyên tính toán giới hạn; (2) Tỷ lệ mất cân bằng cực đoan (Extreme Imbalance) khi các sự kiện độc hại thực sự chỉ chiếm một phần rất nhỏ (thường dưới 0.01%) so với khối lượng khổng lồ các hoạt động quản trị và vận hành hợp lệ; (3) Tính biến đổi phân phối liên tục (Distribution Drift) xuất phát từ các bản cập nhật phần mềm, sự thay đổi chính sách người dùng và các kỹ thuật tấn công liên tục biến hóa."
    ])
    add_p([
        "Không gian dữ liệu nhật ký trong môi trường mạng doanh nghiệp được tổng hợp từ ba nguồn telemetry chính ",
        make_citation_element([9, 10]),
        ": Nhóm thứ nhất là nhật ký kiểm toán hệ điều hành máy chủ và điểm cuối (Host & Endpoint Audit Logs) như Linux Auditd, Windows Event Logs / Sysmon và eBPF, cung cấp chi tiết ở mức hạt nhân về các lệnh gọi hệ thống (system calls), hành vi tạo tiến trình (Process Creation - Sysmon Event ID 1), nạp thư viện động (Image Load - Sysmon Event ID 7), truy vết thao tác tệp tin ",
        make_citation_element([8, 9]),
        ", sửa đổi cấu hình registry (RegSetValue - Sysmon Event ID 13), cùng các thao tác mở và kết nối socket mạng (connect, accept - Sysmon Event ID 3) ",
        make_citation_element([6, 7]),
        ". Nhóm thứ hai là nhật ký luồng mạng (Network Flow & Protocol Logs), được thu thập từ Zeek, Suricata hoặc NetFlow/IPFIX, cung cấp siêu dữ liệu kết nối giữa các nút mạng, giao dịch DNS, chứng chỉ TLS/SSL và thông lượng gói tin. Nhóm thứ ba là nhật ký ứng dụng và dịch vụ (Application & Service Logs), phát sinh từ máy chủ web (Nginx, Apache), cơ sở dữ liệu, dịch vụ phân tán (HDFS) cùng hệ thống điều phối container (Kubernetes Audit Logs) ",
        make_citation_element([6, 7]),
        "."
    ])
    add_p([
        "Tính dị thể sâu sắc của dữ liệu đặt ra bài toán khoa học về việc lựa chọn đơn vị quan sát (Unit of Observation) phù hợp cho mô hình học biểu diễn, như được hệ thống hóa trực quan trong ",
        make_ref_element("BK_FIG_1_001", "Hình 1.1"),
        " và chi tiết hóa trong ",
        make_ref_element("BK_TBL_1_001", "Bảng 1.1"),
        ":"
    ])

    # FIGURE 1.1
    add_figure_image(fig1_path, width_inches=6.2)
    add_figure_caption(doc, target_p, "1.1", "Phân cấp các đơn vị quan sát và sự đánh đổi giữa ngữ cảnh và chi phí tính toán trong biểu diễn dữ liệu log (Nguồn: Tác giả tổng hợp)", bookmark_name="BK_FIG_1_001", seq_id=1)

    # TABLE 1.1
    add_table_caption(doc, target_p, 1, "So sánh chi tiết đặc tính kỹ thuật giữa các mức biểu diễn dữ liệu log hệ thống", bookmark_name="BK_TBL_1_001")
    t1_headers = ["Mức biểu diễn", "Đặc tính kỹ thuật", "Bảo toàn ngữ cảnh", "Chi phí tính toán"]
    t1_rows = [
        ["Từ tố (Token Level)", "Chuỗi con, từ khóa tĩnh, tham số rời rạc, mã lỗi hex", "Cực thấp (Cục bộ)", "Thấp (Tuyến tính O(N))"],
        ["Sự kiện (Event Level)", "Bản ghi đơn lẻ, định danh Event ID, nhãn mốc thời gian", "Thấp (Từng thời điểm)", "Thấp O(N)"],
        ["Chuỗi / Phiên (Sequence)", "Chuỗi trật tự thời gian theo Session/Process ID", "Trung bình (Chuỗi cục bộ)", "Trung bình O(N·L)"],
        ["Thực thể (Entity Level)", "Hành vi gom cụm theo Host, User hoặc Cây tiến trình", "Khá (Phạm vi thực thể)", "Khá O(N·log N)"],
        ["Đồ thị nguồn gốc (Graph)", "Đồ thị luồng phụ thuộc đa thực thể (Process, File, Socket)", "Rất cao (Toàn cục hệ thống)", "Cao O(|V| + |E|)"]
    ]
    t1_widths = [1800, 3200, 2200, 2100]
    insert_clean_table(doc, target_p, t1_headers, t1_rows, t1_widths, font_size_pt=13)

    add_p([
        "Sự phân cấp này thể hiện sự đánh đổi trực tiếp (Fundamental Trade-off): các đơn vị quan sát ở mức thấp (từ tố, sự kiện) có ưu thế về tốc độ xử lý nhưng đánh mất mối liên kết nhân quả dài hạn giữa các thực thể; ngược lại, mức đồ thị nguồn gốc phản ánh toàn diện quan hệ phụ thuộc nhân quả giữa tiến trình, tệp tin và luồng mạng nhưng đòi hỏi chi phí tính toán và bộ nhớ rất lớn để duy trì trạng thái đồ thị ",
        make_citation_element([2, 18]),
        ". Do đó, một mô hình biểu diễn đặc trưng tối ưu cần có khả năng nắm bắt được thông tin ngữ cảnh đa thực thể mà vẫn duy trì chi phí tính toán khả thi trong môi trường xử lý dòng (Streaming Environment) ",
        make_citation_element([11]),
        "."
    ], first_line_indent=False)
    add_p([
        "Bên cạnh tốc độ và tính dị thể, hiện tượng trôi dạt phân phối (Distribution Drift) là một trong những nguyên nhân hàng đầu khiến các mô hình học máy bị suy giảm hiệu năng nghiêm trọng sau khi triển khai thực tế ",
        make_citation_element([2]),
        ". Về mặt toán học, trôi dạt trong biểu diễn log xuất hiện dưới bốn hình thái chủ đạo ",
        make_citation_element([22, 23]),
        ":"
    ])
    add_p([
        "Tần suất xuất hiện của các mẫu sự kiện log thay đổi theo thời gian do sự biến động của khối lượng công việc (Workload) hoặc chu kỳ vận hành của doanh nghiệp: ",
        latex_to_clean_omml(r"P_t(X_{\text{template}}) \neq P_{t+1}(X_{\text{template}})"),
        "."
    ], bold_prefix="• Covariate Drift: ")
    add_p([
        "Hành vi của hệ thống bị biến đổi do nâng cấp phần mềm, vá lỗi hệ điều hành hoặc cập nhật phiên bản ứng dụng, dẫn đến sự thay đổi phân phối điều kiện của các quan sát đối với trạng thái an ninh: ",
        latex_to_clean_omml(r"P_t(X \mid Y) \neq P_{t+1}(X \mid Y)"),
        "."
    ], bold_prefix="• Concept Drift: ")
    add_p([
        "Kẻ tấn công chủ động thay đổi chiến thuật, kỹ thuật và quy trình tấn công (TTPs), hoặc sử dụng các kỹ thuật lẩn tránh (Evasion Techniques), tấn công bắt chước (Mimicry Attacks) nhằm ngụy trang hành vi độc hại thành các chuỗi sự kiện quản trị hợp lệ: ",
        latex_to_clean_omml(r"P_t(Y \mid X) \neq P_{t+1}(Y \mid X)"),
        "."
    ], bold_prefix="• Adversarial Drift: ")
    add_p([
        "Không gian vector tiềm ẩn z bị suy giảm năng lực phân tách do dữ liệu đầu vào trôi dạt khỏi vùng phân phối huấn luyện ban đầu: ",
        latex_to_clean_omml(r"P_t(\mathbf{z} \mid X) \neq P_{t+1}(\mathbf{z} \mid X)"),
        "."
    ], bold_prefix="• Representation Drift: ")

    # --- 1.1.2 ---
    add_h3("Hành vi tấn công đa giai đoạn và ánh xạ đa nhãn MITRE ATT&CK")
    add_p(
        "Trong lịch sử nghiên cứu phát hiện xâm nhập, các tiếp cận ban đầu thường dựa trên giả định chuỗi tấn công đơn tuyến (Cyber Kill Chain) hoặc các mô hình xích Markov tuyến tính để xâu chuỗi tuần tự các giai đoạn tấn công từ Thâm nhập ban đầu (Initial Access), Thực thi (Execution), Duy trì (Persistence), Leo thang đặc quyền (Privilege Escalation) cho đến Đánh cắp dữ liệu (Exfiltration)."
    )
    add_p([
        "Tuy nhiên, trên thực tế vận hành và theo khung phân loại chuẩn hóa của ma trận tri thức an ninh MITRE ATT&CK (Phiên bản Enterprise v19.1, 28/04/2026) ",
        make_citation_element([1]),
        ", các chiến dịch tấn công có chủ đích (APT) mang bản chất phi tuyến tính sâu sắc ",
        make_citation_element([9]),
        ": (1) Nhảy cóc giai đoạn (Step Skipping): Kẻ tấn công có thể khai thác trực tiếp lỗ hổng thực thi mã từ xa để trích xuất dữ liệu ra ngoài mà không cần thiết lập cơ chế duy trì hay di chuyển ngang; (2) Lặp vòng kỹ thuật (Tactic Looping & Interleaving): Kỹ thuật thu thập thông tin nội bộ (Discovery) thường được lặp lại nhiều lần xen kẽ giữa các bước leo thang đặc quyền và chiếm đoạt thông tin xác thực; (3) Phân nhánh tiến trình song song (Parallel Branching): Kẻ tấn công có thể khởi tạo đồng thời nhiều luồng tiến trình con độc lập trên các tiến trình hợp lệ khác nhau nhằm phân tán sự theo dõi của hệ thống phòng thủ."
    ])
    add_p([
        "Do đó, chuyên đề xác lập nguyên tắc: Ma trận MITRE ATT&CK được mô hình hóa thành một Không gian Bằng chứng Hành vi Đa chiều (Multi-label Behavioral Evidence Space), được minh họa trực quan trong ",
        make_ref_element("BK_FIG_1_002", "Hình 1.2"),
        ":"
    ])
    add_display_equation(r"\mathcal{Y} \subseteq \{0, 1\}^{|\mathcal{T}|}")

    # FIGURE 1.2
    add_figure_image(fig2_path, width_inches=6.2)
    add_figure_caption(
        doc,
        target_p,
        "1.2",
        [
            "Mô hình Không gian Bằng chứng Hành vi Đa chiều MITRE ATT&CK và các đặc trưng phi tuyến tính trong tấn công APT (Nguồn: Tác giả tổng hợp dựa trên MITRE ATT&CK ",
            make_citation_element([1]),
            " và Inam et al. ",
            make_citation_element([9]),
            ")"
        ],
        bookmark_name="BK_FIG_1_002",
        seq_id=2
    )

    add_p([
        "Trong đó mỗi chiến thuật (Tactic) đại diện cho mục tiêu chiến thuật của kẻ tấn công (trả lời câu hỏi 'Tại sao') và mỗi kỹ thuật (Technique) đại diện cho phương thức thực thi cụ thể (trả lời câu hỏi 'Làm thế nào'); một chuỗi sự kiện hoặc cây tiến trình có thể đồng thời kích hoạt nhiều nhãn chiến thuật và kỹ thuật tại cùng một thời điểm quan sát ",
        make_citation_element([1]),
        "."
    ], first_line_indent=False)
    add_p([
        "Về mặt dữ liệu thực nghiệm, việc mô hình hóa hành vi tấn công đòi hỏi phải phân định chính xác đặc tính gán nhãn và mức độ hạt (Label Granularity) của từng bộ dữ liệu chuẩn ",
        make_citation_element([18]),
        ": (1) DARPA Transparent Computing (TC E3/E5) ",
        make_citation_element([28]),
        " cung cấp dữ liệu kiểm toán hệ thống mức hạt nhân chi tiết với các kịch bản APT thực tế được gán nhãn ở mức độ hạt tiến trình/luồng phụ thuộc (Fine-grained Ground Truth); (2) LANL Unified Host and Network Dataset ",
        make_citation_element([29]),
        " phản ánh môi trường doanh nghiệp quy mô lớn với hàng tỷ sự kiện xác thực và luồng mạng, gán nhãn theo sự kiện và cửa sổ thời gian; (3) HDFS và BGL Datasets ",
        make_citation_element([3, 6]),
        " đại diện cho nhật ký hệ thống phân tán và siêu máy tính, được gán nhãn bất thường ở mức khối dữ liệu (Block-level) hoặc mức dòng log đơn lẻ."
    ])
    add_p([
        "Đặc biệt, sự xuất hiện của nhiễu từ hành vi quản trị viên (Admin-Noise) là một thách thức then chốt ",
        make_citation_element([9, 18]),
        ". Các quản trị viên hệ thống thường xuyên sử dụng PowerShell, SSH, WMI và các công cụ dòng lệnh nội bộ tương tự như kẻ tấn công APT, tạo ra sự trùng lặp lớn trong không gian đặc trưng hành vi và gây ra nhiều cảnh báo giả nếu mô hình chỉ học các mẫu bề mặt mà không nắm bắt được ngữ cảnh nhân quả sâu ",
        make_citation_element([24]),
        "."
    ])

    # --- 1.1.3 ---
    add_h3("Các mức biểu diễn dữ liệu và Hợp đồng Biểu diễn (Representation Contract)")
    add_p([
        "Để khắc phục tình trạng các nghiên cứu tiền nhiệm thường đồng nhất việc trích xuất đặc trưng với thuật toán phát hiện cụ thể, chuyên đề đề xuất một khung khái niệm chính thức mang tên Hợp đồng Biểu diễn (Representation Contract), ký hiệu là ",
        latex_to_clean_omml(r"\mathcal{C}_{\text{representation}}"),
        ", đóng vai trò như một bản đặc tả ràng buộc toán học chặt chẽ giữa dữ liệu thô và không gian vector tiềm ẩn z:"
    ])
    add_display_equation(r"\mathcal{C}_{\text{representation}} = \langle \mathcal{P}_{\text{preserve}}, \mathcal{I}_{\text{invariant}}, \mathcal{E}_{\text{exclude}} \rangle")
    add_p(
        "Khung Hợp đồng Biểu diễn thiết lập ba nhóm điều kiện ràng buộc bắt buộc mà không gian vector biểu diễn z phải đồng thời thỏa mãn:"
    )
    add_p(
        "Tập hợp các thuộc tính và tương quan an ninh bắt buộc phải được bảo toàn nguyên vẹn trong không gian vector biểu diễn. Bao gồm: ngữ nghĩa sâu của các tham số dòng lệnh quan trọng, trật tự thời gian và mối quan hệ phụ thuộc nhân quả giữa các thực thể hệ thống (Process-File-Socket), cấu trúc tô-pô cục bộ của đồ thị nguồn gốc và các đặc trưng nhận diện chiến thuật/kỹ thuật MITRE ATT&CK.",
        bold_prefix="1. Nhóm Điều kiện Bảo toàn (PRESERVE - P_preserve): "
    )
    add_p(
        "Tập hợp các biến đổi hình thức hoặc biến động môi trường mà không gian vector z phải bất biến tuyệt đối. Bao gồm: sự biến đổi cú pháp vô hại của các chuỗi văn bản log (định dạng dấu thời gian, khoảng trắng, thứ tự các trường không mang nghĩa an ninh), sự xáo trộn mã định danh tiến trình ngẫu nhiên (PID/PPID), tên tài khoản người dùng hoặc địa chỉ IP cục bộ không ảnh hưởng đến bản chất luồng thực thi, và các biến động định kỳ về khối lượng lưu lượng nền (Workload fluctuations).",
        bold_prefix="2. Nhóm Điều kiện Bất biến (INVARIANT - I_invariant): "
    )
    add_p([
        "Tập hợp các thông tin giả tạo và biến số ngoài miền bắt buộc phải bị triệt tiêu hoàn toàn khỏi không gian vector z nhằm ngăn ngừa rủi ro học đường tắt (Shortcut Learning) ",
        make_citation_element([2]),
        ". Bao gồm: các mẫu định dạng đặc thù của môi trường thử nghiệm (Testbed-specific artifacts), các biến số gây rò rỉ phân vùng (Partition-leakage variables như dấu thời gian tuyệt đối của máy thí nghiệm, hostname cố định của môi trường lab) ",
        make_citation_element([25, 27]),
        ", và các đặc trưng tương quan giả (Spurious correlations) có thể khiến mô hình đạt độ chính xác ảo trên tập kiểm thử nhưng thất bại hoàn toàn khi triển khai thực tế."
    ], bold_prefix="3. Nhóm Điều kiện Triệt tiêu (EXCLUDE - E_exclude): ")
    add_p([
        "Các đặc tả chi tiết của Hợp đồng Biểu diễn được hệ thống hóa trong ",
        make_ref_element("BK_TBL_1_002", "Bảng 1.2"),
        ":"
    ])

    # TABLE 1.2
    add_table_caption(doc, target_p, 2, "Đặc tả Hợp đồng Biểu diễn (Representation Contract) cho vector đặc trưng z", bookmark_name="BK_TBL_1_002")
    t2_headers = ["Thành phần hợp đồng", "Mục tiêu ràng buộc", "Đối tượng áp dụng cụ thể trong dữ liệu log", "Cơ chế kiểm chứng"]
    t2_rows = [
        ["PRESERVE (Bảo toàn)", "Bảo tồn đầy đủ tín hiệu ngữ cảnh an ninh phân biệt", "Ngữ nghĩa tham số động, thứ tự nhân quả, tô-pô đồ thị, nhãn ATT&CK", "Linear Probing, Mutual Information"],
        ["INVARIANT (Bất biến)", "Bất biến trước biến đổi hình thức và trôi dạt vô hại", "Định dạng timestamp, PID ngẫu nhiên, hoán vị cú pháp vô hại", "Invariance Loss, Data Augmentation"],
        ["EXCLUDE (Triệt tiêu)", "Loại bỏ hoàn toàn rò rỉ phân vùng và biến số học đường tắt", "Artifacts môi trường lab, hostname cố định, thông tin định danh nhạy cảm (PII)", "Snooping Test, DP Guarantees"]
    ]
    t2_widths = [1800, 2400, 3100, 2000]
    insert_clean_table(doc, target_p, t2_headers, t2_rows, t2_widths, font_size_pt=13)

    add_p([
        "Đi đôi với Hợp đồng Biểu diễn được quy định trong ",
        make_ref_element("BK_TBL_1_002", "Bảng 1.2"),
        ", chuyên đề thiết lập nguyên tắc phân định ranh giới ba tầng phương pháp luận (Three-Tier Methodological Separation) ",
        make_citation_element([2, 22]),
        ", như được minh họa tổng quát trong ",
        make_ref_element("BK_FIG_1_003", "Hình 1.3"),
        ":"
    ])

    # FIGURE 1.3
    add_figure_image(fig3_path, width_inches=6.2)
    add_figure_caption(doc, target_p, "1.3", "Khung phân định ranh giới ba tầng phương pháp luận và vị trí trọng tâm của không gian vector biểu diễn z (Nguồn: Tác giả đề xuất)", bookmark_name="BK_FIG_1_003", seq_id=3)

    add_p(
        "Đảm nhiệm tiền xử lý dữ liệu thô, phân tích cú pháp sơ bộ, chuẩn hóa kiểu dữ liệu trường và làm sạch dữ liệu. Tầng này không gánh vác nhiệm vụ học biểu diễn ngữ cảnh an ninh sâu.",
        bold_prefix="1. Tầng 1 — Trích xuất đặc trưng cơ sở (Feature Extraction): "
    )
    add_p([
        "Thiết lập ánh xạ ",
        latex_to_clean_omml(r"f_\theta: X \to \mathbf{z} \in \mathbb{R}^d"),
        " từ cấu trúc chuỗi sự kiện và đồ thị nguồn gốc sang không gian vector tiềm ẩn. Toàn bộ năng lực bảo toàn ngữ nghĩa an ninh và tính bất biến được đóng gói trọn vẹn bên trong vector z ",
        make_citation_element([2]),
        "."
    ], bold_prefix="2. Tầng 2 — Học không gian biểu diễn (Representation Learning — Trọng tâm Chuyên đề): ")
    add_p(
        "Đánh giá chất lượng của vector biểu diễn z thông qua các bộ thăm dò tuyến tính đóng băng tham số (Frozen Linear Probing):",
        bold_prefix="3. Tầng 3 — Phát hiện và phân loại hạ nguồn (Downstream Detection): "
    )
    add_display_equation(r"\hat{y} = \sigma(\mathbf{W}^\top \mathbf{z} + \mathbf{b})")
    add_p([
        "Trong đó ma trận trọng số W và vector bias b được huấn luyện trên không gian vector z cố định (frozen parameters θ), bảo đảm bộ phân loại hạ nguồn không làm thay nhiệm vụ trích xuất đặc trưng của Tầng 2 ",
        make_citation_element([2]),
        "."
    ], first_line_indent=False)

    # --- 1.2 ---
    add_h2("1.2. Phân tích so sánh các nhóm phương pháp hiện đại")
    add_p([
        "Nhằm định vị chính xác các đóng góp kỹ thuật và cơ sở lý luận của chuyên đề, Mục này tiến hành khảo sát và đánh giá có hệ thống ba nhóm phương pháp trích xuất đặc trưng log hiện đại trong tài liệu khoa học ",
        make_citation_element([6, 9, 10]),
        ": (1) Nhóm phương pháp thống kê và cú pháp; (2) Nhóm phương pháp ngữ nghĩa chuỗi (Semantic–Sequential); (3) Nhóm phương pháp học biểu diễn đồ thị nguồn gốc (Graph Representation Learning)."
    ])

    # --- 1.2.1 ---
    add_h3("Phương pháp thống kê và cú pháp: Event Count, Frequency, Entropy và Template Features")
    add_p([
        "Nhóm tiếp cận sớm nhất và phổ biến nhất trong thực tế công nghiệp dựa trên việc phân tích cấu trúc cú pháp của log thông qua các bộ log parser tự động (như Drain, Spell, LenMa, AEL) để tách biệt phần văn bản tĩnh (Log Template) và phần tham số biến động ",
        make_citation_element([6, 7]),
        ". Sau khi chuyển đổi các thông điệp văn bản thô thành các mã sự kiện rời rạc ",
        latex_to_clean_omml(r"\mathcal{E} = \{e_1, e_2, \dots, e_M\}"),
        ", các phương pháp thống kê xây dựng vector đặc trưng cho mỗi cửa sổ quan sát thời gian W (Time-based Window) hoặc cửa sổ số lượng (Count-based Window) ",
        make_citation_element([6]),
        ":"
    ])
    add_display_equation(r"\mathbf{x} = [c(e_1), c(e_2), \dots, c(e_M)]^\top \in \mathbb{R}^M")
    add_p([
        "Trong đó ",
        latex_to_clean_omml(r"c(e_i)"),
        " là số lần xuất hiện của sự kiện ",
        latex_to_clean_omml(r"e_i"),
        ". Bên cạnh đếm tần suất đơn thuần, các trọng số TF-IDF (Term Frequency - Inverse Document Frequency) hoặc độ hỗn loạn thông tin Shannon (Information Entropy) cũng được áp dụng nhằm nhấn mạnh các sự kiện hiếm gặp ",
        make_citation_element([6]),
        ":"
    ], first_line_indent=False)
    add_display_equation(r"H(W) = -\sum_{i=1}^M p(e_i) \log_2 p(e_i)")
    add_p([
        "Trên không gian vector tần suất này, các mô hình phát hiện bất thường kinh điển được triển khai rộng rãi: (1) Principal Component Analysis - PCA (Xu et al., SOSP 2009) ",
        make_citation_element([31]),
        " phân rã không gian vector ",
        latex_to_clean_omml(r"\mathbb{R}^M"),
        " thành không gian con chuẩn tắc ",
        latex_to_clean_omml(r"\mathcal{S}_n"),
        " và không gian con phần dư ",
        latex_to_clean_omml(r"\mathcal{S}_r"),
        ", nhận diện bất thường khi năng lượng chiếu lên phần dư vượt ngưỡng kiểm định:"
    ])
    add_display_equation(r"\mathbf{x}_a = (\mathbf{I} - \mathbf{P} \mathbf{P}^\top) \mathbf{x}, \quad \|\mathbf{x}_a\|^2 > \gamma_\alpha")
    add_p([
        "(2) Invariant Mining (Lou et al., USENIX ATC 2010) ",
        make_citation_element([32]),
        " tự động khai phá các phương trình bất biến tuyến tính phản ánh mối quan hệ bảo toàn logic giữa các bước thực thi:"
    ])
    add_display_equation(r"\mathbf{A} \mathbf{x} = \mathbf{0}")
    add_p([
        "Mặc dù sở hữu ưu điểm về hiệu năng tính toán (độ phức tạp tuyến tính O(N)), nhóm phương pháp thống kê và cú pháp bộc lộ hai điểm nghẽn phương pháp luận quan trọng ",
        make_citation_element([6, 8]),
        ": (1) Mất mát ngữ nghĩa an ninh do trừu tượng hóa tham số: các bộ log parser dựa trên biểu thức chính quy thường thay thế các tham số biến động (địa chỉ IP, đường dẫn tệp tin, tham số dòng lệnh) bằng ký tự đại diện <*> khiến nhiều thông tin an ninh mang tính phân biệt cao bị lược bỏ; (2) Lan truyền và khuếch đại sai số cú pháp (Parser Error Propagation): khi gặp các định dạng log mới chưa từng xuất hiện (unseen logs), parser có thể phân tách không chính xác, dẫn đến hiện tượng sinh ra các mẫu sự kiện giả lập hoặc gộp nhầm các sự kiện khác biệt, làm xáo trộn cấu trúc không gian vector x."
    ])

    # --- 1.2.2 ---
    add_h3("Phương pháp semantic–sequential: Embeddings, Self-Supervised Learning, Transformer và Parsing-Free")
    add_p([
        "Nhằm khắc phục sự cứng nhắc của các vector đếm tần suất và tận dụng trật tự xuất hiện của các sự kiện, nhóm phương pháp Semantic–Sequential mô hình hóa luồng log tương tự như các chuỗi ngôn ngữ tự nhiên, tích hợp các kỹ thuật nhúng từ (Word Embeddings) và học tự giám sát (Self-Supervised Learning) để nắm bắt phụ thuộc ngữ cảnh dài hạn ",
        make_citation_element([3, 4, 5, 33, 34]),
        "."
    ])
    add_p([
        "Khởi đầu cho hướng nghiên cứu này là mô hình DeepLog ",
        make_citation_element([3]),
        ". DeepLog sử dụng mạng bộ nhớ dài-ngắn hạn (LSTM) để mô hình hóa chuỗi Event ID như một bài toán dự báo phần tử tiếp theo (Next-Event Prediction). Tại mỗi bước thời gian t, mô hình ước lượng phân phối xác suất có điều kiện của sự kiện tiếp theo ",
        latex_to_clean_omml(r"P(e_t \mid e_{t-k}, \dots, e_{t-1})"),
        ". Nếu sự kiện thực tế không nằm trong tập g sự kiện có xác suất cao nhất được mô hình dự đoán, hệ thống sẽ phát tín hiệu cảnh báo bất thường:"
    ])
    add_display_equation(r"\hat{\mathcal{E}}_t = \operatorname{arg\,top-}g_{e \in \mathcal{E}} P(e \mid e_{t-k}, \dots, e_{t-1})")
    add_p([
        "Song song đó, DeepLog xây dựng mô hình LSTM thứ hai dựa trên phân phối chuẩn nhiều chiều để kiểm tra sự bất thường về giá trị tham số số học (Parameter Value Anomaly) ",
        make_citation_element([3]),
        "."
    ], first_line_indent=False)
    add_p([
        "Các công trình kế tiếp đã nâng cấp cơ chế biểu diễn ngữ nghĩa: (1) LogAnomaly (Meng et al., IJCAI 2019) ",
        make_citation_element([33]),
        " đề xuất Template2Vec, trích xuất vector ngữ nghĩa cho từng mẫu log thông qua Word2Vec/FastText kết hợp trọng số d-IDF, giúp nhận biết sự tương đồng giữa các thông điệp có cấu trúc từ ngữ tương đương; (2) Logsy (Nedelkoski et al., ICDM 2020) ",
        make_citation_element([34]),
        " sử dụng hàm mất mát phân loại ngoại lai (Outlier Classification Loss) trên dữ liệu log từ các hệ thống phụ trợ để định hình biên giới phân tách cho lớp bình thường; (3) LogBERT (Guo et al., IJCNN 2021) ",
        make_citation_element([4]),
        " khai thác Transformer hai chiều với hai tác vụ học tự giám sát: Dự đoán sự kiện log bị che (Masked Log Event Prediction) và Dự đoán phân bố khối lượng log (Volume Anomaly Prediction); (4) Nhóm tiếp cận không dùng Parser (Parsing-Free - NeuralLog ",
        make_citation_element([5]),
        ") bỏ qua bước phân tích cú pháp bằng cách sử dụng các mô hình ngôn ngữ tiền huấn luyện (BERT, RoBERTa) để trực tiếp mã hóa chuỗi văn bản log thô thành các vector nhúng ngữ nghĩa liên tục."
    ])
    add_p([
        "Mặc dù đạt kết quả tốt trên các tập dữ liệu phần mềm, nhóm phương pháp Semantic–Sequential đối mặt với ba rào cản khi áp dụng vào an ninh mạng ",
        make_citation_element([2, 18]),
        ": (1) Nguy cơ rò rỉ và thiên lệch từ dữ liệu tiền huấn luyện (Pretraining-Data Advantage): các mô hình sử dụng Transformer tiền huấn luyện trên kho văn bản tổng quát có nguy cơ tận dụng tri thức ngoài miền; khi đánh giá trong điều kiện phân vùng nghiêm ngặt, năng lực phân tách thực tế cần được kiểm chứng cẩn trọng; (2) Chi phí tài nguyên tính toán: độ phức tạp tính toán bậc hai của cơ chế Self-Attention O(L^2) theo độ dài cửa sổ L đòi hỏi tài nguyên tính toán đáng kể trong môi trường lưu lượng lớn; (3) Giới hạn phạm vi quan sát đơn luồng: mô hình chuỗi chủ yếu theo dõi các sự kiện trên một dòng thời gian cục bộ, gặp khó khăn khi mô hình hóa trực tiếp các mối liên hệ phụ thuộc đan xen đa tiến trình, đa thực thể và vượt ranh giới máy chủ."
    ])

    # --- 1.2.3 ---
    add_h3("Đồ thị nguồn gốc và Graph Representation Learning")
    add_p([
        "Để khắc phục hạn chế về phạm vi quan sát của mô hình chuỗi, hướng tiếp cận dựa trên đồ thị nguồn gốc hệ thống (System Provenance Graph) mô hình hóa toàn bộ lịch sử thực thi và tương tác trong hệ điều hành dưới dạng một đồ thị có hướng, không đồng nhất và gán nhãn thời gian ",
        make_citation_element([9, 10]),
        ":"
    ])
    add_display_equation(r"\mathcal{G} = (\mathcal{V}, \mathcal{E}, \mathcal{T}_v, \mathcal{T}_e, \phi, \psi, \tau)")
    add_p([
        "Trong đó V là tập các đỉnh đại diện cho các thực thể hệ thống thuộc tập kiểu ",
        latex_to_clean_omml(r"\mathcal{T}_v = \{\text{Process, File, Socket, Registry, User, Host}\}"),
        "; E là tập các cạnh có hướng mô tả các tương tác luồng phụ thuộc thuộc tập kiểu ",
        latex_to_clean_omml(r"\mathcal{T}_e = \{\text{fork, execve, read, write, connect, bind, send, recv}\}"),
        "; và τ gán nhãn mốc thời gian xảy ra tương tác ",
        make_citation_element([9]),
        "."
    ], first_line_indent=False)
    add_p([
        "Các hệ thống phát hiện xâm nhập dựa trên đồ thị nguồn gốc (PIDS) tiêu biểu bao gồm: UNICORN ",
        make_citation_element([11]),
        " phân tích cấu trúc đồ thị nguồn gốc thời gian thực bằng cách băm cây con Weisfeiler-Lehman (WL-subtree hashing) và mô hình hóa trạng thái hệ thống qua mô hình phân cụm đồ thị động; KAIROS ",
        make_citation_element([12]),
        " học biểu diễn đồ thị thời gian dựa trên mạng nơ-ron đồ thị (GNN) kết hợp cơ chế chú ý theo thời gian; NODLINK ",
        make_citation_element([13]),
        " mã hóa quan hệ tương tác trực tuyến giữa các nút; MAGIC ",
        make_citation_element([14]),
        " áp dụng mô hình Masked Graph Autoencoder tự giám sát; và ORTHRUS ",
        make_citation_element([15]),
        " phân tích đa mức độ chi tiết để quy kết nguồn gốc tấn công."
    ])
    add_p([
        "Tuy nhiên, việc triển khai GNN trên đồ thị nguồn gốc quy mô thực tế đối mặt với ba rào cản nền tảng ",
        make_citation_element([9, 21]),
        ": (1) Hiện tượng bùng nổ phụ thuộc (Dependency Explosion): qua thời gian vận hành, các tiến trình dịch vụ tồn tại lâu dài (như sshd, systemd, web server) liên kết với hàng triệu tệp tin và socket, khiến đồ thị phát triển dày đặc và làm mờ nhạt dấu vết tấn công; (2) Hiện tượng nghẽn cổ chai thông tin và suy giảm phân tách (Over-smoothing và Over-squashing) ",
        make_citation_element([21]),
        " khi truyền thông điệp qua nhiều lớp GNN trên đồ thị lớn; (3) Chi phí duy trì trạng thái đồ thị và độ trễ tính toán lớn, đòi hỏi các kỹ thuật cắt tỉa (Pruning) ",
        make_citation_element([9, 20]),
        " hoặc đơn giản hóa kiến trúc ",
        make_citation_element([16, 30]),
        " để vận hành thực tế."
    ])
    add_p([
        "Bảng so sánh tổng hợp dưới đây đối chiếu toàn diện ba nhóm phương pháp biểu diễn đặc trưng log hiện đại theo năm tiêu chí kỹ thuật cốt lõi:",
        make_ref_element("BK_TBL_1_003", "Bảng 1.3"),
        ":"
    ])

    # TABLE 1.3
    add_table_caption(doc, target_p, 3, "So sánh đối chiếu ba nhóm phương pháp biểu diễn đặc trưng log hiện đại", bookmark_name="BK_TBL_1_003")
    t3_headers = ["Tiêu chí kỹ thuật", "Thống kê & Cú pháp (Event Count, Invariant)", "Ngữ nghĩa Chuỗi (DeepLog, LogBERT, NeuralLog)", "Đồ thị Nguồn gốc (UNICORN, MAGIC, KAIROS)"]
    t3_rows = [
        ["Đơn vị biểu diễn", "Mẫu log / Tần suất từ tố", "Chuỗi Event ID / Mã nhúng từ", "Đồ thị không đồng nhất đa thực thể"],
        ["Mô hình toán học", "PCA, Invariant Matrix, TF-IDF", "LSTM, Transformer, BERT", "GNN, Dynamic Graph, WL Hashing"],
        ["Phạm vi ngữ cảnh", "Rất hẹp (Đếm trong cửa sổ)", "Hẹp (Chuỗi tiến trình cục bộ)", "Rất rộng (Toàn hệ thống đa thực thể)"],
        ["Độ phức tạp tính toán", "Tuyến tính O(N) — Rất thấp", "Bậc hai O(L^2) — Trung bình", "Phụ thuộc đồ thị O(|V|+|E|) — Rất cao"],
        ["Điểm nghẽn phương pháp luận", "Mất tham số động, lỗi parser", "Bỏ sót quan hệ đa thực thể, leak", "Bùng nổ phụ thuộc, Over-squashing"]
    ]
    t3_widths = [1800, 2400, 2600, 2500]
    insert_clean_table(doc, target_p, t3_headers, t3_rows, t3_widths, font_size_pt=12)

    add_p([
        "Tổng hợp phân tích so sánh chỉ ra rằng: cả ba nhóm phương pháp hiện tại đều sở hữu những ưu thế và giới hạn riêng biệt (",
        make_ref_element("BK_TBL_1_003", "Bảng 1.3"),
        "). Mối quan hệ tương tác và sự chuyển dịch từ các giới hạn phương pháp luận sang năm khoảng trống nghiên cứu trọng tâm được tổng hợp trực quan trên bản đồ quan hệ trong ",
        make_ref_element("BK_FIG_1_004", "Hình 1.4"),
        ":"
    ], first_line_indent=False)

    # FIGURE 1.4
    add_figure_image(fig4_path, width_inches=6.2)
    add_figure_caption(doc, target_p, "1.4", "Bản đồ đối chiếu ba nhóm phương pháp biểu diễn log và nguồn gốc hình thành năm khoảng trống nghiên cứu cốt lõi (Nguồn: Tác giả tổng hợp)", bookmark_name="BK_FIG_1_004", seq_id=4)

    # --- 1.3 ---
    add_h2("1.3. Các khoảng trống nghiên cứu cốt lõi")
    add_p([
        "Từ kết quả khảo sát và phân tích đối chiếu ba nhóm phương pháp biểu diễn đặc trưng log tại Mục 1.2, có thể nhận thấy rằng mặc dù các kỹ thuật thống kê, mô hình chuỗi ngữ nghĩa và học biểu diễn đồ thị nguồn gốc đã đạt được nhiều bước tiến quan trọng, việc ứng dụng chúng vào môi trường phát hiện tấn công mạng doanh nghiệp thực tế vẫn đối mặt với những rào cản nền tảng chưa được giải quyết thấu đáo ",
        make_citation_element([2, 6, 9, 18]),
        ". Nhằm thiết lập cơ sở khoa học vững chắc và định hình phạm vi nghiên cứu, chuyên đề tổng kết năm khoảng trống nghiên cứu cốt lõi (Research Gaps) tương ứng với năm câu hỏi nghiên cứu (Research Questions - RQ) định hướng cho toàn bộ các đề xuất phương pháp luận tiếp theo."
    ])

    # --- 1.3.1 ---
    add_h3("Khoảng trống 1: Mất mát ngữ nghĩa an ninh trong quá trình trừu tượng hóa tham số")
    add_p([
        "Dữ liệu thực nghiệm từ các nghiên cứu tiền nhiệm chỉ ra rằng các bộ phân tích cú pháp (Log Parsers) đóng vai trò then chốt trong việc giảm chiều không gian văn bản log thành các mẫu định dạng tĩnh ",
        make_citation_element([6, 7]),
        ". Tuy nhiên, cơ chế trừu tượng hóa tham số phụ thuộc vào từng thuật toán phân tích cú pháp; nhiều pipeline tiền xử lý thay thế các tham số biến động như chuỗi dòng lệnh, đường dẫn tệp tin, địa chỉ IP hoặc mã lỗi hex bằng ký tự đại diện <*> ",
        make_citation_element([8, 18]),
        ". Quá trình này vô tình loại bỏ các tín hiệu phân biệt an ninh quan trọng nhất—nơi chứa đựng dấu vết của các kỹ thuật tấn công LotL hoặc các tham số thực thi độc hại ",
        make_citation_element([6]),
        "."
    ])
    add_p(
        "Làm thế nào để xây dựng một cơ chế biểu diễn đặc trưng bảo toàn được ngữ nghĩa an ninh phong phú của các tham số biến động mà không làm bùng nổ không gian từ vựng và vẫn duy trì được tính khái quát hóa cao?",
        bold_prefix="• Câu hỏi Nghiên cứu 1 (RQ1): "
    )

    # --- 1.3.2 ---
    add_h3("Khoảng trống 2: Bất đồng bộ và suy thoái trong gióng hàng biểu diễn đa góc nhìn")
    add_p([
        "Các công trình nghiên cứu hiện đại đã chỉ ra rằng dữ liệu log sở hữu tính đa góc nhìn tự nhiên: góc nhìn chuỗi phản ánh trật tự diễn tiến thời gian cục bộ ",
        make_citation_element([3, 4]),
        ", trong khi góc nhìn đồ thị nguồn gốc mô hình hóa cấu trúc tương tác nhân quả đa thực thể ",
        make_citation_element([9, 10, 11]),
        ". Tuy nhiên, các giải pháp hiện nay thường chỉ khai thác đơn lẻ một góc nhìn hoặc kết hợp muộn (Late Fusion) ở tầng phân loại, dẫn đến tình trạng bất đồng bộ và suy giảm năng lực biểu diễn khi luồng sự kiện xảy ra đan xen song song trên nhiều tiến trình."
    ])
    add_p(
        "Làm thế nào để thiết lập một cơ chế gióng hàng biểu diễn (Representation Alignment) đồng bộ giữa góc nhìn chuỗi sự kiện và góc nhìn đồ thị nguồn gốc, giúp tận dụng tối đa sức mạnh bổ trợ lẫn nhau mà không làm gia tăng đột biến độ phức tạp tính toán?",
        bold_prefix="• Câu hỏi Nghiên cứu 2 (RQ2): "
    )

    # --- 1.3.3 ---
    add_h3("Khoảng trống 3: Rò rỉ thông tin quy trình, học đường tắt và trôi dạt biểu diễn")
    add_p([
        "Nhiều mô hình học máy hiện đại đạt kết quả thực nghiệm rất cao trên các tập dữ liệu công khai nhưng suy giảm nghiêm trọng khi thử nghiệm trong môi trường mới ",
        make_citation_element([2]),
        ". Nguyên nhân cốt lõi là hiện tượng học đường tắt (Shortcut Learning) và rò rỉ phân vùng (Partition Leakage), khi mô hình học các mẫu tương quan giả tạo đặc thù của môi trường thử nghiệm (như dải địa chỉ IP cố định, định dạng dấu thời gian lab) thay vì học bản chất hành vi an ninh ",
        make_citation_element([2, 18]),
        ". Đồng thời, sự biến đổi tự nhiên của môi trường hệ thống theo thời gian dẫn đến hiện tượng trôi dạt biểu diễn (Representation Drift)."
    ])
    add_p(
        "Làm thế nào để định nghĩa và thực thi các ràng buộc bất biến (Invariance Constraints) trong quá trình học biểu diễn nhằm triệt tiêu các đặc trưng đường tắt và duy trì tính ổn định của không gian vector đặc trưng trước sự trôi dạt phân phối?",
        bold_prefix="• Câu hỏi Nghiên cứu 3 (RQ3): "
    )

    # --- 1.3.4 ---
    add_h3("Khoảng trống 4: Gán nhãn mức thô, phân bổ bằng chứng yếu và nhiễu quản trị viên")
    add_p([
        "Trong các bộ dữ liệu APT thực tế, việc gán nhãn chi tiết cho từng dòng log là cực kỳ tốn kém và không khả thi; hầu hết các tập dữ liệu chỉ cung cấp nhãn mức thô (Coarse-grained Labels) ở mức phiên làm việc, cửa sổ thời gian hoặc cây tiến trình ",
        make_citation_element([18, 28, 29]),
        ". Hơn nữa, dấu vết của các kỹ thuật tấn công thường xuất hiện rất thưa thớt (Weak Evidence) và bị chìm ngập hoàn toàn trong các luồng hoạt động quản trị viên hợp lệ (Admin-Noise) ",
        make_citation_element([9, 18]),
        "."
    ])
    add_p(
        "Làm thế nào để học được không gian biểu diễn có khả năng khuếch đại và phân tách các bằng chứng tấn công thưa thớt từ dữ liệu gán nhãn mức thô mà không bị chi phối bởi nhiễu hành vi quản trị?",
        bold_prefix="• Câu hỏi Nghiên cứu 4 (RQ4): "
    )

    # --- 1.3.5 ---
    add_h3("Khoảng trống 5: Đánh đổi giữa bảo toàn liên kết an ninh và rủi ro quyền riêng tư")
    add_p([
        "Để phát hiện và điều tra các cuộc tấn công APT kéo dài, mô hình học biểu diễn đòi hỏi phải bảo toàn các liên kết định danh thực thể (tài khoản người dùng, tên máy chủ, địa chỉ IP mạng nội bộ) ",
        make_citation_element([9, 10]),
        ". Tuy nhiên, việc lưu trữ và chia sẻ các biểu diễn vector chứa thông tin định danh trực tiếp vi phạm nghiêm trọng các quy định về bảo vệ dữ liệu và tiềm ẩn nguy cơ bị tấn công suy diễn thành viên (Membership Inference) ",
        make_citation_element([25]),
        " hoặc tái cấu trúc thông tin nhạy cảm (Model Inversion) ",
        make_citation_element([26, 27]),
        "."
    ])
    add_p(
        "Làm thế nào để thiết kế cơ chế bảo vệ quyền riêng tư có thể định lượng được (như Differential Privacy) trên không gian vector biểu diễn mà vẫn duy trì được các liên kết ngữ cảnh trọng yếu phục vụ cho phát hiện và quy kết đe dọa?",
        bold_prefix="• Câu hỏi Nghiên cứu 5 (RQ5): "
    )

    # =========================================================================
    # CHAPTER 2: PHƯƠNG PHÁP BIỂU DIỄN ĐẶC TRƯNG LOG ĐA GÓC NHÌN BẢO TOÀN NGỮ CẢNH AN TOÀN
    # =========================================================================
    add_h1("Chương 2. PHƯƠNG PHÁP BIỂU DIỄN ĐẶC TRƯNG LOG ĐA GÓC NHÌN BẢO TOÀN NGỮ CẢNH AN TOÀN")

    # --- 2.1 ---
    add_h2("2.1. Hình thức hóa bài toán và giới hạn xử lý dòng")

    # --- 2.1.1 ---
    add_h3("2.1.1. Biểu diễn đa góc nhìn và ranh giới Extractor–Detector")
    add_p([
        "Để thiết lập nền tảng lý thuyết vững chắc và giải quyết căn bản các khoảng trống phương pháp luận đã được xác lập trong Chương 1, Chương này đề xuất khung phương pháp biểu diễn đặc trưng log đa góc nhìn bảo toàn ngữ cảnh an toàn, vận hành trực tuyến trong môi trường xử lý dòng (Streaming Environment). Quá trình biểu diễn bắt đầu bằng việc hình thức hóa luồng nhật ký kiểm toán hệ thống thành một chuỗi sự kiện được gán nhãn thời gian: ",
        latex_to_clean_omml(r"\mathcal{L}_{1:t} = \langle e_1, e_2, \dots, e_t \rangle"),
        ", trong đó mỗi sự kiện ",
        latex_to_clean_omml(r"e_i = (t_i, v_i, \tau_i, \mathbf{p}_i)"),
        " là một bộ tứ bao gồm mốc thời gian xuất hiện sự kiện ",
        latex_to_clean_omml(r"t_i \in \mathbb{R}^+"),
        ", định danh thực thể chủ quản ",
        latex_to_clean_omml(r"v_i \in \mathcal{V}"),
        " (tiến trình, người dùng, máy chủ), kiểu hành vi kiểm toán ",
        latex_to_clean_omml(r"\tau_i \in \mathcal{T}"),
        " (như fork, execve, read, write, connect, registry_set), và tập hợp các tham số động có cấu trúc ",
        latex_to_clean_omml(r"\mathbf{p}_i = \{(k_j, v_j)\}_{j=1}^m"),
        " (đường dẫn tệp tin, tham số dòng lệnh thực thi, cổng mạng và địa chỉ IP kết nối) ",
        make_citation_element([9]),
        "."
    ])
    add_p([
        "Bài toán học biểu diễn đặc trưng log được phát biểu chính thức là việc tìm kiếm một ánh xạ tham số hóa ",
        latex_to_clean_omml(r"f_\theta"),
        " biến đổi lịch sử sự kiện quan sát được ",
        latex_to_clean_omml(r"\mathcal{L}_{1:t}"),
        " thành một vector biểu diễn tiềm ẩn liên tục ",
        latex_to_clean_omml(r"\mathbf{z}_t \in \mathbb{R}^d"),
        ":"
    ])
    add_display_equation(r"\mathbf{z}_t = f_\theta(\mathcal{S}_t), \quad \mathcal{S}_t = \text{Update}(\mathcal{S}_{t-1}, e_t)")
    add_p([
        "Trong đó ",
        latex_to_clean_omml(r"\mathcal{S}_t"),
        " là trạng thái bộ nhớ nội bộ hữu hạn tại thời điểm t, được cập nhật theo cơ chế nhân quả đơn lượt (Strictly Causal Single-pass Update) khi tiếp nhận sự kiện mới ",
        latex_to_clean_omml(r"e_t"),
        ". Quá trình học tham số ",
        latex_to_clean_omml(r"\theta"),
        " được thực hiện hoàn toàn theo cơ chế tự giám sát (Self-Supervised Learning) trên dữ liệu nền tảng lịch sử mà không cần nhãn tấn công hay thông tin phân loại hạ nguồn."
    ], first_line_indent=False)
    add_p([
        "Chuyên đề xác lập ranh giới nguyên tắc bất biến giữa Bộ trích xuất biểu diễn (Feature Extractor) và Bộ phát hiện hạ nguồn (Downstream Detector) ",
        make_citation_element([2, 22]),
        ": (1) Bộ trích xuất ",
        latex_to_clean_omml(r"f_\theta"),
        " chỉ đảm nhiệm việc ánh xạ cấu trúc tương tác và ngữ cảnh chuỗi vào không gian vector hình học ",
        latex_to_clean_omml(r"\mathbb{R}^d"),
        " thỏa mãn Hợp đồng Biểu diễn ",
        latex_to_clean_omml(r"\mathcal{C}_{\text{representation}}"),
        "; (2) Bộ phát hiện hạ nguồn ",
        latex_to_clean_omml(r"\hat{\mathbf{y}}_t = \sigma(\mathbf{W}^\top \mathbf{z}_t + \mathbf{b})"),
        " là một bộ dò tuyến tính độc lập được đánh giá theo giao thức đóng băng tham số (Frozen Linear Probe), đảm bảo mọi kết quả phát hiện và quy kết chiến thuật/kỹ thuật MITRE ATT&CK ",
        make_citation_element([1]),
        " phản ánh trung thực chất lượng thông tin nội tại của không gian vector z mà không phụ thuộc vào năng lực học bù của bộ phân loại."
    ])
    add_p(
        "Trên cơ sở định nghĩa bài toán và ranh giới phương pháp luận, chuyên đề hình thức hóa năm giả thuyết khoa học cốt lõi (Hypotheses H1–H5 — Đóng góp Độc lập của Chuyên đề):"
    )
    add_p([
        "Việc bảo toàn và mã hóa liên tục các token tham số động ",
        latex_to_clean_omml(r"X_{\text{param}}"),
        " (chuỗi lệnh, đường dẫn tệp tin, địa chỉ IP) vào không gian vector z mang lại lượng thông tin tương hỗ an ninh cao hơn có ý nghĩa so với việc trừu tượng hóa bằng mẫu regex tĩnh ",
        latex_to_clean_omml(r"X_{\text{template}}"),
        ", ngăn ngừa hiện tượng sụp đổ biểu diễn (Representation Collapse) giữa hành vi quản trị hợp lệ và hành vi LotL độc hại:"
    ], bold_prefix="• Giả thuyết H1 (Parameter Semantic Fidelity — Đóng góp Độc lập): ")
    add_display_equation(r"\mathcal{I}(X_{\text{param}}; \mathbf{z}) > \mathcal{I}(X_{\text{template}}; \mathbf{z})")
    add_p([
        "Việc gióng hàng đồng bộ giữa góc nhìn chuỗi sự kiện thời gian ",
        latex_to_clean_omml(r"\mathbf{z}_t^{(\text{seq})}"),
        " và góc nhìn đồ thị nguồn gốc cấu trúc ",
        latex_to_clean_omml(r"\mathbf{z}_t^{(\text{graph})}"),
        " thông qua hàm mục tiêu tương phản InfoNCE tối ưu hóa biểu diễn thống nhất ",
        latex_to_clean_omml(r"\mathbf{z}_t"),
        " giúp triệt tiêu điểm mù đơn luồng mà không gây bùng nổ phụ thuộc toàn cục:"
    ], bold_prefix="• Giả thuyết H2 (Dual-View Synchronous Alignment — Đóng góp Độc lập): ")
    add_display_equation(r"\mathcal{L}_{\text{align}}(\mathbf{z}_t^{(\text{seq})}, \mathbf{z}_t^{(\text{graph})}) = -\log \frac{\exp(\text{sim}(\mathbf{z}_t^{(\text{seq})}, \mathbf{z}_t^{(\text{graph})})/\tau)}{\sum_k \exp(\text{sim}(\mathbf{z}_t^{(\text{seq})}, \mathbf{z}_k^{(\text{graph})})/\tau)}")
    add_p([
        "Không gian biểu diễn tuân thủ các ràng buộc bất biến (Invariance Constraints) ",
        make_citation_element([22, 23]),
        " sẽ duy trì khoảng cách hình học ổn định trước các biến đổi cú pháp vô hại ",
        latex_to_clean_omml(r"T \in \mathcal{T}_{\text{benign}}"),
        ", đồng thời triệt tiêu thông tin tương hỗ đối với các đặc trưng đường tắt ",
        latex_to_clean_omml(r"S \in \mathcal{E}_{\text{exclude}}"),
        " gây trôi dạt biểu diễn:"
    ], bold_prefix="• Giả thuyết H3 (Invariant Robustness & Anti-Drift — Đóng góp Độc lập): ")
    add_display_equation(r"\|\mathbf{z}(T(X)) - \mathbf{z}(X)\|_2 \le \epsilon_{\text{inv}}, \quad \mathcal{I}(\mathbf{z}; S) \le \epsilon_{\text{short}}")
    add_p([
        "Cơ chế tổng hợp đa thể hiện có trọng số chú ý (Attention-based Multiple Instance Learning) ",
        make_citation_element([24]),
        " tự động phân bổ trọng số chú ý ",
        latex_to_clean_omml(r"\alpha_k"),
        " tập trung vào các bằng chứng tấn công thưa thớt, phân tách hiệu quả dấu vết APT khỏi nhiễu nền quản trị viên (Admin-Noise) trên dữ liệu nhãn mức thô:"
    ], bold_prefix="• Giả thuyết H4 (Weak Evidence & Admin-Noise Disentanglement — Đóng góp Độc lập): ")
    add_display_equation(r"\mathbf{z}_t = \sum_{k=1}^K \alpha_k \mathbf{h}_k, \quad \alpha_k = \frac{\exp(\mathbf{w}^\top \tanh(\mathbf{V} \mathbf{h}_k))}{\sum_{j=1}^K \exp(\mathbf{w}^\top \tanh(\mathbf{V} \mathbf{h}_j))}")
    add_p([
        "Ánh xạ làm sạch định danh kết hợp cơ chế nhiễu vi phân ",
        latex_to_clean_omml(r"(\epsilon, \delta)"),
        "-Differential Privacy ",
        make_citation_element([27]),
        " triệt tiêu hoàn toàn nguy cơ suy diễn thành viên ",
        make_citation_element([25]),
        " và tái cấu trúc thông tin định danh nhạy cảm ",
        make_citation_element([26]),
        " mà vẫn bảo toàn độ chính xác quy kết chiến thuật/kỹ thuật MITRE ATT&CK trên vector z."
    ], bold_prefix="• Giả thuyết H5 (Privacy-Preserving Utility Trade-off — Đóng góp Độc lập): ")

    # --- 2.1.2 ---
    add_h3("2.1.2. Độ phức tạp xử lý dòng với trạng thái hữu hạn")
    add_p([
        "Trong môi trường giám sát an ninh SOC thời gian thực, luồng telemetry liên tục đổ về với lưu lượng hàng trăm nghìn sự kiện mỗi giây. Việc duy trì toàn bộ lịch sử tương tác của hệ thống trên bộ nhớ là bất khả thi về mặt tài nguyên và dẫn đến hiện tượng rò rỉ bộ nhớ (Memory Exhaustion). Do đó, chuyên đề thiết lập mô hình xử lý dòng với trạng thái hữu hạn nghiêm ngặt (Bounded-State Streaming Model), được quy định bởi ngân sách bộ nhớ tối đa ",
        latex_to_clean_omml(r"M_{\text{max}}"),
        " và tập thực thể hoạt động ",
        latex_to_clean_omml(r"\mathcal{V}_{\text{active}}(t)"),
        ":"
    ])
    add_display_equation(r"\mathcal{V}_{\text{active}}(t) = \{v \in \mathcal{V} \mid t - t_{\text{last}}(v) \le \tau_{\text{TTL}}\}, \quad |\mathcal{S}_t| \le M_{\text{max}}")
    add_p([
        "Trong đó ",
        latex_to_clean_omml(r"\tau_{\text{TTL}}"),
        " là ngưỡng thời gian sống (Time-To-Live). Một thực thể v (tiến trình, socket, tệp tin) không phát sinh bất kỳ tương tác nào trong khoảng thời gian ",
        latex_to_clean_omml(r"\tau_{\text{TTL}}"),
        " sẽ bị loại bỏ (Eviction) khỏi bộ đệm trạng thái hoạt động ",
        latex_to_clean_omml(r"\mathcal{S}_t"),
        ". Khi kích thước trạng thái đạt ngưỡng ngân sách ",
        latex_to_clean_omml(r"M_{\text{max}}"),
        ", chính sách Least Recently Updated (LRU) kết hợp cơ chế tóm tắt Count-Min Sketch được kích hoạt để nén thông tin tần suất và giải phóng tài nguyên bộ nhớ."
    ], first_line_indent=False)
    add_p([
        "Đối với cấu trúc đồ thị nguồn gốc, để giải quyết hiện tượng bùng nổ cạnh (Edge Explosion) phát sinh từ các tiến trình dịch vụ tồn tại lâu dài ",
        make_citation_element([9, 11]),
        ", chuyên đề áp dụng quy luật suy giảm trọng số cạnh theo hàm mũ thời gian kết hợp cơ chế mốc ngắt thời gian sự kiện (Event-Time Watermark) ",
        make_citation_element([13, 20]),
        ":"
    ])
    add_display_equation(r"\omega(e_{uv}, t) = \exp(-\lambda (t - t_{uv})), \quad t_{\text{wm}} = \max_{1 \le i \le t}(t_{\text{event}}(e_i)) - \delta_{\text{delay}}")
    add_p([
        "Trong đó ",
        latex_to_clean_omml(r"\lambda > 0"),
        " là hệ số suy giảm thời gian; các cạnh có trọng số ",
        latex_to_clean_omml(r"\omega(e_{uv}, t) < \theta_{\text{prune}}"),
        " sẽ tự động được cắt tỉa khỏi đồ thị hoạt động. Mốc thời gian ",
        latex_to_clean_omml(r"t_{\text{wm}}"),
        " thiết lập ranh giới điều hòa cho các sự kiện đến muộn hoặc xáo trộn thứ tự (Out-of-Order Events) trong phạm vi trễ cho phép ",
        latex_to_clean_omml(r"\delta_{\text{delay}}"),
        "; các sự kiện trễ hơn ",
        latex_to_clean_omml(r"\delta_{\text{delay}}"),
        " được chuyển tiếp vào bộ đệm tái điều chỉnh (Reconciliation Buffer) mà không phá vỡ tính liên tục của luồng suy luận."
    ], first_line_indent=False)
    add_p([
        "Khi hệ thống gặp hiện tượng đột biến lưu lượng (Traffic Spike), cơ chế kiểm soát áp lực ngược (Backpressure Control) dựa trên thuật toán Token-Bucket điều tiết tốc độ nạp dữ liệu, đồng thời kích hoạt chính sách loại bỏ ưu tiên (Priority Shedding) đối với các luồng nhật ký ứng dụng có điểm an ninh sơ cấp thấp, bảo đảm thông lượng xử lý dòng luôn nằm trong giới hạn năng lực tính toán của máy chủ giám sát."
    ])
    add_p([
        "Phân tích độ phức tạp thuật toán xác lập rằng mô hình dòng đề xuất đạt độ phức tạp không gian ",
        latex_to_clean_omml(r"\mathcal{O}(|\mathcal{V}_{\text{active}}| \cdot d + |\mathcal{E}_{\text{window}}|)"),
        " và độ phức tạp thời gian cho mỗi sự kiện ",
        latex_to_clean_omml(r"\mathcal{O}(\text{deg}_{\text{max}} + d)"),
        ". Đây là minh chứng lý thuyết khẳng định tính khả thi của việc duy trì ngữ cảnh an ninh APT dài hạn trong điều kiện ràng buộc trạng thái hữu hạn."
    ])

    # --- 2.1.3 ---
    add_h3("2.1.3. Kiến trúc tổng thể và giao diện vào/ra")
    add_p([
        "Kiến trúc tổng thể của khung biểu diễn đặc trưng log đa góc nhìn dòng được thiết kế theo nguyên tắc phân tách độc lập giữa hai mặt phẳng vận hành: Mặt phẳng Huấn luyện (Training Plane) và Mặt phẳng Suy luận Dòng (Streaming Inference Plane), được minh họa chi tiết trong ",
        make_ref_element("BK_FIG_2_001", "Hình 2.1"),
        ":"
    ])

    # FIGURE 2.1
    add_figure_image(fig5_path, width_inches=6.2)
    add_figure_caption(
        doc,
        target_p,
        "2.1",
        "Kiến trúc hai mặt phẳng (Training Plane & Inference Plane) của khung biểu diễn đặc trưng log đa góc nhìn dòng (Nguồn: Tác giả đề xuất)",
        bookmark_name="BK_FIG_2_001",
        seq_id=5
    )

    add_p(
        "Vận hành ngoại tuyến (Offline) trên kho dữ liệu telemetry lịch sử D_train. Mặt phẳng này thực hiện trích xuất song song hai góc nhìn (chuỗi sự kiện và đồ thị nguồn gốc), tối ưu hóa hàm mất mát tương phản gióng hàng đa góc nhìn L_align kết hợp các số hạng điều hòa bất biến L_reg (VICReg/Barlow Twins) để tìm kiếm bộ trọng số tối ưu θ*. Quá trình huấn luyện không sử dụng nhãn tấn công và không can thiệp vào luồng vận hành thời gian thực.",
        bold_prefix="1. Mặt phẳng Huấn luyện (Training Plane — Offline Self-Supervised): "
    )
    add_p([
        "Vận hành trực tuyến (Online) trên luồng sự kiện đang diễn ra ",
        latex_to_clean_omml(r"e_t \in \mathcal{L}_{1:t}"),
        ". Mặt phẳng này tiếp nhận trọng số đã đóng băng hoàn toàn ",
        latex_to_clean_omml(r"\theta^*"),
        ", cập nhật trạng thái bộ nhớ hữu hạn ",
        latex_to_clean_omml(r"\mathcal{S}_t"),
        " theo các chính sách TTL, Eviction và Watermarking, sau đó trích xuất vector đóng gói ",
        latex_to_clean_omml(r"\mathbf{z}_t = f_{\theta^*}(\mathcal{S}_t) \in \mathbb{R}^d"),
        " theo cơ chế đơn lượt đơn hướng (Strictly Causal — Zero Lookahead), cung cấp đầu vào tức thời cho các bộ dò tuyến tính đóng băng phục vụ phát hiện và điều tra nguồn gốc an ninh."
    ], bold_prefix="2. Mặt phẳng Suy luận Dòng (Streaming Inference Plane — Online Causal): ")

    # =========================================================================
    # CONCLUSION & BIBLIOGRAPHY
    # =========================================================================
    add_h1("Kết luận")
    add_p(
        "Chuyên đề đã hoàn thành việc xây dựng cơ sở lý luận, tổng quan hệ thống hóa các phương pháp trích xuất đặc trưng log hiện đại, và xác lập khung hình thức hóa toán học cho bài toán biểu diễn đặc trưng đa góc nhìn bảo toàn ngữ cảnh an toàn trong môi trường xử lý dòng. Các đóng góp cốt lõi bao gồm: (1) Phân định ranh giới phương pháp luận ba tầng và xác lập Hợp đồng Biểu diễn (Representation Contract); (2) Tổng kết năm khoảng trống nghiên cứu nền tảng (RQ1–RQ5); (3) Hình thức hóa toán học bài toán biểu diễn dòng với trạng thái hữu hạn và xác lập năm giả thuyết khoa học cốt lõi (H1–H5); (4) Thiết lập kiến trúc hai mặt phẳng phân tách rõ ràng giữa huấn luyện tự giám sát ngoại tuyến và suy luận dòng trực tuyến nhân quả."
    )
    add_p(
        "Những kết quả lý thuyết và kiến trúc được thiết lập trong Chuyên đề tạo tiền đề khoa học trực tiếp cho việc triển khai chi tiết các mô-đun thuật toán biểu diễn chuỗi, đồ thị nguồn gốc và cơ chế gióng hàng tương phản trong các giai đoạn nghiên cứu tiếp theo."
    )

    add_h1("Tài liệu tham khảo")
    print("[5/6] Creating native Word BIBLIOGRAPHY field...")

    # Build dynamic Bibliography paragraph
    bib_p = doc.add_paragraph(style="Normal")
    bib_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    bib_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    bib_p.paragraph_format.space_before = Pt(0)
    bib_p.paragraph_format.space_after = Pt(4)
    bib_p.paragraph_format.first_line_indent = Cm(0)

    fld_xml_parts = [f'<w:fldSimple {nsdecls("w")} w:instr="BIBLIOGRAPHY \\l 1033 ">']
    for s in sources:
        fld_xml_parts.append(
            '  <w:p>\n'
            '    <w:pPr>\n'
            '      <w:pStyle w:val="Normal"/>\n'
            '      <w:spacing w:line="360" w:lineRule="auto" w:before="0" w:after="80"/>\n'
            '      <w:ind w:left="0" w:firstLine="0"/>\n'
            '    </w:pPr>\n'
            '    <w:r>\n'
            '      <w:rPr>\n'
            '        <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>\n'
            '        <w:sz w:val="28"/>\n'
            '      </w:rPr>\n'
            f'      <w:t>{escape(s.title)}</w:t>\n'
            '    </w:r>\n'
            '  </w:p>'
        )
    fld_xml_parts.append('</w:fldSimple>')
    bib_p._p.append(parse_xml('\n'.join(fld_xml_parts)))

    # Save to temp docx
    temp_file = target_path.parent / (target_path.stem + ".temp.docx")
    doc.save(str(temp_file))

    # Inject customXml/item1.xml into zip package
    sources_xml_data = generate_perfect_sources_xml(sources).encode("utf-8")
    updated_file = target_path.parent / (target_path.stem + ".updated.docx")
    
    with zipfile.ZipFile(str(temp_file), "r") as zin:
        with zipfile.ZipFile(str(updated_file), "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == "customXml/item1.xml":
                    zout.writestr(item, sources_xml_data)
                else:
                    zout.writestr(item, zin.read(item.filename))
            if "customXml/item1.xml" not in zin.namelist():
                zout.writestr("customXml/item1.xml", sources_xml_data)

    temp_file.unlink(missing_ok=True)
    shutil.copyfile(updated_file, target_path)
    print(f"[SUCCESS] Saved and updated DOCX: {target_path}")

    # =========================================================================
    # STEP 6: AUTOMATE MICROSOFT WORD DESKTOP TO INSERT NATIVE DIAGRAMS, UPDATE ALL DYNAMIC FIELDS & EXPORT PDF
    # =========================================================================
    print("[6/6] Launching Desktop Microsoft Word to render Native Diagrams, update all dynamic fields and export PDF...")
    word = None
    doc_com = None
    pdf_path = target_path.parent / (target_path.stem + ".pdf")
    pythoncom.CoInitialize()
    try:
        word = win32.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0  # wdAlertsNone
        
        abs_target = os.path.abspath(str(target_path))
        doc_com = word.Documents.Open(abs_target)

        # Update dynamic fields
        for fld in doc_com.Fields:
            try:
                fld.Update()
            except Exception:
                pass

        for toc in doc_com.TablesOfContents:
            try:
                toc.Update()
            except Exception:
                pass

        for tof in doc_com.TablesOfFigures:
            try:
                tof.Update()
            except Exception:
                pass

        # Save the fully resolved and updated docx
        doc_com.Save()
        print(f"[SUCCESS] Microsoft Word updated and saved: {abs_target}")

        # Export as PDF
        abs_pdf = os.path.abspath(str(pdf_path))
        doc_com.ExportAsFixedFormat(abs_pdf, 17)  # 17 = wdExportFormatPDF
        print(f"[SUCCESS] Exported PDF: {abs_pdf}")

    except Exception as e:
        print(f"[ERROR] Microsoft Word COM Automation error: {e}")
    finally:
        if doc_com:
            try:
                doc_com.Close(SaveChanges=True)
            except Exception:
                pass
            del doc_com
        if word:
            try:
                word.Quit()
            except Exception:
                pass
            del word
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass

    return str(pdf_path)


if __name__ == "__main__":
    build_master_thesis_document()
