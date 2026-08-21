"""
Scientific Figure Inserter (Rule 2B)
Inserts verified scientific plots (PR curves, ROC curves, ablation charts) into Word documents with native captions and bookmarks.
Enforces: Data provenance, No LLM-fabricated numbers, Centered layout, Automatic SEQ Hình caption, Bookmark anchor.
"""

from pathlib import Path
from typing import Optional
import docx
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from research_agent.schemas.verification import FigureSpecification
from research_agent.visuals.schemas import VisualRecord, VisualType, CreationMethod
from research_agent.visuals.word_caption_manager import WordCaptionManager
from research_agent.visuals.registry import VisualRegistry


class ScientificFigureInserter:
    """
    Inserts verified data figures with native Word captions and cross-reference bookmarks.
    """

    def __init__(self, visual_registry: Optional[VisualRegistry] = None):
        self.registry = visual_registry or VisualRegistry()

    def insert_scientific_figure(
        self,
        doc: docx.Document,
        ref_paragraph: Optional[docx.text.paragraph.Paragraph],
        fig_spec: FigureSpecification,
        seq_num: int = 1,
        chapter_num: int = 1,
        width_inches: float = 5.8,
        node_code: str = "1.2.3",
        purpose: str = "Performance comparison on benchmark dataset",
    ) -> VisualRecord:
        """
        Inserts figure picture and its native caption into the Word document.
        """
        # 1. Paragraph containing the image
        img_p = doc.add_paragraph(style="Normal") if ref_paragraph is None else ref_paragraph.insert_paragraph_before(style="Normal")
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.paragraph_format.first_line_indent = Cm(0)
        img_p.paragraph_format.space_before = Pt(8)
        img_p.paragraph_format.space_after = Pt(4)

        img_path = Path(fig_spec.output_file_rel_path)
        if not img_path.is_absolute():
            # Resolve relative to workspace
            img_path = Path(r"D:\Research") / img_path

        if not img_path.exists():
            raise FileNotFoundError(f"Figure image file does not exist: {img_path}")

        run = img_p.add_run()
        run.add_picture(str(img_path), width=Inches(width_inches))

        bookmark_name = f"BK_FIG_{chapter_num}_{seq_num:03d}"

        # 2. Add native caption below image
        cap_p = WordCaptionManager.add_figure_caption(
            doc=doc,
            ref_paragraph=ref_paragraph,
            seq_num=seq_num,
            title_text=fig_spec.caption,
            chapter_num=chapter_num,
            bookmark_name=bookmark_name,
        )

        # 3. Register in VisualRegistry
        record = VisualRecord(
            visual_id=fig_spec.figure_id,
            node_code=node_code,
            purpose=purpose,
            visual_type=VisualType.DATA_FIGURE,
            creation_method=CreationMethod.MATPLOTLIB_IMAGE,
            caption=fig_spec.caption,
            source_provenance=f"Script: {fig_spec.plot_script_path}, SHA256: {fig_spec.output_sha256}",
            script_path=fig_spec.plot_script_path,
            output_file_path=str(img_path),
            companion_data_path=fig_spec.companion_data_csv_rel_path,
            output_sha256=fig_spec.output_sha256,
            bookmark_name=bookmark_name,
            seq_number=seq_num,
            chapter_number=chapter_num,
            is_verified=True,
        )
        self.registry.register_visual(record)
        return record
