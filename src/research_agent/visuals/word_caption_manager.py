"""
Word Native Caption Manager (Rule 4)
Handles insertion and formatting of native Microsoft Word captions with SEQ fields.
Enforces:
- Figure captions: placed BELOW figure, Label = "Hình"
- Table captions: placed ABOVE table, Label = "Bảng"
- Automatic SEQ fields with chapter prefix (e.g., 'Hình 1.1', 'Bảng 1.1')
- Paragraph formatting: Style 'Caption', Centered, Times New Roman 14pt, keepWithNext on tables.
"""

from typing import Optional, Union
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


class WordCaptionManager:
    """
    Manages native Word dynamic captions (SEQ Bảng and SEQ Hình).
    """

    @staticmethod
    def add_table_caption(
        doc: docx.Document,
        ref_paragraph: Optional[docx.text.paragraph.Paragraph],
        seq_num: int,
        title_text: str,
        chapter_num: Optional[int] = 1,
        bookmark_name: Optional[str] = None,
    ) -> docx.text.paragraph.Paragraph:
        """
        Inserts a native Word Table Caption paragraph ABOVE the table with SEQ Bảng field and keepWithNext.
        """
        cap_p = doc.add_paragraph(style="Caption") if ref_paragraph is None else ref_paragraph.insert_paragraph_before(style="Caption")
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_before = Pt(12)
        cap_p.paragraph_format.space_after = Pt(4)
        cap_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        cap_p.paragraph_format.keep_with_next = True

        cap_xml = (
            f'<w:pPr {nsdecls("w")}>\n'
            '  <w:pStyle w:val="Caption"/>\n'
            '  <w:jc w:val="center"/>\n'
            '  <w:spacing w:before="240" w:after="80" w:line="360" w:lineRule="auto"/>\n'
            '  <w:keepNext/>\n'
            '</w:pPr>'
        )
        cap_p._p.remove(cap_p._p.pPr)
        cap_p._p.append(parse_xml(cap_xml))

        # Optional bookmark start
        if bookmark_name:
            bm_start_xml = f'<w:bookmarkStart {nsdecls("w")} w:id="101" w:name="{bookmark_name}"/>'
            cap_p._p.append(parse_xml(bm_start_xml))

        prefix = f"Bảng {chapter_num}." if chapter_num else "Bảng "
        r1 = cap_p.add_run(prefix)
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(14)
        r1.bold = True

        seq_xml = (
            f'<w:fldSimple {nsdecls("w")} w:instr="SEQ Bảng \\* ARABIC ">\n'
            '  <w:r>\n'
            '    <w:rPr>\n'
            '      <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>\n'
            '      <w:b/>\n'
            '      <w:sz w:val="28"/>\n'
            '    </w:rPr>\n'
            f'    <w:t>{seq_num}</w:t>\n'
            '  </w:r>\n'
            '</w:fldSimple>'
        )
        cap_p._p.append(parse_xml(seq_xml))

        if bookmark_name:
            bm_end_xml = f'<w:bookmarkEnd {nsdecls("w")} w:id="101"/>'
            cap_p._p.append(parse_xml(bm_end_xml))

        r2 = cap_p.add_run(f": {title_text}")
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(14)
        r2.bold = True

        return cap_p

    @staticmethod
    def add_figure_caption(
        doc: docx.Document,
        ref_paragraph: Optional[docx.text.paragraph.Paragraph],
        seq_num: int,
        title_text: str,
        chapter_num: Optional[int] = 1,
        bookmark_name: Optional[str] = None,
    ) -> docx.text.paragraph.Paragraph:
        """
        Inserts a native Word Figure Caption paragraph BELOW the figure with SEQ Hình field.
        """
        cap_p = doc.add_paragraph(style="Caption") if ref_paragraph is None else ref_paragraph.insert_paragraph_before(style="Caption")
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_before = Pt(6)
        cap_p.paragraph_format.space_after = Pt(12)
        cap_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        cap_xml = (
            f'<w:pPr {nsdecls("w")}>\n'
            '  <w:pStyle w:val="Caption"/>\n'
            '  <w:jc w:val="center"/>\n'
            '  <w:spacing w:before="120" w:after="240" w:line="360" w:lineRule="auto"/>\n'
            '</w:pPr>'
        )
        cap_p._p.remove(cap_p._p.pPr)
        cap_p._p.append(parse_xml(cap_xml))

        # Optional bookmark start
        if bookmark_name:
            bm_start_xml = f'<w:bookmarkStart {nsdecls("w")} w:id="102" w:name="{bookmark_name}"/>'
            cap_p._p.append(parse_xml(bm_start_xml))

        prefix = f"Hình {chapter_num}." if chapter_num else "Hình "
        r1 = cap_p.add_run(prefix)
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(14)
        r1.bold = True

        seq_xml = (
            f'<w:fldSimple {nsdecls("w")} w:instr="SEQ Hình \\* ARABIC ">\n'
            '  <w:r>\n'
            '    <w:rPr>\n'
            '      <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>\n'
            '      <w:b/>\n'
            '      <w:sz w:val="28"/>\n'
            '    </w:rPr>\n'
            f'    <w:t>{seq_num}</w:t>\n'
            '  </w:r>\n'
            '</w:fldSimple>'
        )
        cap_p._p.append(parse_xml(seq_xml))

        if bookmark_name:
            bm_end_xml = f'<w:bookmarkEnd {nsdecls("w")} w:id="102"/>'
            cap_p._p.append(parse_xml(bm_end_xml))

        r2 = cap_p.add_run(f": {title_text}")
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(14)
        r2.bold = True

        return cap_p
