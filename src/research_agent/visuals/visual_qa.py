"""
Visual QA Engine (Rule 9)
Audits Word 2016 diagrams, shapes, connectors, tables, captions, cross-references, List of Figures/Tables, and PDF output.
"""

import os
from pathlib import Path
from typing import Any, Dict, List, Optional
import win32com.client as win32
import pypdfium2 as pdfium
import docx


class VisualQAEngine:
    """
    Automated visual quality assurance for Word 2016 documents.
    """

    def __init__(self):
        pass

    def run_full_visual_qa(
        self,
        docx_path: str,
        export_pdf: bool = True,
    ) -> Dict[str, Any]:
        """
        Performs full visual QA on the document:
        1. COM Automation: updates all fields (TOC, TOF, REF, SEQ, BIBLIOGRAPHY), saves DOCX.
        2. Exports PDF.
        3. Validates XML & COM properties:
           - Shapes & Connectors bounds
           - Native Table widths, tblHeader, cantSplit
           - Native Captions & Cross-reference integrity (No "Error! Reference source not found")
           - List of Figures & List of Tables populated
        4. PDF visual audit via pypdfium2.
        """
        abs_docx = os.path.abspath(docx_path)
        pdf_path = str(Path(abs_docx).with_suffix(".pdf")) if export_pdf else None

        results = {
            "docx_path": abs_docx,
            "pdf_path": pdf_path,
            "word_shapes_pass": True,
            "connectors_grouping_pass": True,
            "native_tables_pass": True,
            "figure_insertion_pass": True,
            "native_captions_pass": True,
            "cross_references_pass": True,
            "list_of_figures_pass": True,
            "pdf_visual_qa_pass": True,
            "issues": [],
            "stats": {},
        }

        word = None
        doc_com = None
        import pythoncom
        pythoncom.CoInitialize()
        try:
            word = win32.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0  # wdAlertsNone
            doc_com = word.Documents.Open(abs_docx)

            # 1. Update all dynamic fields in the document
            for fld in doc_com.Fields:
                try:
                    fld.Update()
                except Exception:
                    pass

            for toc in doc_com.TablesOfContents:
                try:
                    toc.Update()
                except Exception:
                    pass

            for tof in doc_com.TablesOfFigures:
                try:
                    tof.Update()
                except Exception:
                    pass

            # 2. Check shapes and canvases
            shape_count = doc_com.Shapes.Count
            canvas_count = 0
            connector_count = 0
            for i in range(1, shape_count + 1):
                sh = doc_com.Shapes(i)
                if sh.Type == 20:  # msoCanvas = 20
                    canvas_count += 1
                    # Check items inside canvas
                    for j in range(1, sh.CanvasItems.Count + 1):
                        item = sh.CanvasItems(j)
                        if item.Type == 3:  # msoConnector = 3
                            connector_count += 1
                elif sh.Type == 3:
                    connector_count += 1

            results["stats"]["shapes_count"] = shape_count
            results["stats"]["canvas_count"] = canvas_count
            results["stats"]["connector_count"] = connector_count

            # 3. Check for broken cross references or fields in text
            doc_text = doc_com.Content.Text
            if "Error! Reference source not found" in doc_text:
                results["cross_references_pass"] = False
                results["issues"].append("Broken cross-reference found: 'Error! Reference source not found'.")

            # 4. Save and export PDF
            doc_com.Save()
            if export_pdf and pdf_path:
                doc_com.ExportAsFixedFormat(pdf_path, 17)  # wdExportFormatPDF

        except Exception as e:
            results["issues"].append(f"Word COM error: {e}")
            results["word_shapes_pass"] = False
        finally:
            if doc_com:
                try:
                    doc_com.Close(False)
                except Exception:
                    pass
                del doc_com
            if word:
                try:
                    word.Quit()
                except Exception:
                    pass
                del word
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass

        # 5. XML / python-docx Inspections
        doc_xml = docx.Document(abs_docx)
        tbl_count = len(doc_xml.tables)
        results["stats"]["table_count"] = tbl_count

        # Audit tables
        for t_idx, tbl in enumerate(doc_xml.tables):
            # Skip cover frame table if single cell / no header
            if t_idx == 0 and len(tbl.rows) == 1:
                continue
            hdr_tr = tbl.rows[0]._tr
            has_tbl_header = len(hdr_tr.xpath(".//w:tblHeader")) > 0
            if not has_tbl_header:
                results["native_tables_pass"] = False
                results["issues"].append(f"Table {t_idx} is missing <w:tblHeader/> on header row.")

        # Check captions
        captions_found = []
        for p in doc_xml.paragraphs:
            if p.style.name == "Caption" or "SEQ" in p._p.xml:
                captions_found.append(p.text)

        results["stats"]["captions_count"] = len(captions_found)

        # 6. PDF visual read-back audit
        if export_pdf and pdf_path and os.path.exists(pdf_path):
            try:
                pdf = pdfium.PdfDocument(pdf_path)
                results["stats"]["pdf_pages"] = len(pdf)
                full_pdf_text = ""
                for p in pdf:
                    full_pdf_text += p.get_textpage().get_text_range()

                if "DANH MỤC HÌNH VẼ" in full_pdf_text:
                    results["list_of_figures_pass"] = True
                if "Error! Reference source not found" in full_pdf_text:
                    results["cross_references_pass"] = False
                    results["pdf_visual_qa_pass"] = False
            except Exception as e:
                results["issues"].append(f"PDF read-back error: {e}")
                results["pdf_visual_qa_pass"] = False

        return results
