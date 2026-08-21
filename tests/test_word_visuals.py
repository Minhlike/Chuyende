"""
Automated Test Suite for Word 2016 Scientific Visuals Engine
(Diagrams, Data Figures, Native Tables, Captions, Cross-References & Visual QA)
"""

import os
from pathlib import Path
import pytest
import docx
import pandas as pd

from research_agent.core.enums import TableType, FigureType
from research_agent.schemas.verification import TableSpecification, FigureSpecification
from research_agent.visuals.schemas import (
    VisualType,
    CreationMethod,
    VisualNecessityReason,
    ShapeNodeSpec,
    ConnectorSpec,
    DiagramSpecification,
)
from research_agent.visuals.necessity_gate import VisualNecessityGate
from research_agent.visuals.registry import VisualRegistry
from research_agent.visuals.word_diagram_builder import WordDiagramBuilder
from research_agent.visuals.word_table_builder import WordTableBuilder
from research_agent.visuals.word_caption_manager import WordCaptionManager
from research_agent.visuals.word_cross_reference_manager import WordCrossReferenceManager
from research_agent.visuals.scientific_figure_inserter import ScientificFigureInserter
from research_agent.visuals.visual_qa import VisualQAEngine
from research_agent.verification.figures.builder import FigureBuilder
from research_agent.skills.registry import ResearchSkillRegistry


class TestWordVisualsEngine:
    """
    Comprehensive tests for all 10 rules of the Word 2016 Scientific Visuals specification.
    """

    @pytest.fixture(autouse=True)
    def setup_test_env(self, tmp_path):
        self.test_dir = tmp_path
        self.docx_path = str(self.test_dir / "test_visuals.docx")
        self.pdf_path = str(self.test_dir / "test_visuals.pdf")

    def test_01_visual_necessity_gate(self):
        """Rule 8: Visual Necessity Gate rejects fluff and accepts necessary architectural visuals."""
        # 1. Reject vague purpose
        eval_fail1 = VisualNecessityGate.evaluate(
            visual_id="FIG-000099",
            visual_type=VisualType.CONCEPTUAL_DIAGRAM,
            purpose="to make page look nice",
            clarity_statement="just a nice decoration",
            alternative_prose_deficiency="none",
        )
        assert not eval_fail1.is_necessary
        assert "vague" in eval_fail1.rejection_reason.lower() or "deficient" in eval_fail1.rejection_reason.lower()

        # 2. Accept valid architectural diagram
        eval_pass = VisualNecessityGate.evaluate(
            visual_id="FIG-000001",
            visual_type=VisualType.CONCEPTUAL_DIAGRAM,
            purpose="Illustrate multi-view representation learning and frozen linear probing architecture",
            clarity_statement="Disentangles raw telemetry ingestion, dual sequence/graph encoding, and frozen evaluation probing.",
            alternative_prose_deficiency="Prose alone cannot clearly convey the concurrent data flows and dimensional projection into latent vector z.",
            primary_reason=VisualNecessityReason.ARCHITECTURE,
        )
        assert eval_pass.is_necessary
        assert eval_pass.primary_reason == VisualNecessityReason.ARCHITECTURE

    def test_02_native_table_builder_and_caption(self):
        """Rule 3 & Rule 4: Native Word Table with repeat header, cantSplit, and SEQ Bảng caption above."""
        doc = docx.Document()
        
        # Build Table Spec
        df = pd.DataFrame({
            "Phương pháp": ["PCA", "DeepLog", "UNICORN", "Khung đề xuất"],
            "Độ phức tạp": ["O(N)", "O(N * L^2)", "O(|V| + |E|)", "O(N)"],
            "Bảo toàn tham số": ["Thấp", "Trung bình", "Cao", "Cao (Kiểm soát)"],
            "Góc nhìn": ["Thống kê", "Chuỗi ngữ nghĩa", "Đồ thị nguồn gốc", "Đa góc nhìn (Chuỗi + Đồ thị)"]
        })
        spec = TableSpecification(
            table_id="TBL-000001",
            table_type=TableType.COMPUTED_TABLE,
            title="So sánh phương pháp",
            caption="So sánh các phương pháp biểu diễn đặc trưng log chủ đạo",
            columns=list(df.columns),
            rows_data=df.values.tolist(),
            output_sha256="aabbccddeeff",
        )

        # 1. Add Table Caption ABOVE Table
        WordCaptionManager.add_table_caption(
            doc=doc,
            ref_paragraph=None,
            seq_num=1,
            title_text=spec.caption,
            chapter_num=1,
            bookmark_name="BK_TBL_001"
        )

        # 2. Insert Table
        tbl = WordTableBuilder.insert_table(
            doc=doc,
            ref_paragraph=None,
            spec=spec,
            font_size_pt=14.0
        )

        doc.save(self.docx_path)

        # Verify XML invariants
        doc_read = docx.Document(self.docx_path)
        assert len(doc_read.tables) == 1
        t = doc_read.tables[0]
        assert len(t.rows) == 5
        # Check tblHeader and cantSplit
        hdr_xml = t.rows[0]._tr.xml
        assert "tblHeader" in hdr_xml
        assert "cantSplit" in hdr_xml
        # Check Caption paragraph has SEQ Bảng
        cap_p = doc_read.paragraphs[0]
        assert "Bảng 1." in cap_p.text
        assert "SEQ Bảng" in cap_p._p.xml

    def test_03_scientific_figure_inserter(self):
        """Rule 2B & Rule 4: Data Figure plotted with companion data and inserted with SEQ Hình caption below."""
        fig_builder = FigureBuilder(output_dir=self.test_dir)
        curves_data = [
            {"name": "Proposed Multi-View", "recalls": [0.0, 0.5, 0.9, 1.0], "precisions": [1.0, 0.98, 0.95, 0.88], "pr_auc": 0.962},
            {"name": "DeepLog Baseline", "recalls": [0.0, 0.5, 0.8, 1.0], "precisions": [1.0, 0.90, 0.75, 0.60], "pr_auc": 0.812},
        ]
        fig_spec = fig_builder.plot_pr_curve(
            figure_id="FIG-000001",
            title="Đường cong Precision-Recall trên DARPA TC",
            caption="Đường cong Precision-Recall đánh giá hiệu năng phát hiện trên bộ dữ liệu DARPA TC",
            curves_data=curves_data,
        )

        doc = docx.Document()
        inserter = ScientificFigureInserter()
        record = inserter.insert_scientific_figure(
            doc=doc,
            ref_paragraph=None,
            fig_spec=fig_spec,
            seq_num=1,
            chapter_num=1,
            width_inches=5.5,
            purpose="Quantitative performance comparison",
        )

        doc.save(self.docx_path)

        # Verify insertion
        doc_read = docx.Document(self.docx_path)
        # 1 image paragraph, 1 caption paragraph
        assert len(doc_read.paragraphs) >= 2
        cap_p = doc_read.paragraphs[1]
        assert "Hình 1." in cap_p.text
        assert "SEQ Hình" in cap_p._p.xml
        assert record.bookmark_name == "BK_FIG_1_001"

    def test_04_cross_reference_manager(self):
        """Rule 5: Cross-Reference Manager generates native Word REF fields without plain text hardcoding."""
        doc = docx.Document()
        p = doc.add_paragraph()
        WordCrossReferenceManager.append_cross_reference_to_paragraph(
            paragraph=p,
            prefix_text="Như mô tả trong ",
            bookmark_name="BK_FIG_1_001",
            fallback_target_text="Hình 1.1",
            suffix_text=", kiến trúc biểu diễn kết hợp cả hai góc nhìn.",
            font_size_pt=14.0
        )
        doc.save(self.docx_path)

        doc_read = docx.Document(self.docx_path)
        p_read = doc_read.paragraphs[0]
        assert "Như mô tả trong" in p_read.text
        assert "kiến trúc biểu diễn kết hợp" in p_read.text
        assert "REF BK_FIG_1_001" in p_read._p.xml

    def test_05_skill_registry_registration(self):
        """Rule 10: All 6 required visual skills are registered and executable by name in ResearchSkillRegistry."""
        registry = ResearchSkillRegistry()
        required_skills = [
            "word-diagram-builder",
            "word-table-builder",
            "scientific-figure-inserter",
            "word-caption-manager",
            "word-cross-reference-manager",
            "visual-qa",
        ]
        for skill_name in required_skills:
            skill = registry.get_skill(skill_name)
            assert skill is not None, f"Skill '{skill_name}' was not found in ResearchSkillRegistry."

    def test_06_word_com_diagram_and_visual_qa(self):
        """Rule 1, Rule 6 & Rule 9: Word COM Drawing Canvas, Shapes, Connectors, TOC/TOF Update, and PDF Visual QA."""
        # Create a rich Word document with TOC, TOF, Native Diagram, Native Table, Native Cross-references
        doc = docx.Document()

        # TOC & TOF headings
        doc.add_paragraph("MỤC LỤC", style="Normal")
        doc.add_paragraph("DANH MỤC HÌNH VẼ", style="Normal")
        doc.add_paragraph("DANH MỤC BẢNG", style="Normal")

        # Heading 1
        h1 = doc.add_paragraph("CHƯƠNG 1. TỔNG QUAN HỆ THỐNG BIỂU DIỄN", style="Heading 1")

        # Body paragraph with cross references
        p_body = doc.add_paragraph()
        WordCrossReferenceManager.append_cross_reference_to_paragraph(
            paragraph=p_body,
            prefix_text="Kiến trúc tổng thể được biểu diễn trong ",
            bookmark_name="BK_FIG_1_001",
            fallback_target_text="Hình 1.1",
            suffix_text=" và các tham số kỹ thuật được liệt kê trong ",
            font_size_pt=14.0
        )
        WordCrossReferenceManager.append_cross_reference_to_paragraph(
            paragraph=p_body,
            prefix_text="",
            bookmark_name="BK_TBL_1_001",
            fallback_target_text="Bảng 1.1",
            suffix_text=".",
            font_size_pt=14.0
        )

        # Diagram anchor paragraph & caption below
        p_diag_anchor = doc.add_paragraph()
        WordCaptionManager.add_figure_caption(
            doc=doc,
            ref_paragraph=None,
            seq_num=1,
            title_text="Kiến trúc biểu diễn vector z và phân loại",
            chapter_num=1,
            bookmark_name="BK_FIG_1_001"
        )

        # Native Table
        df = pd.DataFrame({
            "Thuộc tính": ["Độ trễ", "Bộ nhớ", "F1-Score"],
            "Giá trị": ["1.2 ms", "128 MB", "98.4%"]
        })
        spec_tbl = TableSpecification(
            table_id="TBL-000002",
            table_type=TableType.COMPUTED_TABLE,
            title="Thông số kỹ thuật",
            caption="Các chỉ số đo lường tài nguyên và hiệu năng",
            columns=list(df.columns),
            rows_data=df.values.tolist(),
            output_sha256="112233445566",
        )
        WordCaptionManager.add_table_caption(
            doc=doc,
            ref_paragraph=None,
            seq_num=1,
            title_text=spec_tbl.caption,
            chapter_num=1,
            bookmark_name="BK_TBL_1_001"
        )
        WordTableBuilder.insert_table(doc=doc, ref_paragraph=None, spec=spec_tbl, font_size_pt=14.0)

        doc.save(self.docx_path)

        # Open in Word COM, add Native Diagram Canvas, update fields, save and export PDF
        import win32com.client as win32
        import pythoncom
        pythoncom.CoInitialize()
        word = win32.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc_com = word.Documents.Open(os.path.abspath(self.docx_path))

        # Add diagram spec
        diag_spec = DiagramSpecification(
            diagram_id="FIG-000001",
            title="Kiến trúc biểu diễn đa góc nhìn",
            caption="Sơ đồ luồng xử lý và biểu diễn vector z",
            canvas_width_pt=450.0,
            canvas_height_pt=120.0,
            nodes=[
                ShapeNodeSpec(
                    shape_id="N1",
                    shape_type="ROUNDED_RECTANGLE",
                    label="Nhật ký thô\n(Telemetry Logs)",
                    left_pt=10,
                    top_pt=30,
                    width_pt=120,
                    height_pt=60,
                    is_bold=True,
                ),
                ShapeNodeSpec(
                    shape_id="N2",
                    shape_type="ROUNDED_RECTANGLE",
                    label="Mã hóa đa góc nhìn\n(Vector z)",
                    left_pt=165,
                    top_pt=30,
                    width_pt=120,
                    height_pt=60,
                    is_bold=True,
                ),
                ShapeNodeSpec(
                    shape_id="N3",
                    shape_type="ROUNDED_RECTANGLE",
                    label="Phân loại hạ nguồn\n(Frozen Probe)",
                    left_pt=320,
                    top_pt=30,
                    width_pt=120,
                    height_pt=60,
                    is_bold=True,
                ),
            ],
            connectors=[
                ConnectorSpec(
                    connector_id="C1",
                    source_shape_id="N1",
                    target_shape_id="N2",
                    connector_type="STRAIGHT",
                    start_connection_site=3,
                    end_connection_site=1,
                    arrow_head=True,
                ),
                ConnectorSpec(
                    connector_id="C2",
                    source_shape_id="N2",
                    target_shape_id="N3",
                    connector_type="STRAIGHT",
                    start_connection_site=3,
                    end_connection_site=1,
                    arrow_head=True,
                ),
            ],
            group_shapes=True,
            chapter_num=1,
        )

        diag_builder = WordDiagramBuilder()
        diag_res = diag_builder.build_diagram_in_docx(
            doc_com=doc_com,
            target_range=doc_com.Paragraphs(4).Range,
            spec=diag_spec,
            caption_seq=1,
            chapter_num=1,
        )
        assert diag_res["success"]

        doc_com.Save()
        doc_com.Close(False)
        del doc_com
        word.Quit()
        del word
        pythoncom.CoUninitialize()

        # Run Visual QA Engine
        qa_engine = VisualQAEngine()
        qa_res = qa_engine.run_full_visual_qa(docx_path=self.docx_path, export_pdf=True)

        assert qa_res["word_shapes_pass"], f"Word Shapes QA failed: {qa_res['issues']}"
        assert qa_res["pdf_visual_qa_pass"], f"PDF Visual QA failed: {qa_res['issues']}"

    def test_07_frozen_chapter_1_hash_check(self):
        """Validates that Chapter 1 code and text block remain strictly frozen with immutable SHA-256 hash."""
        import hashlib
        code_path = Path(r"D:\Research\src\research_agent\composition\build_word_visual_qa.py")
        assert code_path.exists(), f"File {code_path} not found"
        lines = code_path.read_text(encoding="utf-8").splitlines()
        s_idx = next((i for i, l in enumerate(lines) if 'add_h1("TỔNG QUAN VỀ PHƯƠNG PHÁP TRÍCH XUẤT' in l), None)
        e_idx = next((i for i, l in enumerate(lines) if 'add_h1("PHƯƠNG PHÁP BIỂU DIỄN ĐẶC TRƯNG LOG' in l), None)
        assert s_idx is not None and e_idx is not None, "Chapter 1 boundaries not found"
        ch1_block = "\n".join(lines[s_idx:e_idx])
        computed_hash = hashlib.sha256(ch1_block.encode("utf-8")).hexdigest()
        CANONICAL_HASH = "6097fb1f051573adb21ce65a3466b41d50ba9b8bb9a526c96563e41705c81d5e"
        assert computed_hash == CANONICAL_HASH, f"Chapter 1 code mutated! Expected {CANONICAL_HASH}, got {computed_hash}"

