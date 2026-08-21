"""
Helper module to build native Microsoft Word OMML equations and format Word documents.
"""

from xml.sax.saxutils import escape
import docx
from docx.oxml import parse_xml


def omath(text: str):
    """Create an inline OMML element."""
    escaped = escape(text)
    xml_str = (
        f'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        f'  <m:r><m:t>{escaped}</m:t></m:r>'
        f'</m:oMath>'
    )
    return parse_xml(xml_str)


def omath_para(text: str):
    """Create a display OMML paragraph element."""
    escaped = escape(text)
    xml_str = (
        f'<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        f'  <m:oMath>'
        f'    <m:r><m:t>{escaped}</m:t></m:r>'
        f'  </m:oMath>'
        f'</m:oMathPara>'
    )
    return parse_xml(xml_str)


if __name__ == "__main__":
    doc = docx.Document()
    p = doc.add_paragraph("Kiểm tra vector biểu diễn: ")
    p._p.append(omath("z ∈ ℝ^d"))

    p2 = doc.add_paragraph()
    p2._p.append(omath_para("P_t(Y | X) ≠ P_{t+1}(Y | X)"))

    doc.save("D:/Research/test_math_builder.docx")
    print("Test passed cleanly!")
